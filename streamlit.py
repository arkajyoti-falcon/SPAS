
import os
import io
import streamlit as st
from streamlit.components.v1 import html
import pandas as pd
import json
import re
import string
from typing import Any, Dict, List, Optional, Tuple
import base64
import copy
import tempfile
import time
from io import BytesIO
from datetime import date, datetime
from typing import Optional, Tuple
from dataclasses import dataclass
from typing import Dict, List
from pathlib import Path
from collections import Counter, defaultdict
from functools import lru_cache
import importlib.util
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image
from groq import Groq
from dotenv import load_dotenv
from docxcompose.composer import Composer
import pdfplumber
import ezdxf
import openpyxl
import requests
import convertapi
import pythoncom
from docx2pdf import convert as docx2pdf_convert
import sys
from sentence_transformers import SentenceTransformer, util as st_util
from bert_score import score as bert_score
# Feedback rules management removed: functionality deprecated/disabled

load_dotenv()

# Import the standalone system-description module (we will reuse its deterministic
# DXF/costing extraction + GROQ prompt flow to produce the System Description)
from Supportive_Functions import st_sys_desc as sd_sys

# Import BOM generation functions from bom.py for Mechanical equipment
try:
    from Supportive_Functions import bom as bom_module
except ImportError:
    bom_module = None

# Ensure Supportive_Functions modules are importable
sys.path.insert(0, str(Path(__file__).parent / "Supportive_Functions"))
from Supportive_Functions import combine_old as combine_old_mod
from Supportive_Functions.dxf_extractor import create_dxf_summary_for_embedding

# Import agentY functions for process flow generation
from Supportive_Functions.agentY import (
    generate_initial_flow as agentY_generate_initial_flow,
    generate_second_flow_with_chunks,
    iterative_refinement,
    clean_generated_flow ,
    validate_flow_quality,
    evaluate_process_flow as agentY_evaluate_process_flow,
    fix_empty_output_chutes as agentY_fix_empty_output_chutes
)

from Supportive_Functions.proposal_facts import ProposalFacts, get_counts_source_of_truth
from Supportive_Functions.proposal_facts_extractor import extract_proposal_facts
from Supportive_Functions.proposal_context import ProposalContext, build_proposal_context

import torch
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ==================== FEATURE FLAGS ====================
ENABLE_CONTEXT_UNIFICATION = True
ENABLE_CONSISTENCY_AUDIT = True
ENABLE_ALIAS_MAP = True


@lru_cache(maxsize=1)
def load_cubizone_app():
    """Lazy-load the Cubizone builder module for embedding inside this app."""
    module_path = Path(__file__).parent / "Cubizone" / "V1" / "main.py"
    if not module_path.exists():
        raise FileNotFoundError(f"Cubizone module missing at {module_path}")

    spec = importlib.util.spec_from_file_location("cubizone_v1_main", module_path)
    if spec is None or spec.loader is None:
        raise ImportError("Unable to prepare Cubizone module spec")

    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module

def log_context_counts(ctx: ProposalContext):
    """Log normalized counts and sources for auditability."""
    if not ctx:
        return
    def src_line(k):
        m = ctx.get(k)
        return f"{k}={m.value if m.value is not None else 'Not specified'} (source={m.source})"
    total_chutes = sum([(ctx.get(k).value or 0) for k in [
        "gravity_chutes","mini_gravity_chutes","collection_chutes","rejection_chutes","dispersion_chutes","bulk_chutes","direct_bagging_chutes"
    ]])
    logger.info("Normalized counts: " + " | ".join([
        src_line("feedlines"),
        src_line("telescopic_belt_conveyors"),
        f"total_chutes={total_chutes if total_chutes else 'Not specified'}",
        f"pph={ctx.pph if ctx.pph is not None else 'Not specified'}",
    ]))

def validate_context_counts(ctx: ProposalContext) -> List[str]:
    issues = []
    fl = ctx.get("feedlines").value
    if fl in (0, None):
        issues.append("Feedlines/Induct Lines are not specified; numbers will not be invented.")
    total_chutes = ctx.get("total_chutes").value or 0
    if total_chutes == 0:
        issues.append("Total chutes are not specified; using safe phrasing.")
    if ctx.pph in (None, 0):
        issues.append("Throughput (PPH) is not specified; remove or soften throughput claims.")
    return issues


# ==================== SECTION OUTPUT SANITIZER ====================
def sanitize_section_output(text: str) -> str:
    """Remove LLM meta-artifacts and style objects from generated section text.
    
    Removes:
    - Code fences (```...```)
    - Lines starting with: Reasoning, Corrected, Explanation, Notes, ---
    - Model preamble like "Here is the corrected..."
    - Stringified style objects like _ParagraphStyle('List Bullet')
    - If multiple blocks exist, keeps the largest proposal-like block
    """
    if not text:
        return text
    
    # Remove code fences (```markdown, ```text, ``` etc)
    text = re.sub(r'```(?:\w+)?\s*\n?', '', text, flags=re.IGNORECASE)
    
    # Remove stringified Python style objects (e.g., _ParagraphStyle('List Bullet') id: 2101683332640)
    text = re.sub(r"_\w*Style\(['\"][^'\"]*['\"]\)\s*(?:id:\s*\d+)?", '', text, flags=re.IGNORECASE)
    
    # Remove meta-heading lines
    meta_patterns = [
        r'^#+\s*(Reasoning|Corrected|Explanation|Notes).*$',  # ### Reasoning, ## Corrected, etc.
        r'^(Reasoning|Corrected System Description|Corrected Executive Summary|Corrected Cover Letter|Corrected Process Flow|Explanation|Notes)\s*:?.*$',
        r'^Here\s+is\s+the\s+(corrected|revised|updated|fixed).*$',
        r'^Below\s+is\s+the\s+(corrected|revised|updated|fixed).*$',
        r'^The\s+(corrected|revised|updated|fixed)\s+(?:text|version|section).*$',
        r'^---+\s*$',  # Horizontal rules used as separators
    ]
    for pattern in meta_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove empty lines at start/end and collapse multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    
    # If text has multiple distinct blocks separated by meta-dividers, extract largest proposal block
    # Split by common meta-block patterns
    blocks = re.split(r'\n(?:---+|###?\s*(?:Reasoning|Corrected|Output|Result))\s*\n', text, flags=re.IGNORECASE)
    if len(blocks) > 1:
        # Score each block: prefer blocks with high alphabetic ratio, low meta-keyword ratio
        def score_block(block: str) -> float:
            words = re.findall(r'\b[A-Za-z]+\b', block)
            if not words:
                return 0
            alpha_ratio = len([w for w in words if len(w) > 2]) / len(words) if words else 0
            meta_keywords = ['reasoning', 'corrected', 'explanation', 'notes', 'error', 'fixed']
            meta_count = sum(1 for w in words if w.lower() in meta_keywords)
            meta_ratio = meta_count / len(words) if words else 1
            return len(block) * alpha_ratio * (1 - meta_ratio * 5)
        
        best_block = max(blocks, key=score_block)
        text = best_block.strip()
    
    return text


def enforce_counts(text: str, context: ProposalContext) -> str:
    """Deterministically replace count mentions near anchors with context values.
    
    Only replaces when context has confirmed (non-None) values.
    Uses tight regex patterns to avoid false replacements.
    """
    if not text or not context:
        return text
    
    result = text
    
    # Helper to get confirmed value
    def get_confirmed(key: str) -> Optional[int]:
        meta = context.get(key)
        return meta.value if meta.value is not None else None
    
    # Feedlines / Induct lines
    feedlines = get_confirmed("feedlines")
    if feedlines is not None:
        # Pattern: "X automatic induct lines", "X feedlines", "X induct lines", "X feed lines"
        patterns = [
            (r'(\d+)\s*(nos\.?)?\s*(fully\s*)?(automatic\s*)?induct\s*lines?', f'{feedlines} automatic induct lines'),
            (r'(\d+)\s*(nos\.?)?\s*feed\s*lines?', f'{feedlines} feedlines'),
            (r'(\d+)\s*(nos\.?)?\s*feedlines?', f'{feedlines} feedlines'),
            (r'(\d+)\s*(nos\.?)?\s*induction\s*points?', f'{feedlines} induction points'),
        ]
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    
    # Throughput PPH
    pph = context.pph
    if pph is not None:
        # Pattern: "X PPH", "X parcels per hour", "throughput of X"
        patterns = [
            (r'(\d[\d,]*)\s*PPH\b', f'{pph:,} PPH'),
            (r'(\d[\d,]*)\s*parcels?\s*per\s*hour', f'{pph:,} parcels per hour'),
            (r'throughput\s+of\s+(\d[\d,]*)', f'throughput of {pph:,}'),
            (r'(\d[\d,]*)\s*pph\b', f'{pph:,} PPH'),
        ]
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    
    # Chutes by type - only replace if that specific type is confirmed
    chute_types = {
        'direct_bagging_chutes': ['direct bagging chutes?', 'direct-bagging chutes?'],
        'gravity_chutes': ['gravity chutes?'],
        'rejection_chutes': ['rejection chutes?'],
        'collection_chutes': ['collection chutes?'],
        'dispersion_chutes': ['dispersion chutes?'],
        'bulk_chutes': ['bulk chutes?'],
        'mini_gravity_chutes': ['mini gravity chutes?', 'mini-gravity chutes?'],
    }
    
    for key, type_patterns in chute_types.items():
        val = get_confirmed(key)
        if val is not None and val > 0:
            for tp in type_patterns:
                # Pattern: "X direct bagging chutes"
                pattern = rf'(\d+)\s*(nos\.?)?\s*{tp}'
                # Extract the type name without regex metacharacters for replacement
                type_name = tp.replace('?', '').replace('s$', 's')
                replacement = f'{val} {type_name}'
                result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    
    # Total chutes
    total_chutes = sum([(get_confirmed(k) or 0) for k in chute_types.keys()])
    if total_chutes > 0:
        # Pattern: "total of X chutes", "X total chutes", "X output chutes", "total chutes: X"
        patterns = [
            (r'total\s+of\s+(\d+)\s*(nos\.?)?\s*chutes?', f'total of {total_chutes} chutes'),
            (r'(\d+)\s*(nos\.?)?\s*total\s+chutes?', f'{total_chutes} total chutes'),
            (r'(\d+)\s*(nos\.?)?\s*output\s+chutes?', f'{total_chutes} output chutes'),
            (r'total\s+chutes?\s*:?\s*(\d+)', f'total chutes: {total_chutes}'),
        ]
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    
    return result


def process_section_output(text: str, context: Optional[ProposalContext] = None) -> str:
    """Full post-processing pipeline for LLM-generated sections.
    
    Applies: sanitize -> enforce_counts -> sanitize (final cleanup)
    """
    if not text:
        return text
    
    # Step 1: Remove meta artifacts
    text = sanitize_section_output(text)
    
    # Step 2: Enforce counts deterministically (if context available)
    if context:
        text = enforce_counts(text, context)
    
    # Step 3: Final sanitize pass (in case enforcement introduced artifacts)
    text = sanitize_section_output(text)
    
    return text


# ==================== MANUAL COMPONENT EDITOR HELPERS ====================
def find_component_by_id(components: List[Dict], target_id: str) -> Optional[Dict]:
    """Find a component node by its ID recursively."""
    for comp in components:
        if comp.get("id") == target_id:
            return comp
        if "children" in comp:
            result = find_component_by_id(comp["children"], target_id)
            if result:
                return result
    return None


def delete_component_by_id(components: List[Dict], target_id: str) -> bool:
    """Delete a component node by its ID recursively."""
    for idx, comp in enumerate(components):
        if comp.get("id") == target_id:
            components.pop(idx)
            return True
        if "children" in comp:
            if delete_component_by_id(comp["children"], target_id):
                return True
    return False


def extract_components_from_dxf_summary_with_ai(dxf_summary: str, dxf_json: dict) -> List[Dict]:
    """Use Groq AI to intelligently parse and structure DXF components hierarchically.
    
    Maps DXF categories to proposal sections intelligently based on actual content.
    """
    
    # Build the system prompt with understanding of DXF categories
    system_prompt = """You are an expert at parsing DXF sorter system data and organizing it into a proposal structure.

IMPORTANT: Map DXF component categories to proposal sections like this:

**DXF -> Proposal Mapping:**
- VDS_BUFFER (count > 0) -> "VDS Loop" under Infeed System
- AUTO_INDUCT + OPERATOR_STATION -> Induct section with sub-components
- CHUTE with "irregular" -> "Irregular Takeaway" section
- CHUTE without "irregular" -> Output Chutes section
- PTL (count > 0) -> "PTL set up for Palletizing" under Output Chutes
- BAG_SYSTEM (count > 0) -> "Bag Takeaway Conveyor" under Output Chutes
- COLLECTION -> Collection Chutes under Output Chutes

Return a JSON array with intelligent structure based on what's present in the DXF data."""

    # Build smart user prompt with category analysis
    category_summary = dxf_json.get('category_summary', {})
    categorized_components = dxf_json.get('categorized_components', {})
    cbs_type = dxf_json.get('cbs_type', 'Loop CBS')
    has_vds = dxf_json.get('has_vds', False)
    induction_type = dxf_json.get('induction_type', 'Manual')
    
    user_prompt = f"""Analyze this DXF component data and create the proposal structure:

**DXF Categories Present:**
{json.dumps(category_summary, indent=2)}

**Component Details:**
{json.dumps(categorized_components, indent=2)}

**System Configuration:**
- CBS Type: {cbs_type}
- Has VDS: {has_vds}
- Induction Type: {induction_type}

**Mapping Rules:**
1. If VDS_BUFFER exists (count={category_summary.get('VDS_BUFFER', 0)}): Add "VDS Loop" to Infeed System
2. If AUTO_INDUCT exists (count={category_summary.get('AUTO_INDUCT', 0)}): Add to Induct with count
3. If OPERATOR_STATION exists (count={category_summary.get('OPERATOR_STATION', 0)}): Add "Orientation Conveyor" to Induct
4. Check CHUTE components for "irregular" - if found, create "Irregular Takeaway" section
5. If PTL exists (count={category_summary.get('PTL', 0)}): Add "PTL set up for Palletizing" to Output Chutes
6. If BAG_SYSTEM exists (count={category_summary.get('BAG_SYSTEM', 0)}): Add "Bag Takeaway Conveyor" to Output Chutes
7. If COLLECTION exists (count={category_summary.get('COLLECTION', 0)}): Add "Non-sort Collection Chute" to Output Chutes

Return a JSON array with sections that match the original proposal structure. Only include what's actually present in the DXF data.

Example structure:
[
  {{
    "id": "infeed_system",
    "name": "Infeed System",
    "type": "section",
    "present": true,
    "children": [
      {{"id": "infeed_conveyor", "name": "Infeed Conveyor", "count": 1, "type": "component"}},
      {{"id": "vds_loop", "name": "VDS Loop", "count": {category_summary.get('VDS_BUFFER', 0)}, "type": "component"}}
    ]
  }},
  {{
    "id": "induct",
    "name": "Induct",
    "type": "section",
    "present": true,
    "children": [
      {{"id": "orientation_conveyor", "name": "Orientation Conveyor", "count": {category_summary.get('OPERATOR_STATION', 0)}, "type": "component"}},
      {{"id": "weighing_conveyor", "name": "Weighing Conveyor", "count": 1, "type": "component"}},
      {{"id": "buffer_conveyor", "name": "Buffer Conveyor", "count": 1, "type": "component"}},
      {{"id": "intelligent_merge", "name": "Intelligent Merge Conveyor", "count": 1, "type": "component"}}
    ]
  }},
  {{
    "id": "sorter",
    "name": "{cbs_type} Cross Belt Sorter",
    "type": "section",
    "present": true,
    "children": []
  }},
  {{
    "id": "irregular_takeaway",
    "name": "Irregular Takeaway",
    "type": "section",
    "present": true,
    "children": [
      {{"id": "irregular_takeaway_conveyor", "name": "Irregular Takeaway Conveyor", "count": 1, "type": "component"}},
      {{"id": "irregular_collection_chutes", "name": "Irregular Collection Chutes", "count": 6, "type": "component"}}
    ]
  }},
  {{
    "id": "output_chutes",
    "name": "Output Chutes",
    "type": "section",
    "present": true,
    "children": [
      {{"id": "sliding_chutes", "name": "Sliding Chutes", "count": 17, "type": "component"}},
      {{"id": "rejection_chute", "name": "Rejection Chute", "count": 1, "type": "component"}},
      {{"id": "bagging_ptl", "name": "Bagging Type PTL", "count": {category_summary.get('PTL', 0)}, "type": "component"}},
      {{"id": "bag_takeaway_conveyor", "name": "Bag Takeaway Conveyor", "count": {category_summary.get('BAG_SYSTEM', 0)}, "type": "component"}},
      {{"id": "collection_chute", "name": "Non-sort Collection Chute", "count": {category_summary.get('COLLECTION', 0)}, "type": "component"}},
      {{"id": "ptl_palletizing", "name": "PTL set up for Palletizing", "count": {category_summary.get('PTL', 0)}, "type": "component"}}
    ]
  }}
]

Return ONLY valid JSON, no explanations."""

    # DEBUG: Print DXF info being sent to AI
    print("\n" + "="*80)
    print("DXF INFO BEING SENT TO AI FOR COMPONENT CONFIGURATION")
    print("="*80)
    print(f"DXF Summary:\n{dxf_summary}")
    print(f"\nDXF JSON Data:")
    print(json.dumps(dxf_json, indent=2))
    print(f"\nAI Prompt:\n{user_prompt}")
    print("="*80 + "\n")

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            max_tokens=4000
        )
        
        result_text = response.choices[0].message.content.strip()
        
        # Clean up the response - extract JSON if wrapped in markdown
        if "```json" in result_text:
            result_text = result_text.split("```json")[1].split("```")[0].strip()
        elif "```" in result_text:
            result_text = result_text.split("```")[1].split("```")[0].strip()
        
        components = json.loads(result_text)
        return components
        
    except Exception as e:
        st.warning(f"AI parsing failed, using fallback extraction: {str(e)}")
        return extract_components_from_dxf_summary_fallback(dxf_summary, dxf_json)


def extract_components_from_dxf_summary_fallback(dxf_summary: str, dxf_json: dict) -> List[Dict]:
    """Fallback extraction if AI parsing fails.

    Builds a proposal-like component structure deterministically from DXF JSON,
    matching the original sections you expect.
    """
    components: List[Dict] = []
    category_summary = dxf_json.get('category_summary', {}) if dxf_json else {}
    categorized_components = dxf_json.get('categorized_components', {}) if dxf_json else {}

    # 1. Infeed System
    infeed_children: List[Dict] = []
    # Always include Infeed Conveyor (at least 1)
    infeed_children.append({
        "id": "infeed_conveyor",
        "name": "Infeed Conveyor",
        "count": max(1, int(category_summary.get('CONVEYOR_INFEED', 1))),
        "present": True,
        "type": "component",
    })
    # VDS Loop if present
    vds_count = int(category_summary.get('VDS_BUFFER', 0) or 0)
    if vds_count > 0:
        infeed_children.append({
            "id": "vds_loop",
            "name": "VDS Loop",
            "count": vds_count,
            "present": True,
            "type": "component",
        })
    components.append({
        "id": "infeed_system",
        "name": "Infeed System",
        "count": 0,
        "present": True,
        "type": "section",
        "children": infeed_children,
    })

    # 2. Induct
    induct_children: List[Dict] = []
    # Orientation Conveyor from operator stations
    op_count = int(category_summary.get('OPERATOR_STATION', 0) or 0)
    if op_count > 0:
        induct_children.append({
            "id": "orientation_conveyor",
            "name": "Orientation Conveyor",
            "count": op_count,
            "present": True,
            "type": "component",
        })
    # Weighing / Buffer / Intelligent Merge (defaults to 1 if used)
    induct_children.extend([
        {"id": "weighing_conveyor", "name": "Weighing Conveyor", "count": 1, "present": True, "type": "component"},
        {"id": "buffer_conveyor", "name": "Buffer Conveyor", "count": 1, "present": True, "type": "component"},
        {"id": "intelligent_merge", "name": "Intelligent Merge Conveyor", "count": 1, "present": True, "type": "component"},
    ])
    components.append({
        "id": "induct",
        "name": "Induct",
        "count": 0,
        "present": True,
        "type": "section",
        "children": induct_children,
    })

    # 3. Sorter
    sorter_name = f"{dxf_json.get('cbs_type', 'Loop CBS')} Cross Belt Sorter" if dxf_json else "Loop Cross Belt Sorter"
    components.append({
        "id": "sorter",
        "name": sorter_name,
        "count": 0,
        "present": True,
        "type": "section",
        "children": [],
    })

    # 4. Irregular Takeaway (derive from CHUTE irregular subtype names)
    irregular_count = 0
    chute_details = (categorized_components.get('CHUTE') or {})
    for key, count in chute_details.items():
        # Sum counts for keys that indicate irregular chutes
        if ('irregular' in key.lower()) or ('irchute' in key.lower()):
            # Value may be dict or int depending on schema; handle both
            try:
                irregular_count += int(count if isinstance(count, int) else sum(count.values()))
            except Exception:
                irregular_count += int(count if isinstance(count, int) else 0)
    if irregular_count > 0:
        components.append({
            "id": "irregular_takeaway",
            "name": "Irregular Takeaway",
            "count": 0,
            "present": True,
            "type": "section",
            "children": [
                {"id": "irregular_takeaway_conveyor", "name": "Irregular Takeaway Conveyor", "count": 1, "present": True, "type": "component"},
                {"id": "irregular_collection_chutes", "name": "Irregular Collection Chutes", "count": irregular_count, "present": True, "type": "component"},
            ],
        })

    # 5. Output Chutes
    total_chutes = int(category_summary.get('CHUTE', 0) or 0)
    sliding_count = max(0, total_chutes - irregular_count)
    output_children: List[Dict] = []
    if sliding_count > 0:
        output_children.append({
            "id": "sliding_chutes",
            "name": "Sliding Chutes",
            "count": sliding_count,
            "present": True,
            "type": "component",
        })
    # Rejection chute (assume at least 1 when chutes exist)
    if total_chutes > 0:
        output_children.append({
            "id": "rejection_chute",
            "name": "Rejection Chute",
            "count": 1,
            "present": True,
            "type": "component",
        })
    # Bagging PTL and Bag Takeaway
    ptl_count = int(category_summary.get('PTL', 0) or 0)
    bag_count = int(category_summary.get('BAG_SYSTEM', 0) or 0)
    if ptl_count > 0:
        output_children.append({
            "id": "bagging_ptl",
            "name": "Bagging Type PTL",
            "count": ptl_count,
            "present": True,
            "type": "component",
        })
        output_children.append({
            "id": "ptl_palletizing",
            "name": "PTL set up for Palletizing",
            "count": ptl_count,
            "present": True,
            "type": "component",
        })
    if bag_count > 0:
        output_children.append({
            "id": "bag_takeaway_conveyor",
            "name": "Bag Takeaway Conveyor",
            "count": bag_count,
            "present": True,
            "type": "component",
        })
    # Collection / Non-sort
    collection_count = int(category_summary.get('COLLECTION', 0) or 0)
    if collection_count > 0:
        output_children.append({
            "id": "collection_chute",
            "name": "Non-sort Collection Chute",
            "count": collection_count,
            "present": True,
            "type": "component",
        })
    if output_children:
        components.append({
            "id": "output_chutes",
            "name": "Output Chutes",
            "count": 0,
            "present": True,
            "type": "section",
            "children": output_children,
        })

    # 6. Recirculation (if present)
    if dxf_json and dxf_json.get('has_recirculation', False):
        components.append({
            "id": "recirculation",
            "name": "Recirculation and Manual Feedline",
            "count": 1,
            "present": True,
            "type": "component",
        })

    # Ensure we always return at least the base sections if extraction is empty
    if not components:
        components = [
            {"id": "infeed_system", "name": "Infeed System", "count": 0, "present": True, "type": "section", "children": []},
            {"id": "induct", "name": "Induct", "count": 0, "present": True, "type": "section", "children": []},
            {"id": "sorter", "name": "Loop Cross Belt Sorter", "count": 0, "present": True, "type": "section", "children": []},
            {"id": "output_chutes", "name": "Output Chutes", "count": 0, "present": True, "type": "section", "children": []},
        ]

    return components


def extract_components_from_dxf_summary(dxf_summary: str) -> List[Dict]:
    """Convert DXF summary text into a structured component tree.
    
    This function is kept for backward compatibility but now uses AI extraction.
    """
    # Try to get the full dxf_json from session state
    dxf_json = st.session_state.get("dxf_json_extracted", {})
    
    if dxf_json:
        return extract_components_from_dxf_summary_with_ai(dxf_summary, dxf_json)
    else:
        # Fallback to simple parsing if dxf_json not available
        return extract_components_from_dxf_summary_fallback(dxf_summary, {})


def detect_sorter_type_from_project_name(project_name: str) -> str:
    """Detect sorter type from project name. Returns 'Linear' or 'Loop'."""
    if not project_name:
        return "Loop"
    project_lower = project_name.lower()
    if "linear" in project_lower:
        return "Linear"
    return "Loop"


def render_editable_component_table(components: List[Dict], project_name: str = "") -> List[Dict]:
    """
    Render a dynamic component editor based on AI-extracted sections.
    
    Features:
    - Dynamically renders sections from AI extraction (not hardcoded 4 sections)
    - Shows all components that AI identified from DXF
    - Allows editing quantities and names
    - Clean, minimal interface
    """
    
    def render_subcomponent_row(comp: Dict, section_id: str, comp_idx: int) -> Dict:
        """Render a clean component row with minimal styling."""
        comp_key = f"{section_id}_comp_{comp_idx}"
        
        col1, col2, col3 = st.columns([3, 0.8, 0.2])
        
        with col1:
            name = st.text_input(
                "Component",
                value=comp.get("name", ""),
                key=f"{comp_key}_name",
                label_visibility="collapsed",
                placeholder="Component name"
            )
            comp["name"] = name
        
        with col2:
            count = st.number_input(
                "Qty",
                min_value=1,
                max_value=999,
                value=max(1, int(comp.get("count", 1))),  # Ensure at least 1
                step=1,
                key=f"{comp_key}_count",
                label_visibility="collapsed"
            )
            comp["count"] = count
            comp["unit"] = "Nos"  # Fixed unit
        
        with col3:
            # Delete button with minimal styling
            if st.button(
                "×",
                key=f"{comp_key}_delete",
                help="Delete",
                use_container_width=False
            ):
                return None  # Signal deletion
        
        return comp
    
    def render_section_editor(section_id: str, section_name: str, section_data: Dict) -> None:
        """Render a dynamic section with its components."""
        section_components = section_data.get("children", [])
        
        with st.expander(f"▾ {section_name}", expanded=False):
            st.markdown("<div style='padding: 4px 0;'></div>", unsafe_allow_html=True)
            
            # Check if this section has children to render
            if section_components or section_data.get("type") == "section":
                st.markdown("<div style='margin: 8px 0;'></div>", unsafe_allow_html=True)
                
                # Components header
                col1, col2, col3 = st.columns([3, 0.8, 0.15])
                with col1:
                    st.markdown("<span style='font-weight: 600; font-size: 11px; color: #666;'>Component</span>", unsafe_allow_html=True)
                with col2:
                    st.markdown("<span style='font-weight: 600; font-size: 11px; color: #666;'>Qty</span>", unsafe_allow_html=True)
                with col3:
                    st.markdown("<span></span>", unsafe_allow_html=True)
                
                st.markdown("<div style='border-top: 1px solid #e0e0e0; margin: 4px 0;'></div>", unsafe_allow_html=True)
                
                # Render existing components
                updated_comps = []
                for idx, comp in enumerate(section_components):
                    rendered = render_subcomponent_row(comp, section_id, idx)
                    if rendered is not None:
                        updated_comps.append(rendered)
                
                # Update section with new list
                section_data["children"] = updated_comps
                
                st.markdown("<div style='margin: 8px 0;'></div>", unsafe_allow_html=True)
    
    # Process AI-extracted components dynamically
    # Render sections in order
    for section in components:
        if section.get("type") == "section":
            section_id = section.get("id")
            section_name = section.get("name")
            render_section_editor(section_id, section_name, section)
    
    return components


def convert_components_to_dxf_context(components: List[Dict]) -> Dict:
    """
    Convert the new component structure to a context that can be used by 
    process flow and system description generation.
    
    This function extracts the configured components and creates a structured
    data that represents the system configuration for proposal generation.
    """
    context = {
        "configured_components": {},
        "system_configuration": {
            "has_vds": "Unknown",
            "induction_type": "Unknown",
            "sorter_type": "Unknown"
        },
        "sections_included": []
    }
    
    # Process each main section
    for comp in components:
        section_id = comp.get("id")
        section_name = comp.get("name")
        is_present = comp.get("present", True)
        children = comp.get("children", [])
        
        if is_present:
            context["sections_included"].append(section_name)
            
            # Extract sub-components for each section
            section_comps = []
            for child in children:
                if child.get("present", True):
                    section_comps.append({
                        "name": child.get("name", ""),
                        "count": child.get("count", 1),
                        "unit": child.get("unit", "Nos"),
                        "id": child.get("id")
                    })
            
            context["configured_components"][section_id] = {
                "section_name": section_name,
                "components": section_comps,
                "component_count": len(section_comps)
            }
    
    # Build a summary string for the prompt
    context["components_summary"] = _build_components_summary(context["configured_components"])
    
    return context


def _build_components_summary(components_dict: Dict) -> str:
    """Build a text summary of configured components for use in prompts."""
    summary_lines = []
    
    section_order = {
        "infeed_system": "1. Infeed System",
        "induct": "2. Induct",
        "sorter": "3. Sorter",
        "output_chutes": "4. Output Chutes"
    }
    
    for section_id in ["infeed_system", "induct", "sorter", "output_chutes"]:
        if section_id in components_dict:
            section_data = components_dict[section_id]
            section_name = section_order.get(section_id, section_data["section_name"])
            
            summary_lines.append(f"{section_name}")
            for comp in section_data["components"]:
                count_str = f"{comp['count']} {comp['unit']}" if comp['count'] > 0 else ""
                if count_str:
                    summary_lines.append(f"  - {comp['name']}: {count_str}")
                else:
                    summary_lines.append(f"  - {comp['name']}")
            summary_lines.append("")
    
    return "\n".join(summary_lines)


@st.dialog("Component Configuration", width="large")
def show_component_editing_dialog(dxf_file) -> None:
    """Show a modal dialog for editing extracted DXF components."""
    
    # Clean header
    st.markdown(
        "<div style='margin-bottom: 12px;'>"
        "<h3 style='margin: 0 0 4px 0; color: #1a1a1a; font-size: 16px; font-weight: 600;'>Configure System Components</h3>"
        "<p style='margin: 0; color: #777; font-size: 12px;'>Select components for each section. Sorter type auto-detected from project name.</p>"
        "</div>",
        unsafe_allow_html=True
    )
    
    st.markdown("<div style='height: 1px; background-color: #e0e0e0; margin-bottom: 12px;'></div>", unsafe_allow_html=True)
    
    # Initialize components - always refresh from latest DXF data
    if "dialog_components" not in st.session_state or st.session_state.get("force_refresh_components", False):
        dxf_summary = st.session_state.get("dxf_components_extracted", "")
        dxf_json = st.session_state.get("dxf_json_extracted", {})
        
        if dxf_json:
            st.session_state.dialog_components = extract_components_from_dxf_summary_with_ai(dxf_summary, dxf_json)
        else:
            st.session_state.dialog_components = extract_components_from_dxf_summary(dxf_summary)
        
        st.session_state.force_refresh_components = False
    
    # Get project name from session state for sorter type detection
    project_name = st.session_state.get("project_name", "")
    
    # Render enhanced component table with expanders
    st.markdown(
        "<div style='padding: 12px 0;'>"
        "<h4 style='margin: 0 0 8px 0; color: #1a1a1a; font-size: 14px;'>System Configuration</h4>"
        "<p style='margin: 0; color: #666; font-size: 12px;'>Configure system components for your proposal</p>"
        "</div>",
        unsafe_allow_html=True
    )
    
    st.session_state.dialog_components = render_editable_component_table(
        st.session_state.dialog_components,
        project_name=project_name
    )
    
    st.markdown("<div style='margin: 20px 0;'></div>", unsafe_allow_html=True)
    
    # Action buttons
    st.markdown(
        "<div style='border-top: 1px solid #e0e0e0; padding-top: 16px;'></div>",
        unsafe_allow_html=True
    )
    
    col_btn1, col_btn2, col_spacer = st.columns([1, 1, 2])
    
    with col_btn1:
        if st.button(
            "Confirm",
            key="confirm_components",
            use_container_width=True,
            type="primary"
        ):
            # Save confirmed components to session state
            st.session_state.manual_components_confirmed = st.session_state.dialog_components
            
            # Convert to context format for use in process flow and system description
            components_context = convert_components_to_dxf_context(st.session_state.dialog_components)
            st.session_state.manual_components_context = components_context
            
            # Store the component configuration for propagation to other sections
            st.session_state.configured_components_data = {
                "components": st.session_state.dialog_components,
                "context": components_context,
                "timestamp": datetime.now().isoformat()
            }
            
            st.session_state.show_component_dialog = False
            st.session_state.use_manual_components = True
            st.success("Configuration saved.")
            st.rerun()
    
    with col_btn2:
        if st.button(
            "Cancel",
            key="cancel_components",
            use_container_width=True
        ):
            st.session_state.show_component_dialog = False
            st.session_state.use_manual_components = False
            st.rerun()


# ==================== SENTENCE MODEL CACHE ====================
@st.cache_resource
def load_sentence_model():
    """Load SentenceTransformer model for semantic similarity scoring."""
    return SentenceTransformer('all-MiniLM-L6-v2')

sentence_model = load_sentence_model()

# ==================== DOCX PREVIEW HELPER FUNCTIONS ====================

def extract_docx_content_with_images(doc_bytes: BytesIO) -> list:
    """
    Extract all content from DOCX including text, images, and tables.
    Returns a list of content items that can be split into pages.
    """
    try:
        doc_bytes.seek(0)
        doc = Document(doc_bytes)
        content_items = []
        image_data = {}
        
        # First, extract all images from relationships
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    # Convert to base64
                    import base64
                    img_base64 = base64.b64encode(image_bytes).decode('utf-8')
                    # Determine image type
                    content_type = rel.target_part.content_type
                    if 'png' in content_type:
                        img_type = 'png'
                    elif 'jpeg' in content_type or 'jpg' in content_type:
                        img_type = 'jpeg'
                    else:
                        img_type = 'png'
                    image_data[rel_id] = f"data:image/{img_type};base64,{img_base64}"
                except:
                    pass
        
        # Process paragraphs and tables in document order
        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else ""
            
            # Check for inline images in this paragraph
            # Look for drawing elements which contain images
            drawing_elements = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if not drawing_elements:
                # Also check for pictures (older format)
                drawing_elements = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
            
            # Check for blip elements (actual image references)
            blip_elements = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            
            for blip in blip_elements:
                # Get the embed relationship ID
                embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                embed_id = blip.get(embed_attr)
                if embed_id and embed_id in image_data:
                    content_items.append({
                        'type': 'image',
                        'data': image_data[embed_id]
                    })
            
            if text:
                import html
                text = html.escape(text)
                content_items.append({
                    'type': 'paragraph',
                    'text': text,
                    'style': style_name
                })
        
        # Process tables
        for table in doc.tables:
            table_html = '<table>'
            for i, row in enumerate(table.rows):
                table_html += '<tr>'
                for cell in row.cells:
                    import html
                    cell_text = html.escape(cell.text.strip())
                    if i == 0:
                        table_html += f'<th>{cell_text}</th>'
                    else:
                        table_html += f'<td>{cell_text}</td>'
                table_html += '</tr>'
            table_html += '</table>'
            content_items.append({
                'type': 'table',
                'html': table_html
            })
        
        return content_items
    except Exception as e:
        return [{'type': 'error', 'text': f'Error extracting document: {str(e)}'}]


def split_content_into_pages(content_items: list, items_per_page: int = 15) -> list:
    """
    Split content items into pages.
    Returns a list of pages, where each page is a list of content items.
    """
    pages = []
    current_page = []
    item_count = 0
    
    for item in content_items:
        # Images and tables take more space
        if item['type'] == 'image':
            weight = 5
        elif item['type'] == 'table':
            weight = 3
        else:
            weight = 1
        
        # Check if adding this item would exceed page limit
        if item_count + weight > items_per_page and current_page:
            pages.append(current_page)
            current_page = []
            item_count = 0
        
        current_page.append(item)
        item_count += weight
    
    # Add remaining content
    if current_page:
        pages.append(current_page)
    
    return pages if pages else [[]]


def render_page_content(content_items: list) -> str:
    """Render a list of content items as HTML."""
    html_parts = []
    
    for item in content_items:
        if item['type'] == 'paragraph':
            style = item.get('style', '').lower()
            text = item['text']
            
            if 'heading 1' in style or style == 'Heading 1':
                html_parts.append(f'<h2>{text}</h2>')
            elif 'heading 2' in style or style == 'Heading 2':
                html_parts.append(f'<h3>{text}</h3>')
            elif 'heading 3' in style or style == 'Heading 3':
                html_parts.append(f'<h4>{text}</h4>')
            elif 'list' in style or text.startswith(('•', '-', '–', '*')):
                clean_text = text.lstrip('•-–* ')
                html_parts.append(f'<li>{clean_text}</li>')
            else:
                html_parts.append(f'<p>{text}</p>')
        
        elif item['type'] == 'image':
            html_parts.append(f'<div class="image-container"><img src="{item["data"]}" alt="Document Image"></div>')
        
        elif item['type'] == 'table':
            html_parts.append(item['html'])
        
        elif item['type'] == 'error':
            html_parts.append(f'<p class="error">{item["text"]}</p>')
    
    return '\n'.join(html_parts)


def get_docx_page_count(doc_bytes: BytesIO) -> int:
    """Estimate page count from DOCX (approximate)."""
    try:
        doc_bytes.seek(0)
        doc = Document(doc_bytes)
        # Rough estimate: ~40 paragraphs per page
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        # Approx 3000 characters per page
        return max(1, total_chars // 3000)
    except:
        return 0


def get_docx_stats(doc_bytes: BytesIO) -> Dict[str, Any]:
    """Get statistics about the DOCX document."""
    try:
        doc_bytes.seek(0)
        doc = Document(doc_bytes)
        
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        para_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        image_count = 0
        
        # Count images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        
        return {
            "words": word_count,
            "paragraphs": para_count,
            "tables": table_count,
            "images": image_count,
            "pages_est": max(1, word_count // 400)  # ~400 words per page
        }
    except:
        return {"words": 0, "paragraphs": 0, "tables": 0, "images": 0, "pages_est": 0}


def render_docx_preview(doc_bytes: BytesIO, filename: str):
    """Render actual DOCX preview by converting to PDF and displaying in iframe."""
    import tempfile
    import base64
    import streamlit.components.v1 as components
    
    try:
        # Save DOCX to temporary file
        doc_bytes.seek(0)
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
            tmp_docx.write(doc_bytes.read())
            tmp_docx_path = tmp_docx.name
        
        # Convert DOCX to PDF using ConvertAPI
        convertapi.api_secret = os.getenv('CONVERTAPI_SECRET', 'secret_aTLbINeIh6U9hVsh')
        
        # Convert to PDF
        result = convertapi.convert('pdf', {'File': tmp_docx_path})
        
        # Save PDF to temporary file
        pdf_path = tmp_docx_path.replace('.docx', '.pdf')
        result.save_files(pdf_path)
        
        # Read PDF and convert to base64
        with open(pdf_path, 'rb') as pdf_file:
            pdf_bytes = pdf_file.read()
        
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        
        # Clean up temp files
        try:
            os.remove(tmp_docx_path)
            os.remove(pdf_path)
        except:
            pass
        
        # Display PDF in iframe with embedded viewer
        pdf_display_html = f'''
        <div style="background: #f5f5f5; border: 1px solid #ddd; border-radius: 4px 4px 0 0; padding: 8px 16px; display: flex; justify-content: space-between; align-items: center;">
            <span style="font-weight: 600; color: #333; font-size: 13px;">{filename}</span>
            <span style="color: #666; font-size: 11px; background: #fff; padding: 2px 8px; border-radius: 2px; border: 1px solid #ccc;">Document Preview</span>
        </div>
        <iframe 
            src="data:application/pdf;base64,{pdf_base64}" 
            width="100%" 
            height="700" 
            style="border: 1px solid #ddd; border-top: none;"
            type="application/pdf">
        </iframe>
        '''
        
        components.html(pdf_display_html, height=750, scrolling=False)
        
    except Exception as e:
        # Fallback to text-based preview if PDF conversion fails
        st.warning(f"Could not convert to PDF for preview: {str(e)}. Showing text-based preview instead.")
        
        # Use fallback HTML preview
        doc_bytes.seek(0)
        content_items = extract_docx_content_with_images(doc_bytes)
        pages = split_content_into_pages(content_items, items_per_page=12)
        total_pages = len(pages)
        
        pages_html = []
        for i, page_content in enumerate(pages):
            page_html = render_page_content(page_content)
            pages_html.append(page_html)
        
        import json
        pages_json = json.dumps(pages_html)
        
        iframe_html = f'''
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                * {{ margin: 0; padding: 0; box-sizing: border-box; }}
                body {{ background: #525659; font-family: "Segoe UI", Arial, sans-serif; min-height: 100vh; display: flex; flex-direction: column; }}
                .toolbar {{ background: linear-gradient(180deg, #f8f8f8 0%, #e8e8e8 100%); border-bottom: 1px solid #ccc; padding: 8px 20px; display: flex; justify-content: space-between; align-items: center; position: sticky; top: 0; z-index: 100; }}
                .filename {{ font-weight: 600; color: #333; font-size: 14px; }}
                .nav-controls {{ display: flex; align-items: center; gap: 8px; }}
                .nav-btn {{ background: linear-gradient(180deg, #fff 0%, #f0f0f0 100%); border: 1px solid #999; border-radius: 4px; padding: 6px 14px; font-size: 13px; cursor: pointer; color: #333; font-weight: 500; }}
                .nav-btn:hover:not(:disabled) {{ background: linear-gradient(180deg, #f0f0f0 0%, #e0e0e0 100%); }}
                .nav-btn:disabled {{ opacity: 0.5; cursor: not-allowed; }}
                .page-info {{ font-size: 13px; color: #555; min-width: 100px; text-align: center; }}
                .viewer {{ flex: 1; overflow-y: auto; padding: 30px; display: flex; justify-content: center; }}
                .page {{ background: #fff; width: 8.5in; min-height: 11in; padding: 1in; box-shadow: 0 4px 20px rgba(0,0,0,0.4); font-family: "Calibri", Arial, sans-serif; font-size: 11pt; line-height: 1.6; color: #000; }}
                .page h2 {{ color: #060c71; border-bottom: 2px solid #f9d20e; padding-bottom: 8px; margin: 24px 0 12px; font-size: 18pt; }}
                .page h3 {{ color: #2a3bb8; margin: 18px 0 8px; font-size: 14pt; }}
                .page h4 {{ color: #3d4fc9; margin: 14px 0 6px; font-size: 12pt; }}
                .page p {{ margin-bottom: 10px; text-align: justify; }}
                .page table {{ width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 10pt; }}
                .page th {{ background: #060c71; color: white; padding: 8px; border: 1px solid #060c71; text-align: left; }}
                .page td {{ padding: 8px; border: 1px solid #ddd; }}
                .image-container {{ text-align: center; margin: 15px 0; }}
                .image-container img {{ max-width: 100%; max-height: 400px; }}
            </style>
        </head>
        <body>
            <div class="toolbar">
                <span class="filename">{filename} (Text Preview)</span>
                <div class="nav-controls">
                    <button class="nav-btn" id="prevBtn" onclick="prevPage()">Previous</button>
                    <span class="page-info" id="pageInfo">Page 1 of {total_pages}</span>
                    <button class="nav-btn" id="nextBtn" onclick="nextPage()">Next</button>
                </div>
            </div>
            <div class="viewer"><div class="page" id="pageContent">Loading...</div></div>
            <script>
                const pages = {pages_json};
                let currentPage = 0;
                const totalPages = pages.length;
                function updatePage() {{
                    document.getElementById('pageContent').innerHTML = pages[currentPage];
                    document.getElementById('pageInfo').textContent = 'Page ' + (currentPage + 1) + ' of ' + totalPages;
                    document.getElementById('prevBtn').disabled = currentPage === 0;
                    document.getElementById('nextBtn').disabled = currentPage === totalPages - 1;
                    document.querySelector('.viewer').scrollTop = 0;
                }}
                function nextPage() {{ if (currentPage < totalPages - 1) {{ currentPage++; updatePage(); }} }}
                function prevPage() {{ if (currentPage > 0) {{ currentPage--; updatePage(); }} }}
                document.addEventListener('keydown', function(e) {{
                    if (e.key === 'ArrowRight') nextPage();
                    else if (e.key === 'ArrowLeft') prevPage();
                }});
                updatePage();
            </script>
        </body>
        </html>
        '''
        components.html(iframe_html, height=750, scrolling=False)

# ==================== EVALUATION & VALIDATION FUNCTIONS ====================

def clean_generated_flow(flow: str) -> str:
    """Remove meta-commentary and debugging output from generated flow."""
    # Remove everything after "Note:" or "Changes Made:"
    flow = re.split(r'\n(?:Note:|Changes Made:|The revised flow|DXF Constraints:|However,|By making|To further)', flow)[0]
    
    # Remove leading/trailing whitespace
    flow = flow.strip()
    
    # Remove any "=== REVISED FLOW ===" headers
    flow = re.sub(r'=== REVISED FLOW.*?===\s*', '', flow)
    
    return flow


def validate_flow_quality(flow: str) -> Tuple[bool, str]:
    """
    Validate that the generated flow is not gibberish or corrupted.
    Returns: (is_valid, error_message)
    """
    if not flow or len(flow.strip()) < 100:
        return False, "Flow is too short or empty"
    
    # Check for excessive repetition (gibberish detection)
    words = flow.split()
    if len(words) > 20:
        # Count repeated words in sequence
        max_repeat = 1
        current_repeat = 1
        for i in range(1, len(words)):
            if words[i] == words[i-1]:
                current_repeat += 1
                max_repeat = max(max_repeat, current_repeat)
            else:
                current_repeat = 1
        
        if max_repeat > 5:  # Same word repeated more than 5 times
            return False, f"Detected gibberish: word '{words[i]}' repeated {max_repeat} times"
    
    # Check for presence of key sections
    required_sections = ["Process Flow", "Infeed", "CBS", "Output"]
    found_sections = sum(1 for section in required_sections if section.lower() in flow.lower())
    
    if found_sections < 2:
        return False, "Missing key sections in flow"
    
    # Check for excessive non-alphanumeric characters
    alpha_ratio = sum(c.isalnum() or c.isspace() for c in flow) / len(flow)
    if alpha_ratio < 0.85:
        return False, "Too many special characters (possible corruption)"
    
    return True, ""


def split_into_sentences(text: str) -> List[str]:
    """Split text into sentences."""
    sentences = re.split(r'[.!?]+', text)
    return [s.strip() for s in sentences if s.strip()]


def compute_bert_scores(original: str, generated: str) -> Tuple[float, float, float]:
    """Compute BERT precision, recall, F1 scores."""
    if not original.strip() and not generated.strip():
        return 1.0, 1.0, 1.0
    if not original.strip() or not generated.strip():
        return 0.0, 0.0, 0.0
    
    P, R, F1 = bert_score(
        [generated],
        [original],
        lang="en",
        rescale_with_baseline=False,
    )
    return float(P[0]), float(R[0]), float(F1[0])


def compute_structural_coherence(original: str, generated: str) -> float:
    """
    Compute structural coherence score (0-100).
    
    Evaluates:
    - Sentence order similarity (45% weight)
    - Paragraph structure (35% weight)
    - Bullet point usage
    - Transition smoothness (20% weight)
    """
    o_sents = split_into_sentences(original)
    g_sents = split_into_sentences(generated)
    
    if not o_sents or not g_sents:
        return 0.0
    
    # Encode sentences
    o_emb = sentence_model.encode(o_sents, convert_to_tensor=True)
    g_emb = sentence_model.encode(g_sents, convert_to_tensor=True)
    
    # 1. Sentence order similarity
    min_len = min(len(o_sents), len(g_sents))
    if min_len == 0:
        return 0.0
    
    sim_pairs = st_util.cos_sim(o_emb[:min_len], g_emb[:min_len])
    order_sim = float(sim_pairs.diag().mean().item())
    
    # 2. Paragraph structure similarity
    def split_paragraphs(t):
        return [p.strip() for p in re.split(r"\n\s*\n", t) if p.strip()]
    
    o_paras = split_paragraphs(original)
    g_paras = split_paragraphs(generated)
    para_sim = 1.0 - min(1.0, abs(len(o_paras) - len(g_paras)) / max(len(o_paras), 1))
    
    # 3. Bullet point similarity
    def count_bullets(t):
        return sum(
            1 for line in t.split("\n")
            if re.match(r"^\s*[\-\*\•\da-z]+[\.\)]\s+", line.strip())
        )
    
    o_bullets = count_bullets(original)
    g_bullets = count_bullets(generated)
    if max(o_bullets, g_bullets) == 0:
        bullet_sim = 1.0
    else:
        bullet_sim = 1.0 - min(1.0, abs(o_bullets - g_bullets) / max(o_bullets, g_bullets))
    
    structure_sim = para_sim * 0.6 + bullet_sim * 0.4
    
    # 4. Transition smoothness (adjacent sentence similarity)
    if len(g_sents) > 1:
        adj_sims = []
        for i in range(len(g_sents) - 1):
            sims = float(st_util.cos_sim(g_emb[i], g_emb[i + 1]).item())
            adj_sims.append(sims)
        transition_sim = sum(adj_sims) / len(adj_sims)
    else:
        transition_sim = 1.0
    
    # Weighted combination
    coherence = 0.45 * order_sim + 0.35 * structure_sim + 0.20 * transition_sim
    return float(max(0.0, min(1.0, coherence)) * 100.0)


def analyze_style_differences(reference: str, generated: str, dxf_json: dict) -> List[str]:
    """
    Generate actionable style and tone feedback by comparing with reference.
    Focus on what matters: structure, language, flow - NOT numbers.
    """
    feedback = []
    
    ref_lower = reference.lower()
    gen_lower = generated.lower()
    
    # 1. Check section structure
    ref_sections = re.findall(r'^([A-Z][A-Za-z\s]+):\s*[-–]?', reference, re.MULTILINE)
    gen_sections = re.findall(r'^([A-Z][A-Za-z\s]+):\s*[-–]?', generated, re.MULTILINE)
    
    if len(ref_sections) != len(gen_sections):
        feedback.append(f"Section count mismatch: Reference has {len(ref_sections)} sections, yours has {len(gen_sections)}. Match the reference structure.")
    
    # 2. Check for bullet point style (a., b., c. vs - or •)
    ref_has_letters = bool(re.search(r'^\s*[a-z]\.\s+', reference, re.MULTILINE))
    gen_has_letters = bool(re.search(r'^\s*[a-z]\.\s+', generated, re.MULTILINE))
    ref_has_dashes = bool(re.search(r'^\s*[-\•]\s+', generated, re.MULTILINE))
    
    if ref_has_letters and not gen_has_letters:
        feedback.append("Use lettered sub-points (a., b., c.) for output chutes, matching the reference style")
    elif not ref_has_letters and ref_has_dashes:
        feedback.append("Remove dash bullets, use plain sub-points like the reference")
    
    # 3. Check for key phrases from reference
    key_phrases = [
        ("dumped in bulk", "Use 'dumped in bulk' for infeed arrival (from reference)"),
        ("ascend to a higher level", "Use 'ascend to a higher level' for vertical movement (from reference)"),
        ("picks and positions", "Use 'picks and positions' for operator action (from reference)"),
        ("efficiently sorts", "Use 'efficiently sorts' for CBS operation (from reference)"),
        ("discharged into", "Use 'discharged into' for chute output (from reference)"),
        ("utilizing the data", "Use 'utilizing the data provided by [Client]' (from reference)"),
    ]
    
    for phrase, suggestion in key_phrases:
        if phrase in ref_lower and phrase not in gen_lower:
            feedback.append(suggestion)
    
    # 4. Check for unwanted technical terms
    technical_terms = ["VDS_BUFFER", "AUTO_INDUCT", "OPERATOR_STATION", "BAG_SYSTEM"]
    found_terms = [term for term in technical_terms if term in generated]
    if found_terms:
        feedback.append(f"Remove technical category names: {', '.join(found_terms)}. Use natural language instead")
    
    # 5. Check tone - formal vs casual
    casual_indicators = ["the system features", "there are", "is designed to"]
    formal_indicators = ["shipments are", "parcels enter", "the operator picks"]
    
    casual_count = sum(1 for phrase in casual_indicators if phrase in gen_lower)
    formal_count = sum(1 for phrase in formal_indicators if phrase in ref_lower)
    
    if formal_count > 3 and casual_count > 2:
        feedback.append("Match reference tone: use active voice ('shipments are', 'parcels enter') instead of passive constructions")
    
    # 6. Check paragraph breaks
    ref_para_count = len(re.split(r'\n\s*\n', reference))
    gen_para_count = len(re.split(r'\n\s*\n', generated))
    
    if abs(ref_para_count - gen_para_count) > 2:
        feedback.append(f"Paragraph structure: Reference has {ref_para_count} paragraphs, yours has {gen_para_count}. Match the reference pacing")
    
    # 7. Check for client name usage
    client_name = dxf_json.get('client', 'the client')
    if client_name and client_name.lower() in ref_lower and client_name.lower() not in gen_lower:
        feedback.append(f"Include client name '{client_name}' in sorting logic description, like the reference")
    
    # 8. Check section order
    if ref_sections and gen_sections:
        # Compare first 3 sections
        for i in range(min(3, len(ref_sections), len(gen_sections))):
            if ref_sections[i].lower().strip() != gen_sections[i].lower().strip():
                feedback.append(f"Section order: Position {i+1} should be '{ref_sections[i]}' (reference has this order)")
                break
    
    return feedback


def generate_ai_feedback(reference: str, generated: str, dxf_json: dict, current_score: float, target_score: float) -> List[str]:
    """
    Use AI to generate intelligent, context-aware feedback for improvement.
    Focuses on structural coherence and style matching.
    """
    
    # Extract structural information
    ref_sections = re.findall(r'^([A-Z][A-Za-z\s]+):\s*[-–]?', reference, re.MULTILINE)
    gen_sections = re.findall(r'^([A-Z][A-Za-z\s]+):\s*[-–]?', generated, re.MULTILINE)
    
    ref_has_letters = bool(re.search(r'^\s*[a-z]\.\s+', reference, re.MULTILINE))
    gen_has_letters = bool(re.search(r'^\s*[a-z]\.\s+', generated, re.MULTILINE))
    
    ref_para_count = len(re.split(r'\n\s*\n', reference))
    gen_para_count = len(re.split(r'\n\s*\n', generated))
    
    # Get first few sentences from reference as examples
    ref_sentences = [s.strip() for s in re.split(r'[.!?]+', reference) if len(s.strip()) > 20][:8]
    
    system_prompt = f"""You are an expert technical writing evaluator specializing in Cross-Belt Sorter (CBS) process flow documentation.

## YOUR TASK

Analyze the generated flow against the reference and provide 3-5 actionable feedback items that will improve the **Structural Coherence score**.

**Current Structural Coherence:** {current_score:.1f}/100
**Target Score:** {target_score}/100
**Gap:** {target_score - current_score:.1f} points

## WHAT IS STRUCTURAL COHERENCE?

Structural Coherence (0-100) measures:
1. **Section Order Match** (45% weight) - Are sections in the same order?
2. **Formatting Match** (35% weight) - Bullets, paragraphs, numbering
3. **Flow Smoothness** (20% weight) - Transitions and sentence rhythm

## STRUCTURAL ANALYSIS

**Reference has {len(ref_sections)} sections:** {', '.join(ref_sections[:5])}
**Generated has {len(gen_sections)} sections:** {', '.join(gen_sections[:5])}

**Reference formatting:**
- Uses lettered sub-points (a., b., c.): {'YES' if ref_has_letters else 'NO'}
- Paragraph count: {ref_para_count}

**Generated formatting:**
- Uses lettered sub-points (a., b., c.): {'YES' if gen_has_letters else 'NO'}
- Paragraph count: {gen_para_count}

## FEEDBACK GUIDELINES

Generate feedback that:
1. **Prioritizes structural issues** - section order, formatting, structure (these have highest impact on score)
2. **Is specific and actionable** - "Change X to Y" not "Improve X"
3. **References the reference document** - "Reference uses..." or "Match reference by..."
4. **Focuses on score improvement** - explain HOW the change will improve coherence
5. **Limits to 3-5 items** - most critical issues only
6. **No generic advice** - every item must be specific to this comparison

## FEEDBACK PRIORITY ORDER

**Priority 1: Structural Issues** (Fix these first - highest score impact)
- Section order mismatch
- Missing or extra sections
- Section title differences

**Priority 2: Formatting Issues** (Medium score impact)
- Bullet point style (a., b., c. vs • vs -)
- Paragraph structure mismatch
- Sub-point formatting

**Priority 3: Flow Issues** (Lower score impact, but important)
- Missing transition phrases
- Sentence structure differences
- Key phrase omissions

## OUTPUT FORMAT

Return feedback as a JSON array of strings. Each item should be:
- One sentence or two max
- Specific and actionable
- Focused on structural/formatting changes

Example format:
[
  "Section order mismatch: Move 'Inducts' section before 'Auto Induct Line' to match reference structure (positions 2 and 3 are swapped)",
  "Use lettered sub-points (a., b., c.) for Output Chutes section instead of bullet points (•) to match reference formatting",
  "Add 'Bag Takeaway Conveyor' section at the end - reference has this as final section but yours is missing it"
]

**CRITICAL:** Return ONLY the JSON array, no other text."""

    user_prompt = f"""Compare these two flows and generate 3-5 actionable feedback items to improve structural coherence.

## REFERENCE FLOW (TARGET STYLE)
```
{reference[:1200]}
```

Reference example sentences (for style):
{chr(10).join(f'• {s}' for s in ref_sentences[:5])}

---

## GENERATED FLOW (TO BE IMPROVED)
```
{generated[:1200]}
```

---

## DXF CONTEXT
Client: {dxf_json.get('client', 'Unknown')}
CBS Type: {dxf_json.get('cbs_type', 'Unknown')}
Total Components: {dxf_json.get('total_components', 0)}

---

Generate 3-5 specific, actionable feedback items that will increase Structural Coherence from {current_score:.1f} to {target_score}.

Focus on:
1. Section order/structure differences
2. Formatting mismatches (bullets, paragraphs)
3. Missing/extra sections

Return ONLY a JSON array of feedback strings."""

    try:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        result = combine_old_mod.call_groq(messages, temp=0.3, max_tok=800)
        
        # Try to parse as JSON
        # Clean up the response - remove markdown code blocks if present
        result_clean = result.strip()
        result_clean = re.sub(r'^```json\s*', '', result_clean)
        result_clean = re.sub(r'^```\s*', '', result_clean)
        result_clean = re.sub(r'\s*```$', '', result_clean)
        result_clean = result_clean.strip()
        
        feedback_list = json.loads(result_clean)
        
        if isinstance(feedback_list, list) and len(feedback_list) > 0:
            # Limit to 5 items max
            return feedback_list[:5]
        else:
            logger.warning("AI feedback did not return a valid list")
            return ["Continue refining the flow structure to match reference formatting"]
            
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse AI feedback as JSON: {e}")
        logger.error(f"Raw response: {result}")
        # Fallback: try to extract lines that look like feedback
        lines = [line.strip() for line in result.split('\n') if line.strip() and len(line.strip()) > 20]
        if lines:
            return lines[:5]
        return ["AI feedback generation failed - continue with structural improvements"]
        
    except Exception as e:
        logger.error(f"Error generating AI feedback: {e}")
        return [f"Error generating feedback: {str(e)}"]


def evaluate_process_flow(generated: str, reference: str, dxf_json: dict, 
                          target_score: float = 85) -> Dict:
    """
    Comprehensive evaluation of generated process flow.
    Now uses AI-generated feedback instead of static rules.
    
    Returns:
        dict with scores and feedback
    """
    # Clean generated flow first
    generated_clean = clean_generated_flow(generated)
    
    evaluation = {
        "structural_coherence": 0.0,
        "bert_precision": 0.0,
        "bert_recall": 0.0,
        "bert_f1": 0.0,
        "component_coverage": 0.0,
        "style_match": 0.0,
        "feedback": [],
    }
    
    # 1. Structural coherence
    evaluation["structural_coherence"] = compute_structural_coherence(reference, generated_clean)
    
    # 2. BERT scores
    P, R, F1 = compute_bert_scores(reference, generated_clean)
    evaluation["bert_precision"] = P * 100
    evaluation["bert_recall"] = R * 100
    evaluation["bert_f1"] = F1 * 100
    
    # 3. Component coverage check (keep this as is - it's data validation)
    dxf_cats = dxf_json.get("category_summary", {})
    generated_lower = generated_clean.lower()
    
    coverage_keywords = {
        "AUTO_INDUCT": ["feedline", "automatic", "auto induct", "automatically induct"],
        "OPERATOR_STATION": ["operator", "manual", "positions", "manually position"],
        "VDS_BUFFER": ["vds", "buffer", "distribution"],
        "CHUTE": ["chute", "output", "discharged"],
        "PTL": ["ptl", "put to light", "put-to-light"],
        "BAG_SYSTEM": ["bag", "bagging", "takeaway"],
        "RECIRCULATION": ["recirculation", "recirculate"],
        "CONVEYOR_INFEED": ["infeed", "conveyor", "loaded onto"],
    }
    
    covered = 0
    total_components = 0
    missing_components = []
    
    for cat, count in dxf_cats.items():
        if count > 0 and cat in coverage_keywords:
            total_components += 1
            keywords = coverage_keywords[cat]
            if any(kw in generated_lower for kw in keywords):
                covered += 1
            else:
                # Only flag if it's a major component
                if count > 5 or cat in ["CHUTE", "CBS_SORTER"]:
                    missing_components.append(f"{cat.replace('_', ' ').title()} ({count} units)")
    
    evaluation["component_coverage"] = (covered / total_components * 100) if total_components > 0 else 100
    
    # Add missing components to feedback if any
    if missing_components:
        evaluation["feedback"].append(f"Missing coverage: {', '.join(missing_components[:3])}")
    
    # 4. AI-GENERATED STYLE FEEDBACK (REPLACING STATIC RULES)
    # Only generate feedback if score is below target
    if evaluation["structural_coherence"] < target_score:
        ai_feedback = generate_ai_feedback(reference, generated_clean, dxf_json, evaluation["structural_coherence"], target_score)
        evaluation["feedback"].extend(ai_feedback)
    
    # 5. Style match score (based on feedback count and structural coherence)
    # If structural coherence is high and no feedback, style match should be high
    if evaluation["structural_coherence"] >= target_score and len(evaluation["feedback"]) == 0:
        evaluation["style_match"] = 100.0
    else:
        # Style match is correlation of structural coherence and low feedback count
        feedback_penalty = max(0, len(evaluation["feedback"]) - 1) * 5
        evaluation["style_match"] = max(0, min(100, evaluation["structural_coherence"] - feedback_penalty))
    
    return evaluation

# ==================== GROQ CLIENT SETUP ====================
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
CONVERTAPI_SECRET = os.getenv("CONVERTAPI_SECRET")

if not GROQ_API_KEY:
    st.sidebar.error("⚠️ GROQ_API_KEY not found in .env file!")
else:
    groq_client = Groq(api_key=GROQ_API_KEY)

if CONVERTAPI_SECRET:
    convertapi.api_credentials = CONVERTAPI_SECRET

# ==================== PROPOSAL SECTION MAPPING FOR FEEDBACK ====================
# Only these 5 sections can be edited via feedback
PROPOSAL_SECTIONS = {
    "cover_letter": {
        "name": "Cover Letter",
        "keywords": ["cover letter", "letter", "introduction", "greeting", "dear", "kind attention", "executives", "executive names", "attention"],
        "description": "The formal cover letter introducing the proposal"
    },
    "executive_summary": {
        "name": "Executive Summary",
        "keywords": ["executive summary", "summary", "overview", "highlights", "key points", "executive"],
        "description": "High-level overview of the proposal"
    },
    "proposed_system": {
        "name": "Proposed System Description",
        "keywords": ["proposed system", "solution", "process flow", "system design", "layout", "proposed"],
        "description": "Description of the proposed sorting system and process flow"
    },
    "system_description": {
        "name": "System Description",
        "keywords": ["system description", "technical description", "components", "subsystems", "induct", "merge", "divert", "conveyor"],
        "description": "Detailed technical system description with component details"
    },
    "technical_details": {
        "name": "Proposed System Technical Details",
        "keywords": ["technical details", "bom", "bill of materials", "mechanical", "equipment list", "technical", "specifications"],
        "description": "Technical specifications and Bill of Materials"
    }
}

# ==================== OUTPUT LOGGING SYSTEM ====================
# Backend logging for tracking all proposal generation activities

OUTPUT_FOLDER = Path(__file__).parent / "OUTPUT"
OUTPUT_LOG_FILE = OUTPUT_FOLDER / "proposal_log.xlsx"

def initialize_output_folder():
    """
    Initialize the OUTPUT folder structure with required subfolders.
    Creates the folder structure if it doesn't exist.
    """
    subfolders = [
        "DXF Upload",
        "Costing Sheet",
        "Throughput Calc Sheet",
        "Generated Proposal"
    ]
    
    # Create main OUTPUT folder
    OUTPUT_FOLDER.mkdir(exist_ok=True)
    
    # Create subfolders
    for subfolder in subfolders:
        (OUTPUT_FOLDER / subfolder).mkdir(exist_ok=True)
    
    # Initialize Excel log file if it doesn't exist
    if not OUTPUT_LOG_FILE.exists():
        columns = [
            "Serial No.",
            "Date",
            "Client Name",
            "Client Executives",
            "Project Name",
            "Offer Reference No",
            "Meeting Date",
            "PPH Rate",
            "IPP Rate",
            "Sender Name",
            "Sender Title",
            "Contact Name",
            "Contact Email",
            "Contact Phone",
            "DXF File Name",
            "Costing Sheet Name",
            "Throughput Calc Sheet Name",
            "Layout PNG Name",
            "User Original Feedback",
            "AI Enhanced Feedback",
            "Section Affected",
            "Initial Output",
            "Output Post Feedback"
        ]
        df = pd.DataFrame(columns=columns)
        df.to_excel(OUTPUT_LOG_FILE, index=False, engine='openpyxl')
    
    return True

def save_uploaded_file_to_output(uploaded_file, subfolder: str) -> str:
    """
    Save an uploaded file to the appropriate subfolder in OUTPUT.
    Returns the saved filename or empty string if no file.
    """
    if uploaded_file is None:
        return ""
    
    try:
        # Create timestamped filename to avoid overwrites
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_name = uploaded_file.name
        name_parts = original_name.rsplit('.', 1)
        if len(name_parts) == 2:
            new_filename = f"{name_parts[0]}_{timestamp}.{name_parts[1]}"
        else:
            new_filename = f"{original_name}_{timestamp}"
        
        # Save to subfolder
        save_path = OUTPUT_FOLDER / subfolder / new_filename
        save_path.parent.mkdir(exist_ok=True)
        
        # Write file content
        uploaded_file.seek(0)
        with open(save_path, 'wb') as f:
            f.write(uploaded_file.read())
        uploaded_file.seek(0)  # Reset for further use
        
        return new_filename
    except Exception as e:
        logger.error(f"Error saving uploaded file to output: {e}")
        return ""

def save_generated_proposal(docx_buffer: BytesIO, pdf_buffer: BytesIO, client_name: str, project_name: str) -> Tuple[str, str]:
    """
    Save the generated proposal DOCX and PDF to the Generated Proposal subfolder.
    Returns tuple of (docx_filename, pdf_filename).
    """
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_client = re.sub(r'[^\w\s-]', '', client_name).strip().replace(' ', '_')
        safe_project = re.sub(r'[^\w\s-]', '', project_name).strip().replace(' ', '_')
        
        base_name = f"Proposal_{safe_client}_{safe_project}_{timestamp}"
        
        # Save DOCX
        docx_filename = f"{base_name}.docx"
        docx_path = OUTPUT_FOLDER / "Generated Proposal" / docx_filename
        docx_buffer.seek(0)
        with open(docx_path, 'wb') as f:
            f.write(docx_buffer.read())
        docx_buffer.seek(0)
        
        # Save PDF
        pdf_filename = f"{base_name}.pdf"
        pdf_path = OUTPUT_FOLDER / "Generated Proposal" / pdf_filename
        pdf_buffer.seek(0)
        with open(pdf_path, 'wb') as f:
            f.write(pdf_buffer.read())
        pdf_buffer.seek(0)
        
        return docx_filename, pdf_filename
    except Exception as e:
        logger.error(f"Error saving generated proposal: {e}")
        return "", ""

def log_proposal_generation(
    client_name: str,
    client_executives: str,
    project_name: str,
    offer_reference: str,
    meeting_date: str,
    pph_rate: str,
    ipp_rate: str,
    sender_name: str,
    sender_title: str,
    contact_name: str,
    contact_email: str,
    contact_phone: str,
    dxf_filename: str,
    costing_filename: str,
    throughput_filename: str,
    layout_png_filename: str
) -> int:
    """
    Log the initial proposal generation to the Excel file.
    Returns the serial number assigned to this entry.
    """
    try:
        initialize_output_folder()
        
        # Read existing log
        if OUTPUT_LOG_FILE.exists():
            df = pd.read_excel(OUTPUT_LOG_FILE, engine='openpyxl')
        else:
            df = pd.DataFrame()
        
        # Determine next serial number
        if len(df) > 0 and "Serial No." in df.columns:
            next_serial = int(df["Serial No."].max()) + 1 if pd.notna(df["Serial No."].max()) else 1
        else:
            next_serial = 1
        
        # Create new entry
        new_entry = {
            "Serial No.": next_serial,
            "Date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Client Name": client_name,
            "Client Executives": client_executives,
            "Project Name": project_name,
            "Offer Reference No": offer_reference,
            "Meeting Date": meeting_date,
            "PPH Rate": pph_rate,
            "IPP Rate": ipp_rate,
            "Sender Name": sender_name,
            "Sender Title": sender_title,
            "Contact Name": contact_name,
            "Contact Email": contact_email,
            "Contact Phone": contact_phone,
            "DXF File Name": dxf_filename,
            "Costing Sheet Name": costing_filename,
            "Throughput Calc Sheet Name": throughput_filename,
            "Layout PNG Name": layout_png_filename,
            "User Original Feedback": "",
            "AI Enhanced Feedback": "",
            "Section Affected": "",
            "Initial Output": "",
            "Output Post Feedback": ""
        }
        
        # Append to dataframe
        new_df = pd.DataFrame([new_entry])
        df = pd.concat([df, new_df], ignore_index=True)
        
        # Save back to Excel
        df.to_excel(OUTPUT_LOG_FILE, index=False, engine='openpyxl')
        
        return next_serial
    except Exception as e:
        logger.error(f"Error logging proposal generation: {e}")
        return -1

def log_feedback_entry(
    serial_no: int,
    user_feedback: str,
    ai_enhanced_feedback: str,
    section_affected: str,
    initial_output: str,
    output_post_feedback: str
):
    """
    Log a feedback entry. If serial_no matches an existing entry with empty feedback fields,
    update that entry. Otherwise, create a new entry with the same base info.
    """
    try:
        if not OUTPUT_LOG_FILE.exists():
            logger.warning("Log file does not exist. Cannot log feedback.")
            return
        
        df = pd.read_excel(OUTPUT_LOG_FILE, engine='openpyxl')
        
        # Find the base entry (most recent with this serial number and empty feedback)
        matching_rows = df[df["Serial No."] == serial_no]
        
        if len(matching_rows) > 0:
            # Check if the last matching row has empty feedback - if so, update it
            last_idx = matching_rows.index[-1]
            if pd.isna(df.loc[last_idx, "User Original Feedback"]) or df.loc[last_idx, "User Original Feedback"] == "":
                # Update existing row
                df.loc[last_idx, "User Original Feedback"] = user_feedback
                df.loc[last_idx, "AI Enhanced Feedback"] = ai_enhanced_feedback
                df.loc[last_idx, "Section Affected"] = section_affected
                df.loc[last_idx, "Initial Output"] = initial_output[:5000] if initial_output else ""  # Limit length
                df.loc[last_idx, "Output Post Feedback"] = output_post_feedback[:5000] if output_post_feedback else ""
            else:
                # Create new row with same base info but new feedback
                base_row = df.loc[last_idx].copy()
                base_row["Date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                base_row["User Original Feedback"] = user_feedback
                base_row["AI Enhanced Feedback"] = ai_enhanced_feedback
                base_row["Section Affected"] = section_affected
                base_row["Initial Output"] = initial_output[:5000] if initial_output else ""
                base_row["Output Post Feedback"] = output_post_feedback[:5000] if output_post_feedback else ""
                df = pd.concat([df, pd.DataFrame([base_row])], ignore_index=True)
        
        # Save back to Excel
        df.to_excel(OUTPUT_LOG_FILE, index=False, engine='openpyxl')
        
    except Exception as e:
        logger.error(f"Error logging feedback entry: {e}")

# Initialize output folder on module load
try:
    initialize_output_folder()
except Exception as e:
    logger.warning(f"Could not initialize output folder: {e}")


def rephrase_feedback_to_actionable(feedback: str, section_content: str = None, all_sections_content: dict = None) -> dict:
    """
    Use LLM to rephrase user feedback into clear, actionable instructions.
    This helps standardize various ways users might express the same intent.
    
    Enhanced to handle short/cryptic user inputs by analyzing full section context.
    
    Args:
        feedback: User's feedback string
        section_content: Content of the identified target section
        all_sections_content: Dict of all section contents for comprehensive analysis
    """
    # Build comprehensive context about ALL section contents
    content_context = ""
    
    if section_content:
        content_context = f"""

TARGET SECTION FULL CONTENT:
\"\"\"
{section_content}
\"\"\"
"""
    
    # Add context from all sections if available
    if all_sections_content:
        all_content_preview = "\n\nALL SECTIONS SUMMARY (for context):\n"
        for key, content in all_sections_content.items():
            section_name = key.replace("section_content_", "").replace("_", " ").title()
            preview = content[:500] if content else "[Empty]"
            all_content_preview += f"\n--- {section_name} ---\n{preview}\n"
        content_context += all_content_preview

    system_prompt = f"""You are an expert at understanding and rephrasing user feedback for Cross-Belt Sorter proposal documents.

## YOUR TASK:
1. Understand what the user wants to change (even from very short inputs)
2. Identify the EXACT current value from section content
3. Extract the EXACT new value the user wants
4. Rephrase into clear, actionable instructions

## CRITICAL RULES FOR HANDLING SHORT/CRYPTIC USER INPUTS:

### Rule 1: Component Count Changes
**Pattern:** User says something like "Direct Bagging chute - 42" or "gravity chutes 127" or "feedlines: 8"
**What to do:**
- Look for ANY mention of this component in the section content
- Find the CURRENT count for this component (e.g., "41 direct bagging chutes")
- Set old_value to the CURRENT count (e.g., "41")
- Set new_value to the NEW count from user (e.g., "42")
- Set target to the component name (e.g., "direct_bagging_chutes")
- Create actionable_instruction like "Change direct bagging chutes count from 41 to 42"

### Rule 2: Technical Values
**Pattern:** User says "throughput 50000" or "PPH: 30000" or "capacity 45K"
**What to do:**
- Search for throughput/PPH/capacity mentions in content
- Extract current value (e.g., "27,600 PPH")
- Set old_value to current (e.g., "27,600" or "27600")
- Set new_value to user's value (e.g., "50000")
- Set target to "throughput_pph" or "capacity"

### Rule 3: Reference Numbers/Codes
**Pattern:** User says "offer FA-2025-001" or "reference: ABC-123"
**What to do:**
- Find current offer/reference in content
- Extract exact current value
- Set new value to user's input
- Set target to "offer_reference"

### Rule 4: Names/Entities
**Pattern:** User says "client Amazon" or "executives: John, Jane" or "project XYZ"
**What to do:**
- Search for client names, executive names, project names in content
- Extract current values
- Set new values from user input

## EXAMPLES WITH SHORT INPUTS:

**Example 1: Component Count**
User: "Direct Bagging chute - 42"
Content has: "41 direct bagging chutes for sorted parcels"
→ old_value: "41", new_value: "42", target: "direct_bagging_chutes"
→ actionable_instruction: "Change direct bagging chutes count from 41 to 42"

**Example 2: Throughput**
User: "throughput 50000"
Content has: "system throughput of 27,600 PPH"
→ old_value: "27,600", new_value: "50,000", target: "throughput_pph"
→ actionable_instruction: "Update system throughput from 27,600 PPH to 50,000 PPH"

**Example 3: Gravity Chutes**
User: "gravity chutes 127"
Content has: "There are 86 gravity chutes"
→ old_value: "86", new_value: "127", target: "gravity_chutes"
→ actionable_instruction: "Change gravity chutes count from 86 to 127"

**Example 4: Feedlines**
User: "feedlines: 8"
Content has: "4 automatic induct lines"
→ old_value: "4", new_value: "8", target: "feedlines"
→ actionable_instruction: "Change feedlines count from 4 to 8"
{content_context}

## OUTPUT FORMAT (JSON ONLY):
{{
    "original_feedback": "the exact user input",
    "actionable_instruction": "Clear instruction like 'Change direct bagging chutes from 41 to 42'",
    "action_type": "replace|add|remove|modify|rephrase",
    "entities": {{
        "old_value": "EXACT current value found in content (e.g., '41', '27600', 'Zepto')",
        "new_value": "exact new value from user (e.g., '42', '50000', 'Amazon')",
        "target": "what is being changed (e.g., 'direct_bagging_chutes', 'throughput_pph', 'client_name')"
    }},
    "intent_summary": "One sentence summary of what user wants"
}}

## CRITICAL:
- NEVER use placeholders like [CURRENT_VALUE], [CURRENT_OFFER_REFERENCE]
- ALWAYS find ACTUAL current value from section content
- Handle very short inputs intelligently by analyzing context
- If you cannot find current value, leave old_value as empty string "" but still process new_value
"""

    def api_call():
        return groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"User feedback: {feedback}"},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
            max_tokens=500,
        )
    
    try:
        completion = call_groq_with_retry(api_call)
        result = json.loads(completion.choices[0].message.content)
        return result
    except Exception as e:
        return {
            "original_feedback": feedback,
            "actionable_instruction": feedback,
            "action_type": "modify",
            "entities": {},
            "intent_summary": feedback
        }

def identify_section_with_context(feedback: str, actionable_info: dict, stored_sections: dict) -> dict:
    """
    Use LLM to identify which section the feedback refers to by analyzing:
    1. The rephrased actionable feedback
    2. The actual content of stored sections (to find where changes apply)
    
    Enhanced with comprehensive section content for better identification.
    """
    # Build FULL section previews from stored content (not just 300 chars)
    section_previews = []
    for section_key, section_info in PROPOSAL_SECTIONS.items():
        content_key = f"section_content_{section_key}"
        content = stored_sections.get(content_key, "")
        if content:
            # Use more content for better matching (up to 1500 chars)
            preview = content[:1500].replace('\n', ' ').strip()
            section_previews.append(f"- {section_key} ({section_info['name']}): \"{preview}...\"")
        else:
            section_previews.append(f"- {section_key} ({section_info['name']}): [No content stored]")
    
    sections_context = "\n".join(section_previews)
    
    system_prompt = f"""You are an expert at analyzing user feedback for proposal documents.
Your task is to identify which section of a proposal the user wants to modify.

AVAILABLE SECTIONS (with content previews):
{sections_context}

SECTION DESCRIPTIONS:
- cover_letter: The formal cover letter with greeting, executive names, introduction
- executive_summary: High-level overview with key highlights and project summary
- proposed_system: Description of the proposed sorting system and process flow
- system_description: Detailed technical system description with component details
- technical_details: Technical specifications and Bill of Materials

Analyze the user's intent and the section contents to determine:
1. Which section contains the content they want to change
2. If the change involves specific text/names, check which section has that text

Return ONLY valid JSON:
{{
    "section_key": "cover_letter|executive_summary|proposed_system|system_description|technical_details",
    "section_name": "The Section Name",
    "confidence": "high|medium|low",
    "reasoning": "Explanation of why this section was identified",
    "found_in_content": true/false
}}

If you cannot determine the section, return section_key as "unknown"."""

    user_content = f"""USER FEEDBACK: {feedback}

ACTIONABLE INSTRUCTION: {actionable_info.get('actionable_instruction', feedback)}
ACTION TYPE: {actionable_info.get('action_type', 'modify')}
TARGET: {actionable_info.get('entities', {}).get('target', 'unknown')}
OLD VALUE: {actionable_info.get('entities', {}).get('old_value', 'N/A')}
NEW VALUE: {actionable_info.get('entities', {}).get('new_value', 'N/A')}

Which section should be modified?"""

    def api_call():
        return groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
            max_tokens=500,
        )
    
    try:
        completion = call_groq_with_retry(api_call)
        result = json.loads(completion.choices[0].message.content)
        return result
    except Exception as e:
        return {
            "section_key": "unknown",
            "section_name": "Unknown",
            "confidence": "low",
            "reasoning": f"Error: {str(e)}",
            "found_in_content": False
        }


def regenerate_section_with_feedback(section_key: str, original_content: str, user_feedback: str, actionable_info: dict, context: dict) -> str:
    """
    Regenerate a specific section incorporating the user's feedback.
    Uses actionable_info for precise instructions.
    
    Handles both modifications (change values) and removals (delete entire subsections).
    """
    section_info = PROPOSAL_SECTIONS.get(section_key, {})
    section_name = section_info.get("name", section_key)
    
    # Extract actionable details
    action_type = actionable_info.get("action_type", "modify")
    actionable_instruction = actionable_info.get("actionable_instruction", user_feedback)
    entities = actionable_info.get("entities", {})
    old_value = entities.get("old_value", "")
    new_value = entities.get("new_value", "")
    target = entities.get("target", "")
    
    # Determine if this is a removal action
    is_removal = (
        action_type == "remove" or 
        (new_value and new_value.lower() in ["none", "remove", "delete"]) or
        (not new_value and any(word in actionable_instruction.lower() for word in ["remove", "delete", "eliminate", "no "]))
    )
    
    # Build removal-aware system prompt
    if is_removal and target:
        # For removals, provide specific guidance
        system_prompt = f"""You are an expert proposal writer at Falcon Autotech.
You need to modify the "{section_name}" section by REMOVING a specific component/subsection.

REMOVAL REQUEST:
- Target: {target}
- Instruction: {actionable_instruction}
- User said: {user_feedback}

CRITICAL RULES FOR REMOVAL:
1. FIND the subsection/paragraph that describes {target}
   - Look for: "Manual Induct Station", "manual induct", "operator", etc. (related keywords)
   - May appear in numbered lists (1., 2., 3.) or (a., b., c.) or as standalone paragraphs
2. REMOVE that entire subsection:
   - Delete the heading/title line (e.g., "3. Manual Induct Station:")
   - Delete ALL paragraphs describing it
   - Delete any blank lines immediately after it
3. DO NOT remove any other content
4. Renumber remaining items if in a numbered list (e.g., 1, 2, 3... stays 1, 2 after removing 2)
5. For lettered lists (a., b., c.), keep letters as-is after removal
6. Keep the overall section structure intact and readable
7. Preserve all formatting and style

Return ONLY the modified section with target completely removed. No explanations."""
    else:
        # For modifications
        system_prompt = f"""You are an expert proposal writer at Falcon Autotech.
You need to modify the "{section_name}" section of a proposal based on specific instructions.

MODIFICATION DETAILS:
- Action Type: {action_type}
- Instruction: {actionable_instruction}
{f"- Replace: '{old_value}' → '{new_value}'" if old_value and new_value else ""}

CRITICAL RULES:
1. ONLY make the specific change requested - do NOT modify anything else
2. Keep the exact same format, structure, and LENGTH as the original
3. Preserve ALL formatting: bold, italics, spacing, capitalization, lists, structure
4. When replacing values:
   - Find exact match and replace ONLY that text
   - Do not change text around the value
   - Keep sentence structure identical
   - Keep capitalization and spacing exactly as is
5. Preserve all technical details, dates, and other information not being changed
6. The output should be the complete section, not just the changed part

Context:
- Client: {context.get('client_name', 'N/A')}
- Project: {context.get('project_name', 'N/A')}

Return ONLY the modified section content. No explanations, no markers."""

    user_prompt = f"""ORIGINAL {section_name.upper()}:
{original_content}

REQUESTED CHANGE:
{actionable_instruction}

User feedback: {user_feedback}

Please output the modified {section_name} with ONLY the requested change applied. Keep everything else exactly the same (same formatting, capitalization, structure)."""

    def api_call():
        return groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,  # Lower temperature for more precise edits
            max_tokens=3000,
        )
    
    try:
        completion = call_groq_with_retry(api_call)
        regenerated = completion.choices[0].message.content.strip()
        
        # Clean up any markdown code blocks
        if regenerated.startswith("```"):
            parts = regenerated.split("```")
            if len(parts) >= 2:
                regenerated = parts[1]
                if regenerated.startswith(("text\n", "markdown\n")):
                    regenerated = "\n".join(regenerated.split("\n")[1:])
        
        return regenerated
    except Exception as e:
        return f"Error regenerating section: {str(e)}"


def find_and_replace_section_in_docx(doc, section_name: str, new_content: str) -> bool:
    """
    Find a section by its header name and replace its content.
    Returns True if section was found and updated, False otherwise.
    
    Section identification strategy:
    - Look for paragraphs containing the section name as a header
    - Replace content paragraphs until the next section header
    """
    # Section header patterns (with numbering like "2. Executive Summary")
    section_patterns = {
        "cover_letter": ["Kind Attention", "Dear", "Respected"],
        "executive_summary": ["Executive Summary"],
        "proposed_system": ["Proposed System Description", "Proposed Solution"],
        "system_description": ["System Description"],
        "technical_details": ["Proposed System Technical Details", "Technical Details"]
    }
    
    patterns = section_patterns.get(section_name, [])
    if not patterns:
        return False
    
    # Find the section start
    section_start_idx = None
    section_end_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        
        # Check if this paragraph is the section header
        for pattern in patterns:
            if pattern.lower() in para_text.lower():
                section_start_idx = i
                break
        
        # If we found the start, look for the next major section header
        if section_start_idx is not None and i > section_start_idx:
            # Check for next numbered section (e.g., "3. Company Profile")
            if para_text and len(para_text) > 2:
                # Look for patterns like "N. " or "N " at start
                if para_text[0].isdigit() and (para_text[1] == '.' or para_text[1] == ' '):
                    section_end_idx = i
                    break
                # Also check for page breaks or major style changes
                if para.style and 'Heading' in para.style.name:
                    section_end_idx = i
                    break
    
    if section_start_idx is None:
        return False
    
    # For cover letter, we need special handling as it's before numbered sections
    if section_name == "cover_letter":
        # Replace content in the paragraphs after "Kind Attention"
        content_lines = new_content.split('\n')
        content_idx = 0
        
        for i in range(section_start_idx + 1, min(section_start_idx + 20, len(doc.paragraphs))):
            if content_idx < len(content_lines):
                para = doc.paragraphs[i]
                # Don't overwrite images or tables
                if para.text.strip():
                    para.text = content_lines[content_idx]
                    content_idx += 1
        return True
    
    # For other sections, replace paragraph content
    # This is a simplified approach - full implementation would preserve formatting
    return True


# ==================== CLIENT LOGOS MAPPING ====================
CLIENT_LOGOS = {
    "Zepto": "assests/Images/clients/zepto.png",
    "Flipkart": "assests/Images/clients/flipkart.png",
    "Shiprocket": "assests/Images/clients/shiprocket.png",
    "Amazon": "assests/Images/clients/amazon.png",
    "Delhivery": "assests/Images/clients/delhivery.png",
    "Swiggy": "assests/Images/clients/swiggy.png",
    "Mondial": "assests/Images/clients/mondial.jpg",
    "Zomato": "assests/Images/clients/zomato.png",
}

# ==================== COMPONENT DESCRIPTION IMAGE PATHS ====================
COMPONENT_IMAGES = {
    "cross_belt_sorter": r"assests\Images\CROS_BELT_SORTER.PNG",
    "cbs_carrier": r"assests\Images\CBS_CAREER.PNG",
    "servo_roller": r"assests\Images\SERVO_ROLLER.PNG",
    "chassis": r"assests\Images\CHASIS.PNG",
    "wheel": r"assests\Images\CAREER_WHEEL.PNG",
    "power": r"assests\Images\TRANSMISSION.PNG",
    "linear": r"assests\Images\LINEAR_MOTOR_DRIVE.PNG",
    "friction_wheel": r"assests\Images\FRICTION_WHEEL_DRIVE.PNG",
    "rcoax": r"assests\Images\DATA.PNG",
    "carrier_position": r"assests\Images\CPS.PNG",
}

GLOSSARY_ENTRIES = [
    ("RFQ", "Request For Quotation"),
    ("RFP", "Request For Proposal"),
    ("PPH", "Parcels / Shipments Per Hour"),
    ("ARB", "Actuated Roller Balls"),
    ("DBO", "Damaged Barcode"),
    ("VDS", "Volume Distribution System"),
    ("ICR", "Intelligent Character Recognition"),
    ("MEZZ", "Mezzanine"),
    ("LIM", "Linear Induction Motor"),
    ("LSM", "Linear Synchronous Motor"),
    ("FWD", "Friction Wheel Drive"),
    ("ECDS", "Empty Carrier Detection System"),
    ("AC", "Alternating Current"),
    ("DC", "Direct Current"),
    ("PLC", "Programmable Logic Controller"),
    ("IT", "Information Technology"),
    ("BOQ", "Bill Of Quantity"),
    ("I/O", "Input/ Output"),
    ("PDP", "Power Distribution Panel"),
    ("PC", "Personal Computer"),
    ("UPS", "Uninterrupted Power Supply"),
    ("CBS", "Cross Belt Sorter"),
    ("MDR", "Motor Driven Roller"),
    ("IPP", "Individual Productivity Potential"),
    ("VM", "Virtual Machine"),
    ("MENA", "Middle East North Africa"),
    ("FOC", "Free of Cost"),
    ("CEP", "Courier Express Parcel"),
    ("DAP", "Design Approval Phase"),
    ("DNF", "Data Not Found"),
    ("LGM", "Logic Mismatch"),
    ("RTVC", "Real Time Video Coding"),
    ("NL Shipments", "Non-Large Shipments"),
    ("SL Shipments", "Semi-Large Shipments"),
    ("NO(S)", "Piece(s)"),
]

# ==================== RETRY WRAPPER FOR RATE LIMITS ====================
def call_groq_with_retry(api_call_func, max_retries=5, initial_delay=2):
    """Wrapper to retry GROQ API calls with exponential backoff on rate limits."""
    for attempt in range(max_retries):
        try:
            return api_call_func()
        except Exception as e:
            error_str = str(e)
            # Check if it's a rate limit error
            if "429" in error_str or "rate_limit_exceeded" in error_str.lower():
                if attempt < max_retries - 1:
                    # Extract wait time from error message if available
                    wait_match = re.search(r'try again in ([0-9.]+)s', error_str)
                    if wait_match:
                        wait_time = float(wait_match.group(1)) + 1  # Add 1 second buffer
                    else:
                        wait_time = initial_delay * (2 ** attempt)  # Exponential backoff
                    
                    st.warning(f"Rate limit hit. Waiting {wait_time:.1f}s before retry {attempt + 1}/{max_retries}...")
                    time.sleep(wait_time)
                else:
                    raise  # Re-raise on final attempt
            else:
                raise  # Re-raise non-rate-limit errors immediately
    
    raise RuntimeError(f"Failed after {max_retries} retries")


# ==================== TEXT NORMALIZATION UTILITIES ====================
def normalize_cross_belt_sorter(text: str) -> str:
    """
    Normalize all variations of 'Cross belt sorter' to 'Cross Belt Sorter'.
    Handles case variations like 'cross belt sorter', 'Cross belt Sorter', etc.
    """
    if not text:
        return text
    # Use regex to find all case variations and replace with proper capitalization
    pattern = re.compile(r'cross[\s-]*belt[\s-]*sorter', re.IGNORECASE)
    return pattern.sub('Cross Belt Sorter', text)


# Known client names from reference proposals that might appear in chunks
KNOWN_REFERENCE_CLIENTS = [
    "Amazon", "Noon", "Shadowfax", "Bosta", "Delhivery", "Flipkart",
    "Aramex", "Asendia", "Fastway", "Ekart", "BlueDart", "DTDC",
    "XpressBees", "Ecom Express", "Gati", "Rivigo", "FedEx", "DHL",
    "Ecom", "Xpressbees", "Blue Dart", "Express", "Flipkart India"
]

def replace_client_names_in_flow(text: str, correct_client_name: str) -> str:
    """
    Replace any reference client names with the correct user-provided client name.
    This ensures consistency when process flows are generated using reference chunks
    from previous proposals.
    
    Enhanced with comprehensive pattern matching to catch all client name mentions.
    """
    if not text or not correct_client_name:
        return text
    
    result = text
    correct_client_lower = correct_client_name.lower().strip()
    
    for ref_client in KNOWN_REFERENCE_CLIENTS:
        # Skip if the reference client IS the correct client (case-insensitive)
        if ref_client.lower() == correct_client_lower:
            continue
        
        # Skip if reference client is a substring of correct client or vice versa
        if ref_client.lower() in correct_client_lower or correct_client_lower in ref_client.lower():
            continue
        
        # Pattern 1: "[Client]'s sorting logic" or "[Client] sorting logic"
        pattern1 = re.compile(
            rf"\b{re.escape(ref_client)}(?:'s|'s|s)?\s+sorting\s+logic",
            re.IGNORECASE
        )
        result = pattern1.sub(f"{correct_client_name}'s sorting logic", result)
        
        # Pattern 2: "data provided by [Client]" or "provided by [Client]"
        pattern2 = re.compile(
            rf"provided\s+by\s+{re.escape(ref_client)}(?:'s|'s)?",
            re.IGNORECASE
        )
        result = pattern2.sub(f"provided by {correct_client_name}", result)
        
        # Pattern 3: Direct client name mention in sorting context
        pattern3 = re.compile(
            rf"utilizing.*?{re.escape(ref_client)}(?:'s|'s)?\s+(?:sorting|logic|data)",
            re.IGNORECASE
        )
        result = pattern3.sub(f"utilizing {correct_client_name}'s sorting logic", result)
        
        # Pattern 4: "[Client]'s WCS" or "WCS from [Client]" or "[Client] WCS"
        pattern4 = re.compile(
            rf"\b{re.escape(ref_client)}(?:'s|'s)?\s+(?:WCS|WMS|ERP|system)",
            re.IGNORECASE
        )
        result = pattern4.sub(f"{correct_client_name}'s \\g<0>".replace(ref_client, correct_client_name), result)
        result = re.sub(
            rf"\b{re.escape(ref_client)}(?:'s|'s)?\s+(WCS|WMS|ERP|system)",
            rf"{correct_client_name}'s \1",
            result,
            flags=re.IGNORECASE
        )
        
        # Pattern 5: "data from [Client]" or "information from [Client]"
        pattern5 = re.compile(
            rf"(?:data|information|instructions?|commands?)\s+(?:from|by)\s+{re.escape(ref_client)}(?:'s|'s)?",
            re.IGNORECASE
        )
        result = pattern5.sub(lambda m: m.group(0).replace(ref_client, correct_client_name).replace(ref_client.lower(), correct_client_name), result)
        
        # Pattern 6: "using [Client]'s data" or "using [Client] data"
        pattern6 = re.compile(
            rf"using\s+{re.escape(ref_client)}(?:'s|'s)?\s+(?:data|logic|system)",
            re.IGNORECASE
        )
        result = pattern6.sub(f"using {correct_client_name}'s data", result)
        
        # Pattern 7: Generic "[Client]'s [something]" pattern - be careful, only in sorting/logistics context
        pattern7 = re.compile(
            rf"\b{re.escape(ref_client)}(?:'s|'s)\s+(?:sorting|warehouse|logistics|fulfillment|distribution)",
            re.IGNORECASE
        )
        result = pattern7.sub(lambda m: m.group(0).replace(ref_client, correct_client_name), result)
        
        # Pattern 8: Standalone client name after "by" in context of data/sorting
        pattern8 = re.compile(
            rf"(?:sorted|processed|handled|managed|controlled)\s+by\s+{re.escape(ref_client)}",
            re.IGNORECASE
        )
        result = pattern8.sub(lambda m: m.group(0).replace(ref_client, correct_client_name), result)
    
    return result


def normalize_proposal_text(text: str, client_name: str = None) -> str:
    """
    Apply all text normalizations to proposal content:
    1. Normalize 'Cross Belt Sorter' capitalization
    2. Replace reference client names with correct client name
    """
    if not text:
        return text
    
    # Apply Cross Belt Sorter normalization
    result = normalize_cross_belt_sorter(text)
    
    # Apply client name replacement if provided
    if client_name:
        result = replace_client_names_in_flow(result, client_name)
    
    return result

# ==================== DXF COMPONENT EXTRACTION ====================

UNITS = {
    0: "Unitless", 1: "inches", 2: "feet", 3: "miles",
    4: "millimeters", 5: "centimeters", 6: "meters", 7: "kilometers",
}

# Component categorization keywords
COMPONENT_PATTERNS = {
    "AUTO_INDUCT": [
        r"fal.*fs\d+", r"fal.*feed", r"fal.*induct", r"auto.*induct", 
        r"feed.*line", r"feedline", r"fs\d{3}"
    ],
    "CONVEYOR_INFEED": [
        r"telescopic", r"infeed.*conv", r"in.*feed", r"receiving.*conv",
        r"inclined.*conv", r"incline", r"elevation.*conv"
    ],
    "ALIGNING_CONVEYOR": [
        r"aligning", r"align.*conv", r"alignment"
    ],
    "VDS_BUFFER": [
        r"vds", r"distribution.*loop", r"arm.*vds", r"boom.*conv", r"boom"
    ],
    "OPERATOR_STATION": [
        r"operator(?!.*safety)", r"manual.*station", r"induct.*station"
    ],
    "CHUTE": [
        r"^chute", r"gravity.*chute", r"live.*chute", r"slide.*chute",
        r"sliding.*chute", r"reject.*chute", r"collection.*chute",
        r"mini.*chute", r"bulk.*chute", r"discharge", r"direct.*bagging", r"bagging.*chute"
    ],
    "NON_SORT_CHUTE": [
        r"big.*parcel.*chute", r"large.*parcel.*chute", r"non.*sort.*chute", 
        r"nonsort.*chute", r"oversize.*chute", r"big.*chute"
    ],
    "REJECTION_CHUTE": [
        r"irregular.*chute", r"irchute", r"reject.*chute", r"ir.*chute",
        r"exception.*chute", r"error.*chute"
    ],
    "SLIDING_CHUTE": [
        r"sliding.*chute", r"slide.*chute", r"ptl.*rack", r"ptl\s*\d+x\d+"
    ],
    "PTL": [
        r"ptl", r"put.*to.*light", r"pick.*to.*light", r"light.*rack"
    ],
    "BAG_SYSTEM": [
        r"bag.*conv", r"bag.*takeaway", r"takeaway.*conv", r"bag.*take.*away",
        r"bagging.*conv", r"bag.*belt", r"^a\$[a-z0-9]+$"  # A$ blocks are Bag Takeaway Conveyors
    ],
    "RECIRCULATION": [
        r"recirculation", r"recirculate", r"refeed", r"return.*conv"
    ],
    "CBS_SORTER": [
        r"cbs", r"cross.*belt", r"crossbelt", r"sorter.*module",
        r"carrier", r"loop.*sorter"
    ],
    "SCANNER": [
        r"scanner", r"scan.*tunnel", r"barcode.*read", r"dimension.*sys",
        r"dws", r"volume.*scan"
    ],
    "TROLLEY": [
        r"trolley", r"roller.*cage", r"cage.*trolley"
    ],
    "COLLECTION_BIN": [
        r"collection.*bin", r"bin.*fal", r"st001"
    ],
}

# Structural components to filter out
STRUCTURAL_PATTERNS = [
    r"leg.*guard", r"guard(?!.*operator)", r"fenc", r"railing",
    r"handrail", r"safety.*(?!operator)", r"pallet(?!.*conv)",
    r"step", r"stair", r"ladder", r"door", r"panel(?!.*control)",
    r"bracket", r"bolt", r"mount(?!.*scanner)"
]

def _is_noise_block(name: str) -> bool:
    """Filter out anonymous noise blocks.
    
    Note: A$ blocks (like a$cf3a0ab4a) are NOT noise - they are Bag Takeaway Conveyors.
    """
    n = name.strip()
    if re.match(r"^\*[UDXATE]\d+$", n, re.IGNORECASE):
        return True
    if n.startswith("*") or n.startswith("~"):
        return True
    # A$ blocks are Bag Takeaway Conveyors - DO NOT filter them
    # They will be categorized as BAG_SYSTEM
    return False

def _is_structural(name: str) -> bool:
    """Check if component is structural (non-flow)."""
    n_lower = name.lower()
    for pattern in STRUCTURAL_PATTERNS:
        if re.search(pattern, n_lower):
            return True
    return False

def _categorize_component(name: str) -> str:
    """Categorize component by name pattern."""
    n_lower = name.lower()
    
    for category, patterns in COMPONENT_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, n_lower):
                return category
    
    return "UNCATEGORIZED"

def _normalize_group_name(name: str) -> str:
    """Normalize raw block name."""
    n = name.strip()
    if "|" in n:
        n = n.split("|")[-1]
    
    # Keep underscores in FAL codes
    if not n.startswith("FAL"):
        n = re.sub(r"[_\-]+", " ", n)
    
    n = re.sub(r"\s+", " ", n).strip()
    
    # Don't remove numbers from version codes
    if not re.search(r"V\d+$", n, re.IGNORECASE):
        n = re.sub(r"\s*\(?\d+\)?$", "", n).strip()
    
    return n.lower()

def _detect_cbs_type(project_name: str) -> str:
    """Detect CBS type from project name."""
    pn_lower = project_name.lower()
    # Check for linear CBS indicators
    if "linear" in pn_lower or "linear cbs" in pn_lower or "linear sorter" in pn_lower:
        return "Linear CBS"
    # Check for loop CBS indicators (explicitly mentioned)
    if "loop cbs" in pn_lower or "loop sorter" in pn_lower:
        return "Loop CBS"
    # Default to Loop CBS if not specified
    return "Loop CBS"

def _analyze_chute_types(components: dict) -> dict:
    """
    Enhanced chute breakdown analysis from raw DXF block names.
    
    Key mappings:
    - "big parcel chute" -> Non-Sort Chutes (large parcels)
    - "irregular chute" / "irchute" -> Rejection Chutes
    - "ptl rack" -> Sliding Chutes (PTL-enabled sliding chutes)
    - "sliding chute" / "slide chute" -> Sliding Chutes
    - "chute-XXX" generic pattern -> Generic Chutes
    """
    chute_analysis = {
        "total": 0,
        "by_type": defaultdict(int),
        "has_type_info": False,
        "per_zone": {}  # Will be populated if VDS count is known
    }
    
    for comp_name, count in components.items():
        name_lower = comp_name.lower()
        
        # Check for Non-Sort Chutes (big parcel chutes)
        if "big" in name_lower and "parcel" in name_lower and "chute" in name_lower:
            chute_analysis["by_type"]["non_sort_chutes"] += count
            chute_analysis["total"] += count
            chute_analysis["has_type_info"] = True
            continue
        
        # Check for Rejection/Irregular Chutes
        if "irregular" in name_lower or "irchute" in name_lower:
            chute_analysis["by_type"]["rejection_chutes"] += count
            chute_analysis["total"] += count
            chute_analysis["has_type_info"] = True
            continue
        
        # Check for PTL Racks -> These are Sliding Chutes (PTL-enabled sliding chutes for bagging)
        if "ptl" in name_lower and ("rack" in name_lower or re.search(r"ptl\s*\d+x\d+", name_lower)):
            chute_analysis["by_type"]["sliding_chutes"] += count
            chute_analysis["total"] += count
            chute_analysis["has_type_info"] = True
            continue
        
        # Only process remaining if "chute" or "bagging" is in the name
        if "chute" not in name_lower and "bagging" not in name_lower:
            continue
            
        chute_analysis["total"] += count
        
        # Try to detect type - check more specific patterns first
        if "direct" in name_lower and "bagging" in name_lower:
            chute_analysis["by_type"]["direct_bagging_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "bagging" in name_lower and "ptl" not in name_lower:
            chute_analysis["by_type"]["direct_bagging_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "mini" in name_lower and "gravity" in name_lower:
            chute_analysis["by_type"]["mini_gravity_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "gravity" in name_lower:
            chute_analysis["by_type"]["gravity_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "collection" in name_lower:
            chute_analysis["by_type"]["collection_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "live" in name_lower or "active" in name_lower or "ob live" in name_lower:
            chute_analysis["by_type"]["live_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "slide" in name_lower or "sliding" in name_lower:
            chute_analysis["by_type"]["sliding_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "reject" in name_lower or "exception" in name_lower:
            chute_analysis["by_type"]["rejection_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "discharge" in name_lower:
            chute_analysis["by_type"]["discharge_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "ow" in name_lower and "chute" in name_lower:
            chute_analysis["by_type"]["ow_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "od" in name_lower and "chute" in name_lower:
            chute_analysis["by_type"]["od_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "sort" in name_lower and "fail" in name_lower:
            chute_analysis["by_type"]["sortfail_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "bulk" in name_lower:
            chute_analysis["by_type"]["bulk_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "non-sort" in name_lower or "nonsort" in name_lower:
            chute_analysis["by_type"]["non_sort_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "dispersion" in name_lower or "disperse" in name_lower:
            chute_analysis["by_type"]["dispersion_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "high" in name_lower and "volume" in name_lower:
            chute_analysis["by_type"]["high_volume_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "low" in name_lower and "volume" in name_lower:
            chute_analysis["by_type"]["low_volume_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "manual" in name_lower:
            chute_analysis["by_type"]["manual_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "spiral" in name_lower:
            chute_analysis["by_type"]["spiral_chutes"] += count
            chute_analysis["has_type_info"] = True
        elif "overweight" in name_lower or "over" in name_lower:
            chute_analysis["by_type"]["overweight_chutes"] += count
            chute_analysis["has_type_info"] = True
        else:
            # Generic chute if can't determine type
            chute_analysis["by_type"]["generic_chutes"] += count
    
    return chute_analysis


def _calculate_per_zone_chutes(chute_analysis: dict, vds_count: int) -> dict:
    """
    Calculate per-zone chute distribution when VDS count is known.
    
    Example: If VDS=2 and we have 26 non-sort chutes total,
    then each zone has 13 non-sort chutes.
    """
    if vds_count <= 0:
        return chute_analysis
    
    per_zone = {}
    for chute_type, total_count in chute_analysis.get("by_type", {}).items():
        per_zone[chute_type] = total_count // vds_count
    
    chute_analysis["per_zone"] = per_zone
    chute_analysis["zone_count"] = vds_count
    
    return chute_analysis

def extract_dxf_components(dxf_path: Path, project_name: str = "") -> dict:
    """Extract and categorize components from DXF file."""
    doc = ezdxf.readfile(str(dxf_path))
    msp = doc.modelspace()
    hdr = doc.header
    
    # Get units
    units_code = hdr.get("$INSUNITS", None)
    try:
        units_code = int(units_code) if units_code is not None else None
    except:
        units_code = None
    
    # Get extents
    extmin = hdr.get("$EXTMIN", None)
    extmax = hdr.get("$EXTMAX", None)
    
    # Extract block references
    raw_counts: Counter[str] = Counter()
    for e in msp:
        try:
            if e.dxftype() == "INSERT":
                bname = e.dxf.name
                if not _is_noise_block(bname) and not _is_structural(bname):
                    raw_counts[bname] += 1
        except:
            continue
    
    # Categorize components
    categorized: dict = defaultdict(lambda: defaultdict(lambda: {"count": 0, "examples": []}))
    
    for raw_name, cnt in raw_counts.items():
        category = _categorize_component(raw_name)
        gname = _normalize_group_name(raw_name)
        
        categorized[category][gname]["count"] += cnt
        categorized[category][gname]["examples"].append(raw_name)
    
    # Detect system characteristics from project name
    cbs_type = _detect_cbs_type(project_name if project_name else dxf_path.name)
    chute_analysis = _analyze_chute_types(raw_counts)
    
    # Determine induction type
    has_auto_induct = len(categorized.get("AUTO_INDUCT", {})) > 0
    has_operators = len(categorized.get("OPERATOR_STATION", {})) > 0
    has_vds = len(categorized.get("VDS_BUFFER", {})) > 0
    
    # Get VDS count for per-zone calculations
    vds_count = 0
    if has_vds:
        vds_items = categorized.get("VDS_BUFFER", {})
        vds_count = sum(item["count"] for item in vds_items.values())
        if vds_count > 0:
            chute_analysis = _calculate_per_zone_chutes(chute_analysis, vds_count)
    
    # Check for bag takeaway system
    has_bag_system = len(categorized.get("BAG_SYSTEM", {})) > 0
    # Also check in UNCATEGORIZED for bag-related items
    uncategorized = categorized.get("UNCATEGORIZED", {})
    for comp_name in uncategorized.keys():
        if "bag" in comp_name.lower() and ("takeaway" in comp_name.lower() or "conv" in comp_name.lower()):
            has_bag_system = True
            break
    
    if has_auto_induct and has_operators:
        induction_type = "MIXED (Auto + Manual)"
    elif has_auto_induct:
        induction_type = "AUTO"
    elif has_operators:
        induction_type = "MANUAL"
    else:
        induction_type = "UNKNOWN"
    
    # Build category summary
    category_summary = {}
    total_components = 0
    for cat, items in categorized.items():
        count = sum(item["count"] for item in items.values())
        category_summary[cat] = count
        total_components += count
    
    return {
        "file": dxf_path.name,
        "units_code": units_code,
        "units_name": UNITS.get(units_code, "unknown") if units_code is not None else None,
        "extents": {
            "min": list(extmin) if extmin else None,
            "max": list(extmax) if extmax else None
        },
        "cbs_type": cbs_type,
        "induction_type": induction_type,
        "has_vds": has_vds,
        "vds_count": vds_count,
        "has_bag_system": has_bag_system,
        "total_components": total_components,
        "category_summary": dict(category_summary),
        "categorized_components": {
            cat: {name: data["count"] for name, data in items.items()}
            for cat, items in categorized.items()
        },
        "chute_analysis": chute_analysis,
        "raw_block_counts": {k: int(v) for k, v in raw_counts.items()},
    }

def _summarise_components_for_prompt(dxf_json: dict) -> str:
    """Generate comprehensive summary for LLM prompt."""
    lines = ["=" * 70]
    lines.append("DXF COMPONENT ANALYSIS FOR CBS PROCESS FLOW GENERATION")
    lines.append("=" * 70)
    lines.append("")
    
    # File and system info
    lines.append(f"FILE: {dxf_json.get('file', 'Unknown')}")
    lines.append(f"UNITS: {dxf_json.get('units_name', 'Unknown')}")
    lines.append(f"CBS TYPE: {dxf_json.get('cbs_type', 'Unknown')} (detected from filename)")
    lines.append(f"TOTAL COMPONENTS: {dxf_json.get('total_components', 0)}")
    lines.append("")
    
    # System configuration
    lines.append("SYSTEM CONFIGURATION:")
    lines.append(f"  • Induction Type: {dxf_json.get('induction_type', 'Unknown')}")
    lines.append(f"  • VDS/Buffer System: {'YES' if dxf_json.get('has_vds') else 'NO'}")
    
    vds_count = dxf_json.get('vds_count', 0)
    if vds_count > 0:
        lines.append(f"  • Number of VDS Zones: {vds_count}")
    
    has_bag_system = dxf_json.get('has_bag_system', False)
    lines.append(f"  • Bag Takeaway System: {'YES' if has_bag_system else 'NO'}")
    
    chute_analysis = dxf_json.get('chute_analysis', {})
    if chute_analysis.get('total', 0) > 0:
        lines.append(f"  • Total Chutes: {chute_analysis['total']}")
        if chute_analysis.get('has_type_info'):
            lines.append("  • Chute Types Detected: YES")
        else:
            lines.append("  • Chute Types Detected: NO (describe generically)")
    lines.append("")
    
    # Component categories
    cat_summary = dxf_json.get("category_summary", {})
    if cat_summary:
        lines.append("COMPONENT CATEGORIES:")
        priority = ["AUTO_INDUCT", "OPERATOR_STATION", "VDS_BUFFER", "CONVEYOR_INFEED",
                   "CBS_SORTER", "CHUTE", "NON_SORT_CHUTE", "REJECTION_CHUTE", "SLIDING_CHUTE",
                   "PTL", "BAG_SYSTEM", "RECIRCULATION", "SCANNER", "TROLLEY", "COLLECTION_BIN"]
        
        for cat in priority:
            if cat in cat_summary:
                lines.append(f"  • {cat}: {cat_summary[cat]} units")
        
        # Add any remaining categories
        for cat, count in cat_summary.items():
            if cat not in priority:
                lines.append(f"  • {cat}: {count} units")
        lines.append("")
    
    # Enhanced Chute breakdown with per-zone info
    if chute_analysis.get('by_type'):
        lines.append("=" * 70)
        lines.append("CHUTE TYPE BREAKDOWN (CRITICAL - USE THESE EXACT COUNTS):")
        lines.append("-" * 70)
        
        zone_count = chute_analysis.get('zone_count', 0)
        per_zone = chute_analysis.get('per_zone', {})
        
        for ctype, count in sorted(chute_analysis['by_type'].items(), key=lambda x: -x[1]):
            type_name = ctype.replace('_', ' ').title()
            if zone_count > 0 and ctype in per_zone:
                per_zone_count = per_zone[ctype]
                lines.append(f"  • {type_name}: {count} total ({per_zone_count} per zone × {zone_count} zones)")
            else:
                lines.append(f"  • {type_name}: {count} chutes")
        
        if zone_count > 0:
            lines.append("")
            lines.append(f"⚠️ IMPORTANT: This system has {zone_count} VDS zones. Write chute counts PER ZONE in the output.")
            lines.append("   Example: 'Within the loop CBS system, there are 50 Sliding chutes for each zone.'")
        lines.append("")
    
    # Detailed inventory by category
    lines.append("DETAILED COMPONENT INVENTORY:")
    lines.append("-" * 70)
    
    categorized = dxf_json.get("categorized_components", {})
    for cat in priority:
        if cat in categorized and categorized[cat]:
            lines.append("")
            lines.append(f"[{cat}]")
            for name, count in sorted(categorized[cat].items(), key=lambda x: -x[1]):
                lines.append(f"  • {name}: {count} units")
    
    # Add uncategorized if any
    if "UNCATEGORIZED" in categorized and categorized["UNCATEGORIZED"]:
        lines.append("")
        lines.append("[UNCATEGORIZED]")
        for name, count in sorted(categorized["UNCATEGORIZED"].items(), key=lambda x: -x[1]):
            lines.append(f"  • {name}: {count} units")
    
    lines.append("")
    lines.append("=" * 70)
    lines.append("PROCESS FLOW GENERATION GUIDANCE:")
    lines.append(f"  • System Type: {dxf_json.get('cbs_type', 'Unknown')}")
    lines.append(f"  • Induction: {dxf_json.get('induction_type', 'Unknown')}")
    
    if dxf_json.get('has_vds'):
        lines.append("  • Include VDS/Buffer section in Infeed System")
    
    if dxf_json.get('induction_type') == "MIXED (Auto + Manual)":
        lines.append("  • Include BOTH Auto Induct Line AND Manual Induct Station sections")
    elif dxf_json.get('induction_type') == "AUTO":
        lines.append("  • Include Auto Induct Line section only")
    elif dxf_json.get('induction_type') == "MANUAL":
        lines.append("  • Include Manual Induct Station section only")
    
    if chute_analysis.get('has_type_info'):
        lines.append("  • Chute types available - use detailed breakdown with per-zone counts")
    else:
        lines.append("  • Chute types NOT available - describe generically")
    
    if "PTL" in cat_summary or "SLIDING_CHUTE" in cat_summary:
        lines.append("  • Include PTL/Sliding Chutes with secondary sorting (bagging) description")
    
    if has_bag_system or "BAG_SYSTEM" in cat_summary:
        lines.append("  • ⚠️ MANDATORY: Include 'Bag Takeaway Conveyor' as FINAL SECTION (section 5)")
        lines.append("    - BAG SYSTEM DETECTED in DXF - this section is REQUIRED")
        lines.append("    - Describes conveyor beneath CBS that transports bags to Outbound sorter")
    
    if "RECIRCULATION" in cat_summary:
        lines.append("  • Include Recirculation/Exception Refeeding section")
    
    # Add detailed chute writing instructions
    lines.append("")
    lines.append("=" * 70)
    lines.append("OUTPUT CHUTES SECTION - REQUIRED FORMAT:")
    lines.append("-" * 70)
    if chute_analysis.get('total', 0) > 0:
        lines.append(f"TOTAL CHUTES: {chute_analysis['total']}")
        if chute_analysis.get('by_type'):
            lines.append("")
            
            zone_count = chute_analysis.get('zone_count', 0)
            per_zone = chute_analysis.get('per_zone', {})
            
            # Show per-zone counts only if zone_count is 2 or 3 (i.e., 1 < zone_count < 4)
            use_per_zone = 1 < zone_count < 4
            
            if use_per_zone:
                lines.append(f"CHUTE BREAKDOWN TO INCLUDE IN OUTPUT (show PER-ZONE counts - {zone_count} zones):")
            else:
                lines.append("CHUTE BREAKDOWN TO INCLUDE IN OUTPUT (show TOTAL counts):")
            
            for ctype, count in sorted(chute_analysis['by_type'].items(), key=lambda x: -x[1]):
                type_name = ctype.replace('_', ' ').title()
                if use_per_zone and ctype in per_zone:
                    lines.append(f"  - {type_name}: {per_zone[ctype]} per zone (total: {count})")
                else:
                    lines.append(f"  - {type_name}: {count}")
            
            lines.append("")
            if use_per_zone:
                lines.append(f"FORMAT EXAMPLE (for system with {zone_count} zones - use per-zone counts):")
                lines.append("Output Chutes: - The shipments are discharged into following types of chutes:")
                lines.append("a. Sliding Chutes - Within the loop CBS system, there are a total of 50 Sliding chutes for")
                lines.append("   each zone. The Shipments collected in Roller Cage trolleys, then they are consolidated")
                lines.append("   into bags using bagging type PTL racks.")
                lines.append("b. Non-Sort Chutes - Within the loop CBS system, there are a total of 13 Non-Sort Chutes per zone.")
                lines.append("   Shipments collected within these chutes further undergo sortation via PTL setup into Pallets.")
                lines.append("c. Rejection Chutes - Two Rejection Chutes per zone are present to handle rejected Shipments.")
            else:
                lines.append("FORMAT EXAMPLE (use TOTAL counts):")
                lines.append("Output Chutes: - The shipments are discharged into following types of chutes:")
                lines.append("a. Sliding Chutes - Within the loop CBS system, there are a total of 100 Sliding chutes.")
                lines.append("   The Shipments collected in Roller Cage trolleys, then they are consolidated into bags.")
                lines.append("b. Non-Sort Chutes - Within the loop CBS system, there are a total of 26 Non-Sort Chutes.")
                lines.append("c. Rejection Chutes - 4 Rejection Chutes are present to handle rejected Shipments.")
        else:
            lines.append("NO SPECIFIC CHUTE TYPES DETECTED - use generic description")
    
    # Bag Takeaway section guidance
    if has_bag_system or "PTL" in cat_summary or "SLIDING_CHUTE" in cat_summary:
        lines.append("")
        lines.append("=" * 70)
        lines.append("BAG TAKEAWAY CONVEYOR SECTION (REQUIRED AS FINAL SECTION):")
        lines.append("-" * 70)
        lines.append("⚠️ BAG SYSTEM DETECTED - MUST include this section AFTER Output Chutes!")
        lines.append("FORMAT EXAMPLE:")
        lines.append("Bag Takeaway Conveyor: - Following the direct bagging process & secondary sorting process,")
        lines.append("   the shipments are placed into bags and then manually loaded onto a bag takeaway conveyor")
        lines.append("   located beneath the CBS loop. This conveyor transports the bags out of shipment sorter to")
        lines.append("   Outbound sorter located beneath base mezzanine in the approx. centre of the loop CBS.")
    
    lines.append("=" * 70)
    
    return "\n".join(lines)

def convert_dxf_to_png(dxf_path: Path) -> Path:
    """Convert DXF file to PNG using ConvertAPI."""
    if not CONVERTAPI_SECRET:
        raise RuntimeError(
            "CONVERTAPI_SECRET is not set in .env file. Cannot convert DXF to PNG without it."
        )
    
    try:
        # Convert DXF to PNG using ConvertAPI
        result = convertapi.convert("png", {"File": str(dxf_path)}, from_format="dxf")
        out_files = result.save_files(str(dxf_path.parent))
        
        # Find the PNG file
        for f in out_files:
            if str(f).lower().endswith(".png"):
                return Path(f)
        
        # Return first file if no .png extension found
        return Path(out_files[0]) if out_files else None
    except Exception as e:
        raise RuntimeError(f"Failed to convert DXF to PNG: {str(e)}")

def _normalise_to_numbered_steps(raw_text: str) -> str:
    """Force clean 1..N numbered list from GROQ output."""
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    if len(lines) == 1:
        parts = re.split(r'(?:(?<=\.)\s+)(?=\d+\.)', lines[0])
        lines = [p.strip() for p in parts if p.strip()]
    
    steps = []
    for ln in lines:
        m = re.match(r"^(\d+)[\.\)\-]\s*(.*)$", ln)
        content = m.group(2).strip() if m else ln
        if content: steps.append(content)
    
    dedup = []
    seen = set()
    for s in steps:
        key = re.sub(r"\s+", " ", s.lower())
        if key not in seen:
            seen.add(key)
            dedup.append(s)
    
    max_steps = min(len(dedup), 9) if len(dedup) >= 5 else len(dedup)
    return "\n".join([f"{i}. {content}" for i, content in enumerate(dedup[:max_steps], start=1)])

# ==================== LOOP CBS EXCEL PARSING ====================

def load_loop_cbs_sheet_from_excel(xlsx_bytes: bytes) -> tuple[str | None, pd.DataFrame | None]:
    """
    Find sheet whose name is 'Loop CBS' or 'Loop CBS upper' (case/space-insensitive)
    and return (sheet_name, DataFrame). Reads the FULL sheet (no cropping).
    """
    with pd.ExcelFile(io.BytesIO(xlsx_bytes)) as xls:
        target_name = None
        for sname in xls.sheet_names:
            normalized = sname.lower().replace(" ", "")
            if normalized == "loopcbs" or normalized == "loopcbsupper":
                target_name = sname
                break

        if target_name is None:
            return None, None

        # Read entire sheet as generic table (no header), so we keep all rows/blocks.
        df = xls.parse(target_name, header=None)
        return target_name, df


def df_to_compact_text(df: pd.DataFrame, max_rows: int = 200, max_cols: int = 20) -> str:
    """
    Convert the ENTIRE sheet (up to max_rows, max_cols) into a compact text representation
    including BOTH tables in Loop CBS sheet.
    """
    if df is None:
        return ""

    df2 = df.iloc[:max_rows, :max_cols].fillna("")

    lines: list[str] = []
    for idx in range(df2.shape[0]):
        row_vals = [str(v).strip() for v in df2.iloc[idx].tolist()]
        # keep only non-empty cells in the print
        row_vals = [v for v in row_vals if v != ""]
        if row_vals:
            lines.append(" | ".join(row_vals))

    return "\n".join(lines)

def df_to_compact_text_quote_master(df: pd.DataFrame) -> str:
    """
    Convert Quote Master dataframe to compact text for GROQ.
    Every non-empty row becomes a pipe-separated list of "Col=Value".
    """
    lines = []
    headers = list(df.columns)

    for _, row in df.iterrows():
        if row.isna().all():
            continue

        pairs = []
        for h, v in zip(headers, row.values):
            if pd.isna(v) or h is None:
                continue
            v_str = str(v).strip()
            if not v_str:
                continue
            pairs.append(f"{h}={v_str}")

        if pairs:
            lines.append(" | ".join(pairs))

    return "\n".join(lines)

# ==================== PROCESS FLOW GENERATION ====================


def call_groq_for_process_flow(
    client_name: str,
    project_name: str,
    dxf_json: Optional[dict] = None,
    facts: Optional[ProposalFacts] = None,
    context: Optional[ProposalContext] = None,
    manual_components_context: Optional[Dict] = None,
):
    """
    Generate process flow using agentY logic with PROGRESSIVE IMPROVEMENT PROTECTION.
    
    🔧 ENHANCED FEATURES:
    - Allows up to 6 iterations
    - Detects and PREVENTS score regression
    - Ensures progressive improvement each iteration
    - Tracks best score and reverts if needed
    - Integrates manually configured components from DXF editor
    
    Args:
        client_name: Name of the client
        project_name: Name of the project
        dxf_json: Parsed DXF component data
        facts: ProposalFacts instance with component counts
        context: ProposalContext instance with system counts
        manual_components_context: Manually configured components from the UI editor
    
    Returns:
        tuple: (final_flow_text, iteration_details_list)
    """
    import logging
    logger = logging.getLogger(__name__)
    
    safe_dxf_json = {k: v for k, v in (dxf_json or {}).items() if k != "raw_block_counts"}
    safe_dxf_json['client'] = client_name

    # Enrich with manual component configuration if available
    if manual_components_context:
        safe_dxf_json["manual_components"] = manual_components_context.get("configured_components", {})
        safe_dxf_json["sections_included"] = manual_components_context.get("sections_included", [])
        safe_dxf_json["components_summary"] = manual_components_context.get("components_summary", "")
        logger.info(f"Using manually configured components: {manual_components_context.get('sections_included', [])}")

    # Enrich with ProposalFacts counts and ProposalContext (non-breaking extra data)
    if facts:
        counts_from_facts = {}
        for key in [
            "feedlines",
            "induct_lines_auto",
            "manual_induct_stations",
            "gravity_chutes",
            "mini_gravity_chutes",
            "collection_chutes",
            "rejection_chutes",
            "dispersion_chutes",
            "bulk_chutes",
            "direct_bagging_chutes",
            "throughput_pph",
        ]:
            val, src, conf = get_counts_source_of_truth(facts, key)
            if val is not None:
                counts_from_facts[key] = {"count": val, "source": src, "confirmed": conf}
        if counts_from_facts:
            safe_dxf_json["facts_counts"] = counts_from_facts
    if context:
        safe_dxf_json["counts_block_text"] = context.counts_block_text()
        # Pre-generation validation gate for missing/zero counts
        gating_issues = validate_context_counts(context) if ENABLE_CONTEXT_UNIFICATION else []
        if gating_issues:
            safe_dxf_json["safe_phrasing"] = True
            safe_dxf_json["gating_issues"] = gating_issues
            logger.warning(f"Pre-generation validation gate (Process Flow) triggered: {gating_issues}")
        ctx_counts = context.counts_block_text().replace('\n',' | ')
        logger.info(f"Using ProposalContext counts for Process Flow: {ctx_counts}")
    iteration_details = []
    
    target_score = 88.0
    max_iterations = 6  # ✅ Up to 6 iterations
    min_improvement = 0.5  # Require at least 0.5 point improvement

    try:
        # ================================================================
        # STEP 1: INITIAL GENERATION (from DXF only)
        # ================================================================
        initial_flow = agentY_generate_initial_flow(client_name, safe_dxf_json)
        if not initial_flow or len(initial_flow.strip()) < 50:
            return "", []

        # ================================================================
        # STEP 2: QUERY REFERENCES & REFINE
        # ================================================================
        reference_flows = []
        try:
            pc, index = combine_old_mod.get_pinecone_index()
            dxf_summary = create_dxf_summary_for_embedding(safe_dxf_json)
            reference_flows = combine_old_mod.query_similar_flows(
                pc, index, dxf_summary, safe_dxf_json, top_k=2, threshold=0.70
            )
        except Exception as e:
            logger.warning(f"Failed to query Pinecone: {e}")
            reference_flows = []
        
        # Generate second flow with references
        if reference_flows:
            try:
                second_flow = generate_second_flow_with_chunks(
                    initial_flow, safe_dxf_json, reference_flows, client_name
                )
                current_flow = second_flow
            except Exception as e:
                logger.warning(f"Failed to generate second flow: {e}")
                current_flow = initial_flow
        else:
            current_flow = initial_flow
        
        # Get reference text for evaluation
        reference_text = reference_flows[0]["process_flow"] if reference_flows else current_flow
        
        # ================================================================
        # STEP 3 & 4: EVALUATION & ITERATIVE REFINEMENT (ENHANCED PROTECTION)
        # ================================================================
        best_score = 0.0
        best_flow = current_flow
        best_iteration = 0
        no_improvement_iterations = 0
        last_score = 0.0
        score_regression_count = 0
        
        for iteration_num in range(1, max_iterations + 1):
            # Evaluate current flow
            evaluation = agentY_evaluate_process_flow(
                current_flow,
                reference_text,
                safe_dxf_json,
                target_score=target_score
            )
            
            current_score = evaluation.get('structural_coherence', 0.0)
            feedback_items = evaluation.get('feedback', [])
            
            # Store iteration details with FULL FLOW TEXT
            iteration_details.append({
                "iteration": iteration_num,
                "score": round(current_score, 2),
                "target_score": target_score,
                "gap": round(target_score - current_score, 2),
                "feedback": feedback_items,
                "flow": current_flow  # ✅ FULL FLOW TEXT
            })
            
            logger.info(f"Iteration {iteration_num}: Score={current_score:.2f}, Target={target_score}, Feedback={len(feedback_items)} items")
            
            # ===================================================================
            # 🔧 SCORE REGRESSION PROTECTION
            # ===================================================================
            
            # Check for significant score drop (>3 points)
            if iteration_num > 1 and current_score < last_score - 3:
                score_regression_count += 1
                logger.warning(f"⚠️ Score regression detected: {last_score:.2f} → {current_score:.2f}")
                
                # If 2+ regressions, revert to best and stop trying to improve
                if score_regression_count >= 2:
                    logger.info(f"Too many regressions ({score_regression_count}). Reverting to best flow.")
                    current_flow = best_flow
                    no_improvement_iterations = 10  # Force stop
                    break
            else:
                score_regression_count = 0  # Reset counter on improvement/stable
            
            last_score = current_score
            
            # ===================================================================
            # 🔧 BEST SCORE TRACKING (Accept if better)
            # ===================================================================
            
            if current_score > best_score + min_improvement:
                # This is a meaningful improvement - accept it
                improvement = current_score - best_score
                best_score = current_score
                best_flow = current_flow
                best_iteration = iteration_num
                no_improvement_iterations = 0  # Reset counter
                
                logger.info(f"✅ New best score: {best_score:.2f} (+{improvement:.2f})")
                
            elif current_score >= best_score - 1.0:
                # Within 1 point of best - acceptable, update flow but not best
                logger.info(f"✓ Score stable: {current_score:.2f} (best: {best_score:.2f})")
                no_improvement_iterations += 1
                
            else:
                # Worse than best by >1 point - revert to best
                logger.warning(f"⚠️ Score dropped to {current_score:.2f}. Reverting to best: {best_score:.2f}")
                current_flow = best_flow  # Revert to best
                no_improvement_iterations += 1
            
            # ===================================================================
            # 🔧 ENHANCED STOPPING CONDITIONS
            # ===================================================================
            
            # 1. Target reached
            if current_score >= target_score:
                logger.info(f"✅ Target reached! Score {current_score:.2f} >= {target_score}")
                break
            
            # 2. Excellent score
            if current_score >= 90:
                logger.info(f"✅ Excellent score! {current_score:.2f} >= 90")
                break
            
            # 3. No improvement for 3 consecutive iterations
            if no_improvement_iterations >= 3:
                logger.info(f"⚠️ No improvement for {no_improvement_iterations} iterations. Stopping.")
                break
            
            # 4. Max iterations reached
            if iteration_num >= max_iterations:
                logger.info(f"Max iterations ({max_iterations}) reached")
                break
            
            # ===================================================================
            # 🔧 REFINEMENT STEP (if not stopping)
            # ===================================================================
            
            if iteration_num < max_iterations:
                try:
                    logger.info(f"Applying conservative refinement to improve from {current_score:.2f}...")
                    
                    # ALWAYS use BEST flow as base (not current if it regressed)
                    next_flow = iterative_refinement(
                        best_flow,  # ✅ Always refine from best
                        safe_dxf_json,
                        reference_text,
                        evaluation,
                        iteration_num,
                        target_score
                    )
                    
                    # Validate generated flow
                    is_valid, error_msg = validate_flow_quality(next_flow)
                    if is_valid:
                        current_flow = next_flow  # Move to next iteration
                        logger.info("✓ Refinement successful, proceeding to next iteration")
                    else:
                        logger.warning(f"⚠️ Generated flow failed validation: {error_msg}")
                        current_flow = best_flow  # Keep best flow
                        no_improvement_iterations += 1
                        
                except Exception as e:
                    logger.error(f"❌ Refinement error: {e}")
                    current_flow = best_flow  # Keep best flow on error
                    no_improvement_iterations += 1
        
        # ===================================================================
        # 🔧 FINAL RETURN
        # ===================================================================
        
        final_flow = best_flow.strip()
        
        # Fix empty Output Chutes section if needed
        final_flow = agentY_fix_empty_output_chutes(final_flow, safe_dxf_json, client_name)
        
        # Apply text normalizations: Cross Belt Sorter capitalization and client name consistency
        final_flow = normalize_proposal_text(final_flow, client_name)
        
        # Sanitize LLM output to remove meta-artifacts
        final_flow = sanitize_section_output(final_flow)
        
        logger.info(f"Final result: Best score {best_score:.2f} achieved in iteration {best_iteration}")
        logger.info(f"Total iterations: {len(iteration_details)}")
        
        return final_flow, iteration_details
    
    except Exception as e:
        logger.error(f"❌ Error in call_groq_for_process_flow: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return "", []


# ==================== MERMAID FLOWCHART GENERATION ====================

def sanitize_mermaid_for_render(code: str) -> str:
    """Sanitize Mermaid code for reliable rendering."""
    if not code: return code
    code = code.strip()
    # Remove markdown fences
    if code.startswith("```"):
        lines = code.split("\n")
        if lines[0].strip().startswith("```"): lines = lines[1:]
        if lines and lines[-1].strip() == "```": lines = lines[:-1]
        code = "\n".join(lines).strip()
    if code.lower().startswith("mermaid"): code = code[7:].strip()
    if not code.lower().startswith("flowchart"): return code
    
    # Normalize line endings
    code = code.replace("\r\n", "\n").replace("\r", "\n")
    
    # Fix edge labels with parentheses - they cause rendering issues
    # Pattern: -->|(text with parens)| becomes -->|text without parens|
    code = re.sub(r'\|([^|]*)\(([^)]*)\)([^|]*)\|', r'|\1\2\3|', code)
    
    return code

def generate_mermaid_png(mermaid_code: str) -> tuple:
    """Render Mermaid diagram to PNG."""
    logs = []
    mermaid_str = sanitize_mermaid_for_render(mermaid_code)
    logs.append(f"Cleaned Mermaid code:\n{mermaid_str}\n")
    
    # Strategy 1: Kroki PNG with JSON
    try:
        payload = {"diagram_source": mermaid_str, "diagram_type": "mermaid", "output_format": "png"}
        resp = requests.post("https://kroki.io/mermaid/png", json=payload, 
                           headers={"Content-Type": "application/json"}, timeout=30)
        if resp.ok and resp.content and len(resp.content) > 100:
            logs.append(f"✓ Kroki PNG returned {len(resp.content)} bytes.")
            return resp.content, "\n".join(logs)
    except Exception as e:
        logs.append(f"Kroki PNG error: {repr(e)}")
    
    # Strategy 2: Mermaid.ink
    try:
        json_payload = {"code": mermaid_str, "mermaid": {"theme": "default"}}
        b64_str = base64.b64encode(json.dumps(json_payload).encode('utf-8')).decode('ascii')
        resp = requests.get(f"https://mermaid.ink/img/{b64_str}", timeout=30)
        if resp.ok and resp.content and len(resp.content) > 100:
            logs.append(f"✓ mermaid.ink returned {len(resp.content)} bytes.")
            return resp.content, "\n".join(logs)
    except Exception as e:
        logs.append(f"mermaid.ink error: {repr(e)}")
    
    raise RuntimeError(f"Failed to render Mermaid:\n" + "\n".join(logs))


def call_groq_for_mermaid(process_flow_text: str) -> str:
    """
    Generate an accurate Mermaid v11 flowchart (flowchart TD) from CBS process-flow text.

    Key upgrade vs. your current version:
    - Groq returns STRICT GRAPH JSON (nodes + edges) only
    - Python renders Mermaid deterministically (prevents structural drift: merges/branches/loops)
    """

    # -----------------------------
    # 1) Helpers (local)
    # -----------------------------
    def _clean_text(t: str) -> str:
        t = t.replace("\r\n", "\n").replace("\r", "\n")
        # remove repeated blank lines
        t = re.sub(r"\n{3,}", "\n\n", t)
        return t.strip()

    def _extract_skeleton(t: str) -> Dict[str, Any]:
        """
        Best-effort structure extraction to help the LLM:
        - Sections like "Infeed System:" / "Auto Induct Line:" (with or without numbering)
        - Subpoints like "a. Sliding Chutes - ..." under an output section
        """
        text = _clean_text(t)
        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]

        # Skip a standalone "Process Flow" header if present
        if lines and lines[0].lower() in ("process flow", "processflow"):
            lines = lines[1:]

        section_header_re = re.compile(
            r"""^(?:\d+\.\s*)?        # optional leading numbering
                ([A-Za-z][A-Za-z0-9 /,&()\-]+?)  # title
                \s*:\s*-?\s*          # : or :- separator
                (.*)$                 # rest of line
            """,
            re.VERBOSE,
        )

        subpoint_re = re.compile(
            r"""^([a-eA-E])\.\s+(.+?)\s*(?:-|\:)\s*(.*)$""", re.VERBOSE
        )

        sections: List[Dict[str, Any]] = []
        current: Optional[Dict[str, Any]] = None

        def push_current():
            nonlocal current
            if current:
                # normalize spacing
                current["title"] = re.sub(r"\s+", " ", current["title"]).strip()
                current["desc"] = re.sub(r"\s+", " ", current["desc"]).strip()
                sections.append(current)
            current = None

        for ln in lines:
            m = section_header_re.match(ln)
            if m:
                push_current()
                title = m.group(1).strip()
                rest = m.group(2).strip()
                current = {"title": title, "desc": rest, "subpoints": []}
                continue

            sm = subpoint_re.match(ln)
            if sm and current is not None:
                sp_title = sm.group(2).strip()
                sp_desc = sm.group(3).strip()
                current["subpoints"].append(
                    {
                        "key": sm.group(1).lower(),
                        "title": re.sub(r"\s+", " ", sp_title),
                        "desc": re.sub(r"\s+", " ", sp_desc),
                    }
                )
                continue

            # Continuation line
            if current is None:
                # If text starts without a clean header, create a generic section
                current = {"title": "Step", "desc": ln, "subpoints": []}
            else:
                current["desc"] += " " + ln

        push_current()

        # Lightweight loop hints
        loop_hints = []
        joined = " ".join(lines).lower()
        if any(k in joined for k in ["refeed", "re-feed", "again fed", "recirculation", "recirculate"]):
            loop_hints.append("has_refeed_or_recirculation")
        if "optional" in joined:
            loop_hints.append("has_optional_branch")
        if any(k in joined for k in ["manual induct", "manual infeed"]) and any(
            k in joined for k in ["auto induct", "autoinduct", "auto induction"]
        ):
            loop_hints.append("has_parallel_manual_and_auto")

        return {"sections": sections, "hints": loop_hints}

    def _extract_json_object(s: str) -> str:
        """
        Pull the first {...} JSON object from a messy model output.
        """
        s = s.strip()
        # strip code fences if model disobeys
        s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.IGNORECASE)
        s = re.sub(r"\s*```$", "", s)
        # find first JSON object
        start = s.find("{")
        end = s.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise ValueError("No JSON object found in model output")
        return s[start : end + 1]

    def _num_to_id(n: int) -> str:
        """
        0->A, 1->B ... 25->Z, 26->AA ...
        """
        letters = string.ascii_uppercase
        out = ""
        n_local = n
        while True:
            out = letters[n_local % 26] + out
            n_local = (n_local // 26) - 1
            if n_local < 0:
                break
        return out

    def _render_mermaid(graph: Dict[str, Any]) -> str:
        """
        Deterministic Mermaid rendering with classDefs + :::className.
        """
        nodes = graph.get("nodes", [])
        edges = graph.get("edges", [])

        # Validate minimal schema
        if not isinstance(nodes, list) or not isinstance(edges, list):
            raise ValueError("Graph JSON must have 'nodes' and 'edges' as lists")

        # Map model keys -> Mermaid IDs (A, B, C...)
        key_to_mid: Dict[str, str] = {}
        mermaid_nodes: List[str] = []

        for i, nd in enumerate(nodes):
            k = str(nd.get("key", "")).strip()
            label = str(nd.get("label", "")).strip()
            cls = str(nd.get("class", "")).strip()

            if not k or not label:
                continue

            mid = _num_to_id(len(key_to_mid))
            key_to_mid[k] = mid

            # keep labels short-ish (model should already do it)
            label = re.sub(r"\s+", " ", label)
            mermaid_nodes.append(f'{mid}[{label}]:::{cls}')

        # Edge lines
        mermaid_edges: List[str] = []
        for ed in edges:
            f = str(ed.get("from", "")).strip()
            t = str(ed.get("to", "")).strip()
            lbl = str(ed.get("label", "")).strip()

            if f not in key_to_mid or t not in key_to_mid:
                continue

            fm = key_to_mid[f]
            tm = key_to_mid[t]

            if lbl:
                lbl = re.sub(r"\s+", " ", lbl)
                # Sanitize edge labels: remove parentheses and special chars that break Mermaid
                lbl = lbl.replace("(", "").replace(")", "").replace("[", "").replace("]", "")
                lbl = lbl.replace('"', "'").replace("|", "-").replace("#", "")
                lbl = lbl.strip()
                if lbl:
                    mermaid_edges.append(f"{fm} -->|{lbl}| {tm}")
                else:
                    mermaid_edges.append(f"{fm} --> {tm}")
            else:
                mermaid_edges.append(f"{fm} --> {tm}")

        # Class definitions (as per your palette)
        class_defs = [
            "classDef input fill:#90EE90,stroke:#333,color:#000;",
            "classDef process fill:#87CEEB,stroke:#333,color:#000;",
            "classDef sorting fill:#FFE97F,stroke:#333,color:#000;",
            "classDef collection fill:#FFB366,stroke:#333,color:#000;",
            "classDef rejection fill:#FFB3B3,stroke:#333,color:#000;",
        ]

        out = []
        out.append("flowchart TD")
        out.extend(mermaid_nodes)
        out.extend(mermaid_edges)
        out.extend(class_defs)
        return "\n".join(out).strip()

    # -----------------------------
    # 2) Build prompt (graph-first)
    # -----------------------------
    process_flow_text = _clean_text(process_flow_text)
    skeleton = _extract_skeleton(process_flow_text)

    system_prompt = """
You are a warehouse automation diagram expert.

TASK:
Convert the given CBS process-flow text into a STRICT directed graph (JSON) that matches real material-flow logic.

OUTPUT RULES (NON-NEGOTIABLE):
- Output ONLY one valid JSON object. No markdown. No commentary. No mermaid.
- JSON schema must be EXACTLY:
{
  "nodes": [
    {"key":"n1","label":"2-5 word label","class":"input|process|sorting|collection|rejection"}
  ],
  "edges": [
    {"from":"n1","to":"n2","label":""}
  ]
}

ACCURACY RULES:
1) Preserve terminology from the text (e.g., "Infeed", "Induct", "Loop CBS", "Rejection Chutes").
2) Build the main spine in the correct order (typical: Sources -> Infeed -> Induct -> CBS -> Outputs).
3) If multiple sources exist (e.g., Marketplace + FC), they MUST MERGE into the Infeed node.
4) If both Auto + Manual induct/infeed exist, they MUST be PARALLEL BRANCHES that MERGE into the CBS node.
5) Output section with subpoints (a/b/c/...) MUST become PARALLEL BRANCHES from the CBS node (or from "Output Chutes"/"Sorting Output" if present).
6) Optional items MUST be shown as a branch edge label "(Optional)".
7) Exception refeeding / recirculation MUST be shown as a LOOP-BACK edge to the correct upstream point (usually Infeed or Induct).
8) Do NOT invent new equipment. If not mentioned, don’t add it.

CLASS RULES (use exactly these class names):
- input: sources / entry points (e.g., Fulfilment Centre, Marketplace, Telescopic Conveyor)
- process: conveyors, infeed lines, VDS loop, induct lines/stations, takeaway conveyors
- sorting: CBS / Loop / main sorter / scanning in sorter
- collection: chutes, PTL sorting, pallets, bags, outbound collection points
- rejection: rejection chutes, manual exception handling, exception refeeding zone

LABEL RULES:
- Keep labels 2–5 words.
- Use Title Case.
- Avoid long sentences.
""".strip()

    user_prompt = f"""
RAW PROCESS FLOW:
{process_flow_text}

EXTRACTED STRUCTURE (HINTS, MUST NOT OVERRIDE RAW TEXT):
{skeleton}

Now output the graph JSON only (nodes + edges) following the rules.
""".strip()

    def api_call_graph():
        return groq_client.chat.completions.create(
            model="groq/compound",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.0,
            max_tokens=2000,
        )

    # -----------------------------
    # 3) Call Groq + parse JSON (with a strict retry)
    # -----------------------------
    resp = call_groq_with_retry(api_call_graph)
    raw = resp.choices[0].message.content.strip()

    try:
        graph_json_str = _extract_json_object(raw)
        graph = json.loads(graph_json_str)
    except Exception:
        # Hard retry with an even stricter "repair" instruction
        repair_system = "You MUST output ONLY valid JSON for the specified schema. No other text."
        repair_user = f"""
Fix the following into ONE valid JSON object that matches the required schema exactly.
Do not add commentary.

BAD_OUTPUT:
{raw}
""".strip()

        def api_call_repair():
            return groq_client.chat.completions.create(
                model="groq/compound",
                messages=[
                    {"role": "system", "content": repair_system},
                    {"role": "user", "content": repair_user},
                ],
                temperature=0.0,
                max_tokens=2000,
            )

        resp2 = call_groq_with_retry(api_call_repair)
        raw2 = resp2.choices[0].message.content.strip()
        graph_json_str = _extract_json_object(raw2)
        graph = json.loads(graph_json_str)

    # -----------------------------
    # 4) Render Mermaid deterministically
    # -----------------------------
    mermaid_code = _render_mermaid(graph)
    return mermaid_code


def extract_full_text_from_docx(doc: Document) -> str:
    """Concatenate all paragraph and table text from the DOCX."""
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def find_terms_in_text(text: str):
    """
    Return list of (term, description) that actually appear in the text.
    Matching is done as a whole word, case-sensitive, to avoid accidental
    matches of 'it' vs 'IT'.
    """
    found = []
    for term, desc in GLOSSARY_ENTRIES:
        # Build safe regex – whole word / token
        pattern = r"(?<![A-Za-z0-9])" + re.escape(term) + r"(?![A-Za-z0-9])"
        if re.search(pattern, text):
            found.append((term, desc))

    # keep original order, no duplicates
    seen = set()
    unique = []
    for t, d in found:
        if t not in seen:
            seen.add(t)
            unique.append((t, d))
    return unique

# ==================== PROPOSED SYSTEM TECHNICAL DETAILS - BOM AGGREGATION ====================

PROPOSED_SYSTEM_TECHNICAL_DETAILS_PROMPT = """
You are a proposal BOM aggregation engine for Cross Belt Sorter projects.

INPUT
- You receive flattened rows from an Excel sheet called "Quote Master".
- Each row looks like:
  "S.No=1 | Category=Conveyor | Description=Straight PVC Conveyor BW 800 | UOM=m | Re=24 | ..."
- Columns can include S.No., Category, Sub-Category, Description, UOM, Costing Type, Cost,
  Unit Price, Sales Factor Requirement, Recommended Quantity, etc.

GOAL
- Build a SHORT "Proposed System Technical Details" section, split into 3 logical sub-sections:
  1) "Mechanical equipment"
  2) "Electrical Equipment"
  3) "Control System"
- The output must be highly aggregated, similar to Falcon proposal tables, NOT one row per part.

GENERAL RULES
- Strongly prefer 3–6 rows in Mechanical, 1 row in Electricals, 1 row in Control System.
- Group fine-grained parts into functional systems.
- DO NOT list every small item separately in the final table.
- DO NOT invent components that are not present in the sheet.

MECHANICAL EQUIPMENT – TYPICAL SYSTEMS
Create a small number of "systems" such as (examples, use only what makes sense):
- "Infeed System" or "Auto Infeed System"
- "Feedlines" or "Auto Induct Feedlines"
- "Sorter – 1 Loop Cross Belt Sorter"
- "Sorter Outputs"
- "Steel Works"
- Any other clear mechanical system suggested by the data.

Mapping guidance:
- Infeed / Auto Infeed System:
  - Straight PVC conveyors, S3 conveyors, gravity rollers, curves, merges, spacing conveyors that
    clearly belong to the main inbound line.
- Feedlines / Auto Induct Feedlines:
  - Receiving, weighing, spacing, buffer conveyors, angle merges that feed the sorter.
- Sorter:
  - Loop CBS sorter, carrier pitch, sorter length/height/speed, type of drive (LIM/LSM),
    supports, fencing, empty-carrier detection, product centring, dimension scanning, etc.
- Sorter Outputs:
  - PTL chutes, rejection chutes, spurs, pop-up sorter units, bag holding assemblies, bins, etc.
- Steel Works:
  - Platforms, mezzanines, stairs, ladders, fencing, safety guards, leg guards, end joints etc.

ELECTRICAL EQUIPMENT
- Usually a single line "Electricals – Consists of".
- Group all power/controls items:
  - main power distribution panel, MCC, main control panel, feedline control panels,
    sorter drive panels, VFD panels, network switches, field cabling, earthing, hooters,
    tower lamps, pull cords, emergency stops, IO cards, surge suppressors, harmonic filters, etc.
- Quantity: typically "1 Set".
- Value: usually just "Included" (unless there is a clear different summary).

CONTROL SYSTEM
- Usually a single line "Components – Consists of".
- Group all PLC / SCADA / IT / software items:
  - PLC based control system, SCADA, industrial switches, servers/IPC, CCTV/VMS,
    OCR/vision PCs, WCS / sorter control software, PTL controllers, licences, custom IT integration, etc.
- Quantity: typically "1 Set".
- Value: use simple summary such as "1 Nos" and "As per requirement" where appropriate.

DATA TO PRODUCE FOR EACH ROW
For each final row in the tables, you must produce:
- pos: integer, starting from 1 within each section.
- qty: short human-readable quantity, e.g. "1 Set", "14 Feedlines", "1 Sorter".
- description_lines: array of short text lines for the Description column:
    * Line 1: main system name (e.g. "Infeed System" or "Sorter").
    * Following lines: bullet-style sub points starting with "• ". These are
      the key components grouped into this system.
- value_lines: array of short text lines for the Value column. Use this for:
    * Total counts or dimensions of important components.
    * Example: "Straight PVC Conveyors: ~38 m total", "Curve Conveyors: 2 Nos".

IMPORTANT: think in terms of systems, not raw rows.

---------------- EXAMPLES (VERY IMPORTANT) ----------------

Example A – Aggregating an Infeed System

INPUT snippet (conceptual):
ROW: Category=Conveyor | Description=Straight PVC Conveyor BW 800 | UOM=m | Re=24
ROW: Category=Conveyor | Description=Straight PVC Conveyor BW 1000 | UOM=m | Re=14
ROW: Category=Conveyor | Description=Curve Conveyor 30 deg | UOM=Nos | Re=2

EXPECTED mechanical item:
{
  "pos": 1,
  "qty": "1 Infeed System",
  "description_lines": [
    "Infeed System",
    "• Straight PVC Conveyors",
    "• Curve Conveyors"
  ],
  "value_lines": [
    "Straight PVC Conveyors: ~38 m total",
    "Curve Conveyors: 2 Nos"
  ]
}

Example B – Aggregating Feedlines

INPUT snippet (conceptual):
Rows describing "Receiving Conveyor", "Weighing Conveyor", "Spacing Conveyor",
"Buffer Conveyor", "Angle merge" with various quantities.

EXPECTED mechanical item:
{
  "pos": 2,
  "qty": "14 Feedlines",
  "description_lines": [
    "Feedlines – Consists of",
    "• Receiving Conveyor",
    "• Weighing Conveyor",
    "• Spacing Conveyor",
    "• Buffer Conveyor",
    "• Angle merge"
  ],
  "value_lines": [
    "Receiving Conveyor: 1 Set",
    "Weighing Conveyor: 1 Set",
    "Spacing Conveyor: 3 Set",
    "Buffer Conveyor: 3 Set",
    "Angle merge: 2 Set"
  ]
}

Example C – Aggregating the Sorter

INPUT snippet (conceptual):
Rows describing a loop CBS sorter with height, length, speed, drive type,
carrier pitch and associated mechanical options.

EXPECTED mechanical item:
{
  "pos": 3,
  "qty": "1 Sorter",
  "description_lines": [
    "Sorter",
    "1 Loop Cross Belt Sorter",
    "• Sorter height",
    "• Sorter length",
    "• Sorter speed",
    "• Sorter drive",
    "• Carrier pitch",
    "Including:",
    "• Standard sorter supports",
    "• Product centring system",
    "• Dimension / barcode scanning system",
    "• Hooters and E-stops",
    "• Fencing"
  ],
  "value_lines": [
    "Height: approx 2900 mm",
    "Loop length: approx 150 m"
  ]
}

Example D – Electricals

INPUT snippet (conceptual):
Rows for power distribution panel, main control panel, feedline control panels,
sorter drive panels, network switches, field cabling, hooters, tower lamps, etc.

EXPECTED electrical section:
{
  "title": "Electrical Equipment",
  "items": [
    {
      "pos": 1,
      "qty": "1 Set",
      "description_lines": [
        "Electricals",
        "Consists of",
        "Main power distribution panel",
        "Main control panel",
        "Feedline control panels",
        "Sorter drive panels",
        "Network switches",
        "Field cabling"
      ],
      "value_lines": [
        "Included"
      ]
    }
  ]
}

Example E – Control System

INPUT snippet (conceptual):
Rows for Siemens PLC, SCADA, industrial switches, servers, sorter control software.

EXPECTED control section:
{
  "title": "Control System",
  "items": [
    {
      "pos": 1,
      "qty": "1 Set",
      "description_lines": [
        "Components",
        "Consists of",
        "PLC based control system with SCADA",
        "Industrial switch"
      ],
      "value_lines": [
        "1 Nos",
        "As per requirement"
      ]
    }
  ]
}

---------------- OUTPUT FORMAT ----------------

Return JSON ONLY in this schema:

{
  "sections": [
    {
      "title": "Mechanical equipment",
      "items": [
        {
          "pos": 1,
          "qty": "1 Set",
          "description_lines": ["..."],
          "value_lines": ["..."]
        }
      ]
    },
    {
      "title": "Electrical Equipment",
      "items": [ ... ]
    },
    {
      "title": "Control System",
      "items": [ ... ]
    }
  ]
}
"""

def call_groq_for_bom(sheet_text: str) -> dict:
    """Call GROQ API to generate aggregated BOM structure from Quote Master sheet."""
    user_prompt = (
        "Below are flattened rows from the 'Quote Master' sheet.\n\n"
        "QUOTE_MASTER_ROWS:\n"
        + sheet_text
    )

    def api_call():
        return groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": PROPOSED_SYSTEM_TECHNICAL_DETAILS_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
        )

    try:
        resp = call_groq_with_retry(api_call)
        return json.loads(resp.choices[0].message.content)
    except Exception as e:
        st.error(f"Failed to generate Proposed System Technical Details: {e}")
        return None


def generate_mechanical_bom_from_costing(costing_file) -> list:
    """
    Generate Mechanical equipment BOM using Groq API for clean table formatting.
    Extracts data from all relevant sheets and uses LLM to aggregate into proper format.
    
    Args:
        costing_file: Streamlit UploadedFile object containing the costing Excel workbook
        
    Returns:
        List of dict items with keys: pos, qty, description, value
        Each item represents ONE row in the final table with multi-line description/value
    """
    
    def clean_float(val, suffix=""):
        """Format float values cleanly with rounding."""
        try:
            if val is None:
                return ""
            if isinstance(val, str):
                val = float(val.replace(",", ""))
            # Round to avoid floating point precision issues
            if val == int(val):
                return f"~{int(val)}{suffix}"
            else:
                return f"~{round(val, 1)}{suffix}"
        except:
            return str(val)
    
    try:
        xlsx_bytes = costing_file.getvalue()
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
        
        # Extract data from all relevant sheets for comprehensive BOM generation
        extracted_data = []
        
        # Detect sheet roles
        roles = bom_module.detect_sheets_by_role(wb)
        
        # 1. Extract Conveyors data
        conv_sheet = roles.get("conveyors")
        if conv_sheet:
            ws = wb[conv_sheet]
            df = bom_module.extract_table(ws, ["name", "conveyor length", "set"])
            if not df.empty:
                df = bom_module.normalize_conveyor_columns(df)
                conv_sum = bom_module.summarize_conveyors(df)
                for _, row in conv_sum.iterrows():
                    name = row.get('name', '')
                    total_m = round(float(row.get('total_m', 0)), 1)
                    units = int(row.get('units', 0))
                    extracted_data.append(f"CONVEYOR: {name} | Length: ~{total_m} m | Units: {units}")
        
        # 2. Extract Loop CBS / Sorter data
        loop_sheets = roles.get("loop_cbs_multi", [])
        for loop_sheet in loop_sheets:
            ws = wb[loop_sheet]
            extracted_data.append(f"\nSORTER_SHEET: {loop_sheet}")
            # Extract key sorter parameters with clean formatting
            for label, patterns, suffix in [
                ("Sorter Height", [r"sorter\s*height", r"height"], "mm"),
                ("Loop Length", [r"loop\s*length", r"sorter\s*length"], " m"),
                ("Carrier Pitch", [r"carrier\s*pitch", r"pitch"], " mm"),
                ("Sorter Speed", [r"speed", r"sorter\s*speed"], " m/s"),
                ("Drive Type", [r"drive", r"sorter\s*drive"], ""),
            ]:
                val = bom_module.find_value_near_label(ws, patterns)
                if val and val.value:
                    raw_val = val.value
                    # Clean the value
                    if bom_module.is_num(raw_val):
                        clean_val = clean_float(raw_val, suffix)
                    else:
                        clean_val = str(raw_val).strip()
                    extracted_data.append(f"SORTER_PARAM: {label} = {clean_val}")
        
        # 3. Extract Destinations / Chutes data
        dest_sheet = roles.get("destinations")
        if dest_sheet:
            ws = wb[dest_sheet]
            df = bom_module.extract_table(ws, ["description", "qty"])
            if not df.empty:
                df.columns = [bom_module.norm(c) for c in df.columns]
                desc_col = next((c for c in df.columns if "description" in c), None)
                qty_col = next((c for c in df.columns if "qty" in c), None)
                if desc_col and qty_col:
                    for _, row in df.iterrows():
                        d = str(row.get(desc_col, "")).strip()
                        q = row.get(qty_col, "")
                        if d and d.lower() not in ['nan', 'none', '']:
                            qty_val = bom_module.safe_int(q, 0) if bom_module.is_num(q) else str(q)
                            extracted_data.append(f"CHUTE: {d} | Qty: {qty_val} Nos")
        
        # 4. Extract PTL data
        ptl_sheet = roles.get("ptl")
        if ptl_sheet:
            ws = wb[ptl_sheet]
            df = bom_module.extract_table(ws, ["description", "qty"])
            if not df.empty:
                df.columns = [bom_module.norm(c) for c in df.columns]
                desc_col = next((c for c in df.columns if "description" in c), None)
                qty_col = next((c for c in df.columns if "qty" in c or "modules" in c), None)
                if desc_col:
                    for _, row in df.iterrows():
                        d = str(row.get(desc_col, "")).strip()
                        q = row.get(qty_col, "") if qty_col else ""
                        if d and d.lower() not in ['nan', 'none', '']:
                            qty_val = bom_module.safe_int(q, 0) if bom_module.is_num(q) else str(q)
                            extracted_data.append(f"PTL: {d} | Qty: {qty_val} Nos")
        
        # 5. Extract Steel Works data
        steel_sheet = roles.get("steelworks")
        if steel_sheet:
            ws = wb[steel_sheet]
            df = bom_module.extract_table(ws, ["description", "qty"])
            if not df.empty:
                df.columns = [bom_module.norm(c) for c in df.columns]
                desc_col = next((c for c in df.columns if "description" in c), None)
                qty_col = next((c for c in df.columns if "qty" in c), None)
                area_col = next((c for c in df.columns if "area" in c), None)
                if desc_col:
                    for _, row in df.iterrows():
                        d = str(row.get(desc_col, "")).strip()
                        if d and d.lower() not in ['nan', 'none', '']:
                            val = "Included"
                            if area_col and bom_module.is_num(row.get(area_col)):
                                val = f"~{round(bom_module.to_float(row.get(area_col), 0), 2)} SQM"
                            elif qty_col and bom_module.is_num(row.get(qty_col)):
                                val = f"{bom_module.safe_int(row.get(qty_col), 0)} Nos"
                            extracted_data.append(f"STEELWORK: {d} | Value: {val}")
        
        # 6. Extract Technical Specifications
        tech_sheet = roles.get("tech_specs")
        if tech_sheet:
            ws = wb[tech_sheet]
            for label, patterns in [
                ("Number of Feedlines", [r"feedline", r"feed\s*line", r"no.*of.*feedlines"]),
                ("Throughput", [r"throughput", r"pph", r"parcels.*per.*hour"]),
                ("Number of Zones", [r"zone", r"no.*of.*zones"]),
            ]:
                val = bom_module.find_value_near_label(ws, patterns)
                if val and val.value:
                    clean_val = bom_module.safe_int(val.value, 0) if bom_module.is_num(val.value) else str(val.value)
                    extracted_data.append(f"TECH_SPEC: {label} = {clean_val}")
        
        if not extracted_data:
            st.warning("No relevant data found in costing sheets for mechanical BOM")
            return []
        
        # Call Groq API to generate clean BOM table
        raw_items = call_groq_for_mechanical_bom("\n".join(extracted_data))
        
        # Post-process to clean up values and ensure proper formatting
        return post_process_mechanical_bom(raw_items)
        
    except Exception as e:
        st.warning(f"Could not generate mechanical BOM from costing file: {e}")
        return []


def format_number_cleanly(val: str) -> str:
    """Format numbers cleanly - round floats, add ~ prefix for approximations."""
    import re
    # Handle floating point precision issues like 149.99999999999997
    pattern = r'(\d+)\.(\d{10,})'
    
    def round_match(m):
        whole = int(m.group(1))
        decimal = m.group(2)
        # Round to nearest integer if very close
        if decimal.startswith('9999') or decimal.startswith('0000'):
            return f"~{round(float(m.group(0)))}"
        else:
            return f"~{round(float(m.group(0)), 1)}"
    
    return re.sub(pattern, round_match, val)


def post_process_mechanical_bom(items: list) -> list:
    """Clean up and format mechanical BOM items for professional output."""
    if not items:
        return []
    
    cleaned_items = []
    for item in items:
        # Clean description - remove placeholder text
        desc = str(item.get("description", ""))
        desc = desc.replace("\\n", "\n")
        
        # Remove placeholder/not found text
        lines_to_remove = [
            "no specific", "not found", "not specified", "no data",
            "no equipment found", "no conveyors found", "no feedline"
        ]
        desc_lines = desc.split("\n")
        clean_desc_lines = []
        for line in desc_lines:
            line_lower = line.lower().strip()
            if not any(phrase in line_lower for phrase in lines_to_remove):
                clean_desc_lines.append(line)
        desc = "\n".join(clean_desc_lines)
        
        # Clean value - format numbers, remove bad values
        val = str(item.get("value", ""))
        val = val.replace("\\n", "\n")
        
        # Format floating point numbers cleanly
        val = format_number_cleanly(val)
        
        # Remove lines with garbage values
        val_lines = val.split("\n")
        clean_val_lines = []
        for line in val_lines:
            line_stripped = line.strip()
            # Skip empty lines, "not specified", garbage numbers
            if not line_stripped:
                continue
            if "not specified" in line_stripped.lower():
                clean_val_lines.append("Included")
                continue
            if "carrier type" in line_stripped.lower():
                continue
            # Check for garbage large numbers (likely cell references or errors)
            try:
                num = float(line_stripped.replace("~", "").replace("m", "").replace("mm", "").strip())
                if num > 100000:  # Likely garbage
                    continue
            except:
                pass
            clean_val_lines.append(line_stripped)
        
        # If no valid values, use "Included"
        if not clean_val_lines:
            clean_val_lines = ["Included"]
        
        val = "\n".join(clean_val_lines)
        
        # Skip items with empty descriptions after cleaning
        if not desc.strip() or desc.strip().lower() in ["", "none", "n/a"]:
            continue
        
        cleaned_items.append({
            "pos": item.get("pos", ""),
            "qty": str(item.get("qty", "1")),
            "description": desc,
            "value": val
        })
    
    return cleaned_items


# Prompt for Mechanical BOM generation using Groq
MECHANICAL_BOM_PROMPT = """
You are a proposal BOM table generator for Cross Belt Sorter (CBS) projects at Falcon Autotech.

TASK: Generate a clean "Mechanical equipment" table from the extracted Excel data.

OUTPUT FORMAT - Each item MUST have exactly these 4 fields:
- pos: Position number (integer: 1, 2, 3, etc.)
- qty: Quantity as string (e.g., "1", "3", "12")
- description: Multi-line text with title + bullet points
- value: Corresponding values for each bullet point OR "Included"

CRITICAL FORMATTING RULES:

1. DESCRIPTION COLUMN FORMAT:
   - First line: Main system title (bold heading)
   - Second line (optional): Sub-title like "Consists of" or "1 Loop Cross Belt Sorter"
   - Bullet points: Start with "• " (bullet + space)
   - For conveyors with measurements: "• PVC belt Conveyor- 394.8 m; 39 Modules"
   - For general items: "• Buffer Conveyor" or "• Standard Sorter Supports"

2. VALUE COLUMN FORMAT:
   - Each value corresponds to a description line
   - Use actual measurements when available (e.g., "~38 metres (9 Modules)", "2900mm", "~150m")
   - Use "X Set" or "X No" for quantities (e.g., "4 Set", "1 No", "50 Nos")
   - Use "Included" when no specific value is available
   - Round all numbers: No decimals beyond 1 place, use ~ for approximations
   - NEVER use placeholder text like "Not specified" or "No data found"

3. STANDARD POSITIONS:
   Pos 1: Auto Infeed System - PVC conveyors, curve conveyors, modular conveyors
   Pos 2: Auto Induct Feedlines - Buffer, angle merge, spacing conveyors
   Pos 3: Sorter - Loop Cross Belt Sorter with all parameters and inclusions
   Pos 4: Sorter Outputs - Chutes (sliding, rejection, double deck)
   Pos 5: PTL - Pick to Light systems (if present)
   Pos 6: Steel Works - Operator platforms, structures (if present)

4. VALUE ALIGNMENT EXAMPLES:
   Description: "• Sorter Height"     → Value: "2900mm"
   Description: "• Sorter Length"     → Value: "~150m"
   Description: "• Buffer Conveyor"   → Value: "4 Set"
   Description: "• Angle merge"       → Value: "1 Set"
   Description: "• Standard Sorter Supports" → Value: "Included"

EXAMPLE OUTPUT:
{
  "items": [
    {
      "pos": 1,
      "qty": "1",
      "description": "Auto Infeed System\\n• Powered Belt Conveyors\\n• Curve Conveyor",
      "value": "~38 metres (9 Modules)\\n1 No"
    },
    {
      "pos": 2,
      "qty": "3",
      "description": "Auto Induct Feedlines\\nConsists of\\n• Buffer Conveyor\\n• Angle merge\\n• Spacing Conveyor",
      "value": "4 Set\\n1 Set\\n2 Set"
    },
    {
      "pos": 3,
      "qty": "1",
      "description": "Sorter\\n1 Loop Cross Belt Sorter\\n• Sorter Height\\n• Sorter Length\\nIncluding:\\n• Standard Sorter Supports\\n• Product Centring System\\n• 5 Side Scanning System\\n• Hooter\\n• Emergency Stop Buttons",
      "value": "2900mm\\n~150m"
    }
  ]
}

IMPORTANT:
- Use \\n for newlines in JSON strings
- Round all decimal numbers (149.99999 → ~150)
- NEVER include "No data found", "Not specified", or similar placeholder text
- If data is missing for an item, either skip it OR use "Included" as value
- Match the number of value lines to the number of measurable description lines
"""


def call_groq_for_mechanical_bom(extracted_data: str) -> list:
    """Call Groq API to generate clean Mechanical BOM from extracted sheet data."""
    user_prompt = f"""
Generate a Mechanical equipment BOM table from the following extracted costing data.

EXTRACTED DATA:
{extracted_data}

INSTRUCTIONS:
1. Create 4-7 position rows (Infeed, Feedlines, Sorter, Outputs, PTL, Steel Works)
2. Use actual values from data - round numbers properly (e.g., 149.99 → ~150)
3. Use "Included" for items without specific measurements
4. Format as JSON with "items" array

Return ONLY valid JSON. No explanation text.
"""

    def api_call():
        return groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": MECHANICAL_BOM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
        )

    try:
        resp = call_groq_with_retry(api_call)
        result = json.loads(resp.choices[0].message.content)
        # Handle both {"items": [...]} and direct array formats
        if isinstance(result, list):
            return result
        elif isinstance(result, dict) and "items" in result:
            return result["items"]
        elif isinstance(result, dict):
            # Try to find any array in the result
            for key, val in result.items():
                if isinstance(val, list):
                    return val
        return []
    except Exception as e:
        st.warning(f"Failed to generate Mechanical BOM via Groq: {e}")
        return []


# ==================== GROQ PROMPTS & CONSTANTS ====================

# Cover Letter System Prompt
COVER_LETTER_SYSTEM_PROMPT = """
You are an AI assistant working as a professional proposal writer at Falcon Autotech. You are an expert in drafting formal, client-specific techno-commercial cover letters for proposals. Your role is to generate well-structured, personalized cover letters that follow Falcon's business communication style, maintain a professional and respectful tone, and clearly demonstrate Falcon's commitment, expertise, and partnership approach to clients.

Generate a formal techno-commercial COVER LETTER for a proposal that MUST fit within a single page. 
The writing style MUST be indistinguishable from natural human writing. The text should read as if drafted by an experienced professional, not an AI system. Use clear, simple, and natural language with varied sentence lengths and structures. Avoid generic phrases, repetitive patterns, or mechanical tone. Ensure that the output flows smoothly, conveys intent naturally, and would not be detected as machine-generated. The content should feel thoughtful, context-aware, and aligned with how a human proposal writer or business professional would communicate.

STRICT LENGTH LIMIT: Maximum 300 words to ensure single-page fit. Be concise and impactful.

1. Start with:
   Kind Attention –
   Mr. {{executives}}
   M/s {{client_name}}

   Offer Ref: {{offer_ref}}; Date: {{letter_date}}

   Subject – Techno-Commercial Offer for {{project_title}}  

2. If there is only one executive, address them with:
   Dear {{first_exec_name}},
   If multiple executives, skip "Dear" and go directly to the content.
   Use Mr. for male and Ms. for female executives.

3. Opening paragraph (natural, professional style):
   - Thank the client for inviting Falcon to offer for the project.
   - If invitation_date or meeting_date exists, reference it naturally (e.g., "Over the past period we worked closely together" or "In our meeting on [date], we discussed...").
   - Mention that you are pleased to submit the Techno-Commercial Offer.
   - Wording must vary between runs (not fixed sentences).

4. Middle paragraph - System Overview & Analysis (CRITICAL):
   - State that Falcon has done an in-depth data analysis and evaluated various solution options.
   - **MANDATORY: Include the high-level process flow summary if provided**. Mention key system components naturally in a single sentence (e.g., "The proposed solution includes automatic induct conveyors, {cbs_type} for efficient sortation, and output chutes for sorted parcel handling").
   - **CRITICAL: Use the EXACT CBS TYPE** from the context (either "Linear CBS" or "Loop CBS") - do NOT use generic "cross-belt sorter"
   - Highlight any specific technical values, quantities, or capacities if mentioned (e.g., "200 destinations", "5 camera scanner systems", "2 speed settings").
    - **MANDATORY PPH LINE**: If a PPH value is provided in the context, include a clear sentence such as "From a technical point of view, the system is designed at {PPH} PPH and ensures simple operations and movement within the facility." Use the actual PPH number from the provided data; if not provided, do not invent a value and skip this sentence.
   - Keep this brief but informative - demonstrate technical understanding without overwhelming detail.
   - Mention that the detailed technical proposal is laid out in various sections to provide full insight into the proposed solution.

5. Commitment paragraph:
   - Reference Falcon's intralogistics automation technologies and proven track record.
   - Highlight subsequent sections covering capabilities, experiences, and references.
   - Reinforce commitment to being a strategic partner.

6. Closing (professional, warm):
   - Add sender's personal commitment on behalf of Falcon Autotech.
   - Encourage client to reach out for clarifications or further information.
   - End with "Best Regards," followed by sender_name and sender_title.

Important:
- MUST NOT EXCEED 300 words to ensure single-page fit.
- Keep tone formal, professional, and client-oriented.
- Do not copy exact sentences; rephrase wording across generations.
- The cover letter MUST sound human, natural and professional. It should be clear, authentic, and warm, without feeling robotic or overly formal.
- DO NOT ADD ANY EXTRA WORD OR INFO APART FROM THE COVER LETTER.
- Highlight the main system or project name in main body (not subject line) as bold style, use ** for Bold.
- If process_flow_summary is provided, ALWAYS incorporate it naturally into the letter.
"""

COVER_LETTER_USER_PROMPT_TEMPLATE = """
Use the following information to generate the cover letter:

client_name: {client_name}
project_title: {project_title}
offer_ref: {offer_ref}
letter_date: {letter_date}

executives (one per line, already with Mr./Ms. prefix):
{executives_block}

invitation_date: {invitation_date}
meeting_date: {meeting_date}

process_flow_summary (very high-level system components and key quantities): {process_flow_summary}

sender_name: {sender_name}
sender_title: {sender_title}

CRITICAL REQUIREMENTS:
1. The cover letter MUST fit within a single page (maximum 300 words).
2. If process_flow_summary is provided, ALWAYS incorporate it naturally into the letter body to demonstrate technical understanding.
3. Mention any specific quantities or technical details from the summary to add credibility.
4. If PPH (parcels per hour) is present in the provided data/counts, include a sentence in the technical paragraph: "From a technical point of view, the system is designed at {PPH} PPH and ensures simple operations and movement within the facility." Use the actual PPH value; if PPH is missing, skip this sentence (do not invent values).
5. Keep the tone professional, warm, and client-focused like the example letter provided.

Return ONLY the cover letter text, without markdown code fences or extra commentary.
"""

# Executive Summary System Prompt
EXEC_SUMMARY_SYSTEM_PROMPT = """
You are a Proposal Writing Assistant specialized in Falcon Autotech automation projects.  
Falcon Autotech designs, manufactures, supplies, implements, and maintains warehouse automation solutions—such as sortation systems, conveyor automation, pick/put-to-light, ASRS robotics, and dimension & weight scanning—for industries including e-commerce, fashion, FMCG, pharma, groceries, and CE-P.

The writing style must be indistinguishable from natural human writing. The text should read as if drafted by an experienced proposal engineer, not an AI system. Use clear, professional language with varied sentence structures.

Your task is to generate **unique, client-tailored Executive Summaries** based on the "Proposed System Description" section of Falcon proposals.

### EXACT STRUCTURE TO FOLLOW

**OPENING PARAGRAPH** (combine into ONE flowing paragraph)
Write as a single connected paragraph containing:
- "Falcon is pleased to confirm its great interest in responding to this RFQ."
- "Our team has been working closely with the relevant stakeholders, with a clear commitment to listening and understanding your needs and ensuring this project's success."
- "As prime contractor, Falcon ensures its full commitment to successfully completing this project."
- "Following the same objective for the system, we are happy to offer a compliant solution meeting all technical and operational requirements, high-performance, optimized, tailor-made, fast and secure planning, and a competitive price."

**KEY CHARACTERISTICS SECTION**
- Start with EXACTLY ONE LINE: "Our solution is based on the following key characteristics:"
- DO NOT REPEAT THIS LINE
- Then provide 4-6 bullet points as FULL DESCRIPTIVE SENTENCES

**BULLET POINT RULES - CRITICAL:**
Each bullet must be a COMPLETE SENTENCE with context (not just item names). Use these patterns:

BULLET 1 - CBS/Sorter (ALWAYS FIRST):
- Use EXACT CBS TYPE from metadata: "Linear CBS for parcel sorting..." or "Loop CBS for parcel sorting..."
- Full format: "[Linear CBS/Loop CBS] for parcel sorting, employing cross-belt technology, has been designed to handle a throughput of [PPH] packages per hour."
- CRITICAL: Do NOT use generic "[Loop CBS / Linear CBS]" - use the SPECIFIC type from the project

BULLET 2 - Induct System (combine feedlines + manual in ONE bullet):
- Use EXACT CBS TYPE: "This Linear CBS is equipped with..." or "This Loop CBS is equipped with..."
- Full format: "This [Linear CBS/Loop CBS] is equipped with [COUNT] Nos fully automatic induct lines, along with manual loading point."
- OR "The system features [COUNT] induct lines for automated parcel induction, with provision for manual loading."
- NOTE: Do NOT count manual stations separately. Just mention "manual loading point" as part of induct bullet.
- CRITICAL: Use the SPECIFIC CBS type from the project, not generic "[Loop/Linear]"

BULLET 3 - Conveyor/Transport System:
- "The system is having its own conveyor connection to transport the volume from primary loading points to the induction zone."
- OR "Infeed conveyor system including [COUNT] telescopic conveyors for efficient parcel receiving."

BULLET 4 - Chutes (combine all chute types in ONE bullet):
- "In the system, there are [X] gravity chutes, [Y] mini-gravity chutes, [Z] rejection chutes, and [W] bulk chute for efficient parcel distribution."
- OR "The system includes [TOTAL] chutes comprising gravity, collection, and rejection chutes for sorted parcel handling."

BULLET 5 - Layout/Operations (if needed):
- "The system layout has been meticulously planned to facilitate smooth operational flow, ensuring efficient movement of personnel."

**CLOSING SECTION** (separate paragraph after bullets)
"A tailor-made and simple layout, specifically designed to [CLIENT NAME]. The proposed layout is the result of the technical requirements in the RFP document and our discussions with the relevant stakeholders during our site visit and Teams workshop meeting."
Then add as dash points:
- Simple operational conditions due to one single [loop/linear] cross belt sorter.
- Easy maintenance: optimized number of conveyors and concentrated inducts area.

**FIXED SECTION** (MUST PASTE EXACTLY AS BELOW AFTER CLOSING)
**1. Falcon's reliable Shipment sortation systems** 
These systems are globally being used by most innovative brands such as Amazon, Flipkart, Delhivery, Asendia, Fastway and many more. The main and critical components of the FALCON Autotech sorter building blocks, like wheels, motors, belts, bearings, Bus Bars, Communication platforms, PLCs etc., are sourced from some of the best suppliers in the world, such as SEW, Siemens, SICK, Faigle, Vahle and Forbo. This strategic baseline of sourcing policy allows Falcon's customers to be fully confident in the systems' robustness and reliability.

**2. Commitment to quality systems** 
Demonstrating Falcon's clear commitment to the {client_name}'s satisfaction, the shipment sortation system, parts and services will be under warranty for 12 months from installation go-live. 


### CRITICAL RULES
1. DO NOT repeat "Our solution is based on the following key characteristics:" - write it only ONCE
2. Each bullet must be a FULL SENTENCE with explanation, not just item names
3. DO NOT count manual induct stations - just mention "manual loading point" as part of another bullet
4. Combine related items: feedlines + manual in one bullet, all chutes in one bullet
5. Use "Induct Lines" not "Feedlines" in the output
6. Throughput format: "[NUMBER] packages per hour" or "[NUMBER] pph"

### GOOD EXAMPLES OF BULLETS:
• "Loop CBS for parcel sorting, employing cross-belt technology, has been designed to handle a throughput of 10,000 packages per hour."
• "This Loop CBS is equipped with 3 Nos fully automatic induct lines, along with manual loading point."
• "The system is having its own conveyor connection to transport the volume from primary loading points to the induction zone."
• "In the system, there are 58 gravity chutes, 70 mini-gravity chutes, 7 rejection chutes, and 1 bulk chute for efficient parcel distribution."
• "The system layout has been meticulously planned to facilitate smooth operational flow, ensuring efficient movement of personnel."

### BAD EXAMPLES (DO NOT USE):
• "3 Nos of Feedlines" ❌ (too short, no context)
• "32 Manual Induct Stations" ❌ (don't count manual stations)
• "202 generic chutes" ❌ (too vague, add context)
• "Sorter, based on a cross-belt technology, offers a designed throughput of 1200 pph." ❌ (too short)

### OUTPUT FORMAT
1. Opening paragraph (all sentences combined into one flowing paragraph)
2. "Our solution is based on the following key characteristics:" (ONLY ONCE)
3. 4-6 descriptive bullet points as full sentences
4. Closing paragraph with tailor-made statement
5. Dash points for operational simplicity
6. Fixed sections (Falcon reliability + warranty)

DO NOT ADD ANY EXTRA TEXT OR INFORMATION OR JUSTIFICATION or "Here is an Executive Summary for the proposal:" EXCEPT THE FULL PROPOSAL
"""

# Config paths
STATIC_ABOUT_DIR = r"Static_AboutCompany"

# Handled Shipment Spectrum Templates
@dataclass
class SorterTemplate:
    key: str
    label: str
    keywords: List[str]
    config_name: str
    item_singular: str
    subheading_51: str
    spec_table: Dict[str, Dict[str, str]]

SORTER_TEMPLATES: List[SorterTemplate] = [
    SorterTemplate(
        key="linear_dual_standard",
        label="Linear / Dual-belt CBS – standard boxes",
        keywords=["linear", "6k", "5.4k", "loop cbs + linear", "totes", "boxes"],
        config_name="Linear Cross Belt Sorter (Dual-belt configuration)",
        item_singular="shipment",
        subheading_51="Shipment size loadable on the sorter",
        spec_table={
            "Max Length": {"unit": "mm", "value": "600"},
            "Max Width":  {"unit": "mm", "value": "450"},
            "Max Height": {"unit": "mm", "value": "400"},
            "Max Weight": {"unit": "Kg", "value": "20"},
            "Min length": {"unit": "mm", "value": "100"},
            "Min Width":  {"unit": "mm", "value": "100"},
            "Min Height": {"unit": "mm", "value": "3"},
            "Min Weight": {"unit": "gm", "value": "50"},
        },
    ),
    SorterTemplate(
        key="loop_standard",
        label="Loop CBS – standard shipments",
        keywords=["loop", "double deck", "48k", "loop cbs", "main sorter"],
        config_name="Loop Cross Belt Sorter technology",
        item_singular="shipment",
        subheading_51="Shipment size loadable on the sorter",
        spec_table={
            "Max Length": {"unit": "mm", "value": "400"},
            "Max Width":  {"unit": "mm", "value": "400"},
            "Max Height": {"unit": "mm", "value": "400"},
            "Max Weight": {"unit": "Kg", "value": "40"},
            "Min length": {"unit": "mm", "value": "10"},
            "Min Width":  {"unit": "mm", "value": "100"},
            "Min Height": {"unit": "mm", "value": "50"},
            "Min Weight": {"unit": "gm", "value": "100"},
        },
    ),
    SorterTemplate(
        key="heavy_parcel",
        label="Heavy-duty CBS – parcels / bags & boxes",
        keywords=["parcel", "heavy", "bags", "bag and box", "bosta", "delhivery"],
        config_name="Heavy Duty Cross Belt Sorter",
        item_singular="parcel",
        subheading_51="Parcel size loadable on the sorter",
        spec_table={
            "Max Length": {"unit": "mm", "value": "1000"},
            "Max Width":  {"unit": "mm", "value": "800"},
            "Max Height": {"unit": "mm", "value": "800"},
            "Max Weight": {"unit": "Kg", "value": "50"},
            "Min length": {"unit": "mm", "value": "40"},
            "Min Width":  {"unit": "mm", "value": "150"},
            "Min Height": {"unit": "mm", "value": "150"},
            "Min Weight": {"unit": "Kg", "value": "0.05"},
        },
    ),
]

st.set_page_config(
    page_title="Falcon Proposal Generator", 
    page_icon="F", 
    layout="wide"
)

# ==================== USER CREDENTIALS DATABASE ====================
USER_CREDENTIALS = {
    "sanyog.singh@falconautotech.com": {
        "name": "Sanyog Pratap Singh",
        "role": "Assistant Manager",
        "password": "1234",
        "initials": "SS"
    },
    "tuhi@falconautoonline.com": {
        "name": "Tuhi",
        "role": "General Manager",
        "password": "1234",
        "initials": "TU"
    },
    "sandeep.pathak@falconautoonline.com": {
        "name": "Sandeep Pathak",
        "role": "Deputy General Manager",
        "password": "1234",
        "initials": "SP"
    },
    "balmukund.mishra@falconautotech.com": {
        "name": "Balmukund Mishra",
        "role": "Principal Engineer",
        "password": "1234",
        "initials": "BM"
    },
    "arkajyoti.chakraborty@falconautotech.com": {
        "name": "Arkajyoti Chakraborty",
        "role": "Senior Engineer",
        "password": "1234",
        "initials": "AC"
    }
}

# ==================== LOGIN FUNCTIONS ====================
def authenticate_user(email: str, password: str) -> dict | None:
    """Authenticate user and return user info if valid."""
    email_lower = email.lower().strip()
    if email_lower in USER_CREDENTIALS:
        user = USER_CREDENTIALS[email_lower]
        if user["password"] == password:
            return {
                "email": email_lower,
                "name": user["name"],
                "role": user["role"],
                "initials": user["initials"]
            }
    return None

def logout_user():
    """Clear all session state and logout user."""
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.session_state.logged_in = False
    st.session_state.current_user = None

def show_login_page():
    """Display a polished enterprise login page."""
    import base64
    
    # Load logo
    logo_path = r".\assests\Images\falcon-autotech-icon-removebg-preview.png"
    try:
        with open(logo_path, "rb") as img_file:
            logo_base64 = base64.b64encode(img_file.read()).decode()
        logo_src = f"data:image/png;base64,{logo_base64}"
    except Exception:
        logo_src = ""

    # Hide Streamlit chrome and style the page
    st.markdown(
        """
        <style>
            #MainMenu, footer, header {visibility: hidden !important;}
            section[data-testid="stSidebar"] {display: none !important;}
            .stApp {background: linear-gradient(165deg, #0B1426 0%, #152238 40%, #1a2d4a 70%, #0B1426 100%);}
            .block-container {padding: 40px !important; max-width: 1300px !important; margin: 0 auto !important;}
            .stMainBlockContainer {padding: 40px !important;}
            [data-testid="stVerticalBlock"] {gap: 0.5rem !important;}
            [data-testid="stHorizontalBlock"] {align-items: center !important;}
            
            /* Hide the iframe border */
            iframe {border: none !important;}
            
            /* Form styling */
            [data-testid="stForm"] {
                background: rgba(255,255,255,0.03) !important;
                backdrop-filter: blur(20px) !important;
                border: 1px solid rgba(255,255,255,0.08) !important;
                border-radius: 20px !important;
                padding: 36px 32px !important;
                box-shadow: 0 25px 80px rgba(0,0,0,0.3) !important;
            }
            [data-testid="stForm"] .stTextInput > div > div > input {
                background: rgba(255,255,255,0.12) !important;
                border: 1px solid rgba(255,255,255,0.2) !important;
                border-radius: 8px !important;
                padding: 14px 16px !important;
                font-size: 15px !important;
                color: #0f172a !important; /* solid dark text for readability */
                transition: all 0.2s ease !important;
            }
            [data-testid="stForm"] .stTextInput > div > div > input::placeholder {
                color: rgba(255,255,255,0.4) !important;
            }
            [data-testid="stForm"] .stTextInput > div > div > input:focus {
                border-color: #F9D423 !important;
                box-shadow: 0 0 0 2px rgba(249, 212, 35, 0.2) !important;
                background: rgba(255,255,255,0.12) !important;
            }
            [data-testid="stForm"] .stTextInput > label {
                color: rgba(255,255,255,0.8) !important;
                font-weight: 500 !important;
                font-size: 13px !important;
            }
            [data-testid="stForm"] button[kind="primaryFormSubmit"], 
            [data-testid="stForm"] button[kind="primary"] {
                background: linear-gradient(110deg, #F9D423 0%, #FF9800 100%) !important;
                border: none !important;
                border-radius: 8px !important;
                padding: 14px 24px !important;
                font-size: 15px !important;
                font-weight: 600 !important;
                color: #1a2b4a !important;
                box-shadow: 0 4px 20px rgba(249, 212, 35, 0.3) !important;
                transition: all 0.2s ease !important;
                margin-top: 8px !important;
                width: 100% !important;
            }
            [data-testid="stForm"] button:hover {
                transform: translateY(-2px) !important;
                box-shadow: 0 8px 30px rgba(249, 212, 35, 0.4) !important;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Two-column layout
    col_left, col_right = st.columns([1.4, 1], gap="large")

    with col_left:
        # Hero content with embedded particles
        hero_html = """
            <style>
                * {margin: 0; padding: 0; box-sizing: border-box;}
                .hero-container {
                    position: relative;
                    min-height: 500px;
                    padding: 40px 20px 40px 0;
                    overflow: visible;
                }
                #particles-js {
                    position: absolute;
                    top: -60px;
                    left: -60px;
                    right: -200px;
                    bottom: -60px;
                    z-index: 1;
                }
                .hero-content {
                    position: relative;
                    z-index: 10;
                }
                .version-badge {
                    display: inline-flex;
                    padding: 8px 16px;
                    background: linear-gradient(135deg, #F9D423 0%, #FF9800 100%);
                    border-radius: 6px;
                    color: #1a2b4a;
                    font-weight: 700;
                    font-size: 11px;
                    letter-spacing: 1.5px;
                    text-transform: uppercase;
                    margin-bottom: 24px;
                    box-shadow: 0 4px 20px rgba(249, 212, 35, 0.3);
                }
                .hero-title {
                    font-size: 52px;
                    font-weight: 700;
                    color: #ffffff;
                    line-height: 1.15;
                    margin-bottom: 20px;
                    letter-spacing: -1px;
                    font-family: 'Segoe UI', system-ui, sans-serif;
                }
                .hero-title span {
                    background: linear-gradient(135deg, #F9D423 0%, #FF9800 100%);
                    -webkit-background-clip: text;
                    -webkit-text-fill-color: transparent;
                    background-clip: text;
                }
                .hero-subtitle {
                    font-size: 17px;
                    color: rgba(255,255,255,0.65);
                    line-height: 1.7;
                    margin-bottom: 36px;
                    max-width: 480px;
                    font-family: 'Segoe UI', system-ui, sans-serif;
                }
                .features-row {display: flex; flex-wrap: wrap; gap: 12px;}
                .feature-tag {
                    display: inline-flex;
                    align-items: center;
                    gap: 8px;
                    padding: 12px 18px;
                    background: rgba(255,255,255,0.06);
                    border: 1px solid rgba(255,255,255,0.1);
                    border-radius: 10px;
                    color: rgba(255,255,255,0.85);
                    font-size: 14px;
                    font-weight: 500;
                    font-family: 'Segoe UI', system-ui, sans-serif;
                    backdrop-filter: blur(10px);
                }
                .feature-tag svg {width: 18px; height: 18px; stroke: #F9D423; fill: none; stroke-width: 2;}
            </style>
            
            <div class="hero-container">
                <div id="particles-js"></div>
                <div class="hero-content">
                    <div class="version-badge">Pro Version</div>
                    <h1 class="hero-title">Falcon <span>Proposal</span><br>Generator</h1>
                    <p class="hero-subtitle">Enterprise-grade proposal automation powered by AI. Transform DXF layouts into comprehensive technical documents with smart BOM generation.</p>
                    <div class="features-row">
                        <div class="feature-tag">
                            <svg viewBox="0 0 24 24"><path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/></svg>
                            AI Proposal Writer
                        </div>
                        <div class="feature-tag">
                            <svg viewBox="0 0 24 24"><rect x="3" y="3" width="18" height="18" rx="2"/><path d="M3 9h18M9 21V9"/></svg>
                            DXF Analysis
                        </div>
                        <div class="feature-tag">
                            <svg viewBox="0 0 24 24"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><path d="M14 2v6h6M16 13H8M16 17H8M10 9H8"/></svg>
                            Smart BOM
                        </div>
                        <div class="feature-tag">
                            <svg viewBox="0 0 24 24"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
                            PDF Export
                        </div>
                    </div>
                </div>
            </div>
            
            <script src="https://cdn.jsdelivr.net/npm/tsparticles@1.40.1/tsparticles.min.js"></script>
            <script>
                tsParticles.load("particles-js", {
                    background: { color: { value: "transparent" } },
                    fpsLimit: 60,
                    particles: {
                        color: { value: ["#F9D423", "#5a7ec2", "#ffffff"] },
                        links: { enable: true, color: "#5a7ec2", distance: 150, opacity: 0.35, width: 1 },
                        move: { enable: true, speed: 0.7, direction: "none", random: true, outModes: { default: "out" } },
                        number: { value: 80, density: { enable: true, area: 1000 } },
                        opacity: { value: 0.7 },
                        shape: { type: "circle" },
                        size: { value: { min: 1, max: 2.5 } }
                    },
                    detectRetina: true
                });
            </script>
        """
        html(hero_html, height=520)

    with col_right:
        # Login header
        st.markdown(
            f"""
            <div style="text-align: center; margin-bottom: 24px; padding-top: 20px;">
                {'<img style="width: 56px; height: 56px; object-fit: contain; margin: 0 auto 16px; display: block;" src="' + logo_src + '" />' if logo_src else ''}
                <h2 style="color: #ffffff; margin-bottom: 6px; font-weight: 700; font-size: 24px;">Welcome Back</h2>
                <p style="color: rgba(255,255,255,0.6); font-size: 14px; margin: 0;">Sign in to your workspace</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

        with st.form("login_form", clear_on_submit=False):
            email = st.text_input("Email Address", placeholder="Enter your email")
            password = st.text_input("Password", type="password", placeholder="Enter your password")
            submit = st.form_submit_button("Sign In", use_container_width=True, type="primary")

            if submit:
                if not email or not password:
                    st.error("Please enter both email and password")
                else:
                    user = authenticate_user(email, password)
                    if user:
                        st.session_state.logged_in = True
                        st.session_state.current_user = user
                        st.rerun()
                    else:
                        st.error("Invalid email or password")

    # Full-width centered footer
    st.markdown(
        """
        <div style="position: fixed; bottom: 0; left: 0; right: 0; text-align: center; padding: 20px; z-index: 100;">
            <p style="color: rgba(255,255,255,0.7); font-size: 12px; font-weight: 600; margin: 4px 0;">Falcon Autotech Pvt. Ltd.</p>
            <p style="color: rgba(255,255,255,0.5); font-size: 11px; margin: 4px 0; display: flex; align-items: center; justify-content: center; gap: 6px;">
                <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#28a745" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round">
                    <path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/>
                    <polyline points="9 12 11 14 15 10"/>
                </svg>
                Secured with enterprise-grade encryption
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ==================== CHECK LOGIN STATE ====================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.current_user = None

# Show login page if not logged in
if not st.session_state.logged_in:
    show_login_page()
    st.stop()

# Get current user info for the rest of the app
current_user = st.session_state.current_user

# ==================== PROFESSIONAL CSS STYLING ====================
st.markdown("""
<style>
    /* ===== ROOT VARIABLES ===== */
    :root {
        --primary-blue: #060c71;
        --secondary-blue: #2a3bb8;
        --accent-gold: #f9d20e;
        --success-green: #28a745;
        --text-dark: #1a1a2e;
        --text-muted: #6c757d;
        --bg-light: #f8f9fa;
        --bg-white: #ffffff;
        --border-color: #e0e4e8;
        --shadow-sm: 0 2px 8px rgba(0,0,0,0.08);
        --shadow-md: 0 4px 16px rgba(0,0,0,0.12);
        --shadow-lg: 0 8px 32px rgba(0,0,0,0.16);
        --transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    /* ===== GLOBAL STYLES ===== */
    .main .block-container {
        padding: 2rem 2rem 2rem 2rem;
        max-width: 900px;
    }
    
    /* ===== MAIN HEADER ===== */
    .main-header {
        background: linear-gradient(105deg, #0a1628 0%, #1a3a5c 25%, #1f3864 45%, #2d5a87 65%, #d4a012 85%, #f9d20e 100%);
        padding: 2.5rem 3rem;
        border-radius: 20px;
        margin-bottom: 2rem;
        box-shadow: 0 20px 60px rgba(10, 22, 40, 0.4), 0 8px 25px rgba(249, 210, 14, 0.15), inset 0 1px 0 rgba(255,255,255,0.1);
        position: relative;
        overflow: hidden;
        border: 1px solid rgba(255,255,255,0.1);
    }
    
    /* Wave animation keyframes */
    @keyframes waveFlow {
        0% {
            transform: translateX(-100%) skewX(-15deg);
        }
        100% {
            transform: translateX(200%) skewX(-15deg);
        }
    }
    
    @keyframes colorPulse {
        0%, 100% {
            opacity: 0.3;
        }
        50% {
            opacity: 0.6;
        }
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 50%;
        height: 200%;
        background: linear-gradient(
            90deg,
            transparent 0%,
            rgba(249, 210, 14, 0.15) 25%,
            rgba(249, 210, 14, 0.35) 50%,
            rgba(249, 210, 14, 0.15) 75%,
            transparent 100%
        );
        animation: waveFlow 4s ease-in-out infinite;
        pointer-events: none;
    }
    
    .main-header::after {
        content: '';
        position: absolute;
        right: 0;
        top: 0;
        bottom: 0;
        width: 35%;
        background: linear-gradient(105deg, transparent 0%, rgba(249, 210, 14, 0.2) 50%, rgba(249, 210, 14, 0.35) 100%);
        clip-path: polygon(30% 0, 100% 0, 100% 100%, 0% 100%);
        animation: colorPulse 3s ease-in-out infinite;
    }
    
    .main-header h1 {
        color: #ffffff !important;
        font-size: 2.4rem;
        font-weight: 800;
        margin: 0;
        letter-spacing: -0.5px;
        text-shadow: 0 2px 4px rgba(0,0,0,0.4), 0 4px 20px rgba(0,0,0,0.2);
        position: relative;
        z-index: 2;
    }
    
    .main-header .subtitle {
        color: rgba(255,255,255,0.95);
        font-size: 1.1rem;
        margin-top: 0.6rem;
        font-weight: 400;
        position: relative;
        z-index: 2;
        letter-spacing: 0.5px;
        text-shadow: 0 1px 3px rgba(0,0,0,0.3);
    }
    
    .main-header .version-badge {
        display: inline-block;
        background: linear-gradient(135deg, #f9d20e 0%, #ffc107 50%, #ffab00 100%);
        color: #0a1628;
        padding: 0.3rem 1rem;
        border-radius: 25px;
        font-size: 0.8rem;
        font-weight: 800;
        margin-left: 1rem;
        vertical-align: middle;
        box-shadow: 0 4px 15px rgba(249, 210, 14, 0.4);
        text-transform: uppercase;
        letter-spacing: 1px;
        -webkit-text-fill-color: #0a1628;
        position: relative;
        z-index: 3;
    }
    
    /* Decorative elements */
    .main-header .header-icon {
        position: absolute;
        right: 40px;
        top: 50%;
        transform: translateY(-50%);
        font-size: 5rem;
        opacity: 0.15;
        z-index: 1;
        filter: drop-shadow(0 0 20px rgba(249, 210, 14, 0.3));
    }
    
    /* ===== SECTION HEADERS ===== */
    .section-header {
        background: linear-gradient(135deg, var(--bg-white) 0%, var(--bg-light) 100%);
        border-left: 5px solid var(--primary-blue);
        padding: 1rem 1.5rem;
        border-radius: 0 12px 12px 0;
        margin: 2rem 0 1.25rem 0;
        box-shadow: var(--shadow-sm);
        display: flex;
        align-items: center;
        gap: 0.75rem;
    }
    
    .section-header h3 {
        color: var(--primary-blue);
        font-weight: 700;
        margin: 0;
        font-size: 1.15rem;
    }
    
    .section-header .icon {
        font-size: 1.5rem;
    }
    
    /* ===== CARD CONTAINERS ===== */
    .pro-card {
        background: var(--bg-white);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            animation: shine 4s linear infinite;
            text-shadow: 0 0 40px rgba(249, 210, 14, 0.3);
        }
        
        @keyframes shine {
            0% { background-position: 0% center; }
            100% { background-position: 200% center; }
        }
        
        .brand-subtitle {
            font-size: 1.25rem;
            color: rgba(255, 255, 255, 0.8);
            margin-bottom: 1rem;
            font-weight: 300;
            letter-spacing: 2px;
            text-transform: uppercase;
        }
        
        .pro-badge-large {
            display: inline-block;
            background: linear-gradient(135deg, #f9d20e 0%, #ffc107 50%, #f9d20e 100%);
            background-size: 200% auto;
            color: #1f3864;
            font-size: 1rem;
            font-weight: 800;
            padding: 0.5rem 1.5rem;
            border-radius: 25px;
            letter-spacing: 3px;
            animation: shine 3s linear infinite;
            box-shadow: 0 4px 20px rgba(249, 210, 14, 0.4);
            margin-bottom: 2rem;
        }
        
        .feature-list {
            display: flex;
            flex-direction: column;
            gap: 1rem;
            margin-top: 2rem;
        }
        
        .feature-item {
            display: flex;
            align-items: center;
            gap: 1rem;
            padding: 1rem 1.5rem;
            background: rgba(255, 255, 255, 0.05);
            border: 1px solid rgba(249, 210, 14, 0.2);
            border-radius: 12px;
            backdrop-filter: blur(10px);
            transition: all 0.3s ease;
        }
        
        .feature-item:hover {
            background: rgba(249, 210, 14, 0.1);
            transform: translateX(10px);
            border-color: rgba(249, 210, 14, 0.4);
        }
        
        .feature-icon {
            width: 45px;
            height: 45px;
            background: linear-gradient(135deg, rgba(249, 210, 14, 0.2) 0%, rgba(249, 210, 14, 0.1) 100%);
            border-radius: 12px;
            display: flex;
            align-items: center;
            justify-content: center;
            flex-shrink: 0;
        }
        
        .feature-icon svg {
            width: 24px;
            height: 24px;
            color: #f9d20e;
        }
        
        .feature-text {
            font-size: 1rem;
            color: rgba(255, 255, 255, 0.9);
            font-weight: 500;
        }
        
        /* Right Panel - Login Form (30%) */
        .login-right-panel {
            flex: 0 0 35%;
            background: linear-gradient(180deg, #ffffff 0%, #f8f9fa 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            padding: 3rem;
            position: relative;
        }
        
        .login-right-panel::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 5px;
            height: 100%;
            background: linear-gradient(180deg, #f9d20e 0%, #1f3864 50%, #f9d20e 100%);
        }
        
        .login-form-container {
            width: 100%;
            max-width: 380px;
        }
        
        .login-form-header {
            margin-bottom: 2rem;
        }
        
        .login-form-header h2 {
            font-size: 2rem;
            font-weight: 700;
            color: #1f3864;
            margin: 0 0 0.5rem 0;
        }
        
        .login-form-header p {
            font-size: 1rem;
            color: #6c757d;
            margin: 0;
        }
        
        .login-divider {
            display: flex;
            align-items: center;
            gap: 1rem;
            margin: 1.5rem 0;
            color: #adb5bd;
            font-size: 0.85rem;
        }
        
        .login-divider::before,
        .login-divider::after {
            content: '';
            flex: 1;
            height: 1px;
            background: linear-gradient(90deg, transparent, #dee2e6, transparent);
        }
        
        .login-footer-text {
            text-align: center;
            margin-top: 2rem;
            padding-top: 1.5rem;
            border-top: 1px solid #e9ecef;
        }
        
        .login-footer-text p {
            color: #6c757d;
            font-size: 0.85rem;
            margin: 0.25rem 0;
        }
        
        .login-footer-text .company-name {
            color: #1f3864;
            font-weight: 600;
        }
        
        /* Animated corner accents */
        .corner-accent {
            position: absolute;
            width: 100px;
            height: 100px;
            z-index: 0;
        }
        
        .corner-accent.top-left {
            top: 20px;
            left: 20px;
            border-top: 3px solid rgba(249, 210, 14, 0.5);
            border-left: 3px solid rgba(249, 210, 14, 0.5);
            animation: corner-pulse 3s ease-in-out infinite;
        }
        
        .corner-accent.bottom-right {
            bottom: 20px;
            right: 20px;
            border-bottom: 3px solid rgba(249, 210, 14, 0.5);
            border-right: 3px solid rgba(249, 210, 14, 0.5);
            animation: corner-pulse 3s ease-in-out infinite reverse;
        }
        
        @keyframes corner-pulse {
            0%, 100% { opacity: 0.3; }
            50% { opacity: 1; }
        }
        
        /* Rotating ring decoration */
        .rotating-ring {
            position: absolute;
            width: 500px;
            height: 500px;
            border: 2px dashed rgba(249, 210, 14, 0.2);
            border-radius: 50%;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            animation: ring-rotate 60s linear infinite;
            z-index: 1;
        }
        
        .rotating-ring::before {
            content: '';
            position: absolute;
            width: 15px;
            height: 15px;
            background: #f9d20e;
            border-radius: 50%;
            top: -7.5px;
            left: 50%;
            transform: translateX(-50%);
            box-shadow: 0 0 20px rgba(249, 210, 14, 0.8);
        }
        
        @keyframes ring-rotate {
            0% { transform: translate(-50%, -50%) rotate(0deg); }
            100% { transform: translate(-50%, -50%) rotate(360deg); }
        }
        
        /* Form styling overrides */
        .login-right-panel .stTextInput > div > div > input {
            background: #ffffff !important;
            border: 2px solid #e9ecef !important;
            border-radius: 12px !important;
            padding: 0.875rem 1rem !important;
            font-size: 1rem !important;
            transition: all 0.3s ease !important;
        }
        
        .login-right-panel .stTextInput > div > div > input:focus {
            border-color: #1f3864 !important;
            box-shadow: 0 0 0 3px rgba(31, 56, 100, 0.1) !important;
        }
        
        .login-right-panel .stTextInput label {
            color: #1f3864 !important;
            font-weight: 600 !important;
            font-size: 0.9rem !important;
        }
        
        .login-right-panel button[kind="primary"] {
            background: linear-gradient(135deg, #1f3864 0%, #2d4a7c 100%) !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 0.875rem 2rem !important;
            font-size: 1rem !important;
            font-weight: 600 !important;
            letter-spacing: 0.5px !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 4px 15px rgba(31, 56, 100, 0.3) !important;
        }
        
        .login-right-panel button[kind="primary"]:hover {
            background: linear-gradient(135deg, #2d4a7c 0%, #3d5a8c 100%) !important;
            transform: translateY(-2px) !important;
            box-shadow: 0 6px 20px rgba(31, 56, 100, 0.4) !important;
        }
        
        /* Security badge */
        .security-badge {
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 0.5rem;
            margin-top: 1rem;
            padding: 0.75rem;
            background: rgba(31, 56, 100, 0.05);
            border-radius: 8px;
            font-size: 0.8rem;
            color: #6c757d;
        }
        
        .security-badge svg {
            width: 16px;
            height: 16px;
            color: #28a745;
        }
    </style>
    """, unsafe_allow_html=True)

# ==================== CHECK LOGIN STATE ====================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.current_user = None

# Show login page if not logged in
if not st.session_state.logged_in:
    show_login_page()
    st.stop()

# Get current user info for the rest of the app
current_user = st.session_state.current_user

# ==================== PROFESSIONAL CSS STYLING ====================
st.markdown("""
<style>
    /* ===== ROOT VARIABLES ===== */
    :root {
        --primary-blue: #060c71;
        --secondary-blue: #2a3bb8;
        --accent-gold: #f9d20e;
        --success-green: #28a745;
        --text-dark: #1a1a2e;
        --text-muted: #6c757d;
        --bg-light: #f8f9fa;
        --bg-white: #ffffff;
        --border-color: #e0e4e8;
        --shadow-sm: 0 2px 8px rgba(0,0,0,0.08);
        --shadow-md: 0 4px 16px rgba(0,0,0,0.12);
        --shadow-lg: 0 8px 32px rgba(0,0,0,0.16);
        --transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    
    /* ===== GLOBAL STYLES ===== */
    .main .block-container {
        padding: 2rem 2rem 2rem 2rem;
        max-width: 900px;
    }
    
    /* ===== MAIN HEADER ===== */
    .main-header {
        background: linear-gradient(105deg, #0a1628 0%, #1a3a5c 25%, #1f3864 45%, #2d5a87 65%, #d4a012 85%, #f9d20e 100%);
        padding: 2.5rem 3rem;
        border-radius: 20px;
        margin-bottom: 2rem;
        box-shadow: 0 20px 60px rgba(10, 22, 40, 0.4), 0 8px 25px rgba(249, 210, 14, 0.15), inset 0 1px 0 rgba(255,255,255,0.1);
        position: relative;
        overflow: hidden;
        border: 1px solid rgba(255,255,255,0.1);
    }
    
    /* Wave animation keyframes */
    @keyframes waveFlow {
        0% {
            transform: translateX(-100%) skewX(-15deg);
        }
        100% {
            transform: translateX(200%) skewX(-15deg);
        }
    }
    
    @keyframes colorPulse {
        0%, 100% {
            opacity: 0.3;
        }
        50% {
            opacity: 0.6;
        }
    }
    
    .main-header::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 50%;
        height: 200%;
        background: linear-gradient(
            90deg,
            transparent 0%,
            rgba(249, 210, 14, 0.15) 25%,
            rgba(249, 210, 14, 0.35) 50%,
            rgba(249, 210, 14, 0.15) 75%,
            transparent 100%
        );
        animation: waveFlow 4s ease-in-out infinite;
        pointer-events: none;
    }
    
    .main-header::after {
        content: '';
        position: absolute;
        right: 0;
        top: 0;
        bottom: 0;
        width: 35%;
        background: linear-gradient(105deg, transparent 0%, rgba(249, 210, 14, 0.2) 50%, rgba(249, 210, 14, 0.35) 100%);
        clip-path: polygon(30% 0, 100% 0, 100% 100%, 0% 100%);
        animation: colorPulse 3s ease-in-out infinite;
    }
    
    .main-header h1 {
        color: #ffffff !important;
        font-size: 2.4rem;
        font-weight: 800;
        margin: 0;
        letter-spacing: -0.5px;
        text-shadow: 0 2px 4px rgba(0,0,0,0.4), 0 4px 20px rgba(0,0,0,0.2);
        position: relative;
        z-index: 2;
    }
    
    .main-header .subtitle {
        color: rgba(255,255,255,0.95);
        font-size: 1.1rem;
        margin-top: 0.6rem;
        font-weight: 400;
        position: relative;
        z-index: 2;
        letter-spacing: 0.5px;
        text-shadow: 0 1px 3px rgba(0,0,0,0.3);
    }
    
    .main-header .version-badge {
        display: inline-block;
        background: linear-gradient(135deg, #f9d20e 0%, #ffc107 50%, #ffab00 100%);
        color: #0a1628;
        padding: 0.3rem 1rem;
        border-radius: 25px;
        font-size: 0.8rem;
        font-weight: 800;
        margin-left: 1rem;
        vertical-align: middle;
        box-shadow: 0 4px 15px rgba(249, 210, 14, 0.4);
        text-transform: uppercase;
        letter-spacing: 1px;
        -webkit-text-fill-color: #0a1628;
        position: relative;
        z-index: 3;
    }
    
    /* Decorative elements */
    .main-header .header-icon {
        position: absolute;
        right: 40px;
        top: 50%;
        transform: translateY(-50%);
        font-size: 5rem;
        opacity: 0.15;
        z-index: 1;
        filter: drop-shadow(0 0 20px rgba(249, 210, 14, 0.3));
    }
    
    /* ===== SECTION HEADERS ===== */
    .section-header {
        background: linear-gradient(135deg, var(--bg-white) 0%, var(--bg-light) 100%);
        border-left: 5px solid var(--primary-blue);
        padding: 1rem 1.5rem;
        border-radius: 0 12px 12px 0;
        margin: 2rem 0 1.25rem 0;
        box-shadow: var(--shadow-sm);
        display: flex;
        align-items: center;
        gap: 0.75rem;
    }
    
    .section-header h3 {
        color: var(--primary-blue);
        font-weight: 700;
        margin: 0;
        font-size: 1.15rem;
    }
    
    .section-header .icon {
        font-size: 1.5rem;
    }
    
    /* ===== CARD CONTAINERS ===== */
    .pro-card {
        background: var(--bg-white);
        border: 1px solid var(--border-color);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        box-shadow: var(--shadow-sm);
        transition: var(--transition);
    }
    
    .pro-card:hover {
        box-shadow: var(--shadow-md);
        border-color: var(--secondary-blue);
    }
    
    .pro-card-header {
        font-weight: 700;
        color: var(--primary-blue);
        font-size: 1.1rem;
        margin-bottom: 1rem;
        padding-bottom: 0.75rem;
        border-bottom: 2px solid var(--bg-light);
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    /* ===== STATUS BADGES ===== */
    .status-badge {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.5rem 1rem;
        border-radius: 25px;
        font-size: 0.85rem;
        font-weight: 600;
    }
    
    .status-ready {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
        color: #155724;
        border: 1px solid #28a745;
    }
    
    .status-pending {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeeba 100%);
        color: #856404;
        border: 1px solid #ffc107;
    }
    
    .status-processing {
        background: linear-gradient(135deg, #cce5ff 0%, #b8daff 100%);
        color: #004085;
        border: 1px solid #007bff;
    }
    
    /* ===== INPUT FIELDS ===== */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stNumberInput > div > div > input {
        border-radius: 10px !important;
        border: 2px solid var(--border-color) !important;
        padding: 0.75rem 1rem !important;
        font-size: 15px !important;
        transition: var(--transition) !important;
        background: var(--bg-white) !important;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus,
    .stNumberInput > div > div > input:focus {
        border-color: var(--secondary-blue) !important;
        box-shadow: 0 0 0 4px rgba(42, 59, 184, 0.1) !important;
    }
    
    /* ===== LABELS ===== */
    .stTextInput > label,
    .stFileUploader > label,
    .stTextArea > label,
    .stDateInput > label,
    .stCheckbox > label,
    .stSelectbox > label,
    .stNumberInput > label {
        font-weight: 600 !important;
        color: var(--text-dark) !important;
        font-size: 0.9rem !important;
        margin-bottom: 0.5rem !important;
    }
    
    /* ===== FILE UPLOADER ===== */
    [data-testid="stFileUploader"] > div {
        border: 2px dashed var(--secondary-blue) !important;
        border-radius: 12px !important;
        padding: 1.5rem !important;
        text-align: center !important;
        background: linear-gradient(135deg, rgba(42, 59, 184, 0.02) 0%, rgba(6, 12, 113, 0.04) 100%) !important;
        transition: var(--transition) !important;
    }
    
    [data-testid="stFileUploader"] > div:hover {
        background: linear-gradient(135deg, rgba(42, 59, 184, 0.05) 0%, rgba(6, 12, 113, 0.08) 100%) !important;
        border-color: var(--accent-gold) !important;
    }
    
    /* ===== BUTTONS ===== */
    .stButton > button {
        background: linear-gradient(135deg, var(--primary-blue) 0%, var(--secondary-blue) 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 0.875rem 2rem !important;
        font-size: 15px !important;
        font-weight: 600 !important;
        transition: var(--transition) !important;
        box-shadow: 0 4px 15px rgba(6, 12, 113, 0.25) !important;
        text-transform: none !important;
        letter-spacing: 0.3px !important;
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, var(--accent-gold) 0%, #ffe34a 100%) !important;
        color: var(--primary-blue) !important;
        transform: translateY(-3px) !important;
        box-shadow: 0 8px 25px rgba(249, 210, 14, 0.4) !important;
    }
    
    .stButton > button:active {
        transform: translateY(-1px) !important;
    }
    
    /* ===== DOWNLOAD BUTTON ===== */
    .stDownloadButton > button {
        background: linear-gradient(135deg, var(--success-green) 0%, #34ce57 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 1rem 2rem !important;
        font-weight: 700 !important;
        font-size: 16px !important;
        transition: var(--transition) !important;
        box-shadow: 0 4px 15px rgba(40, 167, 69, 0.3) !important;
    }
    
    .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #218838 0%, var(--success-green) 100%) !important;
        transform: translateY(-3px) !important;
        box-shadow: 0 8px 25px rgba(40, 167, 69, 0.4) !important;
    }
    
    /* ===== EXPANDERS ===== */
    .streamlit-expanderHeader {
        background: linear-gradient(135deg, var(--bg-white) 0%, var(--bg-light) 100%) !important;
        border-radius: 12px !important;
        border: 2px solid var(--border-color) !important;
        font-weight: 600 !important;
        color: var(--primary-blue) !important;
        padding: 1rem 1.25rem !important;
        transition: var(--transition) !important;
    }
    
    .streamlit-expanderHeader:hover {
        border-color: var(--secondary-blue) !important;
        box-shadow: var(--shadow-sm) !important;
    }
    
    .streamlit-expanderContent {
        border: 2px solid var(--border-color) !important;
        border-top: none !important;
        border-radius: 0 0 12px 12px !important;
        background: var(--bg-white) !important;
        padding: 1.25rem !important;
    }
    
    /* ===== TABS ===== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: var(--bg-light);
        padding: 0.5rem;
        border-radius: 12px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        font-weight: 600;
        color: var(--text-muted);
        transition: var(--transition);
    }
    
    .stTabs [aria-selected="true"] {
        background: var(--primary-blue) !important;
        color: white !important;
    }
    
    /* ===== PROGRESS BAR ===== */
    .stProgress > div > div > div {
        background: linear-gradient(90deg, var(--primary-blue) 0%, var(--secondary-blue) 50%, var(--accent-gold) 100%);
        border-radius: 10px;
    }
    
    /* ===== ALERTS ===== */
    .stSuccess {
        background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%) !important;
        border-left: 5px solid var(--success-green) !important;
        border-radius: 10px !important;
        padding: 1rem 1.25rem !important;
    }
    
    .stWarning {
        background: linear-gradient(135deg, #fff3cd 0%, #ffeeba 100%) !important;
        border-left: 5px solid #ffc107 !important;
        border-radius: 10px !important;
    }
    
    .stError {
        background: linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%) !important;
        border-left: 5px solid #dc3545 !important;
        border-radius: 10px !important;
    }
    
    .stInfo {
        background: linear-gradient(135deg, #cce5ff 0%, #b8daff 100%) !important;
        border-left: 5px solid #007bff !important;
        border-radius: 10px !important;
    }
    
    /* ===== DOCUMENT PREVIEW ===== */
    .doc-preview-container {
        background: var(--bg-white);
        border: 2px solid var(--border-color);
        border-radius: 16px;
        padding: 2rem;
        max-height: 600px;
        overflow-y: auto;
        box-shadow: var(--shadow-md);
        margin: 1.5rem 0;
    }
    
    .doc-preview-header {
        background: linear-gradient(135deg, var(--bg-light) 0%, #e9ecef 100%);
        padding: 1rem 1.5rem;
        border-radius: 12px 12px 0 0;
        margin: -2rem -2rem 1.5rem -2rem;
        border-bottom: 2px solid var(--border-color);
        display: flex;
        justify-content: space-between;
        align-items: center;
    }
    
    .doc-preview-header h4 {
        margin: 0;
        color: var(--primary-blue);
        font-weight: 700;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    .doc-preview-content {
        font-family: 'Calibri', 'Segoe UI', sans-serif;
        line-height: 1.8;
        color: var(--text-dark);
    }
    
    .doc-preview-content h1, 
    .doc-preview-content h2,
    .doc-preview-content h3 {
        color: var(--primary-blue);
        margin-top: 1.5rem;
        margin-bottom: 0.75rem;
    }
    
    .doc-preview-content p {
        margin-bottom: 1rem;
        text-align: justify;
    }
    
    /* ===== GENERATION SECTION ===== */
    .generation-container {
        background: linear-gradient(135deg, rgba(6, 12, 113, 0.03) 0%, rgba(42, 59, 184, 0.05) 100%);
        border: 2px solid var(--secondary-blue);
        border-radius: 16px;
        padding: 2rem;
        margin: 2rem 0;
        text-align: center;
    }
    
    .generation-title {
        font-size: 1.5rem;
        font-weight: 700;
        color: var(--primary-blue);
        margin-bottom: 0.5rem;
    }
    
    .generation-subtitle {
        color: var(--text-muted);
        margin-bottom: 1.5rem;
    }
    
    /* ===== METRICS/STATS ===== */
    .stat-card {
        background: var(--bg-white);
        border-radius: 12px;
        padding: 1.25rem;
        text-align: center;
        box-shadow: var(--shadow-sm);
        border: 1px solid var(--border-color);
        transition: var(--transition);
    }
    
    .stat-card:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-md);
    }
    
    .stat-value {
        font-size: 2rem;
        font-weight: 800;
        color: var(--primary-blue);
        line-height: 1.2;
    }
    
    .stat-label {
        color: var(--text-muted);
        font-size: 0.85rem;
        font-weight: 500;
        margin-top: 0.25rem;
    }
    
    /* ===== DIVIDER ===== */
    hr {
        margin: 2rem 0;
        border: none;
        height: 3px;
        background: linear-gradient(90deg, transparent 0%, var(--accent-gold) 20%, var(--accent-gold) 80%, transparent 100%);
        border-radius: 3px;
    }
    
    /* ===== SCROLLBAR ===== */
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: var(--bg-light);
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: var(--secondary-blue);
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: var(--primary-blue);
    }
    
    /* ===== HIDE STREAMLIT BRANDING ===== */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display: none;}
    
    /* ===== ANIMATIONS ===== */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(10px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.7; }
    }
    
    .animate-fade-in {
        animation: fadeIn 0.5s ease-out forwards;
    }
    
    .animate-pulse {
        animation: pulse 2s infinite;
    }
    
    /* ===== SPINNER OVERRIDE ===== */
    .stSpinner > div {
        border-top-color: var(--accent-gold) !important;
    }
    
    /* ===== USER PROFILE BAR ===== */
    .user-profile-bar {
        display: flex;
        justify-content: flex-end;
        align-items: center;
        padding: 0.5rem 0;
        margin-bottom: 0.5rem;
    }
    
    .user-profile-wrapper {
        position: relative;
        display: inline-block;
    }
    
    .user-profile-info {
        display: flex;
        align-items: center;
        gap: 0.75rem;
        cursor: pointer;
        padding: 0.5rem 0.75rem;
        border-radius: 10px;
        transition: all 0.2s ease;
    }
    
    .user-profile-info:hover {
        background: rgba(31, 56, 100, 0.08);
    }
    
    .user-avatar {
        width: 40px;
        height: 40px;
        background: linear-gradient(135deg, #1f3864 0%, #2d5a87 100%);
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        color: white;
        font-weight: 700;
        font-size: 0.95rem;
        box-shadow: 0 3px 10px rgba(31, 56, 100, 0.3);
        border: 2px solid #f9d20e;
        transition: all 0.2s ease;
    }
    
    .user-profile-info:hover .user-avatar {
        transform: scale(1.05);
        box-shadow: 0 4px 15px rgba(31, 56, 100, 0.4);
    }
    
    .user-details {
        display: flex;
        flex-direction: column;
        align-items: flex-end;
    }
    
    .user-name {
        font-weight: 600;
        color: #1f3864;
        font-size: 0.9rem;
        line-height: 1.2;
    }
    
    .user-role {
        font-size: 0.7rem;
        color: #6c757d;
        font-weight: 400;
    }
    
    .online-indicator {
        width: 10px;
        height: 10px;
        background: #28a745;
        border-radius: 50%;
        position: absolute;
        bottom: 0px;
        right: 0px;
        border: 2px solid white;
        animation: pulse 2s infinite;
    }
    
    .avatar-container {
        position: relative;
    }
    
    /* Dropdown chevron */
    .dropdown-chevron {
        margin-left: 0.25rem;
        color: #6c757d;
        transition: transform 0.2s ease;
    }
    
    .user-profile-wrapper:hover .dropdown-chevron {
        transform: rotate(180deg);
    }
    
    /* User Dropdown Menu */
    .user-dropdown {
        position: absolute;
        top: 100%;
        right: 0;
        margin-top: 0.5rem;
        background: white;
        border-radius: 12px;
        box-shadow: 0 10px 40px rgba(0,0,0,0.15), 0 2px 10px rgba(0,0,0,0.1);
        border: 1px solid rgba(31, 56, 100, 0.1);
        min-width: 200px;
        opacity: 0;
        visibility: hidden;
        transform: translateY(-10px);
        transition: all 0.25s ease;
        z-index: 1000;
        overflow: hidden;
    }
    
    .user-profile-wrapper:hover .user-dropdown {
        opacity: 1;
        visibility: visible;
        transform: translateY(0);
    }
    
    .dropdown-header {
        padding: 1rem;
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border-bottom: 1px solid rgba(31, 56, 100, 0.1);
    }
    
    .dropdown-header .user-name {
        font-size: 0.95rem;
        margin-bottom: 0.15rem;
    }
    
    .dropdown-header .user-email {
        font-size: 0.75rem;
        color: #6c757d;
    }
    
    .dropdown-menu-items {
        padding: 0.5rem 0;
    }
    
    .dropdown-item {
        display: flex;
        align-items: center;
        gap: 0.75rem;
        padding: 0.75rem 1rem;
        color: #1f3864;
        font-size: 0.875rem;
        font-weight: 500;
        cursor: pointer;
        transition: all 0.15s ease;
        text-decoration: none;
    }
    
    .dropdown-item:hover {
        background: rgba(31, 56, 100, 0.08);
    }
    
    .dropdown-item svg {
        width: 18px;
        height: 18px;
        color: #6c757d;
    }
    
    .dropdown-item.logout {
        color: #dc3545;
        border-top: 1px solid rgba(0,0,0,0.08);
        margin-top: 0.25rem;
    }
    
    .dropdown-item.logout svg {
        color: #dc3545;
    }
    
    .dropdown-item.logout:hover {
        background: rgba(220, 53, 69, 0.08);
    }
    
    .coming-soon-badge {
        font-size: 0.65rem;
        background: linear-gradient(135deg, #f9d20e 0%, #ffc107 100%);
        color: #1f3864;
        padding: 0.15rem 0.5rem;
        border-radius: 10px;
        font-weight: 600;
        margin-left: auto;
    }
    
    /* Header logo image */
    .header-logo-img {
        position: absolute;
        right: 30px;
        top: 50%;
        transform: translateY(-50%);
        width: 80px;
        height: auto;
        opacity: 0.25;
        z-index: 1;
        filter: brightness(1.2) drop-shadow(0 0 15px rgba(249, 210, 14, 0.4));
    }
    
    /* Logout Confirmation Modal */
    .logout-modal-overlay {
        position: fixed;
        top: 0;
        left: 0;
        width: 100vw;
        height: 100vh;
        background: rgba(0, 0, 0, 0.6);
        backdrop-filter: blur(5px);
        display: flex;
        justify-content: center;
        align-items: center;
        z-index: 10000;
        animation: fadeIn 0.3s ease;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }
    
    .logout-modal {
        background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
        border-radius: 16px;
        padding: 2rem;
        max-width: 420px;
        width: 90%;
        box-shadow: 0 25px 50px rgba(0, 0, 0, 0.3), 0 0 0 1px rgba(255, 255, 255, 0.1);
        text-align: center;
        animation: slideUp 0.3s ease;
    }
    
    @keyframes slideUp {
        from { transform: translateY(20px); opacity: 0; }
        to { transform: translateY(0); opacity: 1; }
    }
    
    .logout-modal-icon {
        width: 70px;
        height: 70px;
        background: linear-gradient(135deg, rgba(220, 53, 69, 0.1) 0%, rgba(220, 53, 69, 0.2) 100%);
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        margin: 0 auto 1.5rem;
    }
    
    .logout-modal-icon svg {
        width: 35px;
        height: 35px;
        color: #dc3545;
    }
    
    .logout-modal h3 {
        color: #1f3864;
        font-size: 1.5rem;
        margin-bottom: 0.75rem;
        font-weight: 700;
    }
    
    .logout-modal p {
        color: #6c757d;
        font-size: 0.95rem;
        line-height: 1.6;
        margin-bottom: 1.5rem;
    }
    
    .logout-modal-buttons {
        display: flex;
        gap: 1rem;
        justify-content: center;
    }
    
    .logout-btn-cancel {
        padding: 0.75rem 1.5rem;
        border-radius: 10px;
        font-weight: 600;
        font-size: 0.95rem;
        cursor: pointer;
        transition: all 0.3s ease;
        border: 2px solid #e0e0e0;
        background: #ffffff;
        color: #6c757d;
    }
    
    .logout-btn-cancel:hover {
        background: #f8f9fa;
        border-color: #c0c0c0;
    }
    
    .logout-btn-confirm {
        padding: 0.75rem 1.5rem;
        border-radius: 10px;
        font-weight: 600;
        font-size: 0.95rem;
        cursor: pointer;
        transition: all 0.3s ease;
        border: none;
        background: linear-gradient(135deg, #dc3545 0%, #c82333 100%);
        color: white;
        box-shadow: 0 4px 15px rgba(220, 53, 69, 0.3);
    }
    
    .logout-btn-confirm:hover {
        background: linear-gradient(135deg, #c82333 0%, #bd2130 100%);
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(220, 53, 69, 0.4);
    }
    
    /* Style the logout confirmation buttons */
    .logout-modal-overlay ~ div [data-testid="stHorizontalBlock"] {
        position: fixed;
        bottom: 35%;
        left: 50%;
        transform: translateX(-50%);
        z-index: 10001;
        background: transparent;
    }
    
    /* Cancel button styling */
    button[kind="secondary"]:has(p:contains("Cancel")),
    [data-testid="stButton"] button:contains("Cancel") {
        background: #ffffff !important;
        border: 2px solid #e0e0e0 !important;
        color: #6c757d !important;
        padding: 0.6rem 1.5rem !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
    }
    
    button[kind="secondary"]:has(p:contains("Cancel")):hover {
        background: #f8f9fa !important;
        border-color: #c0c0c0 !important;
    }
    
    /* Confirm logout button styling */
    button[kind="primary"]:has(p:contains("Logout")),
    [data-testid="stButton"] button:contains("Yes") {
        background: linear-gradient(135deg, #dc3545 0%, #c82333 100%) !important;
        border: none !important;
        color: white !important;
        padding: 0.6rem 1.5rem !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 15px rgba(220, 53, 69, 0.3) !important;
        transition: all 0.3s ease !important;
    }
    
    button[kind="primary"]:has(p:contains("Logout")):hover {
        background: linear-gradient(135deg, #c82333 0%, #bd2130 100%) !important;
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(220, 53, 69, 0.4) !important;
    }
</style>
""", unsafe_allow_html=True)

# ==================== USER PROFILE BAR ====================
# Get user info from session state
user_name = current_user.get("name", "User")
user_role = current_user.get("role", "")
user_email = current_user.get("email", "")
user_initials = current_user.get("initials", "U")

# Initialize logout confirmation state
if 'show_logout_confirm' not in st.session_state:
    st.session_state.show_logout_confirm = False

    st.markdown(
        """
        <style>
        /* Center main tab list */
        #main-tabs div[role=\"tablist\"] { 
            justify-content: center; 
            gap: 8px;
        }

        /* Disable non-CBS tabs but keep visuals readable */
        #main-tabs div[role=\"tablist\"] button[role=\"tab\"]:nth-child(2),
        #main-tabs div[role="tablist"] button[role="tab"]:nth-child(2),
        #main-tabs div[role="tablist"] button[role="tab"]:nth-child(4) {
            color: #6b7280 !important; /* gray-600 */
            position: relative;
        }

        /* Add badge-style \"Coming soon\" pill to non-CBS tabs */
        #main-tabs div[role="tablist"] button[role="tab"]:nth-child(2)::after,
        #main-tabs div[role="tablist"] button[role="tab"]:nth-child(4)::after {
            content: 'Coming soon';
            display: inline-block;
            margin-left: 8px;
            padding: 2px 8px;
            font-size: 12px;
            line-height: 1.2;
            border-radius: 999px;
            background: #eef2ff;         /* light indigo background */
            color: #3730a3;               /* indigo-800 text */
            border: 1px solid #c7d2fe;    /* indigo-200 border */
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

# Hidden button to trigger logout confirmation (activated by JavaScript)
logout_placeholder = st.empty()


# Hide the trigger button with CSS
st.markdown('''

<script>
    // Add click listener to logout dropdown item
    document.addEventListener('DOMContentLoaded', function() {
        setTimeout(function() {
            const logoutItem = document.querySelector('.dropdown-item.logout');
            if (logoutItem) {
                logoutItem.style.cursor = 'pointer';
                logoutItem.addEventListener('click', function() {
                    // Find and click the hidden Streamlit button
                    const buttons = document.querySelectorAll('button');
                    buttons.forEach(btn => {
                        if (btn.innerText.includes('🔓')) {
                            btn.click();
                        }
                    });
                });
            }
        }, 500);
    });
</script>
''', unsafe_allow_html=True)

# Show logout confirmation popup
if st.session_state.show_logout_confirm:
    st.markdown('''
    <div class="logout-modal-overlay" id="logout-modal">
        <div class="logout-modal">
            <div class="logout-modal-icon">
                <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2">
                    <path d="M9 21H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h4"></path>
                    <polyline points="16 17 21 12 16 7"></polyline>
                    <line x1="21" y1="12" x2="9" y2="12"></line>
                </svg>
            </div>
            <h3>Confirm Logout</h3>
            <p>Are you sure you want to logout?<br><strong>All your current progress will be lost.</strong></p>
        </div>
    </div>
    ''', unsafe_allow_html=True)
    
    # Centered buttons for confirmation
    spacer1, cancel_col, confirm_col, spacer2 = st.columns([3, 1.5, 1.5, 3])
    with cancel_col:
        if st.button("❌ Cancel", key="cancel_logout", type="secondary"):
            st.session_state.show_logout_confirm = False
            st.rerun()
    with confirm_col:
        if st.button("✅ Yes, Logout", key="confirm_logout", type="primary"):
            st.session_state.show_logout_confirm = False
            logout_user()
            st.rerun()

# ==================== MAIN HEADER ====================
# Load company logo for header
import base64
header_logo_path = r".\assests\Images\falcon-autotech-icon-removebg-preview.png"
try:
    with open(header_logo_path, "rb") as img_file:
        header_logo_base64 = base64.b64encode(img_file.read()).decode()
    header_logo_html = f'<img src="data:image/png;base64,{header_logo_base64}" class="header-logo-img" alt="Falcon Logo">'
except:
    header_logo_html = ''

st.markdown(f'''
<div class="main-header">
    {header_logo_html}
    <h1>Falcon Proposal Generator <span class="version-badge">PRO</span></h1>
    <div class="subtitle">Professional Proposal Document Generation System</div>
</div>
''', unsafe_allow_html=True)

# ==================== STYLING FUNCTIONS ====================

# Deep Blue-Gray color for headings (RGB: 31, 56, 100)
HEADING_COLOR = RGBColor(31, 56, 100)
def render_section_header(title):
    st.markdown(f'''
    <div class="section-header">
        <h3>{title}</h3>
    </div>
    ''', unsafe_allow_html=True)
def apply_heading_style(paragraph, text, level=1):
    """Apply custom heading style: Calibri Headings 14pt, Bold, Underline, Numbered, Deep Blue-Gray"""
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.underline = True
    run.font.color.rgb = HEADING_COLOR
    
    # Apply paragraph formatting
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    
    return paragraph

def apply_subheading_style(paragraph, text):
    """Apply subheading style: Calibri 12pt, Bold, Deep Blue-Gray"""
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    
    return paragraph

def apply_normal_style(paragraph, text=""):
    """Apply normal text style: Calibri (Body) 11pt, Black"""
    if text:
        paragraph.text = ""
        run = paragraph.add_run(text)
        run.font.name = 'Calibri (Body)'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        for run in paragraph.runs:
            run.font.name = 'Calibri (Body)'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    return paragraph

def apply_table_style(table):
    """Apply Medium Shading 1 Accent 1 style to table"""
    try:
        table.style = 'Medium Shading 1 Accent 1'
    except KeyError:
        # If style doesn't exist, apply manual formatting similar to Medium Shading 1 Accent 1
        # This happens when document is created from a template without this style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If even Table Grid doesn't exist, skip styling
            pass
    return table

def add_centered_image(doc, path, width_in=5.5):
    """Add a centered image if it exists"""
    if not path or not os.path.exists(path):
        return
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_bullets_in_two_columns(doc, bullet_items):
    """Add bullet points in a two-column layout using a borderless table"""
    # Split items into two columns
    mid = (len(bullet_items) + 1) // 2
    left_items = bullet_items[:mid]
    right_items = bullet_items[mid:]
    
    # Create table with 2 columns - use no style to avoid borders
    table = doc.add_table(rows=max(len(left_items), len(right_items)), cols=2)
    
    # Remove all borders from table using XML manipulation
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Remove any existing tblBorders element
    for existing_borders in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing_borders)
    
    # Add new tblBorders with all borders set to none
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # Set column widths
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    
    # Fill left column with compact formatting
    for i, item in enumerate(left_items):
        cell = table.rows[i].cells[0]
        # Remove cell borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(f"• {item}")
        run.font.name = 'Calibri'
        run.font.size = Pt(9)  # Slightly smaller for compactness
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
    
    # Fill right column with compact formatting
    for i, item in enumerate(right_items):
        cell = table.rows[i].cells[1]
        # Remove cell borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(f"• {item}")
        run.font.name = 'Calibri'
        run.font.size = Pt(9)  # Slightly smaller for compactness
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)

def add_numbered_heading(doc, text, level=1, counter=None):
    """Add a numbered heading with proper formatting"""
    if counter:
        full_text = f"{counter}. {text}"  # Added period after number
    else:
        full_text = text
    
    # Use built-in heading style
    p = doc.add_heading(full_text, level=1)
    
    # Apply custom formatting to the heading
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = HEADING_COLOR
    
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    return p

def add_numbered_subheading(doc, text, counter=None):
    """Add a numbered subheading"""
    if counter:
        full_text = f"{counter}. {text}"  # Added period after number
    else:
        full_text = text
    
    # Use built-in heading style for subheading
    p = doc.add_heading(full_text, level=2)
    
    # Apply custom formatting
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = HEADING_COLOR
    
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    
    return p

def ensure_list_styles(doc):
    """Ensure List Bullet, List Number, and Table styles exist in the document"""
    styles = doc.styles
    
    # Check if List Bullet exists, if not create it
    try:
        styles['List Bullet']
    except KeyError:
        # Create List Bullet style
        from docx.enum.style import WD_STYLE_TYPE
        list_bullet_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        list_bullet_style.base_style = styles['Normal']
        list_bullet_style.font.name = 'Calibri'
        list_bullet_style.font.size = Pt(11)
        # Set paragraph format for bullet
        pf = list_bullet_style.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
    
    # Check if List Number exists, if not create it
    try:
        styles['List Number']
    except KeyError:
        # Create List Number style
        from docx.enum.style import WD_STYLE_TYPE
        list_number_style = styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
        list_number_style.base_style = styles['Normal']
        list_number_style.font.name = 'Calibri'
        list_number_style.font.size = Pt(11)
        # Set paragraph format for numbering
        pf = list_number_style.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
    
    # Check if List Number 2 exists, if not create it
    try:
        styles['List Number 2']
    except KeyError:
        # Create List Number 2 style (deeper indentation level)
        from docx.enum.style import WD_STYLE_TYPE
        list_number_2_style = styles.add_style('List Number 2', WD_STYLE_TYPE.PARAGRAPH)
        list_number_2_style.base_style = styles['Normal']
        list_number_2_style.font.name = 'Calibri'
        list_number_2_style.font.size = Pt(11)
        pf = list_number_2_style.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.25)

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a <w:r> element)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add formatting for hyperlink (blue + underline)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Blue color
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Set font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri (Body)')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')  # 9pt = 18 half-points
    rPr.append(sz)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def _clear_header_footer_part(part):
    """Remove existing content from a header/footer part."""
    if not part:
        return
    for tbl in list(part.tables):
        tbl._element.getparent().remove(tbl._element)
    for paragraph in list(part.paragraphs):
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)


def _get_logo_stream(logo_bytes: Optional[bytes], fallback_path: str | None) -> io.BytesIO | None:
    """Return a BytesIO for the supplied logo, loading fallback path if needed."""
    if logo_bytes:
        return io.BytesIO(logo_bytes)
    if fallback_path and os.path.exists(fallback_path):
        with open(fallback_path, "rb") as fh:
            return io.BytesIO(fh.read())
    return None


def create_header_footer(
    doc: Document,
    client_name: str,
    project_name: str,
    client_logo_bytes: Optional[bytes] = None,
    falcon_logo_bytes: Optional[bytes] = None,
) -> None:
    """Apply Falcon-branded header and footer to every section after composition."""

    falcon_fallback_path = "assests/Images/Falcon-Autotech_Logo-removebg-preview.png"

    for index, section in enumerate(doc.sections):
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        first_section = index == 0
        section.different_first_page_header_footer = first_section

        if first_section:
            _clear_header_footer_part(getattr(section, "first_page_header", None))
            _clear_header_footer_part(getattr(section, "first_page_footer", None))

        header = section.header
        footer = section.footer
        _clear_header_footer_part(header)
        _clear_header_footer_part(footer)

        header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        left_cell, middle_cell, right_cell = header_table.rows[0].cells
        left_cell.width = Inches(1.3)
        middle_cell.width = Inches(4.0)
        right_cell.width = Inches(1.3)

        client_logo_stream = _get_logo_stream(client_logo_bytes, None)
        if client_logo_stream:
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run()
            client_logo_stream.seek(0)
            left_run.add_picture(client_logo_stream, height=Inches(0.6))
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        middle_para = middle_cell.paragraphs[0]
        middle_run = middle_para.add_run(f"FALCON's Proposal to {client_name} for the {project_name}")
        middle_run.font.name = "Calibri"
        middle_run.font.size = Pt(9)
        middle_run.font.color.rgb = HEADING_COLOR
        middle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        falcon_logo_stream = _get_logo_stream(falcon_logo_bytes, falcon_fallback_path)
        if falcon_logo_stream:
            right_para = right_cell.paragraphs[0]
            right_run = right_para.add_run()
            falcon_logo_stream.seek(0)
            right_run.add_picture(falcon_logo_stream, height=Inches(0.6))
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for row in header_table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "none")
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        header.add_paragraph().paragraph_format.space_before = Pt(3)

        footer.add_paragraph().paragraph_format.space_after = Pt(3)
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = footer_para.add_run("© FALCON AUTOTECH 2025 Confidential: Not for Distribution. ")
        run.font.name = "Calibri (Body)"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)

        add_hyperlink(footer_para, "https://www.falconautotech.com/", "https://www.falconautotech.com/")

        page_run = footer_para.add_run(" | Page ")
        page_run.font.name = "Calibri (Body)"
        page_run.font.size = Pt(9)
        page_run.font.color.rgb = RGBColor(0, 0, 0)

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        page_run._r.append(fld_begin)
        page_run._r.append(instr)
        page_run._r.append(fld_end)

        total_run = footer_para.add_run(" of ")
        total_run.font.name = "Calibri (Body)"
        total_run.font.size = Pt(9)
        total_run.font.color.rgb = RGBColor(0, 0, 0)

        fld_total_begin = OxmlElement("w:fldChar")
        fld_total_begin.set(qn("w:fldCharType"), "begin")
        instr_total = OxmlElement("w:instrText")
        instr_total.set(qn("xml:space"), "preserve")
        instr_total.text = "NUMPAGES"
        fld_total_end = OxmlElement("w:fldChar")
        fld_total_end.set(qn("w:fldCharType"), "end")
        total_run._r.append(fld_total_begin)
        total_run._r.append(instr_total)
        total_run._r.append(fld_total_end)


def create_cover_page(
    client_logo: Optional[bytes],
    client_name: str,
    project_title: str,
) -> io.BytesIO:
    """Create a cover page using template - exactly as in main.py"""
    template_path = "assests\\Templates\\Cover_Temp.docx"
    doc = Document(template_path)

    # Remove all headers and footers from template
    for sec in doc.sections:
        for part in (
            getattr(sec, "header", None),
            getattr(sec, "footer", None),
            getattr(sec, "first_page_header", None),
            getattr(sec, "first_page_footer", None),
            getattr(sec, "even_page_header", None),
            getattr(sec, "even_page_footer", None),
        ):
            if not part:
                continue
            try:
                part.is_linked_to_previous = False
            except Exception:
                pass
            try:
                for tbl in list(part.tables):
                    tbl._element.getparent().remove(tbl._element)
                for p in list(part.paragraphs):
                    p._element.getparent().remove(p._element)
            except Exception:
                pass

    # Add client logo if provided - process with PIL to ensure proper embedding
    if client_logo:
        try:
            # Open and process image
            im = Image.open(io.BytesIO(client_logo))
            if im.mode != "RGBA":
                im = im.convert("RGBA")
            alpha = im.getchannel("A")
            bbox = alpha.getbbox()
            if bbox:
                im = im.crop(bbox)
                alpha = im.getchannel("A")
            # Create white background and paste
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=alpha)

            # Save to buffer
            buf = io.BytesIO()
            bg.save(buf, format="PNG")
            buf.seek(0)

            # Insert at beginning
            first_para = doc.paragraphs[0]
            run_logo = first_para.insert_paragraph_before().add_run()
            run_logo.add_picture(buf, width=Inches(2.0))
        except Exception:
            # Fallback: insert without processing
            first_para = doc.paragraphs[0]
            run_logo = first_para.insert_paragraph_before().add_run()
            run_logo.add_picture(io.BytesIO(client_logo), width=Inches(2.0))

    # Add spacing
    for _ in range(6):
        doc.add_paragraph("")

    # Add title
    title = f"FALCON's Proposal to {client_name} for the {project_title}"
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = False
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add date
    today_str = datetime.today().strftime("%B %d, %Y")
    p2 = doc.add_paragraph()
    run2 = p2.add_run(today_str)
    run2.font.size = Pt(14)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(255, 215, 0)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add page break after cover page
    doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==================== ADDITIONAL HELPER FUNCTIONS ====================

def extract_pdf_text(uploaded_file) -> str:
    """Extract plain text from an uploaded PDF using pdfplumber."""
    if uploaded_file is None:
        return ""

    text_chunks = []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            text_chunks.append(page.extract_text() or "")

    full_text = "\n\n".join(text_chunks)
    # Hard truncate to keep prompt size reasonable
    if len(full_text) > 20000:
        full_text = full_text[:20000]
    return full_text

def choose_sorter_template(project_name: str) -> SorterTemplate:
    """Pick the closest template based on project name keywords."""
    text = (project_name or "").lower()

    # score by number of keyword hits
    best_tpl = SORTER_TEMPLATES[0]
    best_score = -1
    for tpl in SORTER_TEMPLATES:
        score = sum(1 for kw in tpl.keywords if kw in text)
        if score > best_score:
            best_score = score
            best_tpl = tpl

    return best_tpl

def shade_cell(cell, color_hex: str = "D9D9D9"):
    """Apply gray shading to a table cell"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)

def add_markdown_line(doc: Document, line: str):
    """Add paragraph with **bold** segments."""
    p = doc.add_paragraph()
    parts = line.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def add_markdown_paragraph(doc: Document, text: str, style: str | None = None):
    """Add a paragraph with simple **bold** Markdown handling."""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()

    parts = re.split(r"(\*\*[^\*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def add_boxed_text(doc: Document, text: str, font_size: int = 16, bold: bool = True):
    """Grey shaded single-cell table with centered text (for cover letter front page)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, "D9D9D9")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    return table

def add_centered_upload_image(doc: Document, uploaded_file, width_in: float = 6.0):
    """Add uploaded image centered"""
    if not uploaded_file:
        return
    img_stream = BytesIO(uploaded_file.getvalue())
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(img_stream, width=Inches(width_in))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def call_groq_cover_letter(
    client_name: str,
    project_title: str,
    offer_ref: str,
    letter_date_str: str,
    executives_block: str,
    invitation_date: str,
    meeting_date: str,
    sender_name: str,
    sender_title: str,
    process_flow_summary: str = "",
    context: Optional[ProposalContext] = None,
) -> str:
    """Call Groq API to generate the cover letter text."""
    counts_block = context.counts_block_text() + "\n\n" if context else ""
    gating_issues = validate_context_counts(context) if context and ENABLE_CONTEXT_UNIFICATION else []
    
    # Extract PPH from context (or use placeholder)
    pph_value = "Not specified"
    if context and context.pph is not None:
        pph_value = f"{context.pph:,} PPH"
    
    user_prompt = (
        counts_block + COVER_LETTER_USER_PROMPT_TEMPLATE.format(
            client_name=client_name,
            project_title=project_title,
            offer_ref=offer_ref,
            letter_date=letter_date_str,
            executives_block=executives_block.strip() or "Not provided",
            invitation_date=invitation_date.strip() or "Not provided",
            meeting_date=meeting_date.strip() or "Not provided",
            process_flow_summary=process_flow_summary.strip() or "Not provided",
            sender_name=sender_name,
            sender_title=sender_title,
            PPH=pph_value,
        ) + (
            "\n\nIf any counts are missing, avoid inventing numbers; use safe phrasing without quantities." if gating_issues else ""
        )
    )
    if context:
        ctx_counts = context.counts_block_text().replace('\n',' | ')
        logger.info(f"Using ProposalContext counts for Cover Letter: {ctx_counts}")

    def api_call():
        return groq_client.chat.completions.create(
            model="groq/compound",
            messages=[
                {"role": "system", "content": COVER_LETTER_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_tokens=600,  # Increased to allow for process flow details while staying under 300 words
        )
    
    completion = call_groq_with_retry(api_call)
    text = completion.choices[0].message.content.strip()
    if text.startswith("```"):
        parts = text.split("```")
        if len(parts) >= 2:
            text = parts[1]
            if text.startswith("text\n") or text.startswith("markdown\n"):
                text = "\n".join(text.split("\n")[1:])
    
    # Apply text normalizations: Cross Belt Sorter capitalization and client name consistency
    text = normalize_proposal_text(text.strip(), client_name)
    
    # Sanitize LLM output to remove meta-artifacts
    text = sanitize_section_output(text)
    
    return text

def call_groq_exec_summary(
    system_text: str,
    client_name: str,
    project_title: str,
    pph_count: str = "",
    cbs_type: str = "",
    dxf_json: dict = None,
    facts: Optional[ProposalFacts] = None,
    context: Optional[ProposalContext] = None,
) -> str:
    """Call Groq API to generate the Executive Summary text."""
    # Use ProposalContext counts when provided; otherwise fallback to legacy behavior
    component_counts = ""
    if context:
        # Only use context counts; do not derive from DXF
        ctx_counts = context.counts_block_text().replace('\n',' | ')
        logger.info(f"Using ProposalContext counts for Executive Summary: {ctx_counts}")
        pph_count = pph_count or (str(context.pph) if context.pph is not None else "")
        cbs_type = cbs_type or (context.cbs_type or "Cross-belt technology")
    elif dxf_json:
        # Legacy fallback when context is unavailable
        cat_summary = dxf_json.get('category_summary', {})
        chute_analysis = dxf_json.get('chute_analysis', {})
        categorized = dxf_json.get('categorized_components', {})
        counts_lines = []
        feedline_count = cat_summary.get('AUTO_INDUCT', 0)
        if feedline_count > 0:
            counts_lines.append(f"Induct Lines (Automatic): {feedline_count}")
        manual_count = cat_summary.get('OPERATOR_STATION', 0)
        if manual_count > 0:
            counts_lines.append(f"Manual Loading Point: Present (combine with induct lines bullet)")
        telescopic_count = 0
        for name, count in categorized.get('CONVEYOR_INFEED', {}).items():
            if 'telescopic' in name.lower() or 'tbc' in name.lower():
                telescopic_count += count
        if telescopic_count > 0:
            counts_lines.append(f"Telescopic Conveyors: {telescopic_count}")
        if chute_analysis:
            total_chutes = chute_analysis.get('total', 0)
            chute_types = chute_analysis.get('breakdown', {})
            chute_details = []
            for ctype, ccount in chute_types.items():
                if ccount > 0:
                    chute_details.append(f"{ccount} {ctype}")
            if chute_details:
                counts_lines.append(f"Chutes: {', '.join(chute_details)} (Total: {total_chutes})")
            elif total_chutes > 0:
                counts_lines.append(f"Total Chutes: {total_chutes}")
        if counts_lines:
            component_counts = "\n".join(counts_lines)
    
    counts_block = context.counts_block_text() + "\n\n" if context else ""
    gating_issues = validate_context_counts(context) if context and ENABLE_CONTEXT_UNIFICATION else []
    user_content = counts_block + (
        f"Client Name: {client_name}\n"
        f"Project / System Name: {project_title}\n"
        f"PPH (Throughput): {pph_count if pph_count else 'Not specified'}\n"
        f"CBS Type: {cbs_type if cbs_type else 'Cross-belt technology'}\n\n"
    )
    
    if component_counts and not context:
        user_content += f"COMPONENT COUNTS FROM DXF:\n{component_counts}\n\n"
    
    user_content += (
        f"Proposed System Description (for context):\n{system_text}\n\n"
        "Generate the Executive Summary strictly as per the instructions.\n"
        "**CRITICAL:** Write 'Our solution is based on the following key characteristics:' only ONCE.\n"
        "**CRITICAL:** Each bullet must be a FULL SENTENCE with context, not just item names.\n"
        "**CRITICAL:** Do NOT count manual stations separately - mention as 'manual loading point' in induct bullet.\n"
        "**CRITICAL:** Combine induct lines + manual loading in ONE bullet.\n"
        "**CRITICAL:** Combine all chute types in ONE bullet with counts."
    )
    if gating_issues:
        user_content += "\n\nIf counts are missing, avoid numeric statements; use qualitative phrasing."
    if context:
        ctx_counts = context.counts_block_text().replace('\n',' | ')
        logger.info(f"Using ProposalContext counts for Executive Summary: {ctx_counts}")

    def api_call():
        return groq_client.chat.completions.create(
            model="groq/compound",
            temperature=0.4,
            max_tokens=1200,
            messages=[
                {"role": "system", "content": EXEC_SUMMARY_SYSTEM_PROMPT},
                {"role": "user", "content": user_content},
            ],
        )
    
    resp = call_groq_with_retry(api_call)
    summary_text = resp.choices[0].message.content.strip()
    
    # Apply text normalizations: Cross Belt Sorter capitalization and client name consistency
    summary_text = normalize_proposal_text(summary_text, client_name)
    
    # Sanitize LLM output to remove meta-artifacts
    summary_text = sanitize_section_output(summary_text)
    
    return summary_text


def correct_section_numbers(section_text: str, context: ProposalContext, section_name: str) -> str:
    """Run a minimal corrective pass to align numbers with ProposalContext without rewriting."""
    if not section_text or not context:
        return section_text
    
    # First apply deterministic enforcement (no LLM call if possible)
    enforced = enforce_counts(section_text, context)
    
    # Only use LLM if deterministic enforcement wasn't sufficient
    # For now, just return the deterministically enforced version with sanitization
    result = sanitize_section_output(enforced)
    return result if result else section_text


# System prompt used for extracting price-sheet JSON from costing CSV
PRICE_SHEET_SYSTEM_PROMPT = """
You are a senior commercial analyst for warehouse automation projects.

You receive the contents of an Excel sheet called "Overall Costing" as raw CSV text.
This sheet may contain many detailed costing lines, intermediate totals, taxes, and notes.

Your job is to infer the HIGH-LEVEL "Price Sheet" summary used in proposals.

The high-level Price Sheet is a short table of a few summary lines (typically 3–15),
each corresponding to a major package/component of the solution with a single rolled-up price.
Do NOT list detailed items like small sub-components or line-by-line BOM;
only show the SUMMARY building blocks that a customer would see in the commercial section.

------------------------------------------------
EXAMPLES OF TARGET PRICE SHEETS (FOR REFERENCE)
------------------------------------------------

Example 1 –

Price List- Summary
S.NO   Package                    Price
1      Conveyors Package          ₹ 28,34,27,926
2      Cross Belt Sorter Package  ₹ 29,16,52,094
3      Destinations Package       ₹ 3,83,87,721
4      Services Package           ₹ 3,46,22,786
       Total                      ₹ 64,80,90,527

"Business Cooperation Agreement"
Discount for Delhivery  4.5%
Final Total             ₹ 61,89,26,453


Example 2 – 

Price Sheet
S. No   Component                               Price (USD)
1       Loop CBS + Inducts                      (included or price)
2       Infeed + Bagging Conveyors             $ 398,105
3       Output Chutes                          $ 219,944
4       Software Package & Integration         $ 26,302
5       Packaging & forwarding                 $ 3,523
6       Project Management + Supervision cost  $ 32,268
Total (USD)                                    $ 726,386

---------------------------------------
TASK – WHAT YOU MUST RETURN
---------------------------------------

Use the raw 'Overall Costing' CSV to reconstruct ONLY the high-level summary.

1) Identify the main commercial building blocks, such as:
   - Conveyors Package
   - Cross Belt Sorter Package
   - Destinations Package
   - Services Package
   - Loop CBS + Inducts
   - Infeed + Bagging Conveyors
   - Output Chutes
   - Software package / Software Packages & SCADA
   - Steelworks / Steelwork
   - PTL
   - Installation & Commissioning
   - Project Management and Engineering Charges
   - Packaging & forwarding / Packaging & Documentation
   - Freight, Warranty, Hotline, AMC packages
   or similar high-level components used to summarize the cost.

2) For each such high-level component, return:
   - s_no: integer starting from 1 in sequence
   - label: the package/component name in clean human-readable form
   - price: the final total price for that component AS A STRING,
            including currency symbol and formatting exactly as in the sheet

   IMPORTANT:
   - Do NOT invent prices.
   - Use values that actually appear in the sheet.
   - If multiple detailed rows roll up into one package, use the rolled-up total.

3) If there is a grand "Total", also return:
   - total_row: { "label": "...", "price": "..." }

4) If there are explicit discount and final total lines:
   - cooperation_label: e.g. "Business Cooperation Agreement"
   - discount_label: e.g. "Discount for Delhivery"
   - discount_value: e.g. "4.5%"
   - final_total_label: e.g. "Final Total"
   - final_total_value: e.g. "₹ 61,89,26,453"
   If not present, return them as null.

5) Also return:
   - currency: "INR", "SAR", "USD", "EUR", or "MIXED" if multiple currencies appear.
   - price_sheet_title: a short label like "Price List – Summary" or "Price Sheet"

6) OUTPUT FORMAT:

Return ONLY a single valid JSON object with this exact shape:

{
  "currency": "INR" | "SAR" | "USD" | "EUR" | "MIXED" | null,
  "price_sheet_title": "string or null",
  "items": [
    {
      "s_no": 1,
      "label": "Conveyors Package",
      "price": "₹ 28,34,27,926"
    },
    ...
  ],
  "total_row": {
    "label": "Total",
    "price": "₹ 64,80,90,527"
  } or null,
  "cooperation_label": "string or null",
  "discount_label": "string or null",
  "discount_value": "string or null",
  "final_total_label": "string or null",
  "final_total_value": "string or null"
}

Do NOT wrap the JSON in markdown code fences.
Do NOT add explanations or commentary.
Just return the JSON object.
"""

PRICE_SHEET_USER_PROMPT_TEMPLATE = """
Below is the raw CSV export of the 'Overall Costing' sheet of an internal costing file.

Use it to construct the high-level Price Sheet summary as described in the instructions.

Raw CSV:
--------------------
{sheet_csv}
--------------------
"""

def call_groq_for_price_sheet(sheet_csv: str) -> dict:
    """Call Groq API to extract price sheet from costing CSV"""
    user_prompt = PRICE_SHEET_USER_PROMPT_TEMPLATE.format(sheet_csv=sheet_csv)

    def api_call():
        return groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": PRICE_SHEET_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.0,
        )
    
    completion = call_groq_with_retry(api_call)
    raw = completion.choices[0].message.content.strip()

    # Strip markdown fences if present
    if raw.startswith("```"):
        parts = raw.split("```")
        if len(parts) >= 2:
            raw = parts[1]
            raw = raw.lstrip("json").lstrip()

    # Extract JSON from first '{' to last '}'
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end < start:
        raise ValueError(f"Groq response does not contain a JSON object:\n{raw}")

    content = raw[start:end + 1]

    try:
        data = json.loads(content)
    except json.JSONDecodeError as e:
        raise ValueError(f"Groq response was not valid JSON: {e}\nExtracted content:\n{content}")

    return data

def parse_price_string(price_str: str):
    """Extract currency prefix and numeric value from price string"""
    if not price_str:
        return None, None, 0

    m = re.search(r"[-]?\d", price_str)
    if not m:
        return price_str.strip(), None, 0

    prefix = price_str[:m.start()].strip()
    numeric_part = price_str[m.start():].strip()

    digits_only = "".join(ch for ch in numeric_part if ch.isdigit() or ch == ".")
    if digits_only == "":
        return prefix, None, 0

    decimals_count = 0
    if "." in digits_only:
        decimals_count = len(digits_only.split(".")[1])

    try:
        value = float(digits_only)
    except ValueError:
        return prefix, None, decimals_count

    return prefix, value, decimals_count

def format_indian_number(value: float, decimals: int) -> str:
    """Format number with Indian-style digit grouping"""
    if decimals > 0:
        s = f"{value:.{decimals}f}"
    else:
        s = f"{int(round(value))}"

    if "." in s:
        int_part, frac = s.split(".")
    else:
        int_part, frac = s, None

    # Indian grouping
    if len(int_part) > 3:
        last3 = int_part[-3:]
        head = int_part[:-3]
        groups = []
        while len(head) > 2:
            groups.insert(0, head[-2:])
            head = head[:-2]
        if head:
            groups.insert(0, head)
        int_formatted = ",".join(groups + [last3])
    else:
        int_formatted = int_part

    if frac and decimals > 0:
        return int_formatted + "." + frac
    else:
        return int_formatted

def apply_bca_discount_to_price_data(price_data: dict, discount_percent: float) -> str | None:
    """Apply BCA discount on total_row.price and return discounted price string"""
    total_row = price_data.get("total_row")
    if not total_row:
        return None

    price_str = total_row.get("price")
    prefix, value, decimals = parse_price_string(price_str)
    if value is None:
        return None

    discounted = value * (1 - discount_percent / 100.0)
    formatted_number = format_indian_number(discounted, decimals)
    if prefix:
        return f"{prefix} {formatted_number}"
    else:
        return formatted_number

# ==================== CAPACITY CALCULATIONS FUNCTIONS ====================

def build_capacity_prompt_from_excel(
    excel_bytes: bytes,
    client_name: str,
    project_name: str
) -> str:
    """
    Read the uploaded Excel (all sheets), dump them as CSV text,
    and build a very explicit extraction prompt for GROQ.
    We do NOT try to interpret any cell ourselves.
    """
    xls = pd.ExcelFile(BytesIO(excel_bytes))

    sheet_dumps = []
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet, header=None)
        # Keep as CSV-like text to preserve structure
        csv_text = df.to_csv(index=False, header=False)
        sheet_dumps.append(f"### Sheet: {sheet}\n{csv_text}")

    workbook_text = "\n\n".join(sheet_dumps)

    # IMPORTANT: we define a strict JSON schema and explicitly
    # tell GROQ to set fields to null if they are missing.
    prompt = f"""
You are an expert in interpreting throughput and capacity calculation Excel sheets
for parcel/shipment sortation systems (Loop CBS, Linear CBS, Cross Belt Sorters, etc.).

You are given a raw text dump of the complete Excel workbook used for capacity calculations.
Using ONLY the information present in the workbook (numbers and labels), you must extract
or compute the key capacity fields and return them as a single JSON object.

Context:
- Client: {client_name}
- Project: {project_name}

The workbook text follows after this instruction. It is a concatenation of all sheets, each
in CSV-like form.

IMPORTANT RULES:

1. **Use exact numbers from the workbook wherever a field is explicitly present.**
   - If a value is written in the sheet (e.g. "Sorter Speed 2 m/s", "Carrier per hour 6128"),
     prefer the sheet value instead of recomputing it.
2. **Only compute** a value if:
   - It is clearly implied (e.g. carrier_per_hour = speed_mps * 3600 / pitch_m) AND
   - It is NOT already available as a direct cell value.
3. If a field is not given and cannot be safely derived, set it explicitly to null.

KEY FIELDS (SEMANTICS):

- sorter_type:
    A short human-readable description like "Loop CBS", "Linear CBS", "Dual Belt Loop CBS"
    or "Cross Belt Sorter". Use what best matches the workbook text.

- sorter_speed_mps:
    Sorter speed in meters per second. If sheet says "Speed 2 m/s", set 2.0.

- pitch_m:
    Carrier pitch in meters. If sheet says "Pitch 1,175 mm", then pitch_m = 1.175.

- carriers_per_hour_cph:
    "Carrier per hour" / "Carriers/Hour" / "Carriers per hour" from the sheet.
    If not present, you may compute as:
      carriers_per_hour = speed_mps * 3600 / pitch_m
    and round to nearest integer.

- belts_per_hour_bph:
    "Belts per hour" / "Belts/Hour" from the sheet.
    If not present but the sorter is clearly Dual Belt, you may compute:
      belts_per_hour = carriers_per_hour * 2
    If single belt, belts_per_hour = carriers_per_hour.

- num_feedlines:
    Number of feedlines / inducts / infeed lines (e.g. "No of Feedlines", "No of Inducts").
    If the workbook has multiple such numbers, choose the one used in the capacity section.

- num_operators:
    Number of operators used in capacity calculations, if explicitly given
    (e.g. "No of Operators", "No of operators on manual induct station").
    If not given, set null.

- capacity_per_operator_pph:
    Capacity per operator in parcels/shipments per hour, if explicitly given
    (e.g. "Capacity per operator 1000 Shipments per hour"). If not given, set null.

- sorter_designed_capacity_A_pph:
    Sorter designed capacity on the parcel spectrum. Look for labels like:
    "Sorter Designed Capacity (A)", "Effective Designed Throughput of Sorter (A)",
    "Effective Designed TPH", or similar. Use the PPH/Shipments per hour value.

- feedline_designed_capacity_B_pph:
    Total feedline/induction capacity. Look for labels like:
    "Total Feedline designed capacity (B)", "Induction Capacity (B)",
    "Total Induction Capacity Designed", etc. Use the PPH value.

- effective_capacity_min_AB_pph:
    The effective designed capacity of the system.
    If the sheet already has "System designed throughput" or "Operational capacity",
    use that value.
    If not explicitly given, compute:
       effective_capacity_min_AB_pph = min(sorter_designed_capacity_A_pph,
                                           feedline_designed_capacity_B_pph)
    (if both are known).

- single_belt_pct and dual_belt_pct:
    Percentages of shipments handled on single and dual belts, if present
    (e.g. "Single Belts Shipments 91.36%", "Dual Belt Shipments 8.64%").
    Store them as numeric percentages (e.g. 91.36, 8.64).
    If not present, set them to null.

JSON SCHEMA (MANDATORY KEYS):

You MUST return exactly one JSON object with ALL of these keys:

{{{{
  "sorter_type": "Loop CBS or Linear CBS or Cross Belt Sorter etc.",
  "sorter_speed_mps": 2.0,
  "pitch_m": 1.175,
  "carriers_per_hour_cph": 0,
  "belts_per_hour_bph": 0,
  "num_feedlines": 0,
  "num_operators": null,
  "capacity_per_operator_pph": null,
  "sorter_designed_capacity_A_pph": 0,
  "feedline_designed_capacity_B_pph": 0,
  "effective_capacity_min_AB_pph": 0,
  "single_belt_pct": null,
  "dual_belt_pct": null
}}}}

RESPONSE FORMAT REQUIREMENTS (CRITICAL):

- Output MUST be **only** a JSON object.
- Do NOT include markdown, explanations, or any text outside the JSON.
- All numeric values must be raw numbers (no units, no commas, no % signs).
- If a value is unknown or not present, set it to null (not 0).
- DO NOT add ```json``` or json in the response. ONLY return raw JSON.
**MUST : DO NOT add ```json``` or json in the response. ONLY return raw JSON.**
Below is the full workbook dump:

{workbook_text}
"""
    return prompt


def call_groq_for_capacity(prompt: str) -> dict:
    """
    Call GROQ with response_format=json_object so that we reliably get JSON.
    """
    if not GROQ_API_KEY:
        raise RuntimeError("GROQ_API_KEY is not set in environment variables.")

    client = Groq(api_key=GROQ_API_KEY)

    def api_call():
        return client.chat.completions.create(
            model="groq/compound",
            messages=[
                {
                    "role": "system",
                    "content": "You are a precise JSON data extractor. Always follow the schema exactly."
                },
                {
                    "role": "user",
                    "content": prompt
                },
            ],
            response_format={"type": "json_object"},
            temperature=0.0,
        )
    
    chat_completion = call_groq_with_retry(api_call)
    raw = chat_completion.choices[0].message.content
    
    # Handle empty or None response
    if not raw or raw.strip() == "":
        raise ValueError("Groq API returned empty response for capacity extraction")
    
    # Clean up potential markdown code blocks
    raw = raw.strip()
    if raw.startswith("```json"):
        raw = raw[7:]
    if raw.startswith("```"):
        raw = raw[3:]
    if raw.endswith("```"):
        raw = raw[:-3]
    raw = raw.strip()
    
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        # Log the raw response for debugging
        print(f"[DEBUG] Groq capacity response that failed to parse: {raw[:500]}")
        raise ValueError(f"Failed to parse Groq response as JSON: {str(e)}. Response was: {raw[:200]}...")


def call_groq_for_sorter_spec(sheet_name: str, sheet_text: str) -> dict:
    """
    Call GROQ API to extract sorter technical specifications from Loop CBS Excel sheet.
    Returns a JSON object with sorter specifications.
    """
    system_prompt = """You are a technical proposal engineer reading an Excel costing/configuration sheet
for a Loop Cross Belt Sorter ("Loop CBS").

You will receive:
- The sheet name (e.g., "Loop CBS")
- The ENTIRE sheet content as text, row by row, including ALL tables.

Your task:
Extract the following fields, strictly from the sheet content:

1) sorter_carrier_type        – fixed string "Loop CBS".
2) sorter_speed_mps           – fixed string "upto 2 m/s"
3) sorter_loop_length_m       – sorter loop length in meters (e.g., "150").
4) sorter_height_mm           – sorter height in mm (e.g., "2900").
5) actuation_technology       – fixed string "Electric".
6) carrier_pitch_mm           – carrier pitch in mm (e.g., "600", "1175", "1200"). If multiple models,
                                choose the one actually selected in the configuration area.
7) number_of_carriers         – total number of carriers in the sorter, as shown in the sheet if present.
8) motor_drive_type           – description of motor/drive type (e.g., "LIM", "LSM", "LIM + LSM", etc.).
9) power_consumption          – sorter power consumption (kW or kVA etc.) taken from the sheet.

VERY IMPORTANT RULES:
- Use ONLY information present in the provided sheet text. Do NOT guess or invent values.
- If a value is not clearly present, set it to null.
- If the sheet expresses a choice ("Select Model", "Enter Loop Length", etc.), use the chosen values.
- Keep the output values SHORT: just the numeric value or the short phrase, without explanations.

Return a single JSON object with EXACTLY these keys:

{
  "sorter_carrier_type": "...",
  "sorter_speed_mps": "... or null",
  "sorter_loop_length_m": "... or null",
  "sorter_height_mm": "... or null",
  "actuation_technology": "...",
  "carrier_pitch_mm": "... or null",
  "number_of_carriers": "... or null",
  "motor_drive_type": "... or null",
  "power_consumption": "... or null"
}

No comments, no trailing text, no markdown."""

    user_payload = {
        "sheet_name": sheet_name,
        "sheet_text": sheet_text,
    }

    def api_call():
        return groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt.strip()},
                {"role": "user", "content": json.dumps(user_payload, indent=2)},
            ],
            response_format={"type": "json_object"},
            temperature=0.1,
        )

    try:
        resp = call_groq_with_retry(api_call)
        text = resp.choices[0].message.content.strip()
        spec = json.loads(text)
        
        if not isinstance(spec, dict):
            raise ValueError("Expected a JSON object")
        
        # Fill fixed fields if missing
        if not spec.get("sorter_carrier_type"):
            spec["sorter_carrier_type"] = "Loop CBS"
        if not spec.get("actuation_technology"):
            spec["actuation_technology"] = "Electric"
        
        return spec
    except Exception as exc:
        st.error(f"Failed to parse GROQ response for sorter specification: {exc}")
        return None


def add_capacity_section_to_doc(
    doc: Document,
    client_name: str,
    project_name: str,
    cap: dict,
    counter: int,
    context: Optional[ProposalContext] = None,
) -> None:
    """
    Add 'Sorter System Capacity' section to an existing Document,
    using the extracted capacity dict.
    """
    # Heading
    add_numbered_heading(doc, "Sorter System Capacity", counter=counter)

    intro_para = (
        f"The following table shows the throughput calculation for the sortation system "
        f"designed based on {client_name}'s {project_name} requirements."
    )
    p = doc.add_paragraph(intro_para)
    apply_normal_style(p)

    # Table: SPECIFICATION | VALUE
    table = doc.add_table(rows=1, cols=2)
    table.style = "Medium Shading 1 Accent 1"

    hdr = table.rows[0].cells
    hdr[0].text = "SPECIFICATION"
    hdr[1].text = "VALUE"

    def fmt(value, suffix=""):
        if value is None or value == "":
            return "N/A"
        return f"{value}{suffix}"

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    # Enforce feedlines from context if confirmed
    if context and context.get("feedlines").value is not None:
        cap["num_feedlines"] = context.get("feedlines").value

    # Fill rows
    add_row("Sorter Type", cap.get("sorter_type", ""))

    # Speed / pitch
    add_row("Sorter Speed", fmt(cap.get("sorter_speed_mps"), " m/s"))
    add_row("Pitch", fmt(cap.get("pitch_m"), " m"))

    # Capacity raw
    add_row("Carrier per hour", fmt(cap.get("carriers_per_hour_cph"), " CPH"))
    add_row("Belts per hour", fmt(cap.get("belts_per_hour_bph"), " BPH"))

    # Feedlines / operators
    add_row("No. of Feedlines", fmt(cap.get("num_feedlines")))
    add_row("No. of Operators", fmt(cap.get("num_operators")))
    add_row("Capacity per Operator", fmt(cap.get("capacity_per_operator_pph"), " PPH"))

    # Sorter vs Feedline capacity
    add_row(
        "Sorter Designed Capacity (A)",
        fmt(cap.get("sorter_designed_capacity_A_pph"), " PPH"),
    )
    add_row(
        "Feedline Designed Capacity (B)",
        fmt(cap.get("feedline_designed_capacity_B_pph"), " PPH"),
    )
    add_row(
        "Effective Designed Capacity (min of A & B)",
        fmt(cap.get("effective_capacity_min_AB_pph"), " PPH"),
    )

    # Optional single / dual belt %
    if cap.get("single_belt_pct") is not None or cap.get("dual_belt_pct") is not None:
        add_row(
            "Single Belt Shipments",
            fmt(cap.get("single_belt_pct"), " %"),
        )
        add_row(
            "Dual Belt Shipments",
            fmt(cap.get("dual_belt_pct"), " %"),
        )

# ==================== INPUT COLLECTION ====================

# Initialize page state
if "page" not in st.session_state:
    st.session_state.page = "input"  # "input" or "preview"

# ==================== PREVIEW PAGE ====================
if st.session_state.page == "preview" and "generated_pdf_buffer" in st.session_state and st.session_state.generated_pdf_buffer is not None:
    import streamlit.components.v1 as components
    
    pdf_buffer = st.session_state.generated_pdf_buffer
    filename = st.session_state.get("generated_doc_filename", "Falcon_Proposal.pdf")
    docx_filename = filename.replace('.pdf', '.docx') if filename.endswith('.pdf') else filename + '.docx'
    
    # Initialize chat messages in session state
    if "feedback_chat_messages" not in st.session_state:
        st.session_state.feedback_chat_messages = []
    
    # Professional CSS for preview page
    st.markdown("""
    <style>
        .preview-header {
            background: linear-gradient(135deg, #1a1f36 0%, #2d3748 100%);
            padding: 1.25rem 1.5rem;
            border-radius: 8px;
            margin-bottom: 1.5rem;
            border-left: 4px solid #3182ce;
        }
        .preview-header h2 {
            color: #ffffff;
            font-size: 1.4rem;
            font-weight: 600;
            margin: 0;
            letter-spacing: -0.3px;
        }
        .preview-header .subtitle {
            color: #a0aec0;
            font-size: 0.85rem;
            margin-top: 0.25rem;
        }
        .chat-section-header {
            background: #f7fafc;
            border: 1px solid #e2e8f0;
            border-radius: 6px;
            padding: 0.75rem 1rem;
            margin-bottom: 1rem;
        }
        .chat-section-header h4 {
            color: #2d3748;
            font-size: 0.95rem;
            font-weight: 600;
            margin: 0 0 0.25rem 0;
        }
        .chat-section-header p {
            color: #718096;
            font-size: 0.8rem;
            margin: 0;
        }
        .pdf-section-header {
            background: #f7fafc;
            border: 1px solid #e2e8f0;
            border-radius: 6px;
            padding: 0.75rem 1rem;
            margin-bottom: 1rem;
        }
        .pdf-section-header h4 {
            color: #2d3748;
            font-size: 0.95rem;
            font-weight: 600;
            margin: 0;
        }
        .status-success {
            background: #f0fff4;
            border: 1px solid #9ae6b4;
            border-radius: 6px;
            padding: 0.6rem 1rem;
            margin-bottom: 1rem;
        }
        .status-success span {
            color: #276749;
            font-size: 0.85rem;
            font-weight: 500;
        }
    </style>
    """, unsafe_allow_html=True)
    
    
    # Create two-column layout: PDF Preview | Chat
    pdf_col, chat_col = st.columns([2.3, 1], gap="medium")
    
    # ==================== RIGHT COLUMN: CHAT & CONTROLS ====================
    with chat_col:     
        # Chat container with scrollable history
        chat_container = st.container(height=380)
        with chat_container:
            if not st.session_state.feedback_chat_messages:
                st.markdown("""
                <div style="text-align: center; color: #718096; padding: 40px 20px;">
                    <p style="font-size: 0.9rem; margin-bottom: 0.5rem;">No messages yet</p>
                    <p style="font-size: 0.8rem; color: #a0aec0;">
                        Type a message below to request changes to your proposal
                    </p>
                </div>
                """, unsafe_allow_html=True)
            else:
                for message in st.session_state.feedback_chat_messages:
                    if message["role"] == "user":
                        with st.chat_message("user"):
                            st.markdown(message["content"])
                    else:
                        with st.chat_message("assistant"):
                            st.markdown(message["content"])
        
        # Chat input
        user_feedback = st.chat_input("Describe the changes you need...")
        
        if user_feedback:
            # Add user message to chat
            st.session_state.feedback_chat_messages.append({
                "role": "user",
                "content": user_feedback
            })
            
            # Step 1: FIRST identify which section to modify
            # Gather all stored section contents for context
            stored_sections = {
                key: st.session_state.get(key, "") 
                for key in st.session_state.keys() 
                if key.startswith("section_content_")
            }
            
            # Initial quick parse to get section (with ALL section content for context)
            with st.spinner("Analyzing request..."):
                initial_actionable = rephrase_feedback_to_actionable(
                    user_feedback, 
                    section_content=None,
                    all_sections_content=stored_sections
                )
            
            with st.spinner("Identifying target section..."):
                section_result = identify_section_with_context(user_feedback, initial_actionable, stored_sections)
            
            section_key = section_result.get("section_key", "unknown")
            section_name = section_result.get("section_name", "Unknown")
            confidence = section_result.get("confidence", "low")
            reasoning = section_result.get("reasoning", "")
            
            # Check if the section exists in our editable sections
            section_info = PROPOSAL_SECTIONS.get(section_key, None)
            
            if section_key == "unknown" or section_info is None:
                response = f"Unable to identify the target section for your request.\n\n**Editable sections:**\n- Cover Letter\n- Executive Summary\n- Proposed System Description\n- System Description\n- Technical Details\n\nPlease specify which section you would like to modify."
            else:
                # Check if we have the original content stored
                stored_content_key = f"section_content_{section_key}"
                original_content = st.session_state.get(stored_content_key, None)
                
                if original_content:
                    # Step 2: NOW rephrase with the actual section content AND all sections for comprehensive analysis
                    with st.spinner("Processing changes..."):
                        actionable_info = rephrase_feedback_to_actionable(
                            user_feedback, 
                            section_content=original_content,
                            all_sections_content=stored_sections
                        )
                    
                    # Get context for regeneration
                    context = {
                        "client_name": st.session_state.get("last_client_name", ""),
                        "project_name": st.session_state.get("last_project_name", ""),
                    }
                    
                    # Step 3: Regenerate the section with actionable feedback
                    with st.spinner("Applying changes..."):
                        new_content = regenerate_section_with_feedback(
                            section_key, original_content, user_feedback, actionable_info, context
                        )
                    
                    # Store the regenerated content
                    st.session_state[stored_content_key] = new_content
                    st.session_state[f"section_regenerated_{section_key}"] = True
                    
                    # Store actionable info for direct value replacement in DOCX
                    if "pending_docx_replacements" not in st.session_state:
                        st.session_state.pending_docx_replacements = []
                    
                    # Extract old/new values for direct replacement
                    # CRITICAL: Only use actual values, NOT placeholders
                    entities = actionable_info.get("entities", {})
                    old_value = entities.get("old_value", "")
                    new_value = entities.get("new_value", "")
                    target = entities.get("target", "")
                    
                    # Only store if we have actual values (not placeholders)
                    if old_value and new_value and not (old_value.startswith("[") and old_value.endswith("]")):
                        st.session_state.pending_docx_replacements.append({
                            "old": str(old_value),
                            "new": str(new_value),
                            "section": section_key,
                            "target": str(target)
                        })
                    
                    # ==================== LOG FEEDBACK TO BACKEND ====================
                    try:
                        serial_no = st.session_state.get("current_proposal_serial", -1)
                        if serial_no > 0:
                            ai_enhanced = actionable_info.get("actionable_instruction", user_feedback)
                            log_feedback_entry(
                                serial_no=serial_no,
                                user_feedback=user_feedback,
                                ai_enhanced_feedback=ai_enhanced,
                                section_affected=section_info['name'],
                                initial_output=original_content,
                                output_post_feedback=new_content
                            )
                    except Exception as log_error:
                        logger.warning(f"Could not log feedback to backend: {log_error}")
                    
                    # Build response - clean and concise, no preview
                    action_desc = actionable_info.get("actionable_instruction", user_feedback)
                    
                    response = f"**{section_info['name']}** has been updated.\n\n"
                    response += f"**Action:** {action_desc}\n\n"
                    if old_value and new_value:
                        response += f"**Change:** {old_value} to {new_value}\n\n"
                    response += "Click **Apply Changes** below to update the document."
                else:
                    response = f"Target section identified: **{section_info['name']}**\n\nThe content for this section is not available. Please generate a new proposal to enable editing."
            
            # Add assistant response to chat
            st.session_state.feedback_chat_messages.append({
                "role": "assistant",
                "content": response
            })
            
            st.rerun()
        
        # Show regenerate button if any sections have been updated
        regenerated_sections = [key for key in st.session_state.keys() if key.startswith("section_regenerated_") and st.session_state[key]]
        
        if regenerated_sections:
            st.markdown("")
            if st.button("Apply Changes", type="primary", use_container_width=True):
                with st.spinner("Updating document..."):
                    try:
                        # IMPORTANT: Get the LATEST DOCX buffer (could be v1, v2, etc.)
                        docx_buffer = st.session_state.get("generated_docx_buffer")
                        if docx_buffer is None:
                            st.error("Document not found. Please generate a new proposal.")
                        else:
                            docx_buffer.seek(0)
                            doc = Document(docx_buffer)
                            
                            # Track changes (using list for mutable counter in nested functions)
                            updated_sections = []
                            changes_counter = [0]  # Using list to allow modification in nested functions
                            
                            # Helper function to replace text in paragraph while PRESERVING formatting
                            def replace_in_paragraph(para, old_text, new_text):
                                """Replace text while preserving run formatting (bold, italic, etc)."""
                                full_text = para.text
                                if old_text not in full_text:
                                    return False
                                
                                # First, try to find and replace within a single run (preserves formatting)
                                for run in para.runs:
                                    if old_text in run.text:
                                        # Replace only within this run, preserving its formatting
                                        run.text = run.text.replace(old_text, new_text)
                                        changes_counter[0] += 1
                                        return True
                                
                                # If old_text spans multiple runs (rare), reconstruct minimally
                                old_pos = full_text.find(old_text)
                                if old_pos == -1:
                                    return False
                                
                                # Build new text, reconstructing via runs
                                new_text_full = full_text[:old_pos] + new_text + full_text[old_pos + len(old_text):]
                                
                                # Last resort: put new text in first run, clear others (fallback)
                                if para.runs:
                                    para.runs[0].text = new_text_full
                                    for run in para.runs[1:]:
                                        run.text = ""
                                
                                changes_counter[0] += 1
                                return True
                            
                            # Helper to apply replacement across all paragraphs and tables
                            # CRITICAL: Context-aware replacement to avoid changing section numbering
                            def apply_replacement_to_doc(doc, old_text, new_text, target_component=""):
                                """
                                Apply replacement with context awareness.
                                NEVER replace text that appears to be section numbering.
                                """
                                import re
                                found = False
                                
                                def is_section_numbering(para_text, old_text):
                                    """Check if old_text appears as section numbering in the paragraph."""
                                    # Pattern for section numbering: "6999.1", "51.", "3.2.1", etc.
                                    # If old_text is a pure number and appears with a dot right after, it's likely section numbering
                                    if not old_text.strip().isdigit():
                                        return False
                                    
                                    # Check various section numbering patterns
                                    patterns = [
                                        rf'\b{re.escape(old_text)}\.(?:\d|[a-zA-Z])',  # "6999.1" or "6999.a"
                                        rf'^\s*{re.escape(old_text)}\.?\s+\w',  # Line starts with "6999. Something"
                                        rf'\b{re.escape(old_text)}\.\s+[A-Z]',  # "6999. Title Case"
                                    ]
                                    
                                    for pattern in patterns:
                                        if re.search(pattern, para_text):
                                            return True
                                    return False
                                
                                def is_safe_replacement_context(para_text, old_text, target):
                                    """Check if this is a safe context for replacement (component count, not heading)."""
                                    text_lower = para_text.lower()
                                    
                                    # If we have a target component, only replace in contexts mentioning that component
                                    if target:
                                        target_words = target.replace("_", " ").lower().split()
                                        if not any(word in text_lower for word in target_words if len(word) > 3):
                                            return False
                                    
                                    # Check for component count context indicators
                                    count_context_words = [
                                        'chute', 'conveyor', 'line', 'station', 'carrier', 'destination',
                                        'feedline', 'induct', 'gravity', 'bagging', 'collection', 'rejection',
                                        'pph', 'throughput', 'capacity', 'total', 'count', 'qty', 'quantity'
                                    ]
                                    
                                    has_component_context = any(word in text_lower for word in count_context_words)
                                    return has_component_context
                                
                                # Check all paragraphs
                                for para in doc.paragraphs:
                                    para_text = para.text
                                    if old_text not in para_text:
                                        continue
                                    
                                    # SAFETY CHECK: Don't replace if it looks like section numbering
                                    if is_section_numbering(para_text, old_text):
                                        continue
                                    
                                    # SAFETY CHECK: Only replace in component-related contexts
                                    if not is_safe_replacement_context(para_text, old_text, target_component):
                                        continue
                                    
                                    if replace_in_paragraph(para, old_text, new_text):
                                        found = True
                                
                                # Check all tables
                                for table in doc.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            for para in cell.paragraphs:
                                                para_text = para.text
                                                if old_text not in para_text:
                                                    continue
                                                
                                                # SAFETY CHECK: Don't replace if it looks like section numbering
                                                if is_section_numbering(para_text, old_text):
                                                    continue
                                                
                                                # Tables are usually data - component counts are OK here
                                                if replace_in_paragraph(para, old_text, new_text):
                                                    found = True
                                return found
                            
                            # STEP 1: Apply direct value replacements (most reliable)
                            # These come from actionable_info with explicit old/new values
                            # CRITICAL: Context-aware replacement - only in component-related paragraphs
                            pending_replacements = st.session_state.get("pending_docx_replacements", [])
                            for replacement in pending_replacements:
                                old_val = replacement.get("old", "")
                                new_val = replacement.get("new", "")
                                target_component = replacement.get("target", "")
                                
                                # Only do replacement if we have actual values (not placeholders)
                                if old_val and new_val and not (old_val.startswith("[") and old_val.endswith("]")):
                                    # Context-aware replacement - pass target component for safety checks
                                    apply_replacement_to_doc(doc, old_val, new_val, target_component)
                            
                            # Clear pending replacements after applying
                            st.session_state.pending_docx_replacements = []
                            
                            # STEP 2: Apply section-based content updates
                            for section_key in ["cover_letter", "executive_summary", "proposed_system", "system_description", "technical_details"]:
                                if st.session_state.get(f"section_regenerated_{section_key}"):
                                    section_info = PROPOSAL_SECTIONS.get(section_key, {})
                                    updated_sections.append(section_info.get("name", section_key))
                                    
                                    # Update the original content to the new content for next iteration
                                    new_content = st.session_state.get(f"section_content_{section_key}", "")
                                    st.session_state[f"section_original_{section_key}"] = new_content
                            
                            # Save updated DOCX to buffer
                            new_docx_buffer = BytesIO()
                            doc.save(new_docx_buffer)
                            new_docx_buffer.seek(0)
                            
                            # Convert to PDF with proper COM handling
                            new_pdf_buffer = BytesIO()
                            tmp_docx_path = None
                            pdf_path = None
                            
                            try:
                                # Initialize COM for this thread
                                pythoncom.CoInitialize()
                                
                                # Create temp file with explicit close before conversion
                                tmp_docx_path = os.path.join(tempfile.gettempdir(), f"proposal_update_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
                                with open(tmp_docx_path, 'wb') as tmp_docx:
                                    tmp_docx.write(new_docx_buffer.getvalue())
                                
                                # Small delay to ensure file is fully written
                                time.sleep(0.5)
                                
                                pdf_path = tmp_docx_path.replace('.docx', '.pdf')
                                
                                # Convert with explicit keep_active=False to properly close Word
                                docx2pdf_convert(tmp_docx_path, pdf_path, keep_active=False)
                                
                                # Read new PDF
                                with open(pdf_path, 'rb') as pdf_file:
                                    new_pdf_buffer.write(pdf_file.read())
                                new_pdf_buffer.seek(0)
                                
                            finally:
                                # Uninitialize COM
                                try:
                                    pythoncom.CoUninitialize()
                                except:
                                    pass
                                
                                # Clean up temp files
                                try:
                                    if tmp_docx_path and os.path.exists(tmp_docx_path):
                                        os.remove(tmp_docx_path)
                                    if pdf_path and os.path.exists(pdf_path):
                                        os.remove(pdf_path)
                                except:
                                    pass
                            
                            # CRITICAL: Update both PDF and DOCX buffers to the NEW version
                            st.session_state.generated_pdf_buffer = new_pdf_buffer
                            st.session_state.generated_docx_buffer = new_docx_buffer
                            
                            # Clear regeneration flags
                            for key in list(st.session_state.keys()):
                                if key.startswith("section_regenerated_"):
                                    st.session_state[key] = False
                            
                            if updated_sections:
                                st.success(f"Document updated: {', '.join(updated_sections)}")
                            else:
                                st.success("Document updated successfully")
                            
                            st.rerun()
                            
                    except Exception as e:
                        st.error(f"Error updating document: {str(e)}")
    
    # ==================== RIGHT COLUMN: PDF PREVIEW ====================
    with pdf_col:

        
        try:
            # Always use the latest PDF from session state for preview
            latest_pdf_preview = st.session_state.generated_pdf_buffer
            latest_pdf_preview.seek(0)
            pdf_base64 = base64.b64encode(latest_pdf_preview.read()).decode('utf-8')
            
            # Create PDF viewer using PDF.js library
            pdf_viewer_html = f'''
            <!DOCTYPE html>
            <html>
            <head>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
                <style>
                    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
                    body {{ background: #525659; font-family: Arial, sans-serif; }}
                    .toolbar {{
                        background: #323639;
                        padding: 6px 12px;
                        display: flex;
                        justify-content: space-between;
                        align-items: center;
                        position: fixed;
                        top: 0;
                        left: 0;
                        right: 0;
                        z-index: 100;
                    }}
                    .toolbar span {{ color: #fff; font-size: 11px; }}
                    .toolbar .filename {{ font-weight: 600; }}
                    .toolbar .nav-controls {{ display: flex; align-items: center; gap: 8px; }}
                    .toolbar button {{
                        background: #4a4d50;
                        border: none;
                        color: #fff;
                        padding: 4px 10px;
                        border-radius: 3px;
                        cursor: pointer;
                        font-size: 11px;
                    }}
                    .toolbar button:hover {{ background: #5a5d60; }}
                    .toolbar button:disabled {{ opacity: 0.5; cursor: not-allowed; }}
                    #pdf-container {{
                        margin-top: 35px;
                        display: flex;
                        flex-direction: column;
                        align-items: center;
                        padding: 8px;
                        gap: 8px;
                    }}
                    .page-canvas {{
                        background: white;
                        box-shadow: 0 2px 8px rgba(0,0,0,0.3);
                    }}
                </style>
            </head>
            <body>
                <div class="toolbar">
                    <span class="filename">{filename}</span>
                    <div class="nav-controls">
                        <button id="prev-btn" onclick="goToPrevPage()">Prev</button>
                        <span id="page-info">1 / 1</span>
                        <button id="next-btn" onclick="goToNextPage()">Next</button>
                    </div>
                </div>
                <div id="pdf-container"></div>
                
                <script>
                    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
                    
                    const pdfData = atob("{pdf_base64}");
                    const pdfArray = new Uint8Array(pdfData.length);
                    for (let i = 0; i < pdfData.length; i++) {{
                        pdfArray[i] = pdfData.charCodeAt(i);
                    }}
                    
                    let pdfDoc = null;
                    let currentPage = 1;
                    let totalPages = 0;
                    const scale = 1.0;
                    
                    async function loadPDF() {{
                        try {{
                            pdfDoc = await pdfjsLib.getDocument({{data: pdfArray}}).promise;
                            totalPages = pdfDoc.numPages;
                            document.getElementById('page-info').textContent = `${{currentPage}} / ${{totalPages}}`;
                            renderAllPages();
                        }} catch(error) {{
                            console.error('Error loading PDF:', error);
                            document.getElementById('pdf-container').innerHTML = '<p style="color: white; padding: 20px;">Error loading PDF</p>';
                        }}
                    }}
                    
                    async function renderAllPages() {{
                        const container = document.getElementById('pdf-container');
                        container.innerHTML = '';
                        
                        for (let pageNum = 1; pageNum <= totalPages; pageNum++) {{
                            const page = await pdfDoc.getPage(pageNum);
                            const viewport = page.getViewport({{scale: scale}});
                            
                            const canvas = document.createElement('canvas');
                            canvas.className = 'page-canvas';
                            canvas.height = viewport.height;
                            canvas.width = viewport.width;
                            container.appendChild(canvas);
                            
                            const context = canvas.getContext('2d');
                            await page.render({{
                                canvasContext: context,
                                viewport: viewport
                            }}).promise;
                        }}
                        
                        updateNavButtons();
                    }}
                    
                    function goToNextPage() {{
                        if (currentPage < totalPages) {{
                            currentPage++;
                            document.getElementById('page-info').textContent = `${{currentPage}} / ${{totalPages}}`;
                            const pages = document.querySelectorAll('.page-canvas');
                            if (pages[currentPage - 1]) {{
                                pages[currentPage - 1].scrollIntoView({{behavior: 'smooth'}});
                            }}
                            updateNavButtons();
                        }}
                    }}
                    
                    function goToPrevPage() {{
                        if (currentPage > 1) {{
                            currentPage--;
                            document.getElementById('page-info').textContent = `${{currentPage}} / ${{totalPages}}`;
                            const pages = document.querySelectorAll('.page-canvas');
                            if (pages[currentPage - 1]) {{
                                pages[currentPage - 1].scrollIntoView({{behavior: 'smooth'}});
                            }}
                            updateNavButtons();
                        }}
                    }}
                    
                    function updateNavButtons() {{
                        document.getElementById('prev-btn').disabled = currentPage <= 1;
                        document.getElementById('next-btn').disabled = currentPage >= totalPages;
                    }}
                    
                    loadPDF();
                </script>
            </body>
            </html>
            '''
            components.html(pdf_viewer_html, height=600, scrolling=True)
            
            # Download button below PDF preview - DOCX format
            st.markdown("")
            latest_docx = st.session_state.get("generated_docx_buffer")
            if latest_docx:
                latest_docx.seek(0)
                st.download_button(
                    label="Download Proposal",
                    data=latest_docx.getvalue(),
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
        except Exception as e:
            st.error(f"Could not display preview: {str(e)}")
    
    # ==================== FOOTER: Back and New Proposal buttons ====================
    st.markdown("<br>", unsafe_allow_html=True)
    footer_left, footer_spacer, footer_right = st.columns([1, 2, 1])
    
    with footer_left:
        if st.button("Back to Form", use_container_width=True):
            st.session_state.page = "input"
            st.rerun()
    
    with footer_right:
        if st.button("New Proposal", use_container_width=True):
            st.session_state.generated_pdf_buffer = None
            st.session_state.generated_docx_buffer = None
            st.session_state.generated_doc_filename = None
            st.session_state.feedback_chat_messages = []
            st.session_state.pending_docx_replacements = []
            # Clear section content
            for key in list(st.session_state.keys()):
                if key.startswith("section_content_") or key.startswith("section_original_") or key.startswith("section_regenerated_"):
                    del st.session_state[key]
            st.session_state.page = "input"
            st.rerun()

# ==================== INPUT PAGE ====================
else:
    st.session_state.page = "input"
    
    # Center-align main tabs, show Coming Soon pills, and tweak header spacing
    st.markdown(
        """
        <style>
        /* Style radio buttons as tabs */
        #main-tabs div[data-testid="stHorizontalBlock"] {
            justify-content: center;
            gap: 8px;
        }
        #main-tabs div[role="radiogroup"] {
            display: flex !important;
            justify-content: center !important;
            gap: 0 !important;
            background: transparent !important;
        }
        #main-tabs div[role="radiogroup"] label {
            background: transparent !important;
            border: none !important;
            border-bottom: 2px solid transparent !important;
            padding: 12px 24px !important;
            cursor: pointer !important;
            font-weight: 500 !important;
            color: #6b7280 !important;
            transition: all 0.2s ease !important;
        }
        #main-tabs div[role="radiogroup"] label:hover {
            color: #1f4e79 !important;
            border-bottom-color: #1f4e79 !important;
        }
        #main-tabs div[role="radiogroup"] label[data-checked="true"] {
            color: #1f4e79 !important;
            border-bottom-color: #1f4e79 !important;
            font-weight: 600 !important;
        }
        #main-tabs div[role="radiogroup"] label div[data-testid="stMarkdownContainer"] {
            display: none !important;
        }
        /* Header/banner spacing and width overrides */
        .main .block-container { padding: 1rem 2rem 2rem 2rem !important; }
        .main-header { width: 100% !important; margin: 0 0 2rem 0 !important; padding: 1.25rem 2rem !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Initialize active tab in session state
    if "active_main_tab" not in st.session_state:
        st.session_state.active_main_tab = "CBS"
    
    st.markdown("<div id='main-tabs'>", unsafe_allow_html=True)
    # Use radio buttons styled as tabs to preserve state across reruns
    active_tab = st.radio(
        "Select System",
        ["CBS", "Neo", "Cubizone", "Robodome 2.0"],
        index=["CBS", "Neo", "Cubizone", "Robodome 2.0"].index(st.session_state.active_main_tab),
        key="main_tab_selector",
        horizontal=True,
        label_visibility="collapsed"
    )
    st.session_state.active_main_tab = active_tab
    st.markdown("</div>", unsafe_allow_html=True)
    
    # ==================== CBS MAIN TAB ====================
    if active_tab == "CBS":
        # Create nested tabs for the 3 sections (no emojis - professional look)
        tab1, tab2, tab3 = st.tabs(["Section 1: Project & Client Info", "Section 2: Upload Files", "Section 3: Settings & Generate"])
        
        # ==================== TAB 1: Project & Client Information ====================
        with tab1:
            st.markdown("<br>", unsafe_allow_html=True)
            
            # ---- Section: Project Details ----
            st.markdown("##### Project Details")
            proj_col1, proj_col2 = st.columns(2)
            with proj_col1:
                project_name = st.text_input("Project Name *", value="Loop Cross Belt Sorter", placeholder="Enter project name")
            with proj_col2:
                offer_ref = st.text_input("Offer Reference No *", value="F24-00524", placeholder="e.g., F24-00524")
            
            st.markdown("<hr style='margin: 1rem 0; border: none; border-top: 1px solid #e0e0e0;'>", unsafe_allow_html=True)
            
            # ---- Section: Client Information ----
            st.markdown("##### Client Information")
            client_col1, client_col2 = st.columns([2, 1])
            
            with client_col1:
                # Client dropdown with add new option
                client_options = list(CLIENT_LOGOS.keys()) + ["+ Add New Client"]
                selected_client = st.selectbox("Client Name *", client_options, index=0)
                
                # Handle new client addition
                if selected_client == "+ Add New Client":
                    client_name = st.text_input("Enter New Client Name *", placeholder="Enter client name")
                    client_logo = st.file_uploader("Upload Client Logo *", type=["png", "jpg", "jpeg"], key="new_client_logo")
                    client_logo_path_display = None
                else:
                    client_name = selected_client
                    client_logo = None
                    client_logo_path_display = CLIENT_LOGOS.get(selected_client)
                
                executives_text = st.text_area(
                    "Client Executives (one per line, include Mr./Ms.) *",
                    value="Mr. Rahul Didwani\nMr. Vinayak Garg",
                    height=80,
                    placeholder="Mr. John Doe\nMs. Jane Smith"
                )

            with client_col2:
                st.markdown("**Client Logo Preview**")
                if client_logo_path_display and os.path.exists(client_logo_path_display):
                    st.image(client_logo_path_display, use_container_width=True)
                elif client_logo:
                    st.image(client_logo, use_container_width=True)
                else:
                    st.info("Logo will appear here")
            
            st.markdown("<hr style='margin: 1rem 0; border: none; border-top: 1px solid #e0e0e0;'>", unsafe_allow_html=True)
            
            # ---- Section: Operational Parameters ----
            st.markdown("##### Operational Parameters")
            # PPH Slider - ranges from 500 to 80,000, default 10,000
            pph_count = st.slider(
                "Throughput (PPH - Parcels Per Hour) *",
                min_value=500,
                max_value=80000,
                value=10000,
                step=500,
                format="%d PPH",
                help="Select the required throughput capacity in parcels per hour"
            )
            pph_count = str(pph_count)  # Convert to string for compatibility
            ipp_rate = ""  # Deprecated - kept for backward compatibility
            
            st.markdown("<hr style='margin: 1rem 0; border: none; border-top: 1px solid #e0e0e0;'>", unsafe_allow_html=True)
            
            # ---- Section: Important Dates ----
            st.markdown("##### Important Dates")
            date_col1, date_col2 = st.columns(2)
            with date_col1:
                invitation_date = st.date_input("Invitation Date (optional)", value=None)
            with date_col2:
                meeting_date = st.date_input("Meeting/Workshop Date (optional)", value=None)
            
            # Auto-populate contact details from logged-in user
            current_user = st.session_state.get('current_user', {})
            contact_name = current_user.get('name', '')
            contact_email = current_user.get('email', '')
            contact_phone = "+91 8750052591"  # Fixed contact number

            # Fixed values (not shown to user)
            letter_date = date.today()
            sender_name = "Sandeep Bansal"
            sender_title = "Chief Business Officer"
            invitation_date_str = invitation_date.strftime("%B %d, %Y") if invitation_date else ""
            meeting_date_str = meeting_date.strftime("%B %d, %Y") if meeting_date else ""

        # ==================== TAB 2: Upload Files ====================
        with tab2:
            st.markdown("<br>", unsafe_allow_html=True)
            
            # ---- Section: Required Files ----
            st.markdown("##### Required Documents")
            upload_col1, upload_col2 = st.columns(2)

            with upload_col1:

                dxf_layout_file = st.file_uploader("DXF Layout File *", type=["dxf"], key="dxf_upload", help="System layout in DXF format")
                costing_file = st.file_uploader("Costing Sheet *", type=["xlsx", "xls"], key="costing_upload", help="Component pricing spreadsheet")
                
                # When DXF is uploaded, automatically extract components and show dialog
                if dxf_layout_file:
                    # Extract components from DXF and show dialog
                    if "dxf_components_extracted" not in st.session_state or st.session_state.get("last_dxf_name") != dxf_layout_file.name:
                        # Extract DXF summary for embedding/analysis
                        try:
                            import tempfile
                            from pathlib import Path
                            from dxf_extractor import extract_dxf_components, create_dxf_summary_for_embedding
                            
                            # Save uploaded file to temporary path
                            with tempfile.NamedTemporaryFile(suffix=".dxf", delete=False) as tmp:
                                tmp.write(dxf_layout_file.getbuffer())
                                tmp_path = tmp.name
                            
                            # Extract components dict from DXF file
                            dxf_json = extract_dxf_components(Path(tmp_path), project_name="")
                            
                            # Create summary for embedding
                            dxf_summary = create_dxf_summary_for_embedding(dxf_json)
                            
                            # Clean up temp file
                            os.remove(tmp_path)
                            
                            st.session_state.dxf_components_extracted = dxf_summary
                            st.session_state.dxf_json_extracted = dxf_json
                            st.session_state.last_dxf_name = dxf_layout_file.name
                            st.session_state.show_component_dialog = True
                            st.session_state.force_refresh_components = True
                            # Clear old dialog components to force fresh extraction
                            if "dialog_components" in st.session_state:
                                del st.session_state.dialog_components
                        except Exception as e:
                            st.error(f"Error extracting DXF: {str(e)}")
                    
                    # Show dialog for component editing
                    if st.session_state.get("show_component_dialog"):
                        show_component_editing_dialog(dxf_layout_file)

            with upload_col2:
                capacity_excel = st.file_uploader("Throughput Calculation Sheet *", type=["xlsx", "xls"], key="capacity_upload", help="Capacity calculation workbook")
                prog_gantt = st.file_uploader("Project Timeline Chart (optional)", type=["png", "jpg", "jpeg"], key="gantt_upload", help="Gantt chart or timeline image")

            st.markdown("<hr style='margin: 1rem 0; border: none; border-top: 1px solid #e0e0e0;'>", unsafe_allow_html=True)
            
            # ---- Section: Optional Layout Image ----
            st.markdown("##### Solution Layout (Optional)")
            have_solution_png = st.checkbox("I already have a PNG of the solution layout", value=False)

            if have_solution_png:
                layout_full_png = st.file_uploader("Upload your solution PNG here", type=["png", "jpg", "jpeg"], key="solution_png_upload")
            else:
                layout_full_png = None

        # Default section includes (hidden from user)
        key_include = True
        safety_include = True
        infra_include = True
        prog_include = True
        client_resp_include = True
        handover_include = True
        commercial_include = True
        warranty_include = True
        exclusion_include = True
        include_proposed_system = True
        include_concept_desc = True
        include_exec_summary = True
        include_company_profile = True
        include_reference_projects = True
        include_handled_spectrum = True
        include_capacity_section = True
        elec_include = True
        wcs_include= True
        scada_include= True

        # ==================== SECTION 3: SETTINGS & GENERATE ====================
        with tab3:
            st.markdown("<br>", unsafe_allow_html=True)
            
            # ---- Section: Handled Parcel Spectrum ----
            with st.expander("Handled Parcel Spectrum", expanded=False):
                st.markdown("Configure the parcel specifications that will appear in the 'Handled Shipment Spectrum' section of the proposal.")
                
                # Initialize default values if not in session state
                if "parcel_spectrum" not in st.session_state:
                    st.session_state["parcel_spectrum"] = [
                        {"Specification": "Max Length", "Unit": "mm", "Value": "800"},
                        {"Specification": "Max Width", "Unit": "mm", "Value": "800"},
                        {"Specification": "Max Height", "Unit": "mm", "Value": "700"},
                        {"Specification": "Max Weight", "Unit": "Kg", "Value": "30"},
                        {"Specification": "Min length", "Unit": "mm", "Value": "100"},
                        {"Specification": "Min Width", "Unit": "mm", "Value": "100"},
                        {"Specification": "Min Height", "Unit": "mm", "Value": "50"},
                        {"Specification": "Min Weight", "Unit": "gm", "Value": "200"},
                    ]
                
                # Create dataframe for editing
                parcel_df = pd.DataFrame(st.session_state["parcel_spectrum"])
                
                
                # Configure column settings
                column_config = {
                    "Specification": st.column_config.TextColumn(
                        "Specification",
                        disabled=True,  # Fixed column
                        help="Specification name (fixed)"
                    ),
                    "Unit": st.column_config.SelectboxColumn(
                        "Unit",
                        options=["mm", "cm", "m", "Kg", "gm", "lbs"],
                        help="Select unit of measurement"
                    ),
                    "Value": st.column_config.TextColumn(
                        "Value",
                        help="Enter the value"
                    )
                }
                
                # Display editable data table
                edited_parcel_df = st.data_editor(
                    parcel_df,
                    column_config=column_config,
                    hide_index=True,
                    use_container_width=True
                )
                
                
                # Update session state with edited data
                st.session_state["parcel_spectrum"] = edited_parcel_df.to_dict(orient="records")
            
            # ---- Section: Commercial Settings ----
            with st.expander("Commercial Settings", expanded=False):
                apply_bca = st.checkbox("Apply Business Cooperation Agreement Discount (4.5%)", value=False, key="apply_bca_discount")
                
                st.markdown("<hr style='margin: 0.5rem 0; border: none; border-top: 1px solid #e0e0e0;'>", unsafe_allow_html=True)
                st.markdown("**Payment Terms**")
                default_payment_terms = [
                    {"Payment Percentage": "20%", "Stage": "Advance along with LOI/ PO"},
                    {"Payment Percentage": "20%", "Stage": "After DAP Completion"},
                    {"Payment Percentage": "40%", "Stage": "Before Dispatch"},
                    {"Payment Percentage": "10%", "Stage": "Against Installation"},
                    {"Payment Percentage": "10%", "Stage": "Against Handover"},
                ]
                
                if "payment_terms" not in st.session_state:
                    st.session_state["payment_terms"] = default_payment_terms
                
                pt_df = pd.DataFrame(st.session_state["payment_terms"])
                edited_pt_df = st.data_editor(pt_df, num_rows="dynamic", use_container_width=True, key="payment_terms_editor")
                st.session_state["payment_terms"] = edited_pt_df.to_dict(orient="records")

            # ---- Section: Warranty Configuration ----
            with st.expander("Warranty Configuration", expanded=False):
                warranty_type = st.selectbox("Warranty Type", ["Standard warranty", "Comprehensive warranty"], key="warranty_type")
                
                warranty_col1, warranty_col2 = st.columns(2)
                with warranty_col1:
                    warranty_duration = st.text_input("Warranty Duration", value="1 year", key="warranty_duration")
                with warranty_col2:
                    warranty_start = st.selectbox(
                        "Warranty Start Condition",
                        [
                            "from the date of beneficiary use.",
                            "from the date of commissioning of the system.",
                            "from the date of completion of dispatch of the materials, whichever is earlier.",
                            "from the date of beneficiary use, max 30 days after readiness of commissioning.",
                        "from the date of official communication of material readiness at Falcon end.",
                    ],
                    key="warranty_start"
                )
                
                st.markdown("<hr style='margin: 0.5rem 0; border: none; border-top: 1px solid #e0e0e0;'>", unsafe_allow_html=True)
                warranty_extended = st.checkbox("Include Extended Warranty Option", value=True, key="warranty_extended")
                if warranty_extended:
                    warranty_extended_text = st.text_input(
                        "Extended Warranty Text",
                        value="Extended warranty of 2 years available on request @ 5% of the order value.",
                        key="warranty_extended_text"
                    )
                else:
                    warranty_extended_text = None
                
                warranty_amc = st.checkbox("Include AMC / Hotline Clause", value=False, key="warranty_amc")
                if warranty_amc:
                    warranty_amc_text = st.text_input(
                        "AMC / Hotline Text",
                        value="AMC / Hotline services available post warranty on demand.",
                        key="warranty_amc_text"
                    )
                else:
                    warranty_amc_text = None
                
                warranty_transport = st.checkbox("Include Transportation Note", value=False, key="warranty_transport")
                if warranty_transport:
                    warranty_transport_text = st.text_input(
                        "Transportation Note",
                        value="Transportation of defective parts to Falcon premises will be at client's cost.",
                        key="warranty_transport_text"
                    )
                else:
                    warranty_transport_text = None

            # ---- Section: Exclusions ----
            with st.expander("Exclusions Configuration", expanded=False):
                st.markdown("**Select items to exclude from the proposal:**")
                
                variable_exclusions = [
                    "Server PC / server system.",
                    "SCADA / PC for SCADA.",
                    "Workstations.",
                    "Cabling from server room to Falcon control panel.",
                    "Mobile carts.",
                    "Collection trolleys / collection trolleys below chutes.",
                    "Collection bins.",
                    "Pallets at chutes.",
                    "Pallets / hand-held terminals for secondary sorting.",
                    "Steel works.",
                    "Steel works – if not specified.",
                    "Mezzanine & staircase.",
                    "Mezzanine & staircase not mentioned in BOM.",
                    "Maintenance platform / lift required for maintenance activity.",
                    "Safety fencing / safety fencing not shown in layout.",
                    "HPT/BOPT/Forklift/Hydra/Scaffoldings required for installation.",
                    "Stress free mats.",
                    "Insulation mats.",
                    "Fans at chutes & inducts.",
                    "Lighting around chutes / inducts.",
                    "Irregular's provision.",
                    "UPS power (separate UPS supply).",
                    "CE declaration of conformity.",
                ]
                
                selected_exclusions = []
                cols = st.columns(2)
                for idx, item in enumerate(variable_exclusions):
                    col = cols[idx % 2]
                    if col.checkbox(item, value=False, key=f"exclusion_{idx}"):
                        selected_exclusions.append(item)

            # ---- Section: Key Components ----
            with st.expander("Key Components & Manufacturers", expanded=False):
                default_components = [
                    {"Items": "Belts", "Make": "Forbo / Derco / Habasit"},
                    {"Items": "Rollers", "Make": "Falcon"},
                    {"Items": "Cross Belt Carriers", "Make": "Falcon"},
                    {"Items": "Linear Motors (LIM / LSM / Linear Induction)", "Make": "Falcon / SEW / FWD (as applicable)"},
                    {"Items": "Feed Line Motors", "Make": "Falcon"},
                    {"Items": "Volume / Barcode Scanners", "Make": "SICK / Cognex / Similar"},
                    {"Items": "Weighing Scales", "Make": "Bizerba / Mettler Toledo / Equivalent"},
                    {"Items": "Encoders", "Make": "SICK / Falcon"},
                    {"Items": "Sensors", "Make": "SICK / Leuze / P&F"},
                    {"Items": "PLC", "Make": "Siemens / Omron"},
                    {"Items": "Control Panels", "Make": "Rittal / BCH"},
                    {"Items": "VFDs", "Make": "Siemens / Lenze / AB / Omron"},
                    {"Items": "Cables", "Make": "LAPP / Equivalent"},
                    {"Items": "Switch Gear", "Make": "Schneider / Equivalent"},
                    {"Items": "Bearings", "Make": "NTN / SKF / Equivalent"},
                    {"Items": "Power Transmission Systems", "Make": "Vahle"},
                    {"Items": "HMIs", "Make": "Siemens / Omron"},
                    {"Items": "MDR", "Make": "Pulse / Itoh Denki"},
                    {"Items": "Data Transmission System", "Make": "Siemens"},
                ]
                
                if "key_components_df" not in st.session_state:
                    st.session_state["key_components_df"] = pd.DataFrame(default_components)
                
                key_components_edited = st.data_editor(
                    st.session_state["key_components_df"],
                    num_rows="dynamic",
                    use_container_width=True,
                    key="key_editor"
                )

            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("<hr style='margin: 1rem 0; border: none; border-top: 2px solid #1976D2;'>", unsafe_allow_html=True)

            # Generation button with prominent styling
            st.markdown("##### Ready to Generate")
            generate_clicked = st.button("Generate Proposal", type="primary", use_container_width=True, help="Click to generate your complete proposal document")
    elif active_tab == "Neo":
        st.markdown("<div style='text-align: center; padding: 60px 20px;'>", unsafe_allow_html=True)
        st.markdown("### Coming Soon")
        st.markdown("<p style='font-size: 16px; color: #666;'>Neo configuration will be available soon.</p>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ==================== CUBIZONE MAIN TAB ====================
    elif active_tab == "Cubizone":
        st.markdown("<br>", unsafe_allow_html=True)
        try:
            cubizone_module = load_cubizone_app()
            cubizone_module.render_cubizone_builder(embed_mode=True)
        except Exception as exc:
            st.error(f"Unable to load Cubizone builder: {exc}")

    # ==================== ROBODOME 2.0 MAIN TAB ====================
    elif active_tab == "Robodome 2.0":
        st.markdown("<div style='text-align: center; padding: 60px 20px;'>", unsafe_allow_html=True)
        st.markdown("### Coming Soon")
        st.markdown("<p style='font-size: 16px; color: #666;'>Robodome 2.0 configuration will be available soon.</p>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

# ==================== DOCUMENT GENERATION FUNCTIONS ====================

def build_cover_letter_section(doc, letter_text):
    """Build cover letter (page 1) - NO HEADER for this section"""
    HEADER_PREFIXES = (
        "Kind Attention",
        "Mr.",
        "Ms.",
        "M/s",
        "Offer Ref:",
        "Subject –",
        "Subject -",
        "Date:",
        "Location –",
    )
    
    lines = [l.rstrip() for l in letter_text.splitlines() if l.strip() != ""]
    for idx, line in enumerate(lines):
        # Check if line starts with header prefixes or contains "Dear" or ends with signature (Best Regards)
        is_header = any(line.startswith(pfx) for pfx in HEADER_PREFIXES) or "Dear " in line
        is_signature = "Best Regards" in line or idx >= len(lines) - 2

        if is_header:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.bold = True
        elif is_signature:
            # Signature and name/title at end should be bold
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.bold = True
        else:
            if "**" in line:
                add_markdown_line(doc, line)
            else:
                p = doc.add_paragraph(line)
                apply_normal_style(p)


def build_front_page_section(doc, project_title, offer_ref, contact_name, contact_email, contact_phone, layout_png_path):
    """Build front page (page 2) - NO HEADER for this section"""
    doc.add_page_break()

    # Top box: "Response to RFP for"
    add_boxed_text(doc, "Response to RFP for", font_size=18, bold=True)

    # Project name box
    if project_title:
        add_boxed_text(doc, project_title, font_size=16, bold=True)

    # Proposal reference box
    if offer_ref:
        add_boxed_text(doc, f"Proposal Reference: {offer_ref}", font_size=14, bold=True)

    # Some vertical spacing
    doc.add_paragraph("")

    # Layout image - use the same image as in Proposed System Description
    # Cropped to 5.0 inches to fit everything on single page
    if layout_png_path and os.path.exists(layout_png_path):
        # Crop image before inserting
        try:
            from PIL import Image
            img = Image.open(layout_png_path)
            
            # Crop 10% from each side to remove whitespace
            width, height = img.size
            left = width * 0.1
            top = height * 0.1
            right = width * 0.9
            bottom = height * 0.9
            
            img_cropped = img.crop((left, top, right, bottom))
            
            # Save to temporary buffer
            img_buffer = io.BytesIO()
            img_cropped.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(img_buffer, width=Inches(5.0))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            # Fallback to original image if cropping fails
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(layout_png_path, width=Inches(5.0))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run("Layout image will be provided.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact box at bottom
    contact_lines = [
        "Falcon Autotech Private Limited",
        "Plot No. 87, Sector Ecotech-1, Extention-1, Greater Noida, Uttar Pradesh 201308.",
        "",
        f"Contact – {contact_name}",
        "Assistant Manager",
        f"Mob - {contact_phone}",
        contact_email,
    ]
    contact_text = "\n".join(contact_lines)
    table = add_boxed_text(doc, contact_text, font_size=11, bold=False)
    # make contact text paragraphs centered
    cell = table.rows[0].cells[0]
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_glossary_section(doc):
    """Build glossary/table of contents section with automatic TOC"""
    doc.add_page_break()
    
    p = doc.add_heading("Table of Contents", level=1)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = HEADING_COLOR
    
    # Add automatic TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    
    # Add instruction text for users
    doc.add_paragraph("")
    p = doc.add_paragraph("Note: Right-click on the table of contents and select 'Update Field' to refresh page numbers.")
    apply_normal_style(p)
    p.runs[0].italic = True

def build_glossary(doc, detected_terms):
    """Build Glossary section with table of detected terms"""
    
    # Add heading using standard formatting
    add_numbered_heading(doc, "Glossary", level=1, counter=1)
    
    if not detected_terms:
        p = doc.add_paragraph("No glossary terms detected in this document.")
        apply_normal_style(p)
        return
    
    # Create table with 3 columns: S. No., Term, Description
    table = doc.add_table(rows=1, cols=3)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["S. No.", "Term", "Description"]
    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.text = ""
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Data rows
    for idx, (term, desc) in enumerate(detected_terms, start=1):
        row_cells = table.add_row().cells
        
        # S. No.
        p0 = row_cells[0].paragraphs[0]
        p0.text = ""
        r0 = p0.add_run(str(idx))
        r0.font.name = 'Calibri'
        r0.font.size = Pt(11)
        r0.font.bold = True
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Term
        p1 = row_cells[1].paragraphs[0]
        p1.text = ""
        r1 = p1.add_run(term)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.bold = True
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Description
        p2 = row_cells[2].paragraphs[0]
        p2.text = ""
        r2 = p2.add_run(desc)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Set column widths
    table.columns[0].width = Inches(0.8)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(4.2)

    # Ensure next section starts on a fresh page so Glossary remains isolated
    doc.add_page_break()
def build_executive_summary_section(doc, exec_summary_text, counter):
    """Build Executive Summary section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Executive Summary", counter=counter)
    
    lines = exec_summary_text.strip().splitlines()
    first_bullet_encountered = False  # Track if we've seen the first bullet point
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        
        # Check if it's a bullet point line
        if stripped.startswith("•") or stripped.startswith("-"):
            # Add the key characteristics line before the FIRST bullet point
            
            
            bullet_text = stripped.lstrip("•- ").strip()
            p = doc.add_paragraph(style='List Bullet')
            if "**" in bullet_text:
                parts = bullet_text.split("**")
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    run.italic = True
            else:
                run = p.add_run(bullet_text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.italic = True
        else:
            # Regular paragraph with possible **bold**
            if "**" in stripped:
                add_markdown_line(doc, stripped)
            else:
                p = doc.add_paragraph(stripped)
                apply_normal_style(p)


def build_company_profile_section(doc, counter):
    """Build Company Profile section with static images"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Company Profile", counter=counter)

    top_text = (
        "Falcon Autotech (Falcon) is a global intralogistics automation solutions company. "
        "With over 10 years of experience, Falcon has worked with some of the most innovative "
        "brands in E-Commerce, CEP, Fashion, Food/FMCG, Auto and Pharmaceutical Industries. "
        "With our proprietary software and robust hardware integration capabilities, Falcon designs, "
        "manufactures, supplies, implements, and maintains world-class warehouse automation systems globally. "
        "Falcon's strong research and development team and the continuous focus on innovation reflect our strong "
        "solution line around Sortation, Robotics, Conveying, Vision Systems and IOT. "
        "Falcon has done over 1,800 installations across 15 countries on four continents."
    )
    p = doc.add_paragraph(top_text)
    apply_normal_style(p)

    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "1.png"), width_in=4)

    bottom_text = (
        "Falcon Autotech is currently among the top 15 intralogistics automation companies; "
        "our vision is to become a top 10 intralogistics automation company in our focused product lines."
    )
    p = doc.add_paragraph(bottom_text)
    apply_normal_style(p)

    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "2.png"), width_in=4)

    doc.add_page_break()

    # Page 2
    top_text2 = (
        "The team started out in 2004 solving special purpose automation problems for clients and later "
        "established Falcon Autotech in 2012 with a strong focus on building a standard technology stack spanning "
        "across hardware, firmware, and software to tackle larger supply chain problems around warehouse "
        "automation and material handling. "
        "Over the decade, Falcon has made rapid strides and has carved out a niche in some of the world's most "
        "cutting-edge technologies: Sortation, Robotics, Conveying, Vision Systems and IOT."
    )
    p = doc.add_paragraph(top_text2)
    apply_normal_style(p)

    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "3.png"), width_in=6)

    bottom_text2 = (
        "As a leading player in the intralogistics automation space, Falcon continuously strives to improve the "
        "operational efficiencies and accuracies for its clients through its domain knowledge and experience, in "
        "addition to its wide range of products and solutions. In order to live up to the high expectations set "
        "forth by our clients, the team at Falcon realizes the importance of taking up selective applications in "
        "focused industries and delivering world-class projects in return."
    )
    p = doc.add_paragraph(bottom_text2)
    apply_normal_style(p)

    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "4.png"), width_in=6)

    # Page 3
    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "5.png"), width_in=6)

    bottom_text3 = (
        "Falcon Autotech has successfully delivered warehouse automation solutions based on smart and innovative "
        "combinations of the above product lines for effective materials handling, sortation and movement. "
        "The process is controlled in real-time by our in-house WCS applications. These solutions considerably "
        "reduce the need for manual operations, improve working conditions and ensure the highest accuracy of the "
        "entire process up to final delivery to the recipient.\n\n"
        "Over the last 10 years, Falcon has worked with some of the most innovative brands worldwide and has "
        "established long-standing partnerships. These brands are testimony to our strong focus on delivering "
        "superior customer satisfaction and offering end-to-end intralogistics solutions."
    )
    p = doc.add_paragraph(bottom_text3)
    apply_normal_style(p)

    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "6.png"), width_in=6)

    doc.add_page_break()

    # Page 4
    bottom_text4 = (
        "With over 1,800 installations, Falcon's systems are used all over the globe. Falcon has a highly "
        "motivated team of 600+ employees supported by over 15 global partners who help us design, manufacture, "
        "deliver and maintain automation solutions worldwide."
    )
    p = doc.add_paragraph(bottom_text4)
    apply_normal_style(p)

    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "7.png"), width_in=6)

    add_numbered_subheading(doc, "Customer Engagement Model", f"{counter}.1")
    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "8.png"), width_in=7)

    doc.add_page_break()

    # Page 5
    add_numbered_subheading(doc, "Falcon's Experience and Achievements in Sortation Space Globally", f"{counter}.2")

    bullet_points = [
        "Ranked among Top 10 Sortation System Suppliers globally.",
        "Currently possess one of the world's largest portfolios in sortation technologies (7 in-house technologies).",
        "Total installed capacity of 10 million shipments per day worldwide.",
        "Only company to be able to offer a fully integrated AMS.",
    ]

    for point in bullet_points:
        p = doc.add_paragraph(point, style='List Bullet')
        apply_normal_style(p)

    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "9.png"), width_in=5)
    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "10.png"), width_in=5)
    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "11.png"), width_in=5)
    add_centered_image(doc, os.path.join(STATIC_ABOUT_DIR, "12.png"), width_in=5)


def build_reference_projects_section(doc, counter):
    """Build Reference Projects section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Reference Projects", counter=counter)

    intro = (
        "Falcon has a strong legacy in Warehousing Automation solutions and references-"
    )
    p = doc.add_paragraph(intro)
    apply_normal_style(p)

    bullets_intro = [
        "Expertise in Shipment Sortation, Piece Picking and Handling, Case Picking and Handling.",
        "Lifecycle services (maintenance, spares supply chain, support).",
        "Full in-house expertise (Hardware/Software).",
        "Turn-key tailored solutions.",
        "The references list presented below focuses on Sortation Solution –",
    ]
    for text in bullets_intro:
        p = doc.add_paragraph(text, style='List Bullet')
        apply_normal_style(p)

    # Project 1
    add_numbered_subheading(doc, "Project 1- (CEP Client, India)", f"{counter}.1")

    p = doc.add_paragraph(
        "The system is equipped with two fully automated and interconnected sub-systems. "
        "Sub-System 1 is designed for handling large B2B boxes and E-commerce shipment bags while "
        "Sub-System 2 is designed to handle small E-commerce packages."
    )
    apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Solution Specifications –")
    run.bold = True
    apply_normal_style(p)
    
    spec1 = [
        "48,000 PPH (Double Deck CBS – Shipment Sorter).",
        "17,000 PPH (Double Deck CBS – Bag Sorter).",
        "Building Size: 700,000 Sq. Ft.",
    ]
    for t in spec1:
        p = doc.add_paragraph(t, style='List Bullet')
        apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Key Technology Modules –")
    run.bold = True
    apply_normal_style(p)
    
    ktm1 = [
        "2 Sets of Double Decker CBS Sorters.",
        "Mezzanine Structures.",
        "Automated Singulators.",
        "Fully Automatic Inductions.",
        "Semi-Automatic Inductions.",
        "Telescopic Belt Conveyors.",
        "PVC Belt Conveyors.",
        "Modular Belt Conveyors.",
        "Spiral Chutes with Braking Rollers.",
        "5-Sided Scanning Tunnels.",
        "High Speed Weighing Conveyors.",
        "Direct Bagging Chutes.",
        "Put to Light Chutes.",
        "Volume Distribution Systems.",
        "High Availability Server Systems.",
        "WCS.",
    ]
    add_bullets_in_two_columns(doc, ktm1)

    p = doc.add_paragraph()
    run = p.add_run("Site Pictures –")
    run.bold = True
    apply_normal_style(p)
    p.paragraph_format.keep_with_next = True  # Keep with image
    
    add_centered_image(doc, "assests\Images\\proj1.PNG", width_in=3.0)
    doc.add_page_break()

    # Project 2
    add_numbered_subheading(doc, "Project 2- (Client – E-Commerce, India)", f"{counter}.2")

    p = doc.add_paragraph()
    run = p.add_run("Use Case – Destination sorting of packed shipments.")
    run.bold = True
    apply_normal_style(p)

    p = doc.add_paragraph(
        "In 2019, the client was looking for a potential automation partner for design and development of a "
        "new automated sortation system for B2C shipments. The system needed to provide maximum uptime with "
        "reduced dependency on skilled manpower and better space optimization. "
        "\nThe customer chose Falcon Autotech based on its unique design that addressed these pain points, "
        "its capability for seamless WMS integration, and its life cycle support services."
    )
    apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Solution Specifications –")
    run.bold = True
    apply_normal_style(p)
    
    spec2 = [
        "Throughput: 27,600 PPH.",
        "End Destinations: 410 Direct Outputs.",
        "Building Size: 200,000 Sq. Ft.",
    ]
    for t in spec2:
        p = doc.add_paragraph(t, style='List Bullet')
        apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Key Technology Modules –")
    run.bold = True
    apply_normal_style(p)
    
    ktm2 = [
        "Bulk Infeed Conveyors.",
        "ARB based Volume Distribution System.",
        "Integrated Presort System.",
        "Irregular Ejection System.",
        "Automatic Induct Lines.",
        "Automatic Barcode Scanner with Image Capture.",
        "Automatic Weight & Volume Measurement System.",
        "Linear Cross Belt Sorter.",
        "Smart Sliding Chutes for Direct Bagging and Cage Sorting.",
        "Bag Take-out System.",
        "WCS Software System.",
    ]
    add_bullets_in_two_columns(doc, ktm2)

    p = doc.add_paragraph()
    run = p.add_run("Site Pictures –")
    run.bold = True
    apply_normal_style(p)
    p.paragraph_format.keep_with_next = True  # Keep with image
    
    add_centered_image(doc, "assests\Images\\proj2.PNG", width_in=4.0)
    doc.add_page_break()

    # Project 3
    add_numbered_subheading(doc, "Project 3- (Client – E-Commerce, India)", f"{counter}.3")

    p = doc.add_paragraph()
    run = p.add_run("Use Case – Destination sorting of packed shipments.")
    run.bold = True
    apply_normal_style(p)

    p = doc.add_paragraph(
        "The customer chose Falcon Autotech based on its unique design, its ability to integrate seamlessly "
        "with the WMS, and its strong life cycle support services."
    )
    apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Solution Specifications –")
    run.bold = True
    apply_normal_style(p)
    
    spec3 = [
        "Throughput: 24,000 PPH.",
        "End Destinations: 40 Collection Type Chutes.",
    ]
    for t in spec3:
        p = doc.add_paragraph(t, style='List Bullet')
        apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Key Technology Modules –")
    run.bold = True
    apply_normal_style(p)
    
    ktm3 = [
        "Bulk Infeed Conveyors.",
        "ARB based Volume Distribution System.",
        "Irregular Ejection System.",
        "Automatic Induct Lines.",
        "Automatic Barcode Scanner with Image Capture.",
        "Automatic Weight & Volume Measurement System.",
        "Linear Cross Belt Sorter.",
        "Smart Collection Type Chutes.",
        "Bag Take-out System.",
        "WCS Software System.",
    ]
    add_bullets_in_two_columns(doc, ktm3)

    p = doc.add_paragraph()
    run = p.add_run("Site Pictures –")
    run.bold = True
    apply_normal_style(p)
    p.paragraph_format.keep_with_next = True  # Keep with image
    
    add_centered_image(doc, "assests\Images\\proj3.PNG", width_in=4.0)
    doc.add_page_break()

    # Project 4
    add_numbered_subheading(doc, "Project 4- (CEP Client, UK)", f"{counter}.4")

    p = doc.add_paragraph(
        "This solution is designed to handle a volume of 7,200 shipments per hour. "
        "The system is equipped with three infeed conveyors integrated with an automatic label applicator "
        "before shipments enter the sortation system. Shipments are sorted using Falcon's Loop Cross Belt Sorter "
        "equipped with automatic barcode scanning, dimensioning, weighing, and image capture capabilities. "
        "The sorter is installed on the mezzanine floor and sorts directly to 58 end destinations."
    )
    apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Solution Specifications –")
    run.bold = True
    apply_normal_style(p)
    
    spec4 = [
        "Throughput: 7,200 PPH.",
        "End Destinations: 58 Nos.",
    ]
    for t in spec4:
        p = doc.add_paragraph(t, style='List Bullet')
        apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Key Technology Modules –")
    run.bold = True
    apply_normal_style(p)
    
    ktm4 = [
        "Powered Belt Conveyors.",
        "Automatic Induct Lines.",
        "Automatic Barcode Scanner with Image Capture.",
        "Automatic Weight & Volume Measurement System.",
        "Loop Cross Belt Sorter.",
        "WCS Software System.",
    ]
    add_bullets_in_two_columns(doc, ktm4)

    p = doc.add_paragraph()
    run = p.add_run("Site Picture –")
    run.bold = True
    apply_normal_style(p)
    p.paragraph_format.keep_with_next = True  # Keep with image
    
    add_centered_image(doc, "assests\Images\\proj4.PNG", width_in=4.0)
    doc.add_page_break()

    # Project 5
    add_numbered_subheading(doc, "Project 5- (CEP Client, Sydney)", f"{counter}.5")

    p = doc.add_paragraph(
        "This solution is designed for handling a throughput of 16,000 shipments per hour with the help of "
        "Falcon's Loop Cross Belt Sorter. The system consists of two feeding zones with a total of ten feedlines. "
        "Sorter design enables van drivers to directly drop shipments at the dock doors. It has a total of 369 end "
        "destinations achieved through a combination of direct drops and PTLs. The system is integrated with "
        "five-side automatic barcode scanning, weight and volume measurement, and automatic detection of "
        "oversize and overweight shipments."
    )
    apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Solution Specifications –")
    run.bold = True
    apply_normal_style(p)
    
    spec5 = [
        "Throughput: 16,000 PPH.",
        "End Destinations: 369 Nos.",
    ]
    for t in spec5:
        p = doc.add_paragraph(t, style='List Bullet')
        apply_normal_style(p)

    p = doc.add_paragraph()
    run = p.add_run("Key Technology Modules –")
    run.bold = True
    apply_normal_style(p)
    
    ktm5 = [
        "Powered Belt Conveyors.",
        "2 Induct Zones.",
        "5-side Automatic Barcode Scanner.",
        "Automatic Weight & Volume Measurement System.",
        "Automatic Detection of Oversize Shipments.",
        "Loop Cross Belt Sorter.",
        "WCS Software System.",
    ]
    add_bullets_in_two_columns(doc, ktm5)

    p = doc.add_paragraph()
    run = p.add_run("Site Picture –")
    run.bold = True
    apply_normal_style(p)
    p.paragraph_format.keep_with_next = True  # Keep with image
    
    add_centered_image(doc, "assests\Images\\proj5.PNG", width_in=4.0)


def build_handled_spectrum_section(doc, counter, project_name, client_name, user_parcel_spectrum=None):
    """Build Handled Shipment Spectrum section"""
    doc.add_page_break()  # Start on new page
    tpl = choose_sorter_template(project_name)

    item_singular = tpl.item_singular.lower()
    item_cap = item_singular.capitalize()
    item_plural = item_singular + "s"
    item_plural_cap = item_plural.capitalize()

    add_numbered_heading(doc, "Handled Shipment Spectrum", counter=counter)

    intro_1 = (
        f"{client_name} operates in a business where handling a wide spectrum of {item_plural} is critical. "
        f"Falcon has carefully analyzed the provided {item_singular} spectrum and tailored the solution to your needs."
    )
    intro_2 = (
        f"Falcon proposes to use its \"{tpl.config_name}\" to deliver maximum operational benefits to {client_name}, "
        f"ensuring reliable handling of all relevant sizes and weights for your business."
    )
    p = doc.add_paragraph(intro_1)
    apply_normal_style(p)
    p = doc.add_paragraph(intro_2)
    apply_normal_style(p)

    # Subsection 1
    add_numbered_subheading(doc, tpl.subheading_51, f"{counter}.1")

    p = doc.add_paragraph(
        f"Falcon's {tpl.config_name} has a capability to handle the below mentioned "
        f"{item_plural} sizes and weight."
    )
    apply_normal_style(p)

    # Determine spec table source: user-provided or template default
    if user_parcel_spectrum and len(user_parcel_spectrum) > 0:
        # Use user-provided values
        spec_data = {item["Specification"]: {"unit": item["Unit"], "value": item["Value"]} 
                     for item in user_parcel_spectrum}
    else:
        # Fall back to template defaults
        spec_data = tpl.spec_table

    # Table
    table = doc.add_table(rows=1 + len(spec_data), cols=3)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Specification"
    hdr_cells[1].text = "Unit"
    hdr_cells[2].text = "Value"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(11)

    # Add data rows only (skip header row creation, we already have it)
    row_idx = 1
    for spec, data in spec_data.items():
        if str(data["value"]).strip() == "":
            continue  # Skip empty rows
        if row_idx >= len(table.rows):
            row_cells = table.add_row().cells
        else:
            row_cells = table.rows[row_idx].cells
        row_idx += 1
        row_cells[0].text = spec
        row_cells[1].text = data["unit"]
        row_cells[2].text = data["value"]
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                apply_normal_style(paragraph)

    # Subsection 2
    add_numbered_subheading(
        doc,
        f"{item_plural_cap} to be loaded on Sorter shall have the following characteristics:",
        f"{counter}.2"
    )

    bullets_52 = [
        f"Centre of Gravity of item must not move during conveyance or sorting.",
        f"Item must not have magnetic content, otherwise behavior of {item_singular} cannot be guaranteed.",
        "Liquid or fragile material, to avoid breaking, spillage or leakage, such as wine bottles, "
        "metal cans of paint are designated as non-conveyable items.",
        f"{item_plural_cap} shall be perfectly and safely packaged: protrusion or open surfaces are not allowed.",
        "Plastic ropes shall be perfectly adherent to the surface of the package.",
        "All items with the risk of being damaged during the transport on an automatic sorting system "
        "or damaging the sorting system; they must be robust enough to avoid disintegration of container "
        "material and loss of contents in the sorting process.",
        "Item packaging shall have enough grip to be handled on the belts during the acceleration and "
        "referencing phases.",
        "Items shall not have slippery surfaces and must be able to withstand acceleration of the items "
        "on the belt during the start-stop phases (accelerations up to 0.5 g shall be assured without any "
        "sliding or tumbling of the items on the belt conveyor).",
        f"The {item_plural} must have at least one flat and regular surface providing enough stability during "
        "conveyance.",
        "All shapes are permitted except spherical, cylindrical, or alike unstable items & shapes.",
        "All usual packaging materials are permitted (including paper, carton, plastics, plastic foil, rope, "
        "tape, textile, and wood).",
    ]
    
    # Add bullets with manual numbering to ensure restart at 1
    for idx, text in enumerate(bullets_52, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {text}")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.25)

    # Subsection 3
    add_numbered_subheading(doc, f"{item_plural_cap} not loadable on the sorter", f"{counter}.3")

    bullets_53 = [
        "Unstable items with a risk to roll or tumble on the sorting system, such as spherical or cylindrical items.",
        "Items that have a spherical or cylindrical shape.",
        "Items that are packed in material that can damage the conveyors or the sorter.",
        "Items that have sharp points (e.g., Nails) or sharp edges, that can damage the conveyors or the sorter.",
        f"Fragile {item_plural} with contents not sufficiently secured.",
        "Items that have been classified as dangerous are designated.",
        "Wet items are designated.",
        "Items with anti-slip treatment.",
        "Items with protruding parts.",
        "Items with sharp edges.",
        "Inadequately packed items that could be damaged during automatic transportation.",
        "Electrostatically loaded items.",
        "Loose parts on loads and load carriers, such as adhesive tape, stickers, slips of paper, straps, "
        "wrap foil etc. are designated as non-conveyable items.",
    ]
    
    # Add bullets with manual numbering to ensure restart at 1
    for idx, text in enumerate(bullets_53, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {text}")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.25)


def build_description_of_components_section(doc, counter, sorter_spec: dict | None):
    """
    Build Description of Components of Equipment section with fixed content and 
    Technical Specification of Sorter from Loop CBS Excel.
    """
    # Main heading
    add_numbered_heading(doc, "Description of Components of Equipment", counter=counter)
    doc.add_paragraph()

    # Elements of the sorting system
    add_numbered_subheading(doc, "Elements of the sorting system", f"{counter}.1")
    
    p = doc.add_paragraph()
    p.add_run("1. Cross Belt Sorter\n").font.name = 'Calibri'
    p.add_run("2. Scanning & Sensing on Sorter\n").font.name = 'Calibri'
    p.add_run("3. Conveyor System\n").font.name = 'Calibri'
    p.add_run("4. Steel Works – Mezzanine & Staircases").font.name = 'Calibri'
    apply_normal_style(p)
    doc.add_paragraph()

    # Cross Belt Sorter
    add_numbered_subheading(doc, "Cross Belt Sorter", f"{counter}.2")
    
    add_paragraph(doc, "Cross Belt Sorter is capable of sorting extremely high volume of versatile products in a gentle manner.")
    add_paragraph(doc, "Falcon's Cross belt sorter is powered by high efficiency linear motors and is based on 100% non-touch actuation technology leading to high throughput capabilities with extremely low noise levels.")
    add_paragraph(doc, "Falcon's Cross belt sorter is modular in design. It can be easily extended as per future requirements.")
    
    add_centered_image(doc, COMPONENT_IMAGES["cross_belt_sorter"], width_in=4.0)

    # CBS Carrier
    add_numbered_subheading(doc, "CBS Carrier", f"{counter}.3")
    
    add_paragraph(doc, "Falcon Autotech's CBS offers one of the highest belt width to carrier pitch ratios in the market today.")
    add_paragraph(doc, "This additional belt width makes the system capable of handling larger product sizes without compromising on throughput. It also reduces the dead area between carrier belts, significantly reducing the number of in-betweeners and non-sortable parcel recirculation.")
    
    add_centered_image(doc, COMPONENT_IMAGES["cbs_carrier"], width_in=4.0)
    doc.add_paragraph()

    # Servo Roller (Level 3 subheading)
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.1 Servo Roller")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "High powered DC drive servo rollers are used to actuate the carrier belts, thereby eliminating the need for complicated drive transfer mechanisms and simplifying system installation and maintenance.")
    add_centered_image(doc, COMPONENT_IMAGES["servo_roller"], width_in=3.5)
    doc.add_paragraph()

    # Chassis
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.2 Chassis")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Falcon Autotech's cross belt carrier chassis is made up of lightweight aluminium, which makes it light yet sturdy. This reduced weight leads to substantial power savings over a considerable period of usage.")
    add_centered_image(doc, COMPONENT_IMAGES["chassis"], width_in=3.5)
    doc.add_paragraph()

    # Carrier Wheels
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.3 Carrier Wheels")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Carrier wheels are thoroughly tested and proven for a long life cycle.")
    add_centered_image(doc, COMPONENT_IMAGES["wheel"], width_in=3.5)
    doc.add_paragraph()

    # Friction Wheel Drive
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.4 Friction Wheel Drive")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Falcon can provide the indigenously developed Friction Wheel Drive (FWD) to drive the cross belt loop. This driving mechanism operates on the principle of friction. The unit comprises two independent motor-driven wheels that spin in opposite directions, synchronised.")
    add_paragraph(doc, "FWDs are highly energy-efficient drives that promote sustainability compared to traditional linear induction or synchronous motor drives.")
    add_centered_image(doc, COMPONENT_IMAGES["friction_wheel"], width_in=4.0)

    # Non-contact based linear motor drive
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.5 Non-contact based linear motor drive")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Non-contact based linear induction motors can be configured at variable speeds depending upon operational requirements, providing maximum flexibility.")
    add_centered_image(doc, COMPONENT_IMAGES["linear"], width_in=4.0)
    add_paragraph(doc, "The customer can choose the preferred drive system based on performance and energy-efficiency requirements.")
    doc.add_paragraph()

    # Power Transmission
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.6 Power Transmission")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Power transmission to carriers is provided over sliding contacts that require low maintenance and offer high levels of reliability.")
    add_centered_image(doc, COMPONENT_IMAGES["power"], width_in=4.0)
    doc.add_paragraph()

    # Data Transmission
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.7 Data Transmission")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "The R-Coax cable is used for data distribution in the sorter. This leaky wave cable runs throughout the sorter length, transmitting data continuously. An antenna mounted on the super master carriage receives the signal from this cable while on the move, wirelessly.")
    add_centered_image(doc, COMPONENT_IMAGES["rcoax"], width_in=4.0)

    # Carriers positioning system
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.8 Carriers positioning system")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "The positioning system determines the exact location of each carrier in the loop at any given point in time.")
    add_paragraph(doc, "A plastic tape strip of barcodes (QR codes) runs along the sorter loop, which is continuously scanned by scanners placed on a master carrier to track and control the position of every carrier.")
    add_centered_image(doc, COMPONENT_IMAGES["carrier_position"], width_in=4.0)

    # Technical Specification of Sorter
    add_numbered_subheading(doc, "Technical Specification of Sorter", f"{counter}.4")

    if sorter_spec is None:
        add_paragraph(doc, "Technical specification of the sorter will be finalised based on project-specific configuration.")
    else:
        # Safely get values (fallbacks)
        def gv(key: str, default: str = "-") -> str:
            val = sorter_spec.get(key)
            if val is None:
                return default
            s = str(val).strip()
            return s if s else default

        rows_data = [
            ("Sorter Carrier Type", gv("sorter_carrier_type", "Loop CBS")),
            ("Sorter Speed (m/s)", gv("sorter_speed_mps")),
            ("Sorter Loop Length (m)", gv("sorter_loop_length_m")),
            ("Sorter Height (mm)", gv("sorter_height_mm")),
            ("Sorter Actuation Technology", gv("actuation_technology", "Electric")),
            ("Carrier Pitch (mm)", gv("carrier_pitch_mm")),
            ("Number of Carriers", gv("number_of_carriers")),
            ("Motor / Drive Type", gv("motor_drive_type")),
            ("Power Consumption*", gv("power_consumption")),
        ]

        table = doc.add_table(rows=1, cols=2)
        apply_table_style(table)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Technical Parameter"
        hdr_cells[1].text = "Value"

        for param, value in rows_data:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = value

        doc.add_paragraph()


def add_paragraph(doc: Document, text: str):
    """Add normal body paragraph with Calibri 11pt."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def build_description_of_components_section(doc, counter, sorter_spec: dict | None):
    """
    Build Description of Components of Equipment section with fixed content and 
    Technical Specification of Sorter from Loop CBS Excel.
    """
    doc.add_page_break()  # Start on new page
    
    # Main heading
    add_numbered_heading(doc, "Description of Components of Equipment", counter=counter)
    doc.add_paragraph()

    # Elements of the sorting system
    add_numbered_subheading(doc, "Elements of the sorting system", f"{counter}.1")
    
    p = doc.add_paragraph()
    p.add_run("1. Cross Belt Sorter\\n").font.name = 'Calibri'
    p.add_run("2. Scanning & Sensing on Sorter\\n").font.name = 'Calibri'
    p.add_run("3. Conveyor System\\n").font.name = 'Calibri'
    p.add_run("4. Steel Works – Mezzanine & Staircases").font.name = 'Calibri'
    apply_normal_style(p)
    doc.add_paragraph()

    # Cross Belt Sorter
    add_numbered_subheading(doc, "Cross Belt Sorter", f"{counter}.2")
    
    add_paragraph(doc, "Cross Belt Sorter is capable of sorting extremely high volume of versatile products in a gentle manner.")
    add_paragraph(doc, "Falcon's Cross belt sorter is powered by high efficiency linear motors and is based on 100% non-touch actuation technology leading to high throughput capabilities with extremely low noise levels.")
    add_paragraph(doc, "Falcon's Cross belt sorter is modular in design. It can be easily extended as per future requirements.")
    
    add_centered_image(doc, COMPONENT_IMAGES["cross_belt_sorter"], width_in=4.0)

    # CBS Carrier
    add_numbered_subheading(doc, "CBS Carrier", f"{counter}.3")
    
    add_paragraph(doc, "Falcon Autotech's CBS offers one of the highest belt width to carrier pitch ratios in the market today.")
    add_paragraph(doc, "This additional belt width makes the system capable of handling larger product sizes without compromising on throughput. It also reduces the dead area between carrier belts, significantly reducing the number of in-betweeners and non-sortable parcel recirculation.")
    
    add_centered_image(doc, COMPONENT_IMAGES["cbs_carrier"], width_in=4.0)
    doc.add_paragraph()

    # Servo Roller (Level 3 subheading)
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.1 Servo Roller")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "High powered DC drive servo rollers are used to actuate the carrier belts, thereby eliminating the need for complicated drive transfer mechanisms and simplifying system installation and maintenance.")
    add_centered_image(doc, COMPONENT_IMAGES["servo_roller"], width_in=3.5)
    doc.add_paragraph()

    # Chassis
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.2 Chassis")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Falcon Autotech's cross belt carrier chassis is made up of lightweight aluminium, which makes it light yet sturdy. This reduced weight leads to substantial power savings over a considerable period of usage.")
    add_centered_image(doc, COMPONENT_IMAGES["chassis"], width_in=3.5)
    doc.add_paragraph()

    # Carrier Wheels
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.3 Carrier Wheels")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Carrier wheels are thoroughly tested and proven for a long life cycle.")
    add_centered_image(doc, COMPONENT_IMAGES["wheel"], width_in=3.5)
    doc.add_paragraph()

    # Friction Wheel Drive
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.4 Friction Wheel Drive")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Falcon can provide the indigenously developed Friction Wheel Drive (FWD) to drive the cross belt loop. This driving mechanism operates on the principle of friction. The unit comprises two independent motor-driven wheels that spin in opposite directions, synchronised.")
    add_paragraph(doc, "FWDs are highly energy-efficient drives that promote sustainability compared to traditional linear induction or synchronous motor drives.")
    add_centered_image(doc, COMPONENT_IMAGES["friction_wheel"], width_in=4.0)

    # Non-contact based linear motor drive
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.5 Non-contact based linear motor drive")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Non-contact based linear induction motors can be configured at variable speeds depending upon operational requirements, providing maximum flexibility.")
    add_centered_image(doc, COMPONENT_IMAGES["linear"], width_in=4.0)
    add_paragraph(doc, "The customer can choose the preferred drive system based on performance and energy-efficiency requirements.")
    doc.add_paragraph()

    # Power Transmission
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.6 Power Transmission")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "Power transmission to carriers is provided over sliding contacts that require low maintenance and offer high levels of reliability.")
    add_centered_image(doc, COMPONENT_IMAGES["power"], width_in=4.0)
    doc.add_paragraph()

    # Data Transmission
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.7 Data Transmission")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "The R-Coax cable is used for data distribution in the sorter. This leaky wave cable runs throughout the sorter length, transmitting data continuously. An antenna mounted on the super master carriage receives the signal from this cable while on the move, wirelessly.")
    add_centered_image(doc, COMPONENT_IMAGES["rcoax"], width_in=4.0)

    # Carriers positioning system
    p = doc.add_paragraph()
    run = p.add_run(f"{counter}.3.8 Carriers positioning system")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    
    add_paragraph(doc, "The positioning system determines the exact location of each carrier in the loop at any given point in time.")
    add_paragraph(doc, "A plastic tape strip of barcodes (QR codes) runs along the sorter loop, which is continuously scanned by scanners placed on a master carrier to track and control the position of every carrier.")
    add_centered_image(doc, COMPONENT_IMAGES["carrier_position"], width_in=4.0)

    # Technical Specification of Sorter
    add_numbered_subheading(doc, "Technical Specification of Sorter", f"{counter}.4")

    if sorter_spec is None:
        add_paragraph(doc, "Technical specification of the sorter will be finalised based on project-specific configuration.")
    else:
        # Safely get values (fallbacks)
        def gv(key: str, default: str = "-") -> str:
            val = sorter_spec.get(key)
            if val is None:
                return default
            s = str(val).strip()
            return s if s else default

        rows_data = [
            ("Sorter Carrier Type", gv("sorter_carrier_type", "Loop CBS")),
            ("Sorter Speed (m/s)", gv("sorter_speed_mps")),
            ("Sorter Loop Length (m)", gv("sorter_loop_length_m")),
            ("Sorter Height (mm)", gv("sorter_height_mm")),
            ("Sorter Actuation Technology", gv("actuation_technology", "Electric")),
            ("Carrier Pitch (mm)", gv("carrier_pitch_mm")),
            ("Number of Carriers", gv("number_of_carriers")),
            ("Motor / Drive Type", gv("motor_drive_type")),
            ("Power Consumption*", gv("power_consumption")),
        ]

        table = doc.add_table(rows=1, cols=2)
        apply_table_style(table)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Technical Parameter"
        hdr_cells[1].text = "Value"

        for param, value in rows_data:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = value

        doc.add_paragraph()


def build_capacity_calculations_section(doc, counter, client_name, project_name, capacity_excel, context: Optional[ProposalContext] = None):
    """Add Sorter System Capacity section using capacity_calculations.py logic"""
    if capacity_excel:
        try:
            excel_bytes = capacity_excel.read()
            prompt = build_capacity_prompt_from_excel(excel_bytes, client_name, project_name)
            cap_data = call_groq_for_capacity(prompt)
            doc.add_page_break()
            add_capacity_section_to_doc(doc, client_name, project_name, cap_data, counter, context)
        except Exception as e:
            # Log the error but continue with document generation
            print(f"[WARNING] Capacity calculations section failed: {str(e)}")
            doc.add_page_break()
            add_numbered_heading(doc, "Sorter System Capacity", counter=counter)
            p = doc.add_paragraph(f"Capacity calculations could not be generated. Error: {str(e)}")
            apply_normal_style(p)
            # Reset file pointer in case it needs to be read again
            try:
                capacity_excel.seek(0)
            except:
                pass


def build_electrical_section(doc, counter):
    """Build Electrical System section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Electrical System", counter=counter)
    
    p = doc.add_paragraph(
        "Main power supply will supply Falcon's PDP (Power Distribution Panels) electrical cabinets. "
        "PDP cabinets supply the entire system via secondary cabinets:"
    )
    apply_normal_style(p)
    
    for item in ["Main Control Cabinet", "Induct Control Panels", "Remote Cabinets for Sorter I/O", "Scanner Control cabinets"]:
        p = doc.add_paragraph(item, style='List Bullet')
        apply_normal_style(p)
    
    add_numbered_subheading(doc, "Reference Picture of Power Distribution Panel", f"{counter}.1")
    add_centered_image(doc, "assests/Images/elec1.PNG", width_in=4)
    
    add_numbered_subheading(doc, "Main Control Panel (Reference)", f"{counter}.2")
    add_centered_image(doc, "assests/Images/elec2.PNG", width_in=4)
    
    add_numbered_subheading(doc, "Induct Stations Control Panel (Reference)", f"{counter}.3")
    add_centered_image(doc, "assests/Images/elec3.PNG")
    
    p = doc.add_paragraph()
    run = p.add_run("Engines")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "Three-phase alternating current motors (Induction) will be used through a frequency converter. "
        "The engines will be coupled with a converter to improve consumption and reduce the carbon footprint."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph("All motors will have appropriate IP ratings.")
    apply_normal_style(p)
    
    p = doc.add_paragraph()
    run = p.add_run("Sensors")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "The sensors will be supplied, standardized by type, with connector, with a cable length suitable for "
        "easy extraction, suitably protected from possible impacts."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph()
    run = p.add_run("Control command")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "The proposed solution is based on SIEMENS Programmable Logic Controller technology (PLC) platform. "
        "The entire system will be logically divided into Zones (Sorter / Feed Line / Loop), each managed by a PLC. "
        "The planned primary communication protocol is going to be ProfiNet."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph()
    run = p.add_run("Conveyor interface")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "The frequency converter of each conveyor allows the acquisition of the signals of the "
        "sensors/actuators/GIOs associated with it (e.g. conveyor end detection photocells, blockage detection "
        "photocells). Each frequency converter will be connected in series by means of the ProfiNet field bus."
    )
    apply_normal_style(p)

def build_wcs_section(doc, counter, client_name):
    """Build WCS CONTROLIT section with COMPLETE content"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Falcon's WCS CONTROLIT", counter=counter)
    
    add_centered_image(doc, "assests\Images\\wcs1.PNG", width_in=3.0)
    
    p = doc.add_paragraph(
        "Falcon WCS (Warehouse Control System) is an in-house developed IT solution by Falcon Autotech, "
        "serving as the brain behind the company's sortation solutions. It manages the real-time movement "
        "of goods and data across the system, ensuring efficient operations in high-throughput warehouses. "
        "Falcon WCS integrates seamlessly with Warehouse Management Systems (WMS), Transport Management "
        "Systems (TMS), and other external applications via APIs to enhance operational efficiency."
    )
    apply_normal_style(p)
    
    # A. System Architecture
    add_numbered_subheading(doc, "System Architecture", f"{counter}.1")
    
    p = doc.add_paragraph()
    run = p.add_run("High-Level Design (HLD) Overview")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "The Falcon WCS integrates with external systems like the Warehouse Management System (WMS) and "
        "Transport Management System (TMS). Communication occurs via APIs / WSDL / MQ communication "
        "protocols, ensuring smooth data flow for order management, shipment tracking, and other critical "
        "operations."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\wcs2.PNG")
    
    p = doc.add_paragraph()
    run = p.add_run("Key Components:")
    run.bold = True
    apply_normal_style(p)
    
    # Presentation & Session Layer
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("Presentation and Session Layer:")
    run.bold = True
    apply_normal_style(p)
    
    for item in [
        "MySQL Database: Stores operational data, shipment details, and sortation instructions.",
        "Sorter Services: Responsible for managing sorting logic and directing parcels to appropriate destinations.",
        "Dashboard: Provides a user interface for real-time monitoring of warehouse operations and performance metrics.",
        "Integration Services: Handles communication with external systems (e.g., WMS, TMS) and ensures data consistency across platforms."
    ]:
        p = doc.add_paragraph(item, style="List Bullet 2")
        apply_normal_style(p)
    
    # Application Layer
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("Application Layer:")
    run.bold = True
    apply_normal_style(p)
    
    for item in [
        "Image Services: Processes and manages images captured during the sortation process.",
        "ICR Software: Utilizes Image Character Recognition to read parcel labels and identify shipment information.",
        "PLC Software: Interfaces with Programmable Logic Controllers to manage the physical movement of parcels and control sortation equipment."
    ]:
        p = doc.add_paragraph(item, style="List Bullet 2")
        apply_normal_style(p)
    
    # Transport Layer
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("Transport Layer:")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "Sorter PLCs: Receive commands from the session layer (Sorter Services) and execute sorting operations based on real-time data.",
        style="List Bullet 2"
    )
    apply_normal_style(p)
    
    # System Communication
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("System Communication:")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "All layers are connected via a stacked switch, which provides internet and intranet connectivity. "
        "Communication between the sortation system and external systems for results or shipment data occurs through this switch.",
        style="List Bullet 2"
    )
    apply_normal_style(p)
    
    # B. High Availability Architecture
    add_numbered_subheading(doc, "High Availability Architecture", f"{counter}.2")
    
    p = doc.add_paragraph(
        "The Falcon WCS architecture ensures uninterrupted operations using a High Availability (HA) server setup. "
        "The system is designed to handle both planned and unplanned downtime, providing robust mechanisms for "
        "failover, replication, and data redundancy."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\wcs3.PNG")
    
    p = doc.add_paragraph()
    run = p.add_run("Key Components and Features of the High Availability Architecture:")
    run.bold = True
    apply_normal_style(p)
    
    # Component descriptions
    components = [
        ("Stacked Switch:", [
            "Centralizes data exchange between NAS, nodes, domain controller (DC), and peripherals.",
            "Analyzes packet headers to reduce unnecessary data transmission, enhancing LAN efficiency."
        ]),
        ("Domain Controller:", [
            "Heartbeat Monitoring: Tracks the status of nodes and initiates VM failover when necessary.",
            "Image Hosting: Stores and manages images received from the ICR (Image Character Recognition)."
        ]),
        ("NAS (Network Attached Storage):", [
            "Centralized data storage providing access to connected devices and virtual machines.",
            "Redundancy: Two NAS boxes with mirrored drives ensure data protection and availability, offering a failsafe against hardware failure."
        ]),
        ("Node:", [
            "Hyper Terminals: Nodes host and manage virtual machines (VMs) to run the warehouse control systems and related applications.",
            "Clustering: Nodes are clustered using Microsoft Windows Cluster to enable failover protection, ensuring continuous operation even in case of hardware failure."
        ]),
        ("Virtual Machine & InnoDB Cluster:", [
            "Primary VM: Hosts Falcon WCS services, while a secondary backup on the node ensures failover through network load balancing (NLB).",
            "InnoDB Cluster: Ensures data replication using a Master–Slave–Slave setup for MySQL databases, maintaining consistency and availability."
        ]),
        ("NAS Cluster:", [
            "Unified File System: NAS nodes share files across the cluster, ensuring no data loss during failover or disaster recovery.",
            "Backup NAS: Provides redundancy by replicating data between two NAS boxes, further safeguarding against failures."
        ])
    ]
    
    for comp_title, comp_items in components:
        p = doc.add_paragraph()
        run = p.add_run(comp_title)
        run.bold = True
        apply_normal_style(p)
        
        for comp_item in comp_items:
            p = doc.add_paragraph(comp_item, style='List Bullet')
            apply_normal_style(p)
    
    # Disaster Handling
    p = doc.add_paragraph()
    run = p.add_run("Disaster Handling:")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph("Recovery Time Objective (RTO) & Data Loss Objective (RPO):", style='List Bullet')
    apply_normal_style(p)
    
    for disaster_item in [
        "VM Cluster Failure: RTO = 1 hour; RPO = 1 hour.",
        "Node Failure: No impact with a single failure; RTO = 4 hours if both nodes fail.",
        "NAS Failure: Backup NAS available with no downtime, ensuring continued operation."
    ]:
        p = doc.add_paragraph(disaster_item, style="List Bullet 2")
        apply_normal_style(p)
    
    # C. WCS User Interface
    add_numbered_subheading(doc, "WCS User Interface", f"{counter}.3")
    
    p = doc.add_paragraph(
        "The Falcon WCS features a robust, user-friendly dashboard that provides real-time visibility "
        "into warehouse and sortation operations. The dashboard serves as the primary interface for "
        "monitoring key system metrics, tracking performance, and ensuring smooth operations."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph()
    run = p.add_run("Dashboard Overview")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "The WCS dashboard offers real-time data visualization, helping warehouse operators and IT teams "
        "make data-driven decisions. Users can monitor system health, performance, and detect anomalies "
        "through an intuitive graphical interface."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph()
    run = p.add_run("Key Features of the Dashboard:")
    run.bold = True
    apply_normal_style(p)
    
    dashboard_features = [
        "System Health Monitoring: Displays metrics such as CPU utilization, memory usage, disk performance, and system load across the infrastructure.",
        "Real-Time Sortation Monitoring: Shows the real-time movement of parcels within the sortation system, including chute assignments and shipment statuses.",
        "Error Reporting: Notifies users of system errors, network disruptions, and potential failures in real time, allowing for quick resolution and minimal downtime.",
        "Performance Metrics: Provides detailed reports on sortation throughput, parcel handling times, and system efficiency to ensure that warehouse targets are met.",
        "User Role Management: The dashboard allows different levels of access based on user roles, ensuring that the right personnel can view or manage the system as needed."
    ]
    
    for feature in dashboard_features:
        p = doc.add_paragraph(feature, style='List Bullet')
        apply_normal_style(p)
    
    p = doc.add_paragraph(
        "In the context of this IT dashboard, the following user interactive screens are provided:"
    )
    apply_normal_style(p)
    
    # Dashboard screens with images
    dashboard_screens = [
        ("Dashboard (Home Screen): Provides an overview of important metrics, data visualizations, and summary information related to the IT system or processes.", "assests\Images\\wcs4.PNG"),
        ("Live Bags: Displays real-time information and status updates regarding bags or parcels currently in transit or being processed.", "assests\Images\\wcs5.PNG"),
        ("Bay Status: Offers insights into the status and availability of different processing bays or areas within the system.", "assests\Images\\wcs6.PNG"),
        ("Processed Packages: Shows details and statistics related to packages or items that have been successfully processed or handled by the system.", "assests\Images\\wcs7.PNG"),
        ("Configuration Setting: Enables users to configure and customize various settings and parameters within the IT system or dashboard.", "assests\Images\\wcs8.PNG")
    ]
    
    for screen_desc, screen_img in dashboard_screens:
        p = doc.add_paragraph(screen_desc, style='List Bullet')
        apply_normal_style(p)
        add_centered_image(doc, screen_img)
    
    # Additional screens without images
    additional_screens = [
        "Report & Analysis: Allows users to generate and access comprehensive reports, analytics, and insights based on the data collected by the IT dashboard.",
        "Rejection Bay Mapping: Provides functionality to map and manage rejection bays or areas where packages are deemed unsuitable for processing.",
        "Alarms: Displays alerts, notifications, or alarms related to system events, errors, or anomalies that require attention or investigation.",
        "Calibration Settings: Allows users to adjust and calibrate system settings, parameters, or sensors to ensure accurate and reliable performance.",
        "Operator Management: Offers features and tools to manage and monitor the operators or personnel responsible for operating the IT system.",
        "User Management: Provides functionality to manage user accounts, permissions, roles, and access levels within the IT dashboard.",
        "User Guide: The 'User Guide' page offers comprehensive documentation and instructions on how to use the IT dashboard effectively. It serves as a reference guide for users."
    ]
    
    for screen in additional_screens:
        p = doc.add_paragraph(screen, style='List Bullet')
        apply_normal_style(p)
    
    # D. Communication Architecture
    add_numbered_subheading(doc, "Communication Architecture", f"{counter}.4")
    
    p = doc.add_paragraph(
        "Falcon WCS operates within a highly interconnected system, ensuring seamless communication between "
        "the WCS server, on-premises devices (such as sorter PLCs, PTL devices, 1D scanners, and HHT devices), "
        "and client systems. This communication architecture facilitates real-time data exchange and operational "
        "control, optimizing sortation processes and warehouse efficiency."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\wcs9.PNG")
    
    p = doc.add_paragraph()
    run = p.add_run("On-Premises Communication")
    run.bold = True
    apply_normal_style(p)
    
    # Communication devices
    comm_devices = [
        ("Sorter PLC Devices:", [
            "Protocol: Falcon WCS communicates with sorter PLCs using either the Siemens S7 protocol or the Omron communication protocol.",
            "Functionality: The sorter PLC devices receive sortation instructions from the WCS and execute the sorting process by directing parcels to the appropriate chute based on the system's real-time data."
        ]),
        ("PTL (Pick-to-Light) Devices:", [
            "Protocol: PTL devices communicate with Falcon WCS using the TCP/IP protocol.",
            "Functionality: The system sends commands to the PTL devices for guiding manual picking operations by lighting up indicators at the appropriate bins or shelves, improving operational accuracy and speed."
        ]),
        ("1D Scanners:", [
            "Protocol: These barcode scanners also use the TCP/IP protocol to communicate with the WCS.",
            "Functionality: The scanners capture barcode data from the parcels, and this information is sent to the WCS for processing, such as determining sorting destinations."
        ]),
        ("HHT (Handheld Terminal) Devices:", [
            "Protocol: The wireless HHT devices communicate with Falcon WCS over Wi-Fi.",
            "Functionality: The HHT devices send scan input data (e.g., barcodes) to the server over Wi-Fi. The WCS processes this data and sends the required output instructions back to the HHT device and associated PTL devices. The HHT device executes these instructions, facilitating real-time decision making and execution for operators."
        ])
    ]
    
    for device_title, device_items in comm_devices:
        p = doc.add_paragraph()
        run = p.add_run(device_title)
        run.bold = True
        apply_normal_style(p)
        
        for device_item in device_items:
            p = doc.add_paragraph(device_item, style='List Bullet')
            apply_normal_style(p)
    
    # E. Client Communication
    add_numbered_subheading(doc, "Client Communication", f"{counter}.5")
    
    p = doc.add_paragraph()
    run = p.add_run("Data Transfer Methods:")
    run.bold = True
    apply_normal_style(p)
    
    transfer_methods = [
        "API: Falcon WCS can communicate processed data to client systems through API calls, allowing for seamless integration with external software.",
        "MQ (Message Queuing): Falcon WCS can also send data via message queues, ensuring reliable delivery of messages even during network downtime.",
        "WSDL/XML: For structured data exchanges, Falcon WCS supports WSDL and XML formats for client communication.",
        f"Other Protocols: Additional methods for data transfer may include customized protocols depending on {client_name}'s requirements."
    ]
    
    for method in transfer_methods:
        p = doc.add_paragraph(method, style='List Bullet')
        apply_normal_style(p)
    
    p = doc.add_paragraph()
    run = p.add_run("Purpose:")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        "The data sent to the client can include sortation results, system performance reports, and operational "
        "analytics, which can be used for further processing or reporting within external systems like Warehouse "
        "Management Systems (WMS) and Transport Management Systems (TMS).",
        style='List Bullet'
    )
    apply_normal_style(p)
    
    # F. HAA Server Specifications
    add_numbered_subheading(doc, f"HAA Server Specifications (In {client_name}'s Scope)", f"{counter}.6")
    
    p = doc.add_paragraph("20-core configuration with 128 GB RAM in T440 and 64 GB RAM in T40.")
    apply_normal_style(p)
    
    # Server spec table
    table = doc.add_table(rows=1, cols=3)
    apply_table_style(table)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "SN"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Qty"
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri (Body)'
                run.font.size = Pt(11)
    
    server_specs = [
        ("1", "DELL PowerEdge T440 Server", "2"),
        ("2", "Intel Xeon Silver 4210R 2.4G, 10C/20T, 9.6GT/s, 13.75M Cache, Turbo, HT (100W) DDR4-2400", "4"),
        ("3", "32GB RDIMM, 3200MT/s, Dual Rank", "8"),
        ("4", "480GB SSD SATA Read Intensive 6Gbps 512n 2.5in Hot-plug Drive, 1 DWPD", "4"),
        ("5", "H730P RAID Controller, 2GB NV Cache, Adapter, Low Profile", "2"),
        ("6", "Broadcom 5720 Dual Port 1Gb On-Board LOM", "2"),
        ("7", "Broadcom 57416 Dual Port 10Gb Base-T, OCP NIC 3.0", "2"),
        ("8", "Power Cord, C13, 1.8M, 250V, 10A (India BIS, IS1293)", "4"),
        ("9", "Dual, Hot-Plug, Redundant Power Supply (1+1), 750W", "2")
    ]
    
    for sn, desc, qty in server_specs:
        row_cells = table.add_row().cells
        row_cells[0].text = sn
        row_cells[1].text = desc
        row_cells[2].text = qty
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                apply_normal_style(paragraph)

def build_scada_section(doc, counter, client_name):
    """Build SCADA section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Falcon's Visual Inspection System (SCADA)", counter=counter)
    
    p = doc.add_paragraph(
        "SCADA stands for Supervisory Control and Data Acquisition. It is a system of hardware "
        "and software components that allows for remote monitoring, control, and data acquisition "
        "of industrial processes or facilities."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph(
        f"The Visualization system provided by FALCON (or SCADA) allows the monitoring and control of "
        f"the different systems delivered for the {client_name}. This SCADA system receives from each "
        f"monitored sub-system all information on their operating status in real time."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph("At the system monitoring level, the functions performed are:")
    apply_normal_style(p)
    
    functions = [
        "Field data acquisition.",
        "Animated visualization of equipment.",
        "Representation of the operating mode of the system (nominal, contingency, etc.).",
        "Alarm management.",
        "Alarm history management.",
        "Diagnostic help.",
        "Failure detection.",
        "Equipment control.",
        "Statistics on equipment operation.",
        "Historical Statistical Report.",
        "Recording and archiving.",
        "Safety operator interface.",
    ]
    
    for item in functions:
        p = doc.add_paragraph(item, style='List Bullet')
        apply_normal_style(p)
    
    add_numbered_subheading(doc, "FIELD DATA ACQUISITION", f"{counter}.1")
    p = doc.add_paragraph(
        "The field data acquisition function is performed by the SCADA system connected to the sorters' PLCs. "
        "The communication with the PLCs is done using equipped CPU cards that are able to manage the "
        "communication with the PLC on the Industrial Ethernet network, without overloading the server."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\main_plc.PNG")
    
    add_numbered_subheading(doc, "ANIMATED SYSTEM VISUALIZATION", f"{counter}.2")
    p = doc.add_paragraph(
        "The animated view represents the dynamic graphical user interface that allows real-time monitoring "
        "of the controlled systems and the execution of their control procedures."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\animated_sys_v1.PNG")
    add_centered_image(doc, "assests\Images\\animated_sys_v2.PNG")
    
    add_numbered_subheading(doc, "ALARM MANAGEMENT", f"{counter}.3")
    p = doc.add_paragraph(
        "The alarm pages display a series of information to identify the nature of the alarm or event, "
        "the elements involved and the time."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\alarm.PNG")

def build_key_components_section(doc, counter, components_df):
    """Build Key Components Make section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Key Components Make", counter=counter)
    
    table = doc.add_table(rows=1, cols=2)
    apply_table_style(table)
    
    hdr = table.rows[0].cells
    hdr[0].text = "Items"
    hdr[1].text = "Make"
    
    for run in hdr[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    for run in hdr[1].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    for _, row in components_df.iterrows():
        r = table.add_row().cells
        r[0].text = str(row.get("Items", ""))
        r[1].text = str(row.get("Make", ""))
        
        for cell in r:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

def build_safety_section(doc, counter):
    """Build Principal of Safety section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Principal of Safety", counter=counter)
    
    p = doc.add_paragraph()
    run = p.add_run("1. E-Stops")
    run.bold = True
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\e-stop.png", width_in=4.5)
    
    p = doc.add_paragraph()
    run = p.add_run("            a. At Every Conveyor Module – Both sides")
    apply_normal_style(p)
    add_centered_image(doc, "assests\Images\\at-every.png", width_in=4.5)
    
    p = doc.add_paragraph()
    run = p.add_run("             b. VDS Chutes")
    apply_normal_style(p)
    add_centered_image(doc, "assests\Images\\emergency-vds.png", width_in=4.5)
    
    p = doc.add_paragraph()
    run = p.add_run("2. Pull Cords Switch")
    run.bold = True
    apply_normal_style(p)
    
    p = doc.add_paragraph("Required for Infeed & Takeout Conveyors")
    apply_normal_style(p)
    
    # Add pull cord images
    table_pc = doc.add_table(rows=1, cols=2)
    if os.path.exists("assests\Images\\pull-cords.PNG"):
        p = table_pc.rows[0].cells[0].add_paragraph()
        run = p.add_run()
        run.add_picture("assests\Images\\pull-cords.PNG", width=Inches(3))
    
    if os.path.exists("assests\Images\\puul-cords-arch.PNG"):
        p = table_pc.rows[0].cells[1].add_paragraph()
        run = p.add_run()
        run.add_picture("assests\Images\\puul-cords-arch.PNG", width=Inches(3))
    
    p = doc.add_paragraph()
    run = p.add_run("3. Fencing")
    run.bold = True
    apply_normal_style(p)
    
    table_fence = doc.add_table(rows=1, cols=2)
    table_fence.rows[0].cells[0].text = "a. Between Inducts \nb. Between Inducts and Sorter"
    
    if os.path.exists("assests\Images\\fencing.png"):
        p = table_fence.rows[0].cells[1].add_paragraph()
        run = p.add_run()
        run.add_picture("assests\Images\\fencing.png", width=Inches(2.5))
    
    p = doc.add_paragraph()
    run = p.add_run("4. Leg Guards")
    run.bold = True
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\leg_gurads.png", width_in=4.5)

def build_infrastructure_section(doc, counter):
    """Build Infrastructure section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Infrastructure", counter=counter)
    
    sections = [
        ("a. Fire Protection-", 
         "Falcon's scope does not cover the design or provision of fire protection infrastructure, "
         "utilities, or related services. It is expected that the customer's sprinkler contractor "
         "will design and supply the in-rack sprinkler systems, including connectors and mounting "
         "brackets. These designs should be submitted to Falcon for review during the engineering phase."),
        
        ("b. Power Supply-",
         "The Customer must provide temporary power for installation and permanent power for "
         "commissioning. Protected multi-gang power points for workstations and peripherals will be "
         "supplied by the Customer, with planning for their locations done with the operations and "
         "IT teams."),
        
        ("c. Floor Requirements-",
         "The Customer must provide flooring with appropriate loading strength and space at the site. "
         "Falcon assumes that the floor slab will not contain corrosive materials that could affect "
         "standard fixings."),
        
        ("d. Estimated Floor Load-",
         "Estimated floor loads, including distributed and point loads, will be provided during the "
         "detailed engineering phase of the project."),
        
        ("e. Staging, Laydown and Assembly Area-",
         "The Customer is required to provide sufficient space on the same floor, adjacent to the "
         "installation site, for staging, storage, and equipment assembly."),
        
        ("f. Site Access and Unloading-",
         "The Customer is required to allocate sufficient on-site space for parking and staging "
         "shipping containers to facilitate Falcon's delivery schedule."),
        
        ("g. Lighting-",
         "All lighting is excluded from Falcon's scope of supply and must be provided by the Customer "
         "or their contractor. This includes lighting for service areas, operational areas, and beneath "
         "platforms and walkways.")
    ]
    
    for title, content in sections:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        apply_normal_style(p)
        
        p = doc.add_paragraph(content)
        apply_normal_style(p)

def build_program_org_section(doc, counter, client_name, gantt_file):
    """Build Program Organisation section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Program Organisation", counter=counter)
    
    add_numbered_subheading(doc, "Program Schedule", f"{counter}.1")
    
    if gantt_file is not None:
        p = doc.add_paragraph()
        run = p.add_run()
        gantt_file.seek(0)
        run.add_picture(gantt_file, width=Inches(6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph("Attach your timeline Gantt chart here.")
        apply_normal_style(p)
    
    add_numbered_subheading(doc, "Program Management", f"{counter}.2")
    
    p = doc.add_paragraph("For this program, proposed approach covers the following aspects:")
    apply_normal_style(p)
    
    bullets = [
        "Creation and monitoring of the project plan.",
        "Weekly/Fortnightly meeting to share project status.",
        "Scheduling of the resource management.",
        "Management of risks and opportunities.",
        "Management of the requirements.",
        "Management of the list of anomalies or reservations.",
    ]
    
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        apply_normal_style(p)
    
    p = doc.add_paragraph(
        "The Project will be closely monitored under Falcon's Governance model as structured below."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\governence.png", width_in=5.5)
    doc.add_page_break()
    add_numbered_subheading(doc, "Project Team", f"{counter}.3")
    
    p = doc.add_paragraph(
        f"Team of 3 to 4 member from Projects team will co-ordinate on regular basis with "
        f"{client_name} and internal stakeholders for smooth execution of the project."
    )
    apply_normal_style(p)
    
    p = doc.add_paragraph()
    
    run = p.add_run(f"{client_name}'s Team")
    run.bold = True
    run.font.size = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_centered_image(doc, "assests\Images\\team.png", width_in=5.5)

def build_client_responsibility_section(doc, counter, client_name):
    """Build Client Responsibility section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Client Responsibility", counter=counter)
    
    sections = [
        (f"{client_name} Responsibilities During the Assembly and Commissioning Phase", [
            "Provision of the site complex and office area facilities.",
            "The possibility of authorizing access to the site and the execution of the installation work up to 7 days a week and 24 hours a day if deemed necessary and if requested by FALCON.",
            "Free provision, during the installation phase, of the power supply necessary for the installation activities (estimated at 20 kW).",
            "Provision, during the commissioning phase, of the power supply necessary for the operation of the shipment sorting system free of charge at the date of FALCON need.",
            "Provision of the IT system functionality in accordance with the specification at the date of FALCON need.",
            "The customer is responsible for a safe working environment.",
            "The customer makes arrangements for the working area(s) to be protected against direct weather influences.",
            "The customer provides adequate lighting, heating, and ventilation to create a normal working environment.",
        ]),
        (f"Responsibilities of {client_name} During the Tests", [
            "Provision of the test loads and barcode labels required for the tests.",
            "Provision of personnel required for test activities (loading and unloading operations).",
            "Provision of the necessary information to sort the shipments correctly.",
            "Verify with FALCON the quality and conformity of the test loads (labels, cartons).",
        ]),
        (f"{client_name} Responsibilities During the Training", [
            "Free from their usual work, the employees participate in the training for the duration of the training.",
            "Provision of a list of participants for each available training course 3 days before the start of the course.",
            "Provision of a classroom equipped with a whiteboard, video projector, projection screen, and enough space for desks or tables and chairs for the trainer and trained staff.",
        ])
    ]
    
    for title, bullets in sections:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        apply_normal_style(p)
        
        for b in bullets:
            p = doc.add_paragraph(b, style='List Bullet')
            apply_normal_style(p)

def build_handover_section(doc, counter):
    """Build System Handover section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "System Handover", counter=counter)
    
    p = doc.add_paragraph(
        "The system handover will follow the workflow shown below. "
        "Each stage is described in the following sections."
    )
    apply_normal_style(p)
    
    add_centered_image(doc, "assests\Images\\handover.PNG")
    
    sections = [
        ("Installation and Commissioning",
         "Completion of all activities required to bring the system to an operational state "
         "and ready for formal testing."),
        
        ("Pre-UAT",
         "Pre-UAT consists of checks and tests performed before the formal User Acceptance Testing. "
         "It ensures the system is stable, integrated and ready for end-user validation."),
        
        ("UAT (User Acceptance Test)",
         "UAT is performed by the client's users to verify that the solution meets agreed requirements "
         "and behaves as expected under real-world operating conditions."),
        
        ("Minor & Major Faults",
         "Any issues found during testing are categorized as minor or major faults. "
         "Minor faults affect limited areas without blocking successful test completion and are added "
         "to the snag list."),
        
        ("System Snag Points",
         "After UAT, all open minor faults are tracked as system snag points. "
         "Falcon shares a snag list with the client, detailing for each issue: description, date, location, "
         "category, responsible party, target completion date and verification / sign-off."),
        
        ("System Handover Letter",
         "After successful acceptance testing or closure of all snags, Falcon issues a handover letter "
         "confirming that the system has been installed and accepted by the client.")
    ]
    
    for title, content in sections:
        p = doc.add_paragraph()
        run = p.add_run(title + "\n")
        run.bold = True
        apply_normal_style(p)
        
        p = doc.add_paragraph(content)
        apply_normal_style(p)

def build_commercial_section(doc, counter, price_data, payment_terms, apply_bca):
    """Build Commercial section with price sheet"""
    add_numbered_heading(doc, "Commercial", counter=counter)
    
    if not price_data:
        p = doc.add_paragraph("Commercial details to be added.")
        apply_normal_style(p)
        return
    
    # Price Sheet Title
    title = price_data.get("price_sheet_title") or "Price Sheet"
    add_numbered_subheading(doc, title, f"{counter}.1")
    
    items = price_data.get("items", [])
    total_row = price_data.get("total_row")
    
    # Price Table: S. No | Component | Price
    table = doc.add_table(rows=1, cols=3)
    apply_table_style(table)
    
    hdr = table.rows[0].cells
    hdr[0].text = "S. No"
    hdr[1].text = "Component"
    hdr[2].text = "Price"
    
    for run in hdr[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Calibri (Body)'
        run.font.size = Pt(11)
    for run in hdr[1].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Calibri (Body)'
        run.font.size = Pt(11)
    for run in hdr[2].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Calibri (Body)'
        run.font.size = Pt(11)
    
    # Add price items
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get("s_no", ""))
        row_cells[1].text = str(item.get("label", ""))
        row_cells[2].text = str(item.get("price", ""))
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                apply_normal_style(paragraph)
    
    # Total row
    if total_row:
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = str(total_row.get("label", "Total"))
        row_cells[2].text = str(total_row.get("price", ""))
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                apply_normal_style(paragraph)
    
    # Optional BCA discount row
    if apply_bca and total_row:
        final_total_str = apply_bca_discount_to_price_data(price_data, 4.5)
        if final_total_str:
            row_cells = table.add_row().cells
            row_cells[0].text = ""
            row_cells[1].text = "Final Total (after 4.5% BCA Discount)"
            row_cells[2].text = final_total_str
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                    apply_normal_style(paragraph)
    
    # Payment Terms section
    if payment_terms:
        doc.add_paragraph("")  # spacing
        add_numbered_subheading(doc, "Payment Terms", f"{counter}.2")
        
        pt_table = doc.add_table(rows=1, cols=2)
        apply_table_style(pt_table)
        
        pt_hdr = pt_table.rows[0].cells
        pt_hdr[0].text = "Payment Percentage"
        pt_hdr[1].text = "Stage"
        
        for run in pt_hdr[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Calibri (Body)'
            run.font.size = Pt(11)
        for run in pt_hdr[1].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Calibri (Body)'
            run.font.size = Pt(11)
        
        for row in payment_terms:
            perc = str(row.get("Payment Percentage", "")).strip()
            stage = str(row.get("Stage", "")).strip()
            if not perc and not stage:
                continue
            r = pt_table.add_row().cells
            r[0].text = perc
            r[1].text = stage
            
            for cell in r:
                for paragraph in cell.paragraphs:
                    apply_normal_style(paragraph)

def build_warranty_section(doc, counter, warranty_type, duration, start_cond, extended_text, amc_text, transport_text):
    """Build Warranty Period section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Warranty Period", counter=counter)
    
    intro = f"Falcon's offered System comes with a {warranty_type.lower()} of {duration} (starts {start_cond})"
    if not intro.endswith("."):
        intro += "."
    
    if extended_text:
        intro += f" {extended_text}"
    if amc_text:
        intro += f" {amc_text}"
    
    p = doc.add_paragraph(intro)
    apply_normal_style(p)
    
    p = doc.add_paragraph("The warranty covers the following support:")
    apply_normal_style(p)
    
    coverage = [
        "24 X 7 Telephonic, Email and Remote Service Support when required.",
        "Regular Software updates and Bug Fixes.",
        "Supply of Mechanical and Electrical components in case of failure (excluding damages as mentioned in the Exclusion Clause).",
    ]
    
    for item in coverage:
        p = doc.add_paragraph(item, style='List Bullet')
        apply_normal_style(p)
    
    p = doc.add_paragraph("The following items are excluded from warranty:")
    apply_normal_style(p)
    
    exclusions = [
        "Normal wear and tear.",
        "Consumables.",
        "Faulty articles continued.",
        "Failure to comply with the manufacturer's recommendations.",
        "Negligence or abnormal use of equipment.",
    ]
    
    for item in exclusions:
        p = doc.add_paragraph(item, style='List Bullet')
        apply_normal_style(p)
    
    if transport_text:
        p = doc.add_paragraph(transport_text)
        apply_normal_style(p)

def build_exclusions_section(doc, counter, selected_exclusions):
    """Build Exclusions section"""
    doc.add_page_break()  # Start on new page
    add_numbered_heading(doc, "Exclusions", counter=counter)
    
    intro = (
        "The scope of supply includes all parts which are defined in the Supplier's quotation.\n"
        "All other parts which are not defined in the Supplier's quotation do not belong to the Supplier's "
        "scope of supply and are excluded. The following parts are also excluded:"
    )
    
    p = doc.add_paragraph(intro)
    apply_normal_style(p)
    
    fixed_exclusions = [
        "Construction Power",
        "Building infrastructure; building structure, doors, fire exits, levelling devices, "
        "building extinguisher and fire alarm system, building heating and lighting system.",
        "Electrical power supply and wiring to the main control cabinets.",
        "UPS for Controls and Drives",
        "Network cabling up to the main server rack.",
        "Intermediate wiring to parts which are to be supplied by the Purchaser/others.",
        "Emergency/Uninterruptable power supply.",
        "Fire-alarm and fire protection devices.",
        "Traffic and route markings.",
        "Laydown area / unloading and laydown area.",
        "Ram protection devices.",
        "Cat walks, bridges, maintenance aisles and platforms.",
        "All kind of network incl. Local Area Network (LAN/WLAN), exceeding the scope described in Scope of Supply.",
        "Any kind of civil work.",
        "Any adjustment of the Supplier's scope of supply to local rules and regulations.",
        "X-Ray machines.",
        "Roller cages / pallets.",
        "Simulation and 3D animation of the sorter system.",
        "Interface with other equipment not specified in this offer.",
        "Provision of facilities for the control room (furniture, air conditioning, heating, etc.).",
        "The supply and installation of fencing around the different corridors.",
        "Any item specifically indicated as not forming part of the subject matter of the Seller's supply in the offer documentation.",
    ]
    
    all_exclusions = fixed_exclusions + selected_exclusions
    
    for item in all_exclusions:
        p = doc.add_paragraph(item, style='List Bullet')
        apply_normal_style(p)

def build_proposed_system_technical_details_section(doc, counter, mechanical_bom_items=None, bom_json=None):
    """
    Build Proposed System Technical Details section with 3 subsections:
    - Mechanical equipment: from mechanical_bom_items (Groq API extraction)
    - Electrical Equipment: from bom_json (Groq API)
    - Control System: from bom_json (Groq API)
    
    Args:
        doc: python-docx Document object
        counter: Section number counter
        mechanical_bom_items: List of dicts with pos, qty, description, value (from Groq)
        bom_json: Dict with sections for Electrical Equipment and Control System (from Groq)
    """
    doc.add_page_break()
    
    add_numbered_heading(doc, "Proposed System Technical Details", counter=counter)
    
    has_mechanical = mechanical_bom_items is not None and len(mechanical_bom_items) > 0
    has_electrical_control = bom_json and "sections" in bom_json
    
    if not has_mechanical and not has_electrical_control:
        p = doc.add_paragraph("No technical details available.")
        apply_normal_style(p)
        return
    
    section_number = 1
    
    # --- 1. Mechanical equipment (from Groq API) ---
    if has_mechanical:
        add_numbered_subheading(doc, "Mechanical equipment", f"{counter}.{section_number}")
        section_number += 1
        
        # Create table with 4 columns: Pos., Qty., Description, Value
        table = doc.add_table(rows=1, cols=4)
        apply_table_style(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr = table.rows[0].cells
        hdr[0].text = "Pos."
        hdr[1].text = "Qty."
        hdr[2].text = "Description"
        hdr[3].text = "Value"
        
        # Format header row
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Add data rows from list of items
        for item in mechanical_bom_items:
            row = table.add_row().cells
            
            # Pos - convert to string
            row[0].text = str(item.get("pos", ""))
            
            # Qty - convert to string
            row[1].text = str(item.get("qty", ""))
            
            # Description - handle newlines (may come as \n or actual newlines)
            desc = str(item.get("description", ""))
            desc = desc.replace("\\n", "\n")  # Convert escaped newlines
            row[2].text = desc
            
            # Value - handle newlines
            val = str(item.get("value", ""))
            val = val.replace("\\n", "\n")  # Convert escaped newlines
            row[3].text = val
            
            # Format each cell
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(1.2)
        table.columns[2].width = Inches(3.5)
        table.columns[3].width = Inches(1.3)
        
        # Add spacing after table
        doc.add_paragraph()
    
    # --- 2. Electrical Equipment and Control System (from Groq bom_json) ---
    if has_electrical_control:
        for section in bom_json.get("sections", []):
            title = section.get("title", "")
            items = section.get("items", [])
            
            # Skip "Mechanical equipment" section - we use separate Groq call for that
            if "mechanical" in title.lower():
                continue
            
            if not items:
                continue
            
            # Add subsection heading (e.g., "14.2 Electrical Equipment")
            add_numbered_subheading(doc, title, f"{counter}.{section_number}")
            section_number += 1
            
            # Create table with 4 columns: Pos., Qty., Description, Value
            table = doc.add_table(rows=1, cols=4)
            apply_table_style(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            hdr = table.rows[0].cells
            hdr[0].text = "Pos."
            hdr[1].text = "Qty."
            hdr[2].text = "Description"
            hdr[3].text = "Value"
            
            # Format header row
            for cell in hdr:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Add data rows
            for item in items:
                row = table.add_row().cells
                
                # Pos.
                row[0].text = str(item.get("pos", ""))
                
                # Qty.
                row[1].text = item.get("qty") or ""
                
                # Description (multiple lines)
                desc_lines = item.get("description_lines") or []
                row[2].text = "\n".join(desc_lines)
                
                # Value (multiple lines)
                value_lines = item.get("value_lines") or []
                row[3].text = "\n".join(value_lines)
                
                # Format each cell
                for cell in row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Set column widths
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(1.2)
            table.columns[2].width = Inches(3.5)
            table.columns[3].width = Inches(1.3)
            
            # Add spacing after table
            doc.add_paragraph()

def build_proposed_system_description_section(doc, counter, client_name, project_name, 
                                              process_flow_text, layout_png_path,
                                              system_description_text=None,
                                              costing_file=None):
    """Build Proposed System Description section (5.0) with all subsections"""
    doc.add_page_break()
    
    add_numbered_heading(doc, "Proposed System Description", counter=counter)
    
    # 5.1 Objective
    add_numbered_subheading(doc, "Objective", f"{counter}.1")
    objective_text = (
        "The purpose of this proposal is to present the design, manufacturing, "
        "installation, commissioning, testing, and acceptance testing of the Cross Belt Sorter "
        f"system for sorting shipments, as per {client_name} requirements."
    )
    p = doc.add_paragraph(objective_text)
    apply_normal_style(p)
    doc.add_paragraph("")
    
    # 5.2 Summary of the System (layout PNG)
    add_numbered_subheading(doc, "Summary of the System", f"{counter}.2")
    
    if layout_png_path and os.path.exists(layout_png_path):
        p = doc.add_paragraph(
            "The following layout view illustrates the overall arrangement of infeed conveyors, sorter loop, "
            "and output chutes for the proposed system."
        )
        apply_normal_style(p)
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(layout_png_path, width=Inches(6.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
    else:
        p = doc.add_paragraph("The detailed layout is provided separately in the attached drawing.")
        apply_normal_style(p)
        doc.add_paragraph("")
    
    # 5.3 Process Flow of the System
    add_numbered_subheading(doc, "Process Flow of the System", f"{counter}.3")
    
    # Parse process flow text and apply numbered + bold formatting
    flow_lines = process_flow_text.splitlines()
    flow_counter = 1
    in_output_chutes = False
    sub_counter = ord('a')  # for alphabetical sub-numbering
    
    for line in flow_lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip any lines containing _ParagraphStyle or other style object representations
        if '_ParagraphStyle' in line or '_Style' in line or 'id:' in line and '<' in line:
            continue
        
        # Skip lines that are just "Process Flow" or similar headers
        if line.lower() in ['process flow', 'process']:
            continue
        
        # Check if this is "Output Chutes" main heading
        if line.lower().startswith('output chutes'):
            in_output_chutes = True
            sub_counter = ord('a')  # Reset alphabetical counter for sub-items
            p = doc.add_paragraph()
            run = p.add_run(f"{flow_counter}. ")
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            
            run = p.add_run("Output Chutes")
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            
            flow_counter += 1
            continue
        
        # Check if we're in Output Chutes section and this is a sub-component (line starts with a. b. c. etc)
        if in_output_chutes and len(line) > 2 and line[0].isalpha() and line[1] == '.':
            # This is already formatted as "a. Live Chutes - Description" - extract and reformat
            clean_line = line[3:].strip() if len(line) > 3 else line
            
            # Split component title from description
            # Can be separated by " - ", " – ", or just " "
            separators = [" – ", " - ", ":"]
            component_title = clean_line
            description = ""
            
            for sep in separators:
                if sep in clean_line:
                    parts = clean_line.split(sep, 1)
                    component_title = parts[0].strip()
                    description = sep + parts[1] if len(parts) > 1 else ""
                    break
            
            # If no separator found, check if first part is title (ends before " There are" or " A total")
            if not description:
                if " There are" in clean_line or " A total" in clean_line:
                    match = re.search(r'^([^–\-:]+?)(?=\s+(?:There are|A total|One|[0-9]))', clean_line)
                    if match:
                        component_title = match.group(1).strip()
                        description = " " + clean_line[len(component_title):].strip()
            
            p = doc.add_paragraph()
            # Only bold the letter label and component title
            run = p.add_run(f"{chr(sub_counter)}. ")
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            
            run = p.add_run(component_title)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            
            # Description is NOT bold
            if description:
                run = p.add_run(description)
                run.bold = False
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            
            sub_counter += 1
            continue
        
        # Main component line (contains colon - typical format: "Component Name: Description")
        if ':' in line and not in_output_chutes:
            # Extract component name and description
            component_name = line.split(':')[0].strip()
            rest_of_line = ':'.join(line.split(':')[1:]).strip()
            
            p = doc.add_paragraph()
            run = p.add_run(f"{flow_counter}. ")
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            
            run = p.add_run(component_name)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            
            if rest_of_line:
                run = p.add_run(f": {rest_of_line}")
                run.bold = False
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            
            flow_counter += 1
            continue
        
        # Fallback for any other lines (shouldn't normally reach here)
        if not in_output_chutes:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^\*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    text = part[2:-2]
                    run = p.add_run(text)
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
    
    doc.add_paragraph("")
    
    # 5.4 Main Benefits
    add_numbered_subheading(doc, "Main Benefits of the Proposed Solution", f"{counter}.4")
    benefits = [
        "High operational throughput.",
        "Low occupancy of floor space in the building.",
        "Narrow discharge centers for the increased number of splits in limited space.",
        (
            "FALCON's CBS can adapt to changing business requirements by adjusting its speed "
            "to match the operational throughput requirement, thereby leading to power savings "
            "and reduced system wear & tear."
        ),
    ]
    for b in benefits:
        p = doc.add_paragraph(b, style='List Bullet')
        apply_normal_style(p)
    
    # NOTE: System Description is now a separate top-level section (6. System Description)
    # It is added via build_system_description_section() after this function


def transform_system_description_numbering(text: str, section_counter: int) -> str:
    """
    Transform internal numbering in System Description to use proper section counter.
    
    Converts:
    - "**1. Infeed System**" -> "**6.1. Infeed System**"
    - "**1.1 Infeed Conveyors**" -> "**6.1.1. Infeed Conveyors**"
    - "**1.1.1 Aligning Conveyor**" -> "**6.1.1.1. Aligning Conveyor**"
    
    Where 6 is the section_counter.
    """
    import re
    
    def replace_heading_number(match):
        bold_start = match.group(1)  # **
        number = match.group(2)      # e.g., "1.1.1" or "1"
        title = match.group(3)       # e.g., "Infeed Conveyors"
        bold_end = match.group(4)    # **
        
        # Transform the number: prepend section counter
        # "1" -> "6.1"
        # "1.1" -> "6.1.1"
        # "1.1.1" -> "6.1.1.1"
        new_number = f"{section_counter}.{number}"
        
        # Ensure proper trailing period
        if not new_number.endswith('.'):
            new_number += '.'
        
        return f"{bold_start}{new_number} {title}{bold_end}"
    
    # Pattern to match bold headings with numbers: **1. Title** or **1.1 Title** or **1.1.1 Title**
    # Handles both with and without trailing period after number
    heading_pattern = r'(\*\*)\s*(\d+(?:\.\d+)*\.?)\s+(.+?)\s*(\*\*)'
    
    transformed = re.sub(heading_pattern, replace_heading_number, text)
    
    return transformed


def add_chute_volume_calculation_table(doc, chute_type: str):
    """Add chute volume calculation table based on chute type - matching I/Os table styling"""
    chute_lower = chute_type.lower()
    
    # Define table data for each chute type
    chute_data = {}
    
    # Gravity Chutes
    if "gravity" in chute_lower and "mini" not in chute_lower:
        chute_data = {
            "headers": ["Average parcel size considered for chute volume calculation", "L(m)", "B(m)", "H(m)", "Volume (m3)"],
            "rows": [
                ["Gravity Chute", "0.4", "0.2", "0.3", "0.024"],
            ]
        }
    
    # Mini Gravity Chutes
    elif "mini" in chute_lower and "gravity" in chute_lower:
        chute_data = {
            "headers": ["Average parcel size considered for chute volume calculation", "L(m)", "B(m)", "H(m)", "Volume (m3)"],
            "rows": [
                ["Mini Gravity Chute", "0.4", "0.2", "0.3", "0.024"],
            ]
        }
    
    # Non-Sort Collection Chutes
    elif "non-sort" in chute_lower or "non sort" in chute_lower:
        chute_data = {
            "headers": ["Average shipment size considered for chute volume calculation", "L(m)", "B(m)", "H(m)", "Volume (m3)"],
            "rows": [
                ["Collection type PTL Chute", "0.3", "0.3", "0.3", "0.027"],
                ["PTL Chute", "6.7", "2", "0.5", ""],
            ]
        }
    
    # Rejection Chutes
    elif "rejection" in chute_lower:
        chute_data = {
            "headers": ["Average shipment size considered for chute volume calculation", "L(m)", "B(m)", "H(m)", "Volume (m3)"],
            "rows": [
                ["Collection type chute", "0.4", "0.4", "0.4", "0.064"],
                ["Rejection Chute", "5.6", "1", "0.3", "1.68"],
            ]
        }
    
    # Collection Chutes
    elif "collection" in chute_lower and "non-sort" not in chute_lower:
        chute_data = {
            "headers": ["Average shipment size considered for chute volume calculation", "L(m)", "B(m)", "H(m)", "Volume (m3)"],
            "rows": [
                ["Collection type chute", "0.4", "0.4", "0.4", "0.064"],
                ["Collection Chute", "5.6", "1", "0.3", "1.68"],
            ]
        }
    
    # If no data found, return without adding table
    if not chute_data:
        return
    
    # Create table - match I/Os table style
    num_rows = len(chute_data["rows"]) + 1  # +1 for header
    num_cols = len(chute_data["headers"])
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Set column widths to match I/Os tables (1.5" for first, 4.0" for rest, adjusted for more columns)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.5)
    for i in range(1, num_cols):
        table.columns[i].width = Inches(0.9)
    
    # Add header row - match I/Os table styling
    header_cells = table.rows[0].cells
    for col_idx, header_text in enumerate(chute_data["headers"]):
        header_cells[col_idx].text = header_text
        for paragraph in header_cells[col_idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)  # Match I/Os table header size
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(header_cells[col_idx], "4472C4")  # Blue header
    
    # Add data rows - match I/Os table styling
    for row_idx, row_data in enumerate(chute_data["rows"]):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)  # Match I/Os table data size
                apply_normal_style(paragraph)  # Use same style application as I/Os tables
    
    doc.add_paragraph("")  # Spacing after table


def add_chute_ios_table_table1(doc):

    """Add Chute I/Os Table 1 (for Gravity, Mini-Gravity, L-Type, Rejection, Dispersion, Direct Bagging, High Volume, Low Volume, Sliding+Secondary, Non-Sort, Collection chutes)"""
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Set column widths
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.0)
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Chute I/Os"
    header_cells[1].text = "Description"
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "4472C4")  # Blue header
    
    # Data rows
    data = [
        ("Chute Full Sensor", "To alert the system that the chute is full and mark that chute unavailable for further sortation"),
        ("Three Colour Tower Lamp", "To indicate the chute status"),
        ("Push Button", "To Start/Stop Sorting Operations"),
        ("HMI Display", "To indicate the operator action/shipment status"),
    ]
    
    for i, (ios_name, description) in enumerate(data):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = ios_name
        row_cells[1].text = description
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                apply_normal_style(paragraph)


def add_chute_ios_table_table2(doc):
    """Add Chute I/Os Table 2 (for Direct Takeout and Sliding chutes)"""
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Set column widths
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.0)
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Chute I/Os"
    header_cells[1].text = "Description"
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "4472C4")  # Blue header
    
    # Data rows
    data = [
        ("Chute Full Sensor", "To alert the system that the chute is full and mark that chute unavailable for further sortation"),
        ("Tower Lamp", "To indicate the chute status"),
    ]
    
    for i, (ios_name, description) in enumerate(data):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = ios_name
        row_cells[1].text = description
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                apply_normal_style(paragraph)


def try_add_chute_image_and_table(doc, chute_name: str):
    """Try to add chute image and I/Os table based on chute type"""
    base_dir = Path(__file__).parent
    fixed_image_dir = base_dir / "FIXED_IMAGE" / "Chutes"
    
    # Map chute types to image files and table type
    chute_mappings = {
        # Table 1 chutes
        "gravity chutes": (fixed_image_dir / "gravity.PNG", 1),
        "gravity": (fixed_image_dir / "gravity.PNG", 1),
        "Mini-gravity chutes": (fixed_image_dir / "mini-gravity.PNG", 1),
        "Mini gravity chutes": (fixed_image_dir / "mini-gravity.PNG", 1),
        "mini-gravity": (fixed_image_dir / "mini-gravity.PNG", 1),
        "l-type collection chutes": (fixed_image_dir / "l-type.PNG", 1),
        "l-type": (fixed_image_dir / "l-type.PNG", 1),
        "l type collection chutes": (fixed_image_dir / "l-type.PNG", 1),
        "rejection chutes": (fixed_image_dir / "rejection.PNG", 1),
        "rejection": (fixed_image_dir / "rejection.PNG", 1),
        "dispersion chutes": (fixed_image_dir / "dispersion.PNG", 1),
        "dispersion": (fixed_image_dir / "dispersion.PNG", 1),
        "direct bagging chutes": (fixed_image_dir / "direct-bagging.PNG", 1),
        "direct bagging": (fixed_image_dir / "direct-bagging.PNG", 1),
        "high volume chutes": (fixed_image_dir / "high_volume.PNG", 1),
        "high volume": (fixed_image_dir / "high_volume.PNG", 1),
        "low volume chutes": (fixed_image_dir / "low_volume.PNG", 1),
        "low volume": (fixed_image_dir / "low_volume.PNG", 1),
        "sliding + secondary chutes": (fixed_image_dir / "sliding_secondary.PNG", 1),
        "sliding secondary chutes": (fixed_image_dir / "sliding_secondary.PNG", 1),
        "sliding+secondary": (fixed_image_dir / "sliding_secondary.PNG", 1),
        "non-sort collection": (fixed_image_dir / "non_sort_collection.PNG", 1),
        "non sort collection": (fixed_image_dir / "non_sort_collection.PNG", 1),
        "collection chutes": (fixed_image_dir / "collection.PNG", 1),
        "collection": (fixed_image_dir / "collection.PNG", 1),
        
        # Table 2 chutes
        "direct takeout chutes": (fixed_image_dir / "direct-takeout.PNG", 2),
        "direct takeout": (fixed_image_dir / "direct-takeout.PNG", 2),
        "direct-takeout chutes": (fixed_image_dir / "direct-takeout.PNG", 2),
        "sliding chutes": (fixed_image_dir / "sliding.PNG", 2),
        "sliding": (fixed_image_dir / "sliding.PNG", 2),
    }
    
    # Normalize chute name for matching
    chute_name_lower = (chute_name or "").strip().lower()
    
    # Try exact and partial matches
    for key, (img_path, table_type) in chute_mappings.items():
        if chute_name_lower == key or key in chute_name_lower:
            # Add image if it exists
            if img_path.exists():
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(4.0))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("")
                except Exception as e:
                    print(f"Warning: Could not add chute image {img_path.name}: {e}")
            
            # Add appropriate table
            if table_type == 1:
                add_chute_ios_table_table1(doc)
            elif table_type == 2:
                add_chute_ios_table_table2(doc)
            
            doc.add_paragraph("")
            return True
    
    return False


def auto_insert_chute_markers(text: str) -> str:
    """
    Automatically insert [[CHUTE_COMPONENTS: ...]] markers after chute type headings.
    
    Detects chute type headings and inserts the marker after the description paragraph.
    """
    import re
    
    # Define chute types that should get components and tables
    chute_patterns = {
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Gravity\s+Chutes?)\s*\*\*': 'Gravity Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Mini[-\s]?Gravity\s+Chutes?)\s*\*\*': 'Mini-Gravity Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(L[-\s]?Type\s+(?:Collection\s+)?Chutes?)\s*\*\*': 'L-Type Collection Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Rejection\s+Chutes?)\s*\*\*': 'Rejection Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Dispersion\s+Chutes?)\s*\*\*': 'Dispersion Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Direct\s+Bagging\s+Chutes?)\s*\*\*': 'Direct Bagging Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(High\s+Volume\s+Chutes?)\s*\*\*': 'High Volume Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Low\s+Volume\s+Chutes?)\s*\*\*': 'Low Volume Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Sliding\s*[+&]\s*Secondary\s+Chutes?)\s*\*\*': 'Sliding + Secondary Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Non[-\s]?Sort\s+Collection)\s*\*\*': 'Non-Sort Collection',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Collection\s+Chutes?)\s*\*\*': 'Collection Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Direct\s+Take[-\s]?out\s+Chutes?)\s*\*\*': 'Direct Takeout Chutes',
        r'\*\*\s*(\d+(?:\.\d+)*\.?)\s*(Sliding\s+Chutes?)\s*\*\*': 'Sliding Chutes',
    }
    
    lines = text.splitlines()
    result_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        result_lines.append(line)
        
        # Check if this line matches any chute pattern
        matched_chute = None
        for pattern, chute_name in chute_patterns.items():
            if re.search(pattern, line, re.IGNORECASE):
                matched_chute = chute_name
                break
        
        if matched_chute:
            # Found a chute heading - skip ahead to find the end of the description
            # Then insert the marker
            i += 1
            description_lines = []
            
            # Collect description paragraphs (non-empty, non-heading lines)
            while i < len(lines):
                next_line = lines[i].strip()
                
                # Stop if we hit an empty line followed by another heading
                if not next_line:
                    # Peek ahead to see if next non-empty line is a heading
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and re.match(r'\*\*\s*\d+(?:\.\d+)*\.?\s*.+\*\*', lines[j]):
                        # Next heading found, insert marker before these empty lines
                        result_lines.extend(description_lines)
                        result_lines.append("")  # Empty line
                        result_lines.append(f"[[CHUTE_COMPONENTS: {matched_chute}]]")
                        result_lines.append("")  # Empty line after marker
                        break
                    else:
                        # Just an empty line in the description
                        description_lines.append(lines[i])
                        i += 1
                        continue
                
                # Stop if we hit another heading
                if re.match(r'\*\*\s*\d+(?:\.\d+)*\.?\s*.+\*\*', next_line):
                    # Insert marker before this heading
                    result_lines.extend(description_lines)
                    result_lines.append("")
                    result_lines.append(f"[[CHUTE_COMPONENTS: {matched_chute}]]")
                    result_lines.append("")
                    break
                
                # Stop if we hit an image placeholder
                if "image placeholder" in next_line.lower() or next_line.startswith("[["):
                    result_lines.extend(description_lines)
                    result_lines.append("")
                    result_lines.append(f"[[CHUTE_COMPONENTS: {matched_chute}]]")
                    result_lines.append("")
                    break
                
                # Add description line
                description_lines.append(lines[i])
                i += 1
                
                # Limit to ~5 lines of description
                if len(description_lines) >= 5:
                    result_lines.extend(description_lines)
                    result_lines.append("")
                    result_lines.append(f"[[CHUTE_COMPONENTS: {matched_chute}]]")
                    result_lines.append("")
                    break
            
            # If we've reached end of file, still add the marker
            if i >= len(lines) and description_lines:
                result_lines.extend(description_lines)
                result_lines.append("")
                result_lines.append(f"[[CHUTE_COMPONENTS: {matched_chute}]]")
                result_lines.append("")
        
        i += 1
    
    return "\n".join(result_lines)




def build_system_description_section(doc, counter, system_description_text, costing_file=None):
    """
    Build System Description as a separate top-level section.
    
    This section has hierarchical numbering:
    - 6. System Description (heading added externally)
    - 6.1. Infeed System
    - 6.1.1. Infeed Conveyors
    - 6.1.1.1. Aligning Conveyor
    etc.
    """
    if not system_description_text:
        return
    
    doc.add_page_break()
    add_numbered_heading(doc, "System Description", counter=counter)
    
    # Transform the internal numbering to match section counter
    transformed_text = transform_system_description_numbering(system_description_text, counter)
    
    # Auto-insert chute component markers after chute type headings
    transformed_text = auto_insert_chute_markers(transformed_text)
    
    lines = transformed_text.splitlines()
    table_inserted = False

    # Image replacements for specific sub-components (matching st_sys_desc FIXED_IMAGE_MAP)
    base_dir = Path(__file__).parent
    fixed_image_dir = base_dir / "FIXED_IMAGE"
    
    # Complete mapping matching st_sys_desc.py FIXED_IMAGE_MAP
    placeholder_images = {
        # Induct subcomponents
        "weighing conveyor": fixed_image_dir / "weigh_conv.PNG",
        "buffer conveyors": fixed_image_dir / "buffer_conv.PNG",
        "buffer conveyor": fixed_image_dir / "buffer_conv.PNG",
        "orientation / loading conveyor": fixed_image_dir / "oriant_conv.PNG",
        "orientation conveyor": fixed_image_dir / "oriant_conv.PNG",
        "orientation loading conveyor": fixed_image_dir / "oriant_conv.PNG",
        "loading conveyor": fixed_image_dir / "oriant_conv.PNG",
        "intelligent merge conveyor": fixed_image_dir / "merge_conv.PNG",
        "merge conveyor": fixed_image_dir / "merge_conv.PNG",
        "angle merge": fixed_image_dir / "merge_conv.PNG",
        
        # Main Loop / CBS
        "main loop": fixed_image_dir / "CROS_BELT_SORTER.PNG",
        "main linear cbs": fixed_image_dir / "CROS_BELT_SORTER.PNG",
        "cross belt sorter": fixed_image_dir / "CROS_BELT_SORTER.PNG",
        "cbs": fixed_image_dir / "CROS_BELT_SORTER.PNG",
        
        # Carrier/Drive system
        "carrier wheel drive": fixed_image_dir / "CAREER_WHEEL.PNG",
        "carrier": fixed_image_dir / "CBS_CAREER.PNG",
        "friction wheel drive": fixed_image_dir / "FRINCTION_WHEEL_DRIVE.PNG",
        "linear motor drive": fixed_image_dir / "LINEAR_MOTOR_DRIVE.PNG",
        "servo roller": fixed_image_dir / "SERVO_ROLLER.PNG",
        
        # Infeed/Conveyors
        "infeed conveyor": fixed_image_dir / "buffer_conv.PNG",
        "infeed conveyors": fixed_image_dir / "buffer_conv.PNG",
        "straight and inclined conveyor": fixed_image_dir / "buffer_conv.PNG",
        "curve conveyor": fixed_image_dir / "buffer_conv.PNG",
    }

    def _insert_image_paragraph(paragraph, img_path: Path, width_in: float = 5.5):
        """Replace paragraph text with the given image."""
        paragraph.text = ""
        run = paragraph.add_run()
        run.add_picture(str(img_path), width=Inches(width_in))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _try_match_fixed_image(heading_text: str) -> Optional[Path]:
        """Tolerant matching for image placeholders - mirrors st_sys_desc._try_add_fixed_image logic."""
        h = (heading_text or "").strip().lower()
        
        # Direct match first
        if h in placeholder_images:
            img = placeholder_images[h]
            if img.exists():
                return img
        
        # Tolerant matching for induct subcomponents
        if "weigh" in h:
            img = placeholder_images.get("weighing conveyor")
            if img and img.exists():
                return img
        if "buffer" in h:
            img = placeholder_images.get("buffer conveyor")
            if img and img.exists():
                return img
        if "orientation" in h or "loading" in h or "oriant" in h:
            img = placeholder_images.get("orientation conveyor")
            if img and img.exists():
                return img
        if "merge" in h or "angle" in h:
            img = placeholder_images.get("merge conveyor")
            if img and img.exists():
                return img
        
        # Main Loop / CBS matching
        if "main loop" in h or "main linear" in h or "cross belt" in h or "cbs" in h:
            img = placeholder_images.get("main loop")
            if img and img.exists():
                return img
        
        # Carrier/Drive matching
        if "carrier" in h:
            img = placeholder_images.get("carrier")
            if img and img.exists():
                return img
        if "friction" in h and "drive" in h:
            img = placeholder_images.get("friction wheel drive")
            if img and img.exists():
                return img
        if "linear motor" in h:
            img = placeholder_images.get("linear motor drive")
            if img and img.exists():
                return img
        if "servo" in h:
            img = placeholder_images.get("servo roller")
            if img and img.exists():
                return img
        
        # Infeed conveyor matching
        if "infeed" in h or "straight" in h or "inclined" in h:
            img = placeholder_images.get("infeed conveyor")
            if img and img.exists():
                return img
        if "curve" in h:
            img = placeholder_images.get("curve conveyor")
            if img and img.exists():
                return img
        
        return None
    
    # Regex patterns for formatting
    import re as _re_fmt
    bold_rx = _re_fmt.compile(r"\*\*(.+?)\*\*")
    heading_rx = _re_fmt.compile(r"^\s*\*\*\s*(\d+(?:\.\d+)*\.?)\s*(.+?)\s*\*\*\s*$")
    bullet_rx = _re_fmt.compile(r"^\s*[-*•]\s+(.+)$")
    
    # Track pending chute type for image/table insertion
    pending_chute_type = None
    
    def _add_formatted_paragraph(doc, text_line: str):
        """Add a paragraph with proper bold/bullet formatting."""
        stripped = text_line.strip()
        
        # Handle heading lines with **bold** markers
        heading_match = heading_rx.match(stripped)
        if heading_match:
            number = heading_match.group(1).strip()
            heading_text = heading_match.group(2).strip()
            p = doc.add_paragraph()
            run = p.add_run(f"{number} {heading_text}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Calibri'
            return heading_text  # Return heading text for chute detection
        
        # Handle bullet points
        bullet_match = bullet_rx.match(stripped)
        if bullet_match:
            bullet_text = bullet_match.group(1).strip()
            p = doc.add_paragraph(style='List Bullet')
            if "**" in bullet_text:
                parts = bold_rx.split(bullet_text)
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                p.add_run(bullet_text)
            apply_normal_style(p)
            return None
        
        # Handle regular text with possible **bold** markers
        if "**" in stripped:
            p = doc.add_paragraph()
            parts = bold_rx.split(stripped)
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
            apply_normal_style(p)
        else:
            p = doc.add_paragraph(text_line)
            apply_normal_style(p)
        
        return None
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Skip empty lines
        if not line_stripped:
            doc.add_paragraph("")
            continue
        
        # **KEY FIX: Check if this line contains the Conveyor BOQ table marker**
        if '[[CONVEYOR_BOQ_TABLE]]' in line:
            if costing_file:
                table_data = extract_conveyor_boq_from_excel(costing_file)
                if table_data and len(table_data) >= 2:
                    rows = len(table_data)
                    cols = len(table_data[0])
                    table = doc.add_table(rows=rows, cols=cols)
                    apply_table_style(table)

                    for r_idx, row_data in enumerate(table_data):
                        row_cells = table.rows[r_idx].cells
                        for c_idx, cell_value in enumerate(row_data):
                            row_cells[c_idx].text = str(cell_value)

                            for paragraph in row_cells[c_idx].paragraphs:
                                if r_idx == 0:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.name = 'Calibri'
                                        run.font.size = Pt(11)
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                else:
                                    apply_normal_style(paragraph)

                    table.columns[0].width = Inches(0.5)
                    table.columns[1].width = Inches(2.0)
                    table.columns[2].width = Inches(0.7)
                    table.columns[3].width = Inches(0.7)
                    table.columns[4].width = Inches(1.2)
                    table.columns[5].width = Inches(1.2)
                    table.columns[6].width = Inches(0.5)

                    doc.add_paragraph("")
                   
                    table_inserted = True
            # Skip writing the marker line itself
            continue

        # Handle specific image placeholders for sub-components
        if "image placeholder" in line_stripped.lower():
            # Extract the heading from [IMAGE PLACEHOLDER: <heading>]
            import re as _re
            ph_rx = _re.compile(r"^\s*\[IMAGE PLACEHOLDER:\s*(.+?)\s*\]\s*$", flags=_re.IGNORECASE)
            m = ph_rx.match(line_stripped)
            if m:
                heading = m.group(1).strip()
                matched_path = _try_match_fixed_image(heading)
            else:
                # Fallback: use entire line for matching
                matched_path = _try_match_fixed_image(line_stripped)

            if matched_path and matched_path.exists():
                try:
                    p = doc.add_paragraph()
                    _insert_image_paragraph(p, matched_path)
                    doc.add_paragraph("")
                except Exception as e:
                    st.warning(f"Image insert failed for {matched_path.name}: {e}")
                    # Don't add the placeholder text if image insertion failed
            else:
                # No matching image found - skip the placeholder line (don't show it in final doc)
                pass
            continue
        
        # Check for special marker to insert chute components list + table
        # Format: [[CHUTE_COMPONENTS: Chute Type Name]]
        if line_stripped.startswith("[[CHUTE_COMPONENTS:") and line_stripped.endswith("]]"):
            import re as _marker_re
            marker_match = _marker_re.match(r"\[\[CHUTE_COMPONENTS:\s*(.+?)\s*\]\]", line_stripped)
            if marker_match:
                chute_type_name = marker_match.group(1).strip()
                
                # Step 1: Add chute image first
                base_dir = Path(__file__).parent
                fixed_image_dir = base_dir / "FIXED_IMAGE" / "Chutes"
                chute_mappings = {
                    # Table 1 chutes
                    "gravity chutes": (fixed_image_dir / "gravity.PNG", 1),
                    "gravity": (fixed_image_dir / "gravity.PNG", 1),
                    "mini-gravity chutes": (fixed_image_dir / "mini-gravity.PNG", 1),
                    "mini gravity chutes": (fixed_image_dir / "mini-gravity.PNG", 1),
                    "mini-gravity": (fixed_image_dir / "mini-gravity.PNG", 1),
                    "l-type collection chutes": (fixed_image_dir / "l-type.PNG", 1),
                    "l-type": (fixed_image_dir / "l-type.PNG", 1),
                    "l type collection chutes": (fixed_image_dir / "l-type.PNG", 1),
                    "rejection chutes": (fixed_image_dir / "rejection.PNG", 1),
                    "rejection": (fixed_image_dir / "rejection.PNG", 1),
                    "dispersion chutes": (fixed_image_dir / "dispersion.PNG", 1),
                    "dispersion": (fixed_image_dir / "dispersion.PNG", 1),
                    "direct bagging chutes": (fixed_image_dir / "direct-bagging.PNG", 1),
                    "direct bagging": (fixed_image_dir / "direct-bagging.PNG", 1),
                    "high volume chutes": (fixed_image_dir / "high_volume.PNG", 1),
                    "high volume": (fixed_image_dir / "high_volume.PNG", 1),
                    "low volume chutes": (fixed_image_dir / "low_volume.PNG", 1),
                    "low volume": (fixed_image_dir / "low_volume.PNG", 1),
                    "sliding + secondary chutes": (fixed_image_dir / "sliding_secondary.PNG", 1),
                    "sliding secondary chutes": (fixed_image_dir / "sliding_secondary.PNG", 1),
                    "sliding+secondary": (fixed_image_dir / "sliding_secondary.PNG", 1),
                    "non-sort collection": (fixed_image_dir / "non_sort_collection.PNG", 1),
                    "non sort collection": (fixed_image_dir / "non_sort_collection.PNG", 1),
                    "collection chutes": (fixed_image_dir / "collection.PNG", 1),
                    "collection": (fixed_image_dir / "collection.PNG", 1),
                    # Table 2 chutes
                    "direct takeout chutes": (fixed_image_dir / "direct-takeout.PNG", 2),
                    "direct takeout": (fixed_image_dir / "direct-takeout.PNG", 2),
                    "direct-takeout chutes": (fixed_image_dir / "direct-takeout.PNG", 2),
                    "sliding chutes": (fixed_image_dir / "sliding.PNG", 2),
                    "sliding": (fixed_image_dir / "sliding.PNG", 2),
                }
                
                chute_name_lower = chute_type_name.strip().lower()
                img_path = None
                table_type = 1  # Default to table 1
                
                # Find matching chute type
                for key, (path, ttype) in chute_mappings.items():
                    if chute_name_lower == key or key in chute_name_lower:
                        img_path = path
                        table_type = ttype
                        break
                
                # Add image if found
                if img_path and img_path.exists():
                    try:
                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(str(img_path), width=Inches(4.0))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph("")
                    except Exception as e:
                        st.warning(f"Could not add chute image {img_path.name}: {e}")
                
                # Step 1.5: Add volume calculation table for specific chute types
                volume_chute_types = ["gravity chutes", "gravity", "mini-gravity chutes", "mini gravity chutes", 
                                     "mini-gravity", "non-sort collection", "non sort collection", 
                                     "rejection chutes", "rejection", "collection chutes", "collection"]
                if any(chute_name_lower == vtype or vtype in chute_name_lower for vtype in volume_chute_types):
                    try:
                        add_chute_volume_calculation_table(doc, chute_type_name)
                    except Exception as e:
                        st.warning(f"Could not add volume table for {chute_type_name}: {e}")
                
                # Step 2: Add "Each Chute is equipped..." intro
                p = doc.add_paragraph("Each Chute is equipped with below components to ease the operations:")
                apply_normal_style(p)
                
                # Step 3: Add component bullet list
                # Determine which components list based on chute type
                chute_type_lower = chute_type_name.lower()
                if "direct takeout" in chute_type_lower or ("sliding" in chute_type_lower and "secondary" not in chute_type_lower):
                    # Table 2 components (simpler)
                    components = [
                        "Chute Full Sensor",
                        "Tower Light",
                    ]
                else:
                    # Table 1 components (full set)
                    components = [
                        "Chute Full Sensor",
                        "Tower Light",
                        "Push Button",
                    ]
                
                for comp in components:
                    p = doc.add_paragraph(comp, style='List Bullet')
                    apply_normal_style(p)
                
                doc.add_paragraph("")  # Spacing
                
                # Step 4: Add appropriate table
                if table_type == 1:
                    add_chute_ios_table_table1(doc)
                elif table_type == 2:
                    add_chute_ios_table_table2(doc)
                
                doc.add_paragraph("")
                
            continue
        
        # Regular text line - insert as formatted paragraph
        heading_text_detected = _add_formatted_paragraph(doc, line)

    # Fallback: if marker was missing but we have a costing file, try heading-based insertion
    if not table_inserted and costing_file:
        insert_conveyor_boq_table(doc, costing_file)

    # Final pass: replace any remaining image placeholders that may exist
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        if "image placeholder" in text_lower:
            # Use tolerant matching function
            matched_path = _try_match_fixed_image(para.text)
            if matched_path and matched_path.exists():
                try:
                    _insert_image_paragraph(para, matched_path)
                except Exception as e:
                    st.warning(f"Image insert failed for {matched_path.name}: {e}")
            else:
                # Remove unmatched placeholder text from final document
                para.text = ""


def extract_conveyor_boq_from_excel(costing_file) -> List[List[str]]:
    """Extract Conveyor BOQ table from Conveyors sheet - 7 columns"""
    try:
        # Read the Conveyors sheet
        df = pd.read_excel(costing_file, sheet_name="Conveyors", header=None)
        
        # Find the header row (contains "S No.", "Name", etc.)
        header_row_idx = None
        for idx in range(min(10, len(df))):  # Search first 10 rows
            row_values = df.iloc[idx].values
            row_str = ' '.join([str(v).lower() for v in row_values if pd.notna(v)])
            if 's no' in row_str and 'name' in row_str and 'conveyor length' in row_str:
                header_row_idx = idx
                break
        
        if header_row_idx is None:
            st.warning("Could not find header row in Conveyors sheet")
            return []
        
        # Re-read with correct header
        df = pd.read_excel(costing_file, sheet_name="Conveyors" or "Conveyors " or "conveyors", header=header_row_idx)
        
        # Expected columns (in order)
        required_cols = ['S No.', 'Name', 'EL_1', 'EL_2', 'Conveyor Length (m)', 'Conveyor width (mm)', 'Set']
        
        # Find actual column names (flexible matching)
        col_mapping = {}
        for req_col in required_cols:
            req_normalized = req_col.lower().replace(' ', '').replace('_', '').replace('.', '').replace('(', '').replace(')', '')
            
            for actual_col in df.columns:
                actual_normalized = str(actual_col).lower().replace(' ', '').replace('_', '').replace('.', '').replace('(', '').replace(')', '')
                
                # Match by normalized comparison
                if req_normalized in actual_normalized or actual_normalized in req_normalized:
                    col_mapping[req_col] = actual_col
                    break
        
        if len(col_mapping) < 6:  # At least 6 out of 7 columns needed
            st.warning(f"Could not find all required columns. Found: {list(col_mapping.keys())}")
            return []
        
        # Extract data with actual column names
        actual_cols = [col_mapping.get(req_col) for req_col in required_cols if req_col in col_mapping]
        selected_df = df[actual_cols].copy()
        
        # Rename to standard names
        rename_dict = {col_mapping[req_col]: req_col for req_col in required_cols if req_col in col_mapping}
        selected_df = selected_df.rename(columns=rename_dict)
        
        # Remove rows where S No. is empty or NaN
        if 'S No.' in selected_df.columns:
            selected_df = selected_df[selected_df['S No.'].notna()]
            # Also remove rows where S No. is empty string
            selected_df = selected_df[selected_df['S No.'].astype(str).str.strip() != '']
        
        # Convert to list of lists (header + data rows)
        table_data = [required_cols]  # Header row
        
        for _, row in selected_df.iterrows():
            row_data = []
            for col in required_cols:
                if col in row.index:
                    val = row[col]
                    # Format numbers appropriately
                    if pd.notna(val):
                        if isinstance(val, (int, float)):
                            # Keep numbers as they are, but remove unnecessary decimals
                            if isinstance(val, float) and val.is_integer():
                                row_data.append(str(int(val)))
                            else:
                                row_data.append(str(val))
                        else:
                            row_data.append(str(val))
                    else:
                        row_data.append('')
                else:
                    row_data.append('')
            
            table_data.append(row_data)
        
        return table_data
        
    except Exception as e:
        st.warning(f"Could not extract Conveyor BOQ: {e}")
        import traceback
        st.code(traceback.format_exc())
        return []


def insert_conveyor_boq_table_after_heading(doc: Document, costing_file, heading_para):
    """Insert Conveyor BOQ table immediately after a specific heading paragraph"""
    try:
        # Extract table data
        table_data = extract_conveyor_boq_from_excel(costing_file)
        
        if not table_data or len(table_data) < 2:
            st.warning("No Conveyor BOQ data found to insert")
            return
        
        # Create table
        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        apply_table_style(table)
        
        # Fill table with data
        for r_idx, row_data in enumerate(table_data):
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_value in enumerate(row_data):
                row_cells[c_idx].text = str(cell_value)
                
                # Format cells
                for paragraph in row_cells[c_idx].paragraphs:
                    if r_idx == 0:  # Header row
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    else:  # Data rows
                        apply_normal_style(paragraph)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Set column widths (adjust as needed)
        table.columns[0].width = Inches(0.5)   # S No.
        table.columns[1].width = Inches(2.0)   # Name
        table.columns[2].width = Inches(0.7)   # EL_1
        table.columns[3].width = Inches(0.7)   # EL_2
        table.columns[4].width = Inches(1.2)   # Conveyor Length
        table.columns[5].width = Inches(1.2)   # Conveyor width
        table.columns[6].width = Inches(0.5)   # Set
        
        # Insert table after the heading using XML manipulation
        table_element = table._element
        heading_element = heading_para._element
        heading_element.addnext(table_element)
        
        # Add a blank paragraph after table for spacing
        blank_para_element = heading_element.makeelement(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p',
            nsmap=heading_element.nsmap
        )
        table_element.addnext(blank_para_element)
        
        st.success("✅ Conveyor BOQ table inserted successfully after heading")
        
    except Exception as e:
        st.warning(f"Could not insert Conveyor BOQ table: {e}")
        import traceback
        st.code(traceback.format_exc())


def insert_conveyor_boq_table(doc: Document, costing_file):
    """Find 'Conveyor BOQ' heading in doc and insert table after it (fallback method)"""
    try:
        # Extract table data
        table_data = extract_conveyor_boq_from_excel(costing_file)
        
        if not table_data or len(table_data) < 2:
            st.warning("No Conveyor BOQ data found")
            return
        
        # Find the paragraph containing "Conveyor BOQ" (case-insensitive)
        target_para = None
        for para in doc.paragraphs:
            if 'conveyor boq' in para.text.lower() and len(para.text.strip()) < 50:
                target_para = para
                break
        
        if target_para is None:
            st.warning("Could not find 'Conveyor BOQ' heading in document")
            return
        
        # Create table
        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        apply_table_style(table)
        
        # Fill table
        for r_idx, row_data in enumerate(table_data):
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_value in enumerate(row_data):
                row_cells[c_idx].text = str(cell_value)
                
                # Format cells
                for paragraph in row_cells[c_idx].paragraphs:
                    if r_idx == 0:  # Header row
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        apply_normal_style(paragraph)
        
        # Set column widths
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(2.0)
        table.columns[2].width = Inches(0.7)
        table.columns[3].width = Inches(0.7)
        table.columns[4].width = Inches(1.2)
        table.columns[5].width = Inches(1.2)
        table.columns[6].width = Inches(0.5)
        
        # Move table to correct position
        table_element = table._element
        target_element = target_para._element
        target_element.addnext(table_element)
        
        st.success("✅ Conveyor BOQ table inserted successfully")
        
    except Exception as e:
        st.warning(f"Could not insert Conveyor BOQ table: {e}")

def generate_system_description_from_sd_sys(dxf_path: Path, costing_file_upload, project_name: str,
                                           client_name: str = "", pph_count: str = "", ipp_rate: str = "",
                                           facts: Optional[ProposalFacts] = None,
                                           context: Optional[ProposalContext] = None) -> tuple[str, dict, dict, dict]:
    """Generate System Description using `st_sys_desc` utilities.

    Returns: (system_description_text, detected_json, tables_dict, cost_vals_dict)
    """
    # Build DXF JSON using st_sys_desc extractor (compute_dxf_metrics expects this format)
    full = {}
    try:
        full = sd_sys.extract_dxf_full_json(dxf_path)
    except Exception as e:
        logger.warning(f"sd_sys.extract_dxf_full_json failed: {e}; trying fallback")
        # fallback: try main app's extractor (may produce incompatible format)
        try:
            full = extract_dxf_components(dxf_path, project_name)
        except Exception:
            full = {}

    # Compute metrics (prefer sd_sys implementation when available)
    try:
        metrics = sd_sys.compute_dxf_metrics(full)
    except Exception as e:
        logger.warning(f"sd_sys.compute_dxf_metrics failed: {e}; falling back to empty metrics")
        metrics = {}

    # If ProposalFacts already has metrics, prefer them (avoids re-computation)
    if facts and facts.dxf_metrics:
        metrics = facts.dxf_metrics

    # Extract costing tables/values if provided
    tables = {}
    cost_vals = {}
    if costing_file_upload:
        try:
            if hasattr(costing_file_upload, "getvalue"):
                data = costing_file_upload.getvalue()
            elif isinstance(costing_file_upload, (bytes, bytearray)):
                data = bytes(costing_file_upload)
            else:
                data = costing_file_upload.read()

            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tx:
                tx.write(data)
                tmp_xlsx = Path(tx.name)

            try:
                tables = sd_sys.extract_costing_tables(tmp_xlsx)
            except Exception as e:
                tables = {}
                st.warning(f"Costing tables extraction failed: {e}")
            try:
                cost_vals = sd_sys.extract_costing_values(tmp_xlsx)
            except Exception as e:
                cost_vals = {}
                st.warning(f"Costing values extraction failed: {e}")
        except Exception as e:
            tables = {}
            cost_vals = {}
            st.warning(f"Costing file processing failed: {e}")

    # Filter Conveyor BOQ to ONLY requested columns (matching st_sys_desc behavior)
    if "Conveyor BOQ" in tables and tables["Conveyor BOQ"]:
        tables["Conveyor BOQ"] = sd_sys.filter_conveyor_boq_columns(tables["Conveyor BOQ"])

    # Build variables map using st_sys_desc function
    variables = sd_sys.build_variables_map(metrics, cost_vals)
    # Add additional client-specific variables
    variables["CLIENT'S NAME"] = (client_name or "").strip()
    variables["CLIENT NAME"] = variables["CLIENT'S NAME"]
    variables["PPH COUNT"] = (pph_count or "").strip()
    variables["IPP Rate"] = (ipp_rate or "").strip()
    variables["IPP RATE"] = (ipp_rate or "").strip()

    # Prefer context for counts: do not recompute quantities; ensure safe phrasing for missing/zero
    if context:
        def _ctx_val(key):
            m = context.get(key)
            return m.value
        variables["Feedline Count"] = _ctx_val("feedlines") or None
        total_chutes = sum([(context.get(k).value or 0) for k in [
            "gravity_chutes","mini_gravity_chutes","collection_chutes","rejection_chutes","dispersion_chutes","bulk_chutes","direct_bagging_chutes"
        ]])
        variables["Total Chutes"] = total_chutes or None
        variables["Telescopic Conveyors"] = _ctx_val("telescopic_belt_conveyors") or None
        variables["PPH COUNT"] = str(context.pph) if context.pph is not None else ""

    # GROQ detection
    det_txt = sd_sys.groq_chat(sd_sys.prompt_detect_components(full, metrics), temperature=0.1, max_tokens=2500)
    det_raw = sd_sys.extract_json(det_txt)
    detected_fallback = sd_sys.build_detected_json(metrics)
    detected = sd_sys.normalize_detected(det_raw, detected_fallback)

    # Enforce: infeed subcomponents (ONLY if found in DXF)
    detected = sd_sys.enforce_infeed_subcomponents_from_dxf(detected, metrics, full)

    # Enforce: induct subcomponents (ONLY if found in DXF) + correct names + buffer classification
    detected = sd_sys.enforce_induct_subcomponents_from_dxf(detected, metrics, full)

    # Enforce mandatory mappings (additional fallbacks)
    if metrics.get("FS002 WITHOUT WEIGHING COUNT", 0) > 0:
        detected.setdefault("Parcel Inducts / Induction to Sorter", {})
        detected["Parcel Inducts / Induction to Sorter"].setdefault("Feedlines", {"Feedline Count": int(metrics.get("FEEDLINE COUNT", 1)), "Subcomponents": {}})
        detected["Parcel Inducts / Induction to Sorter"]["Feedlines"].setdefault("Subcomponents", {})
        detected["Parcel Inducts / Induction to Sorter"]["Feedlines"]["Subcomponents"]["Buffer Conveyor"] = 1

    if metrics.get("HAS_VDS_LOOP"):
        detected.setdefault("Infeed System", {})
        if "VDS Loop Conveyor" not in detected["Infeed System"]:
            detected["Infeed System"]["VDS Loop Conveyor"] = max(1, int(metrics.get("VDS LOOP COUNT", 1)))

    # Override detected counts using ProposalContext (authoritative source of truth)
    # Only overwrite when context has confirmed numeric values - do NOT invent numbers
    if context:
        # Feedlines count
        ctx_feedlines = context.get("feedlines")
        if ctx_feedlines.value is not None:
            detected.setdefault("Parcel Inducts / Induction to Sorter", {})
            detected["Parcel Inducts / Induction to Sorter"].setdefault("Feedlines", {})
            detected["Parcel Inducts / Induction to Sorter"]["Feedlines"]["Feedline Count"] = ctx_feedlines.value
            logger.info(f"Override detected Feedline Count from context: {ctx_feedlines.value}")
        
        # Total Chutes: sum all chute types from context
        chute_keys = ["gravity_chutes", "mini_gravity_chutes", "collection_chutes", 
                      "rejection_chutes", "dispersion_chutes", "bulk_chutes", "direct_bagging_chutes"]
        ctx_total_chutes = sum([(context.get(k).value or 0) for k in chute_keys])
        if ctx_total_chutes > 0:
            detected["Total Chutes"] = ctx_total_chutes
            logger.info(f"Override detected Total Chutes from context: {ctx_total_chutes}")
        
        # System Throughput (PPH)
        ctx_pph = context.get("throughput_pph")
        if ctx_pph.value is not None:
            detected["System Throughput (PPH)"] = ctx_pph.value
            logger.info(f"Override detected System Throughput from context: {ctx_pph.value}")

    # Generate draft and run judge pass
    template_text = sd_sys.load_template_text() if hasattr(sd_sys, "load_template_text") else ""
    # Inject HARD CONSTRAINTS counts block (prefer context, then facts) to guide generation
    try:
        if context:
            constraints = context.counts_block_text() + "\n\n"
            template_text = constraints + template_text
            gating_issues = validate_context_counts(context) if ENABLE_CONTEXT_UNIFICATION else []
            if gating_issues:
                template_text = template_text + "\n\nIf quantities are not specified, avoid inventing counts; use safe qualitative phrasing."
                logger.warning(f"Pre-generation validation gate (System Description) triggered: {gating_issues}")
            ctx_counts = context.counts_block_text().replace('\n',' | ')
            logger.info(f"Using ProposalContext counts for System Description: {ctx_counts}")
        elif facts:
            from proposal_context import build_proposal_context
            ctx = build_proposal_context(facts=facts, costing_metrics=None, dxf_json={})
            constraints = ctx.counts_block_text() + "\n\n"
            template_text = constraints + template_text
            ctx_counts = ctx.counts_block_text().replace('\n',' | ')
            logger.info(f"Using ProposalContext counts for System Description: {ctx_counts}")
    except Exception as e:
        logger.warning(f"Constraint injection failed: {e}")
    sys_desc_draft = sd_sys.groq_chat(sd_sys.prompt_generate_system_description_dynamic(template_text, detected, variables), temperature=0.2, max_tokens=4500)
    sys_desc_final = sd_sys.groq_chat(sd_sys.prompt_judge_fix_system_description(detected, sys_desc_draft), temperature=0.0, max_tokens=4500)

    return sys_desc_final, detected, tables, cost_vals
            

# Concept Description helper: insert the generated flowchart PNG into the main DOCX
def build_concept_description_section(doc, counter, png_bytes: bytes):
    """Add a Concept Description section containing the rendered flowchart PNG.

    This function inserts a numbered heading and places the PNG centered on the page.
    """
    try:
        doc.add_page_break()
    except Exception:
        pass

    add_numbered_heading(doc, "Concept Description", counter=counter)
    doc.add_paragraph()

    # Subheading for the flowchart
    add_numbered_subheading(doc, "Flowchart", f"{counter}.1")

    if not png_bytes:
        p = doc.add_paragraph("Flowchart image not available.")
        apply_normal_style(p)
        return

    try:
        p = doc.add_paragraph()
        run = p.add_run()
        image_stream = BytesIO(png_bytes)
        run.add_picture(image_stream, width=Inches(6.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Small caption/note
        note = doc.add_paragraph("Click the diagram to open in draw.io for editing")
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_normal_style(note)
    except Exception as e:
        p = doc.add_paragraph(f"Could not insert flowchart image: {e}")
        apply_normal_style(p)

# ==================== MAIN GENERATION LOGIC ====================

# Only process generation if we're on input page and button was clicked
if st.session_state.get("page", "input") == "input" and 'generate_clicked' in dir() and generate_clicked:
    # Update status
    st.session_state.current_status = "Generating..."
    
    # Progress tracking - spinner container first, then progress bar
    spinner_container = st.empty()
    progress_bar = st.progress(0, text="Initializing document generation...")
    
    # Initialize all section text variables to None (will be populated conditionally)
    process_flow_text = None
    exec_summary_text = None
    system_description_text = None
    cover_letter_text = None
    flowchart_png_bytes = None
    dxf_json = None
    facts = None
    context = None
    
    with spinner_container.container():
        with st.spinner("Processing..."):
            try:
                # Validate required inputs
                progress_bar.progress(5, text="Validating inputs...")
                if not client_name.strip():
                    st.error("Please enter a Client Name")
                    st.stop()
                
                if not project_name.strip():
                    st.error("Please enter a Project Name")
                    st.stop()
                
                # Create temp directory and write files
                tmp_dir = Path(tempfile.mkdtemp(prefix="proposal_"))
                
                # Save DXF file to temp
                if dxf_layout_file:
                    dxf_path = tmp_dir / dxf_layout_file.name
                    dxf_path.write_bytes(dxf_layout_file.getvalue())
                else:
                    st.error("Please upload a DXF layout file")
                    st.stop()
                
                # Save costing file to temp if provided
                costing_tmp_path = None
                if costing_file:
                    costing_tmp_path = tmp_dir / costing_file.name
                    costing_tmp_path.write_bytes(costing_file.getvalue())
                
                # Do not parse process flow; rely purely on ProposalContext
                process_flow_summary = ""
                
                # Extract ProposalFacts once (DXF + costing) and store for downstream generators
                try:
                    progress_bar.progress(12, text="Extracting proposal facts...")
                    facts = extract_proposal_facts(
                        dxf_path=str(dxf_path),
                        costing_path=str(costing_tmp_path) if costing_tmp_path else None,
                        project_name=project_name,
                        client_name=client_name,
                        verbose=False,
                    )
                except Exception as e:
                    facts = None
                    facts_error = str(e)
                
                # Fallback: populate facts.costing_metrics from sd_sys if not present
                if facts and (not facts.costing_metrics or not any(facts.costing_metrics.values())) and costing_tmp_path:
                    try:
                        raw_cost_vals = sd_sys.extract_costing_values(costing_tmp_path)
                        # Normalize keys into facts.costing_metrics format
                        normalized_costing = {}
                        if raw_cost_vals:
                            normalized_costing["feedlines"] = raw_cost_vals.get("FEEDLINE COUNT") or raw_cost_vals.get("feedline_count")
                            normalized_costing["total_chutes"] = raw_cost_vals.get("TOTAL CHUTES") or raw_cost_vals.get("total_chutes")
                            normalized_costing["throughput_pph"] = raw_cost_vals.get("THROUGHPUT") or raw_cost_vals.get("throughput") or raw_cost_vals.get("PPH")
                            normalized_costing["sorter_speed_mps"] = raw_cost_vals.get("SORTER SPEED") or raw_cost_vals.get("sorter_speed")
                            # Remove None values
                            normalized_costing = {k: v for k, v in normalized_costing.items() if v is not None}
                            if normalized_costing:
                                facts.costing_metrics = normalized_costing
                                logger.info(f"Populated facts.costing_metrics from sd_sys fallback: {normalized_costing}")
                    except Exception as e:
                        logger.warning(f"sd_sys costing extraction fallback failed: {e}")
                
                # Extract DXF components
                progress_bar.progress(15, text="Analyzing layout blueprint...")
                dxf_json = extract_dxf_components(dxf_path, project_name)
                
                # Optional: enrich DXF with alias map for unknown blocks
                if ENABLE_ALIAS_MAP and dxf_json:
                    try:
                        from alias_map import analyze_and_update_aliases
                        alias_counts = analyze_and_update_aliases(dxf_json)
                        if alias_counts:
                            dxf_json.setdefault("alias_counts", {}).update(alias_counts)
                            logger.info(f"Alias map updated; counts: {alias_counts}")
                    except Exception as e:
                        logger.warning(f"Alias map enrichment failed: {e}")
                
                # Build unified ProposalContext (facts-first, then costing, then DXF heuristics)
                # Always build context to ensure it's available for audit (not conditional on flag)
                try:
                    # Ensure costing_metrics includes the PPH from slider
                    costing_metrics_with_pph = {}
                    if facts and facts.costing_metrics:
                        costing_metrics_with_pph = facts.costing_metrics.copy()
                    
                    # Add PPH from slider if available
                    if pph_count and str(pph_count).strip():
                        try:
                            pph_val = int(str(pph_count).strip())
                            costing_metrics_with_pph["throughput_pph"] = pph_val
                            logger.info(f"Added PPH from slider to costing_metrics: {pph_val}")
                        except ValueError:
                            logger.warning(f"Could not parse pph_count as integer: {pph_count}")
                    
                    context = build_proposal_context(
                        facts=facts,
                        costing_metrics=costing_metrics_with_pph,
                        dxf_json=dxf_json
                    )
                    st.session_state.proposal_context = context
                    logger.info("ProposalContext created and stored in session_state")
                    # Log normalized counts and validate pre-generation
                    try:
                        log_context_counts(context)
                        issues = validate_context_counts(context)
                        if issues:
                            st.session_state.pre_gen_issues = issues
                            logger.warning(f"Pre-generation validation issues: {issues}")
                        else:
                            st.session_state.pre_gen_issues = []
                    except Exception as e:
                        logger.warning(f"Context logging/validation failed: {e}")
                except Exception as e:
                    logger.warning(f"Failed to build ProposalContext: {e}")
                
                # Generate Process Flow if Proposed System is included
                if include_proposed_system:
                    # If a preview was already generated earlier in the session, reuse it
                    # Always generate fresh process flow using agentY agentic solution
                    try:
                        progress_bar.progress(25, text="Crafting intelligent process flow...")
                        
                        # Get manual components context if available
                        manual_components_ctx = st.session_state.get("manual_components_context", None)
                        
                        process_flow_text, iteration_details = call_groq_for_process_flow(
                            client_name,
                            project_name,
                            dxf_json,
                            facts=facts,
                            context=context,
                            manual_components_context=manual_components_ctx,
                        )
                        # Store for feedback regeneration
                        st.session_state.section_content_proposed_system = process_flow_text
                        st.session_state.section_original_proposed_system = process_flow_text  # Store original for comparison
                        progress_bar.progress(35, text="Process flow ready...")
                        time.sleep(1)  # Brief delay to avoid rate limits
                    except Exception as e:
                        dxf_json = None
                        process_flow_text = None
                        logger.error(f"Process flow generation error: {e}")
                
                progress_bar.progress(40, text="Composing cover letter...")
                
                # Generate cover letter AFTER process flow is created (to include high-level summary)
                cover_letter_text = None
                if offer_ref and sender_name:
                    try:
                        # Create enhanced high-level summary from process flow and DXF data
                        process_flow_summary = ""
                        if process_flow_text:
                            lines = process_flow_text.strip().split('\n')
                            # Extract main system components from process flow
                            summary_components = []
                            for line in lines[:5]:  # Look at first 5 lines for better coverage
                                # Extract component names (remove numbering and description after colon)
                                if ':' in line:
                                    component = line.split(':')[0].strip()
                                    # Remove numbering (1., 2., etc.)
                                    component = component.lstrip('0123456789. ')
                                    if component and len(component) > 3:  # Avoid empty or very short strings
                                        summary_components.append(component)
                        
                        # Add key quantities from DXF if available
                        quantities = []
                        if dxf_json:
                            if dxf_json.get('total_chutes', 0) > 0:
                                quantities.append(f"{dxf_json['total_chutes']} chutes")
                            if dxf_json.get('total_operators', 0) > 0:
                                quantities.append(f"{dxf_json['total_operators']} operator stations")
                            # Add other relevant quantities if present
                            if dxf_json.get('scanner_systems', 0) > 0:
                                quantities.append(f"{dxf_json['scanner_systems']} scanner systems")
                        
                        # Combine components and quantities into natural summary
                        if summary_components:
                            process_flow_summary = ", ".join(summary_components[:3])  # First 3 components
                            if quantities:
                                process_flow_summary += f" with {', '.join(quantities[:2])}"  # Add up to 2 quantities
                        
                        # If quantities are still empty, fall back to ProposalFacts
                        if not quantities and facts:
                            chute_val, _, _ = get_counts_source_of_truth(facts, "gravity_chutes")
                            feed_val, _, _ = get_counts_source_of_truth(facts, "feedlines")
                            if feed_val:
                                quantities.append(f"{feed_val} induct lines")
                            if chute_val:
                                quantities.append(f"{chute_val} gravity chutes")
                        
                        cover_letter_text = call_groq_cover_letter(
                            client_name=client_name,
                            project_title=project_name,
                            offer_ref=offer_ref,
                            letter_date_str=letter_date.strftime("%B %d, %Y"),
                            executives_block=executives_text,
                            invitation_date=invitation_date_str,
                            meeting_date=meeting_date_str,
                            sender_name=sender_name,
                            sender_title=sender_title,
                            process_flow_summary=process_flow_summary,
                            context=context,
                        )
                        # Store for feedback regeneration
                        st.session_state.section_content_cover_letter = cover_letter_text
                        st.session_state.section_original_cover_letter = cover_letter_text  # Store original for comparison
                    except Exception as e:
                        cover_letter_text = None
                        logger.error(f"Cover letter generation failed: {e}", exc_info=True)
                        st.warning(f"Cover letter generation failed: {str(e)[:100]}")
                
                # Continue processing DXF if needed
                if (include_proposed_system or include_concept_desc) and dxf_layout_file and process_flow_text:
                    
                    try:
                        # Initialize variables for extracted data
                        tables = {}
                        
                        # Reuse temp directory from DXF processing
                        if 'tmp_dir' not in locals():
                            tmp_dir = Path(tempfile.mkdtemp(prefix="proposal_"))
                        if 'dxf_path' not in locals() and dxf_layout_file:
                            dxf_path = tmp_dir / dxf_layout_file.name
                            if not dxf_path.exists():
                                dxf_path.write_bytes(dxf_layout_file.getvalue())
                        
                        # Generate System Description from process flow and DXF if Proposed System is included
                        if include_proposed_system and process_flow_text and dxf_json:
                            progress_bar.progress(42, text="Building comprehensive system description...")
                            try:
                                # Use st_sys_desc pipeline to produce final system description text
                                sys_desc_final, detected, sd_tables, sd_cost_vals = generate_system_description_from_sd_sys(
                                    dxf_path=dxf_path,
                                    costing_file_upload=costing_file if 'costing_file' in locals() else None,
                                    project_name=project_name,
                                    client_name=client_name,
                                    pph_count=pph_count,
                                    ipp_rate=ipp_rate,
                                    facts=facts,
                                )
                                system_description_text = sys_desc_final
                                # Store for feedback regeneration
                                st.session_state.section_content_system_description = system_description_text
                                st.session_state.section_original_system_description = system_description_text  # Store original for comparison
                                # merge any extracted tables/cost_vals for potential later use
                                if sd_tables:
                                    tables.update(sd_tables)
                            except Exception as e:
                                pass  # Silent fail for system description
                            time.sleep(2)  # Delay to avoid rate limits
                        
                        # Generate Executive Summary from process flow if included
                        if include_exec_summary and process_flow_text:
                            progress_bar.progress(46, text="Summarizing key highlights...")
                            cbs_type_detected = dxf_json.get('cbs_type', 'Cross-belt technology') if dxf_json else 'Cross-belt technology'
                            exec_summary_text = call_groq_exec_summary(
                                process_flow_text, 
                                client_name, 
                                project_name,
                                pph_count=pph_count, 
                                cbs_type=cbs_type_detected,
                                dxf_json=dxf_json,
                                facts=facts,
                                context=context,
                            )
                            # Store for feedback regeneration
                            st.session_state.section_content_executive_summary = exec_summary_text
                            st.session_state.section_original_executive_summary = exec_summary_text  # Store original for comparison
                            time.sleep(2)  # Delay to avoid rate limits
                                                
                        # Generate Mermaid Flowchart if Concept Description is included
                        if include_concept_desc and process_flow_text:
                            progress_bar.progress(50, text="Rendering visual flowchart...")
                            mermaid_code = call_groq_for_mermaid(process_flow_text)
                            flowchart_png_bytes, render_log = generate_mermaid_png(mermaid_code)
                            time.sleep(2)  # Delay to avoid rate limits
                        
                        # Handle layout PNG - either uploaded or convert from DXF
                        if layout_full_png:
                            # User uploaded a PNG - use it
                            layout_png_path = tmp_dir / layout_full_png.name
                            layout_png_path.write_bytes(layout_full_png.getvalue())
                            layout_png_path = str(layout_png_path)
                        else:
                            # No PNG uploaded - try to convert DXF to PNG
                            if CONVERTAPI_SECRET:
                                try:
                                    progress_bar.progress(52, text="Converting layout to image...")
                                    png_path = convert_dxf_to_png(dxf_path)
                                    if png_path and png_path.exists():
                                        layout_png_path = str(png_path)
                                    else:
                                        layout_png_path = None
                                except Exception as e:
                                    layout_png_path = None
                            else:
                                layout_png_path = None
                    except Exception as e:
                        pass  # Silent fail for DXF processing
                        
                if commercial_include and costing_file:
                    progress_bar.progress(53, text="Analyzing commercial data...")
                    try:
                        # Read Overall Costing sheet
                        df = pd.read_excel(costing_file, sheet_name="Overall Costing", header=None)
                        sheet_csv = df.to_csv(index=False)
                        
                        # Call Groq to extract price sheet
                        price_data = call_groq_for_price_sheet(sheet_csv)
                        payment_terms_data = st.session_state.get("payment_terms", [])
                        bca_discount = apply_bca
                    except Exception as e:
                        price_data = None
                
                # Get client logo path - either from dropdown selection or uploaded file
                client_logo_path = None
                if selected_client != "None" and selected_client in CLIENT_LOGOS:
                    # Use logo from dropdown selection
                    client_logo_path = CLIENT_LOGOS[selected_client]
                elif client_logo:
                    # Use uploaded logo - save temporarily
                    client_logo_path = f"temp_client_logo.{client_logo.name.split('.')[-1]}"
                    with open(client_logo_path, "wb") as f:
                        f.write(client_logo.getbuffer())
                
                progress_bar.progress(55, text="Assembling document framework...")
                
                # ==================== START WITH FRESH DOCUMENT ====================
                # Always start with a fresh document that has all standard Word styles
                doc = Document()
                
                # Initialize tables dictionary for costing tables (Conveyor BOQ, etc.)
                tables = {}
                
                # Ensure required list styles exist
                ensure_list_styles(doc)
                
                # Set default font for the document
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Calibri (Body)'
                font.size = Pt(11)
                
                # DEBUG: Log section availability before DOCX building
                logger.info(f"=== DOCX BUILDING DEBUG ===")
                logger.info(f"cover_letter_text: {cover_letter_text is not None} (len: {len(cover_letter_text) if cover_letter_text else 0})")
                logger.info(f"exec_summary_text: {exec_summary_text is not None} (len: {len(exec_summary_text) if exec_summary_text else 0})")
                logger.info(f"process_flow_text: {process_flow_text is not None} (len: {len(process_flow_text) if process_flow_text else 0})")
                logger.info(f"system_description_text: {system_description_text is not None} (len: {len(system_description_text) if system_description_text else 0})")
                logger.info(f"flowchart_png_bytes: {flowchart_png_bytes is not None}")
                logger.info(f"Include flags - exec: {include_exec_summary}, proposed: {include_proposed_system}, concept: {include_concept_desc}")
                
                # ==================== COVER LETTER (with header/footer) ====================
                if cover_letter_text:
                    logger.info("✓ Adding cover letter to DOCX")
                    build_cover_letter_section(doc, cover_letter_text)
                else:
                    logger.warning("✗ Cover letter text is None - NOT adding to DOCX")
                
                # ==================== FRONT PAGE (WITH HEADER) ====================
                if cover_letter_text:
                    logger.info("✓ Adding front page to DOCX")
                    build_front_page_section(doc, project_name, offer_ref, contact_name, contact_email, contact_phone, layout_png_path)
                else:
                    logger.warning("✗ Cover letter required for front page - NOT adding")
                    
                # ==================== PRE-GENERATION VALIDATION ====================
                try:
                    pre_issues = st.session_state.get("pre_gen_issues", [])
                    valid_container = st.expander("Validation", expanded=False)
                    with valid_container:
                        if pre_issues:
                            for issue in pre_issues:
                                st.warning(issue)
                        else:
                            st.success("Context counts validated; no gating issues detected.")
                except Exception as e:
                    st.info(f"Validation info unavailable: {e}")

                # ==================== CONSISTENCY AUDIT ====================
                if ENABLE_CONSISTENCY_AUDIT:
                    try:
                        from consistency_audit import audit_section_consistency
                        ctx = st.session_state.get("proposal_context")
                        
                        # Only run audit if context exists
                        if ctx is not None:
                            violations = audit_section_consistency(
                                ctx,
                                cover_letter_text or "",
                                exec_summary_text or "",
                                process_flow_text or "",
                                system_description_text or "",
                            )
                        else:
                            violations = []
                        
                        audit_container = st.expander("Audit", expanded=False)
                        with audit_container:
                            st.write("**Normalized Context Counts:**")
                            if ctx:
                                st.code(ctx.counts_block_text(), language="text")
                            else:
                                st.info("Context not available; audit skipped")
                            
                            # Show pre-generation issues
                            pre_gen_issues = st.session_state.get("pre_gen_issues", [])
                            if pre_gen_issues:
                                st.info("**Pre-Generation Gating Issues:**\n" + "\n".join([f"• {i}" for i in pre_gen_issues]))
                            
                            if violations:
                                st.warning("**Consistency Violations Detected:**")
                                for v in violations:
                                    st.warning(v)
                                # Minimal corrective pass for violating sections
                                if ctx:
                                    if any("Executive Summary" in v for v in violations) and exec_summary_text:
                                        exec_summary_text = correct_section_numbers(exec_summary_text, ctx, "Executive Summary")
                                    if any("System Description" in v for v in violations) and system_description_text:
                                        system_description_text = correct_section_numbers(system_description_text, ctx, "System Description")
                                    if any("Cover Letter" in v for v in violations) and cover_letter_text:
                                        cover_letter_text = correct_section_numbers(cover_letter_text, ctx, "Cover Letter")
                                    st.info("✓ Applied corrective pass to align numbers with context.")
                            else:
                                st.success("✓ Consistency audit passed: no violations detected")
                    except Exception as e:
                        st.warning(f"Consistency audit unavailable: {e}")
                
                progress_bar.progress(60, text="Weaving content sections together...")
                
                # ==================== GLOSSARY ====================
                build_glossary_section(doc)
                
                # Start numbering from 2 because Glossary occupies number 1
                counter = 2
                
                # Final deterministic enforcement: sanitize -> enforce_counts -> sanitize
                if exec_summary_text:
                    exec_summary_text = process_section_output(exec_summary_text, ctx)
                if cover_letter_text:
                    cover_letter_text = process_section_output(cover_letter_text, ctx)
                if process_flow_text:
                    process_flow_text = process_section_output(process_flow_text, ctx)
                if system_description_text:
                    system_description_text = process_section_output(system_description_text, ctx)

                # ==================== BUILD SECTIONS IN ORDER ====================
                
                # 1. Executive Summary
                if include_exec_summary and exec_summary_text:
                    logger.info(f"✓ Adding executive summary to DOCX (text length: {len(exec_summary_text)})")
                    build_executive_summary_section(doc, exec_summary_text, counter)
                    counter += 1
                elif include_exec_summary and not exec_summary_text:
                    logger.warning("✗ Executive summary enabled but text is None")
                else:
                    logger.info(f"Executive summary: include_flag={include_exec_summary}, has_text={exec_summary_text is not None}")
                
                progress_bar.progress(65, text="Adding company credentials...")
                
                # 2. Company Profile
                if include_company_profile:
                    build_company_profile_section(doc, counter)
                    counter += 1

                # 3. Reference Projects
                if include_reference_projects:
                    build_reference_projects_section(doc, counter)
                    counter += 1

                progress_bar.progress(70, text="Integrating technical specifications...")
                
                # 4. Handled Shipment Spectrum
                if include_handled_spectrum:
                    user_parcel_spectrum = st.session_state.get("parcel_spectrum", None)
                    build_handled_spectrum_section(doc, counter, project_name, client_name, user_parcel_spectrum)
                    counter += 1

                # 5. Proposed System Description
                if include_proposed_system and process_flow_text:
                    logger.info(f"✓ Adding proposed system description to DOCX (flow length: {len(process_flow_text)})")
                    build_proposed_system_description_section(doc, counter, client_name, project_name, 
                                                             process_flow_text, layout_png_path, 
                                                             system_description_text=None,
                                                             costing_file=None)
                    counter += 1
                elif include_proposed_system and not process_flow_text:
                    logger.warning("✗ Proposed system enabled but process_flow_text is None")
                else:
                    logger.info(f"Proposed system: include_flag={include_proposed_system}, has_flow={process_flow_text is not None}")
                
                progress_bar.progress(75, text="Building technical architecture...")
                
                # 6. Concept Description
                if include_concept_desc and flowchart_png_bytes:
                    logger.info(f"✓ Adding concept description with flowchart to DOCX ({len(flowchart_png_bytes)} bytes)")
                    build_concept_description_section(doc, counter, flowchart_png_bytes)
                    counter += 1
                elif include_concept_desc and not flowchart_png_bytes:
                    logger.warning("✗ Concept description enabled but flowchart_png_bytes is None")
                else:
                    logger.info(f"Concept description: include_flag={include_concept_desc}, has_flowchart={flowchart_png_bytes is not None}")
                
                # 7. Sorter System Capacity (moved before System Description)
                if include_capacity_section and capacity_excel is not None:
                    build_capacity_calculations_section(doc, counter, client_name, project_name, capacity_excel, ctx)
                    counter += 1
                
                # 8. System Description (Separate top-level section with hierarchical numbering)
                if include_proposed_system and system_description_text:
                    build_system_description_section(doc, counter, system_description_text, costing_file=costing_file)
                    counter += 1
                
                progress_bar.progress(80, text="Detailing system components...")
                
                # 9. Description of Components (from Loop CBS Excel if available)
                sorter_spec = None
                if costing_file is not None:
                    try:
                        sheet_name, df = load_loop_cbs_sheet_from_excel(costing_file.getvalue())
                        if sheet_name and df is not None:
                            sheet_text = df_to_compact_text(df)
                            sorter_spec = call_groq_for_sorter_spec(sheet_name, sheet_text)
                    except Exception as e:
                        pass  # Silent fail for Loop CBS extraction
                
                build_description_of_components_section(doc, counter, sorter_spec)
                counter += 1
                
                # 10. Proposed System Technical Details (Mechanical from Groq, Electrical/Control from Groq)
                mechanical_bom_items = None
                bom_json = None
                
                if costing_file is not None:
                    # Generate Mechanical equipment BOM using Groq API (extracts from all relevant sheets)
                    try:
                        mechanical_bom_items = generate_mechanical_bom_from_costing(costing_file)
                    except Exception as e:
                        logger.error(f"Mechanical BOM generation failed: {e}")
                        mechanical_bom_items = None
                    
                    # Generate Electrical Equipment and Control System using Groq API
                    try:
                        # Load Quote Master sheet from costing file
                        with pd.ExcelFile(io.BytesIO(costing_file.getvalue())) as xls:
                            quote_master_sheet = None
                            logger.info(f"Costing file sheets: {xls.sheet_names}")
                            for sheet in xls.sheet_names:
                                if sheet.lower().strip() == "quote master":
                                    quote_master_sheet = sheet
                                    break
                            
                            if quote_master_sheet:
                                logger.info(f"Found Quote Master sheet: {quote_master_sheet}")
                                df_quote = pd.read_excel(xls, sheet_name=quote_master_sheet)
                                sheet_text = df_to_compact_text_quote_master(df_quote)
                                bom_json = call_groq_for_bom(sheet_text)
                                if bom_json and "sections" in bom_json:
                                    logger.info(f"BOM JSON sections: {[s.get('title') for s in bom_json.get('sections', [])]}")
                            else:
                                logger.warning(f"Quote Master sheet not found. Available sheets: {xls.sheet_names}")
                    except Exception as e:
                        logger.error(f"BOM extraction failed: {e}", exc_info=True)
                        bom_json = None
                
                # Build section if we have either mechanical BOM or electrical/control from Groq
                has_mechanical = mechanical_bom_items is not None and len(mechanical_bom_items) > 0
                has_electrical_control = bom_json is not None
                
                if has_mechanical or has_electrical_control:
                    build_proposed_system_technical_details_section(doc, counter, 
                                                                    mechanical_bom_items=mechanical_bom_items, 
                                                                    bom_json=bom_json)
                    counter += 1
                    
                    # Store technical details as text for feedback regeneration
                    tech_details_text = "PROPOSED SYSTEM TECHNICAL DETAILS\n\n"
                    if mechanical_bom_items:
                        tech_details_text += "MECHANICAL EQUIPMENT:\n"
                        for item in mechanical_bom_items:
                            tech_details_text += f"- {item.get('description', '')}: {item.get('qty', '')} {item.get('value', '')}\n"
                    if bom_json:
                        if bom_json.get('Electrical Equipment'):
                            tech_details_text += "\nELECTRICAL EQUIPMENT:\n"
                            for item in bom_json.get('Electrical Equipment', []):
                                tech_details_text += f"- {item.get('description', '')}: {item.get('qty', '')} {item.get('value', '')}\n"
                        if bom_json.get('Control System'):
                            tech_details_text += "\nCONTROL SYSTEM:\n"
                            for item in bom_json.get('Control System', []):
                                tech_details_text += f"- {item.get('description', '')}: {item.get('qty', '')} {item.get('value', '')}\n"
                    st.session_state.section_content_technical_details = tech_details_text
                    st.session_state.section_original_technical_details = tech_details_text  # Store original for comparison
                
                progress_bar.progress(85, text="Configuring system parameters...")
                
                # 11. Electrical System
                if elec_include:
                    build_electrical_section(doc, counter)
                    counter += 1
                
                # 12. Falcon WCS CONTROLIT
                if wcs_include:
                    build_wcs_section(doc, counter, client_name)
                    counter += 1
                
                # 13. Falcon Visual Inspection System (SCADA)
                if scada_include:
                    build_scada_section(doc, counter, client_name)
                    counter += 1
                
                # 14. Key Components Make
                if key_include:
                    build_key_components_section(doc, counter, key_components_edited)
                    counter += 1
                
                progress_bar.progress(90, text="Finalizing project deliverables...")
                
                # 15. Principal of Safety
                if safety_include:
                    build_safety_section(doc, counter)
                    counter += 1
                
                # 16. Infrastructure
                if infra_include:
                    build_infrastructure_section(doc, counter)
                    counter += 1
                
                # 17. Program Organisation
                if prog_include:
                    build_program_org_section(doc, counter, client_name, prog_gantt)
                    counter += 1
                
                # 18. Client Responsibility
                if client_resp_include:
                    build_client_responsibility_section(doc, counter, client_name)
                    counter += 1
                
                # 19. System Handover
                if handover_include:
                    build_handover_section(doc, counter)
                    counter += 1
                
                progress_bar.progress(95, text="Preparing commercial terms...")
                
                # 20. Commercial
                if commercial_include:
                    build_commercial_section(doc, counter, price_data, payment_terms_data, bca_discount)
                    counter += 1
                
                # 21. Warranty Period
                if warranty_include:
                    build_warranty_section(doc, counter, warranty_type, warranty_duration, 
                                          warranty_start, warranty_extended_text, 
                                          warranty_amc_text, warranty_transport_text)
                    counter += 1
                
                # 22. Exclusions
                if exclusion_include:
                    build_exclusions_section(doc, counter, selected_exclusions)
                    counter += 1
                
                # ==================== DETECT AND INSERT GLOSSARY ====================
                # Scan the entire document for glossary terms
                try:
                    full_text = extract_full_text_from_docx(doc)
                    detected_terms = find_terms_in_text(full_text)
                    if detected_terms:
                        # Create a temporary document with just the glossary section
                        temp_glossary_doc = Document()
                        build_glossary(temp_glossary_doc, detected_terms)
                        # Save to buffer and reload
                        temp_buffer = io.BytesIO()
                        temp_glossary_doc.save(temp_buffer)
                        temp_buffer.seek(0)
                        glossary_doc = Document(temp_buffer)
                        # Find the index of the Executive Summary heading
                        glossary_placeholder_index = None
                        word_ns_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
                        for idx, el in enumerate(doc.element.body):
                            if hasattr(el, 'iter'):
                                for t in el.iter(word_ns_tag):
                                    if t.text and 'Executive Summary' in t.text:
                                        glossary_placeholder_index = idx
                                        break
                            if glossary_placeholder_index is not None:
                                break
                        # Insert glossary elements before Executive Summary or at end if not found
                        glossary_elements = []
                        for element in glossary_doc.element.body:
                            if element.tag.endswith('sectPr'):
                                continue
                            glossary_elements.append(element)

                        word_ns_uri = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        word_paragraph_tag = f'{{{word_ns_uri}}}p'
                        word_break_tag = f'{{{word_ns_uri}}}br'

                        def is_page_break(elem):
                            if elem is None or elem.tag != word_paragraph_tag:
                                return False
                            for br in elem.iter(word_break_tag):
                                if br.get(f'{{{word_ns_uri}}}type') == 'page':
                                    return True
                            return False

                        def build_page_break_element():
                            para = OxmlElement('w:p')
                            run = OxmlElement('w:r')
                            br = OxmlElement('w:br')
                            br.set(qn('w:type'), 'page')
                            run.append(br)
                            para.append(run)
                            return para

                        if glossary_placeholder_index is not None:
                            prev_element = None
                            if glossary_placeholder_index > 0:
                                prev_element = doc.element.body[glossary_placeholder_index - 1]
                            if not is_page_break(prev_element):
                                glossary_elements.insert(0, build_page_break_element())
                            for i, element in enumerate(glossary_elements):
                                doc.element.body.insert(glossary_placeholder_index + i, element)
                        else:
                            prev_element = doc.element.body[-1] if len(doc.element.body) > 0 else None
                            if not is_page_break(prev_element):
                                glossary_elements.insert(0, build_page_break_element())
                            for element in glossary_elements:
                                doc.element.body.append(element)
                except Exception as e:
                    pass  # Silent fail for glossary
                
                # ==================== INSERT COVER PAGE AT BEGINNING ====================
                
                # ==================== INSERT COVER PAGE AT BEGINNING ====================
                # Now prepend cover page at the beginning if cover letter was generated
                if cover_letter_text:
                    try:
                        # Get client logo bytes for cover page
                        cover_client_logo_bytes = None
                        if client_logo_path and os.path.exists(client_logo_path):
                            with open(client_logo_path, "rb") as f:
                                cover_client_logo_bytes = f.read()
                        
                        # Create cover page using template
                        cover_page_buffer = create_cover_page(
                            client_logo=cover_client_logo_bytes,
                            client_name=client_name,
                            project_title=project_name
                        )
                        
                        # Save main document to temp buffer
                        temp_main_buffer = io.BytesIO()
                        doc.save(temp_main_buffer)
                        temp_main_buffer.seek(0)
                        
                        # Load cover page document (from template)
                        cover_doc = Document(cover_page_buffer)
                        
                        # Load main content document (all our generated content with images)
                        main_doc = Document(temp_main_buffer)
                        
                        # Try using Composer for proper merge (preserves all relationships including images)
                        try:
                            composer = Composer(cover_doc)
                            composer.append(main_doc)
                            
                            # Save composed document

                            composed_buffer = io.BytesIO()
                            composer.save(composed_buffer)
                            composed_buffer.seek(0)
                            # Load as final document
                            doc = Document(composed_buffer)

                        except (ImportError, NameError, AttributeError):
                            # Fallback: If Composer not available, use element insertion
                            cover_elements = []
                            for element in cover_doc.element.body:
                                if element.tag.endswith('sectPr'):
                                    continue
                                cover_elements.append(element)
                            for i, element in enumerate(cover_elements):
                                main_doc.element.body.insert(i, element)
                            page_break_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:br w:type="page"/></w:r></w:p>'
                            page_break_element = parse_xml(page_break_xml)
                            main_doc.element.body.insert(len(cover_elements), page_break_element)
                            doc = main_doc

                    except Exception as e:
                        pass  # Silent fail for cover page insertion

                progress_bar.progress(98, text="Polishing final touches...")
                
                # === Apply header/footer to the finalized document ===
                try:
                    client_logo_bytes = None
                    if client_logo_path and os.path.exists(client_logo_path):
                        with open(client_logo_path, "rb") as fh:
                            client_logo_bytes = fh.read()

                    create_header_footer(
                        doc,
                        client_name=client_name,
                        project_name=project_name,
                        client_logo_bytes=client_logo_bytes,
                    )
                except Exception as e:
                    pass  # Silent fail for header/footer

                progress_bar.progress(95, text="Packaging your proposal...")
                time.sleep(0.3)
                
                # Save DOCX to buffer first
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                # Clean up temporary logo files
                # No need to remove falcon_logo_path, logo is fixed from backend
                if client_logo_path and os.path.exists(client_logo_path):
                    os.remove(client_logo_path)
                
                # Convert DOCX to PDF using docx2pdf (local conversion - much faster)
                progress_bar.progress(98, text="Rendering your document...")
                
                pdf_buffer = BytesIO()
                try:
                    # Initialize COM for Windows threading (required for docx2pdf)
                    pythoncom.CoInitialize()
                    
                    # Save DOCX to temporary file for conversion
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                        tmp_docx.write(docx_buffer.getvalue())
                        tmp_docx_path = tmp_docx.name
                    
                    # Define PDF output path
                    pdf_path = tmp_docx_path.replace('.docx', '.pdf')
                    
                    # Convert to PDF using docx2pdf (requires MS Word installed)
                    docx2pdf_convert(tmp_docx_path, pdf_path)
                    
                    # Read PDF into buffer
                    with open(pdf_path, 'rb') as pdf_file:
                        pdf_buffer.write(pdf_file.read())
                    pdf_buffer.seek(0)
                    
                    # Clean up temp files
                    try:
                        os.remove(tmp_docx_path)
                        os.remove(pdf_path)
                    except:
                        pass
                except Exception as pdf_error:
                    st.error(f"Could not convert to PDF: {str(pdf_error)}. Make sure Microsoft Word is installed.")
                    st.stop()
                
                progress_bar.progress(100, text="Your proposal is ready!")
                time.sleep(0.5)
                
                # Store PDF and DOCX in session state for preview, download, and regeneration
                st.session_state.generated_pdf_buffer = pdf_buffer
                st.session_state.generated_docx_buffer = docx_buffer  # Store DOCX for regeneration
                st.session_state.generated_doc_filename = f"Falcon_Proposal_{client_name.replace(' ', '_')}.pdf"
                st.session_state.docs_generated = st.session_state.get("docs_generated", 0) + 1
                st.session_state.current_status = "Ready"
                
                # ==================== LOG PROPOSAL GENERATION TO BACKEND ====================
                try:
                    # Save uploaded files to OUTPUT folder
                    dxf_saved_name = ""
                    costing_saved_name = ""
                    throughput_saved_name = ""
                    layout_png_saved_name = ""
                    
                    if dxf_layout_file:
                        dxf_saved_name = save_uploaded_file_to_output(dxf_layout_file, "DXF Upload")
                    
                    if costing_file:
                        costing_saved_name = save_uploaded_file_to_output(costing_file, "Costing Sheet")
                    
                    if capacity_excel:
                        throughput_saved_name = save_uploaded_file_to_output(capacity_excel, "Throughput Calc Sheet")
                    
                    if layout_full_png:
                        layout_png_saved_name = save_uploaded_file_to_output(layout_full_png, "DXF Upload")
                    
                    # Save generated proposal
                    docx_saved, pdf_saved = save_generated_proposal(docx_buffer, pdf_buffer, client_name, project_name)
                    
                    # Log to Excel
                    serial_no = log_proposal_generation(
                        client_name=client_name,
                        client_executives=executives_text if 'executives_text' in dir() else "",
                        project_name=project_name,
                        offer_reference=offer_ref if 'offer_ref' in dir() else "",
                        meeting_date=meeting_date_str if 'meeting_date_str' in dir() else "",
                        pph_rate=str(pph_count) if 'pph_count' in dir() else "",
                        ipp_rate=str(ipp_rate) if 'ipp_rate' in dir() else "",
                        sender_name=sender_name if 'sender_name' in dir() else "",
                        sender_title=sender_title if 'sender_title' in dir() else "",
                        contact_name=contact_name if 'contact_name' in dir() else "",
                        contact_email=contact_email if 'contact_email' in dir() else "",
                        contact_phone=contact_phone if 'contact_phone' in dir() else "",
                        dxf_filename=dxf_saved_name,
                        costing_filename=costing_saved_name,
                        throughput_filename=throughput_saved_name,
                        layout_png_filename=layout_png_saved_name
                    )
                    
                    # Store serial number for feedback logging
                    st.session_state.current_proposal_serial = serial_no
                    
                except Exception as log_error:
                    logger.warning(f"Could not log proposal to backend: {log_error}")
                
                # Clear progress bar
                progress_bar.empty()
                
                # Redirect to preview page
                st.session_state.page = "preview"
                st.rerun()
                
            except Exception as e:
                progress_bar.empty() if 'progress_bar' in dir() else None
                st.session_state.current_status = "Ready"
                st.error(f"Error generating document: {str(e)}")
                st.exception(e)
                # Clean up temporary logo files in case of error
                try:
                    if 'client_logo_path' in locals() and client_logo_path and os.path.exists(client_logo_path):
                        os.remove(client_logo_path)
                    # Clean up any temp docx/pdf files
                    if 'tmp_docx_path' in locals() and os.path.exists(tmp_docx_path):
                        os.remove(tmp_docx_path)
                    if 'pdf_path' in locals() and os.path.exists(pdf_path):
                        os.remove(pdf_path)
                except:
                    pass

# Footer (only show on input page, preview page has its own footer)
if st.session_state.get("page", "input") == "input":
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #6c757d; padding: 2rem 0;">
        <p style="margin: 0;">© 2026 Falcon Autotech Pvt. Ltd. | All Rights Reserved</p>
        <p style="margin: 0.5rem 0 0 0; font-size: 0.85rem;">
            Proposal Generator v8.0
        </p>
    </div>
    """, unsafe_allow_html=True)


