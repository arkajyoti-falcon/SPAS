# app.py
import io
import os
import re
from datetime import datetime
from typing import Optional, List, Dict, Any

import streamlit as st
import pandas as pd
from PIL import Image

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================================================
# Paths / Assets
# =========================================================
COVER_TEMPLATE_PATH = r"assests\Templates\template.docx"     # cover template
FALCON_LOGO_PATH = r"assests\Images\falcon-autotech-icon-removebg-preview.png"      # optional

SOLUTION_IMAGES = {
    "Cubizone-R": r"assests\Images\r.PNG",
    "Cubizone-R Thru": r"assests\Images\thru.PNG",
    "Cubizone-R Eco": r"assests\Images\eco.PNG",
    "Cubizone-R Cross": r"assests\Images\cross.PNG",
    "Cubizone-V Eco": r"assests\Images\v-eco.PNG",
    "Cubizone-V Jumbo": r"assests\Images\v-jumbo.PNG",
}

# ---------------------------------------------------------
# Colors (approx from reference)
# ---------------------------------------------------------
COLOR_HEADING_BLUE = RGBColor(31, 78, 121)
COLOR_TABLE_HEADER = RGBColor(31, 78, 121)
COLOR_TABLE_HEADER_TEXT = RGBColor(255, 255, 255)
COLOR_LIGHT_BAND = RGBColor(217, 225, 242)
COLOR_PRICE_BODY = RGBColor(226, 239, 218)


# =========================================================
# Data: Technical Specs (by type)
# =========================================================
TECH_SPECS: Dict[str, Dict[str, Any]] = {
    "Cubizone-R": {
        "SIZE": [
            ("Maximum Length", "mm", "700"),
            ("Maximum Width", "mm", "700"),
            ("Maximum Height", "mm", "700"),
            ("Minimum Length", "mm", "25"),
            ("Minimum Width", "mm", "25"),
            ("Minimum Height", "mm", "25"),
            ("Measurement Increment", "mm", "1"),
        ],
        "WEIGHT": [
            ("Maximum Weight", "grams", "50000"),
            ("Minimum Weight", "grams", "20"),
            ("Least Count", "grams", "1"),
            ("Throughput", "PPH", "400–600*"),
        ],
        "SHIPMENT SHAPE": [
            ("Regular Boxes", "NA", "YES (Without Flaps)"),
            ("Irregular/Non Boxes", "NA", "YES (With Flaps)"),
            ("Dimension Interval", "Sec", "3"),
            ("Surface Characteristics", "NA", "Stainless Steel"),
            ("Orientation", "NA", "Object to be butted at Axis"),
            ("Spacing", "NA", "NA"),
            ("Operational Temperature", "Celsius", "5 to 55"),
            ("Power", "NA", "230V"),
            ("Barcode Scanner", "NA", "Wired 1D/2D Scanner (Wireless option available)"),
            ("Interface with PC", "NA", "USB/Serial"),
            ("Portable", "NA", "Yes"),
            ("Camera", "NA", "Optional"),
            ("Printer", "NA", "Optional"),
        ],
        "FOOTNOTE": "(*)Throughput depends upon efficiency of the operator",
    },
    "Cubizone-R Eco": {
        "SIZE": [
            ("Maximum Length", "mm", "600"),
            ("Maximum Width", "mm", "600"),
            ("Maximum Height", "mm", "600"),
            ("Minimum Length", "mm", "25"),
            ("Minimum Width", "mm", "25"),
            ("Minimum Height", "mm", "25"),
            ("Measurement Increment", "mm", "1"),
        ],
        "WEIGHT": [
            ("Maximum Weight", "grams", "50000"),
            ("Minimum Weight", "grams", "20"),
            ("Least Count", "grams", "1"),
            ("Throughput", "PPH", "400–600*"),
        ],
        "SHIPMENT SHAPE": [
            ("Regular Boxes", "NA", "YES (Without Flaps)"),
            ("Irregular/Non Boxes", "NA", "YES (With Flaps)"),
            ("Dimension Interval", "Sec", "3"),
            ("Surface Characteristics", "NA", "Stainless Steel"),
            ("Orientation", "NA", "Object to be butted at Axis"),
            ("Spacing", "NA", "NA"),
            ("Operational Temperature", "Celsius", "5 to 55"),
            ("Power", "NA", "230V"),
            ("Barcode Scanner", "NA", "Wired 1D/2D Scanner (Wireless option available)"),
            ("Interface with PC", "NA", "USB/Serial"),
            ("Portable", "NA", "Yes"),
            ("Camera", "NA", "Optional"),
            ("Printer", "NA", "Optional"),
        ],
        "FOOTNOTE": "(*)Throughput depends upon efficiency of the operator",
    },
    "Cubizone-R Thru": {
        "SIZE": [
            ("Maximum Length", "mm", "1200"),
            ("Maximum Width", "mm", "1000"),
            ("Maximum Height", "mm", "1000"),
            ("Minimum Length", "mm", "50"),
            ("Minimum Width", "mm", "50"),
            ("Minimum Height", "mm", "50"),
            ("Measurement Increment", "mm", "1"),
        ],
        "WEIGHT": [
            ("Maximum Weight", "Kilo Grams", "60 / 80 / 100"),
            ("Minimum Weight", "grams", "60"),
            ("Least Count", "grams", "10"),
            ("Throughput", "PPH", "700–900*"),
        ],
        "SHIPMENT SHAPE": [
            ("Regular Boxes", "NA", "YES"),
            ("Dimension Interval", "Sec", "3"),
            ("Surface Characteristics", "NA", "Stainless Steel with Roller Bed"),
            ("Orientation", "NA", "Object to be butted at Axis"),
            ("Spacing", "NA", "NA"),
            ("Operational Temperature", "Celsius", "5 to 55"),
            ("Power", "NA", "230V"),
            ("Barcode Scanner", "NA", "Wired 1D/2D Scanner (Wireless Optional)"),
            ("Interface with PC", "NA", "USB/Serial"),
            ("Portable", "NA", "No"),
            ("Camera", "NA", "Optional"),
            ("Printer", "NA", "Optional"),
        ],
        "FOOTNOTE": "(*)Throughput depends upon efficiency of the operator",
    },
    "Cubizone-R Cross": {
        "SIZE": [
            ("Maximum Length", "mm", "600"),
            ("Maximum Width", "mm", "600"),
            ("Maximum Height", "mm", "600"),
            ("Minimum Length", "mm", "25"),
            ("Minimum Width", "mm", "25"),
            ("Minimum Height", "mm", "25"),
            ("Measurement Increment", "mm", "1"),
        ],
        "WEIGHT": [
            ("Maximum Weight", "grams", "50000"),
            ("Minimum Weight", "grams", "20"),
            ("Least Count", "grams", "1"),
            ("Throughput", "PPH", "800–900*"),
        ],
        "SHIPMENT SHAPE": [
            ("Regular Boxes", "NA", "YES"),
            ("Dimension Interval", "Sec", "3"),
            ("Surface Characteristics", "NA", "Stainless Steel"),
            ("Orientation", "NA", "Object to be butted at Axis"),
            ("Spacing", "NA", "NA"),
            ("Operational Temperature", "Celsius", "5 to 55"),
            ("Power", "NA", "230V"),
            ("Barcode Scanner", "NA", "Wireless 1D/2D Scanner (Wireless Optional)"),
            ("Interface with PC", "NA", "USB/Serial"),
            ("Portable", "NA", "No (Optional)"),
            ("Camera", "NA", "Optional"),
            ("Printer", "NA", "Optional"),
        ],
        "FOOTNOTE": "(*)Throughput depends upon efficiency of the operator",
    },
    "Cubizone-V Eco": {
        "SIZE": [
            ("Maximum Length", "mm", "600"),
            ("Maximum Width", "mm", "600"),
            ("Maximum Height", "mm", "600"),
            ("Minimum Length", "mm", "35"),
            ("Minimum Width", "mm", "35"),
            ("Minimum Height", "mm", "10"),
            ("Measurement Increment", "mm", "1"),
        ],
        "WEIGHT": [
            ("Maximum Weight", "grams", "50000"),
            ("Minimum Weight", "grams", "50"),
            ("Least Count", "grams", "1"),
            ("Throughput", "PPH", "400–600*"),
        ],
        "SHIPMENT SHAPE": [
            ("Regular Boxes", "NA", "Vision Based"),
            ("Dimension Interval", "Sec", "3"),
            ("Surface Characteristics", "NA", "Non Slippery Rubber"),
            ("Orientation", "NA", "Object to be butted at Axis"),
            ("Spacing", "NA", "NA"),
            ("Operational Temperature", "Celsius", "5 to 55"),
            ("Power", "NA", "230V"),
            ("Barcode Scanner", "NA", "Wireless 1D/2D Scanner (Wireless Optional)"),
            ("PC", "NA", "Yes"),
            ("Interface with PC", "NA", "USB/Serial"),
            ("Portable", "NA", "Yes"),
            ("Camera", "NA", "Optional"),
            ("Printer", "NA", "Optional"),
        ],
        "FOOTNOTE": "(*)Throughput depends upon efficiency of the operator",
    },
    "Cubizone-V Jumbo": {
        "SIZE": [
            ("Maximum Length", "mm", "1100"),
            ("Maximum Width", "mm", "800"),
            ("Maximum Height", "mm", "800"),
            ("Minimum Length", "mm", "60"),
            ("Minimum Width", "mm", "60"),
            ("Minimum Height", "mm", "60"),
            ("Measurement Increment", "mm", "1"),
        ],
        "WEIGHT": [
            ("Maximum Weight", "grams", "80000"),
            ("Minimum Weight", "grams", "100"),
            ("Least Count", "grams", "1"),
            ("Throughput", "PPH", "400–600*"),
        ],
        "SHIPMENT SHAPE": [
            ("Regular Boxes", "NA", "Vision Based"),
            ("Dimension Interval", "Sec", "3"),
            ("Surface Characteristics", "NA", "Non Slippery Rubber"),
            ("Orientation", "NA", "Object to be butted at Axis"),
            ("Spacing", "NA", "NA"),
            ("Operational Temperature", "Celsius", "5 to 55"),
            ("Power", "NA", "230V"),
            ("Barcode Scanner", "NA", "Wireless 1D/2D Scanner (Wireless Optional)"),
            ("PC", "NA", "Yes"),
            ("Interface with PC", "NA", "USB/Serial"),
            ("Portable", "NA", "Yes"),
            ("Camera", "NA", "Optional"),
            ("Printer", "NA", "Optional"),
        ],
        "FOOTNOTE": "(*)Throughput depends upon efficiency of the operator",
    },
}


# =========================================================
# Word helpers
# =========================================================
def set_run_font(run, name="Calibri", size=11, bold=False, color: Optional[RGBColor] = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_doc_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Ensure Heading 2/3 exist and set them
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_HEADING_BLUE

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_HEADING_BLUE


def add_field(run, field_text: str):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_text

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_separate)
    r.append(fld_end)


def set_cell_shading(cell, rgb: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def safe_set_table_style(doc: Document, table, preferred="Table Grid"):
    # try preferred
    try:
        doc.styles[preferred]
        table.style = preferred
        return
    except KeyError:
        pass

    # fallback: any available table style
    try:
        for s in doc.styles:
            if s.type == WD_STYLE_TYPE.TABLE:
                table.style = s
                return
    except Exception:
        pass


def safe_add_paragraph_style(doc: Document, text: str, style_name: str):
    """
    Adds paragraph with style if available; else adds without style.
    """
    try:
        doc.styles[style_name]
        return doc.add_paragraph(text, style=style_name)
    except KeyError:
        return doc.add_paragraph(text)


def add_bullet_paragraph(doc: Document, text: str):
    """
    Bullet paragraph that works even when 'List Bullet' style doesn't exist.
    Uses numbering properties directly (Word bullet list).
    """
    p = doc.add_paragraph(text)

    # Add w:numPr to paragraph for bullet list
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")

    numId = OxmlElement("w:numId")
    # 1 is commonly the default bullet numbering in Word templates.
    # If template doesn't have numbering, Word still usually renders bullets on update.
    numId.set(qn("w:val"), "1")

    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

    # font
    if p.runs:
        set_run_font(p.runs[0], size=11)
    return p


def add_header_footer_to_section(doc: Document, section, header_title: str, client_logo: Optional[bytes]):
    header = section.header
    header.is_linked_to_previous = False

    for p in list(header.paragraphs):
        p.clear()

    ht = header.add_table(rows=1, cols=3, width=Inches(7.5))
    safe_set_table_style(doc, ht, "Grid Table 4 - Accent 1")
    ht.autofit = True
    left_cell, mid_cell, right_cell = ht.rows[0].cells

    if client_logo:
        try:
            im = Image.open(io.BytesIO(client_logo))
            if im.mode != "RGBA":
                im = im.convert("RGBA")
            alpha = im.getchannel("A")
            bbox = alpha.getbbox()
            if bbox:
                im = im.crop(bbox)
                alpha = im.getchannel("A")
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=alpha)
            buf = io.BytesIO()
            bg.save(buf, format="PNG")
            buf.seek(0)
            lc_para = left_cell.paragraphs[0]
            lc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lc_para.add_run().add_picture(buf, width=Inches(0.6))
        except Exception:
            pass

    mid_p = mid_cell.paragraphs[0]
    mid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mid_p.add_run(header_title)
    set_run_font(run, size=9, bold=False, color=RGBColor(60, 60, 60))

    if os.path.exists(FALCON_LOGO_PATH):
        rc_para = right_cell.paragraphs[0]
        rc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rc_para.add_run().add_picture(FALCON_LOGO_PATH, width=Inches(0.6))

    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing footer paragraphs
    for p in list(footer.paragraphs):
        p.clear()

    # Footer with centered content
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Line 1: Confidential notice
    r1 = fp.add_run("© FALCON AUTOTECH Confidential: Not for Distribution.")
    set_run_font(r1, size=8, color=RGBColor(120, 120, 120))
    fp.add_run().add_break()
    
    # Line 2: Website
    r2 = fp.add_run("https://www.falconautotech.com/")
    set_run_font(r2, size=8, color=RGBColor(120, 120, 120))
    fp.add_run().add_break()
    
    # Line 3: Page X of Y
    rP = fp.add_run("Page ")
    set_run_font(rP, size=8, color=RGBColor(120, 120, 120))
    add_field(rP, "PAGE")
    rOf = fp.add_run(" of ")
    set_run_font(rOf, size=8, color=RGBColor(120, 120, 120))
    add_field(rOf, "NUMPAGES")


def insert_toc(doc: Document):
    """
    Real TOC field (Heading 2-3). User updates in Word.
    """
    p = doc.add_paragraph()
    r = p.add_run()
    add_field(r, r'TOC \o "2-3" \h \z \u')


# =========================================================
# Utility / parsing
# =========================================================
def normalize_exec_names(raw: str) -> List[str]:
    if not raw:
        return []
    raw = raw.replace("\r", "\n")
    out = []
    for line in raw.split("\n"):
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in line.split(",") if p.strip()]
        out.extend(parts)
    return out


def money_to_int(s: str) -> Optional[int]:
    if s is None:
        return None
    s = str(s).strip()
    if not s:
        return None
    s2 = re.sub(r"[^\d]", "", s)
    if not s2:
        return None
    try:
        return int(s2)
    except Exception:
        return None


def int_to_money(n: int, currency: str = "INR") -> str:
    if currency == "USD":
        return f"${n:,}"
    return f"{n:,}/-"


def convert_price_string(price_str: str, from_currency: str, to_currency: str) -> str:
    """Convert price string between INR and USD"""
    if from_currency == to_currency:
        return price_str
    
    # Extract numeric value
    amount = money_to_int(price_str)
    if amount is None:
        return price_str
    
    if from_currency == "INR" and to_currency == "USD":
        usd_amount = int(amount * 0.011)
        return f"${usd_amount:,}"
    elif from_currency == "USD" and to_currency == "INR":
        inr_amount = int(amount / 0.011)
        return f"{inr_amount:,}/-"
    
    return price_str


# =========================================================
# Cover page (template based)
# =========================================================
def add_cover_page_content(doc: Document, client_logo: Optional[bytes], client_name: str, solution_type: str):
    # Remove headers/footers from template cover
    for sec in doc.sections:
        for part in (
            getattr(sec, "header", None),
            getattr(sec, "footer", None),
            getattr(sec, "first_page_header", None),
            getattr(sec, "first_page_footer", None),
            getattr(sec, "even_page_header", None),
            getattr(sec, "even_page_footer", None),
        ):
            if not part:
                continue
            try:
                part.is_linked_to_previous = False
            except Exception:
                pass
            try:
                for tbl in list(part.tables):
                    tbl._element.getparent().remove(tbl._element)
                for p in list(part.paragraphs):
                    p._element.getparent().remove(p._element)
            except Exception:
                pass

    # Client logo
    if client_logo:
        try:
            im = Image.open(io.BytesIO(client_logo))
            if im.mode != "RGBA":
                im = im.convert("RGBA")
            alpha = im.getchannel("A")
            bbox = alpha.getbbox()
            if bbox:
                im = im.crop(bbox)
                alpha = im.getchannel("A")
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=alpha)
            buf = io.BytesIO()
            bg.save(buf, format="PNG")
            buf.seek(0)

            first_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
            run_logo = first_para.insert_paragraph_before().add_run()
            run_logo.add_picture(buf, width=Inches(2.0))
        except Exception:
            first_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
            run_logo = first_para.insert_paragraph_before().add_run()
            run_logo.add_picture(io.BytesIO(client_logo), width=Inches(2.0))

    for _ in range(6):
        doc.add_paragraph("")

    title = f"Techno-Commercial Offer of Cubizon {solution_type} – {client_name}"
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = False
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    today_str = datetime.today().strftime("%B %d, %Y")
    p2 = doc.add_paragraph()
    run2 = p2.add_run(today_str)
    run2.font.size = Pt(14)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(255, 215, 0)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT


# =========================================================
# Content blocks
# =========================================================
def add_cover_letter(doc: Document, solution_type: str, client_name: str, exec_names: List[str], offer_ref: str,
                     contact_name: str, contact_title: str, contact_email: str, contact_mobile: str):
    today_str = datetime.today().strftime("%d/%m/%Y")
    bold_paragraphs = []

    bold_paragraphs.append(doc.add_paragraph("Kind Attention –"))
    if exec_names:
        for nm in exec_names:
            bold_paragraphs.append(doc.add_paragraph(nm))
    else:
        bold_paragraphs.append(doc.add_paragraph("[Client executive name]"))

    bold_paragraphs.append(doc.add_paragraph(f"M/s {client_name}"))
    bold_paragraphs.append(doc.add_paragraph(f"Offer Ref: {offer_ref}; Date {today_str}."))
    bold_paragraphs.append(doc.add_paragraph(f"Subject – Techno-commercial Offer for Static DWS System – {solution_type}"))

    for para in bold_paragraphs:
        for run in para.runs:
            run.font.bold = True

    doc.add_paragraph("Dear Team,")
    doc.add_paragraph("We are pleased to submit our Techno-commercial Offer in response to your RFQ")
    doc.add_paragraph(
        "As you will note, we have studied your requirements in great depth and, along with the information "
        "collected during the workshops and calls with you, we have put together a detailed technical proposal "
        "laid out in various sections and sequenced to enable you to understand our proposed solution and to "
        "re-enforce our commitment to be your partner in this strategic initiative."
    )
    doc.add_paragraph(
        "In subsequent sections, we have highlighted the capabilities and experiences of Falcon Autotech with "
        "sections on our Automation Technologies."
    )
    doc.add_paragraph(
        "To conclude, I would like to add my personal commitment, on behalf of Falcon Autotech. As we move "
        "through the RFP process, please do not hesitate to contact me and my team. We will be pleased to "
        "assist you for any further information or clarifications that you might have."
    )
    doc.add_paragraph("tails")
    doc.add_paragraph("")

    rows = [
        ("Contact Person", contact_name),
        ("Title", contact_title),
        ("Mobile", contact_mobile),
        ("Email", contact_email),
        ("Company Name", "Falcon Autotech Pvt. Ltd."),
        ("Regd. Office", "57, Nimri Colony, Double Storey Flats, Delhi – 110052"),
        ("Website", "www.falconautotech.com"),
        ("LinkedIn", "www.linkedin.com/company/falconautotech/"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=3)
    safe_set_table_style(doc, tbl, "Grid Table 4 - Accent 1")
    for i, (k, v) in enumerate(rows):
        tbl.cell(i, 0).text = k
        tbl.cell(i, 1).text = ":"
        tbl.cell(i, 2).text = v
        for run in tbl.cell(i, 0).paragraphs[0].runs:
            set_run_font(run, size=11)
        for run in tbl.cell(i, 1).paragraphs[0].runs:
            set_run_font(run, size=11)
        for run in tbl.cell(i, 2).paragraphs[0].runs:
            set_run_font(run, size=11)


def add_reference_image(doc: Document, solution_type: str, section_num: int):
    doc.add_paragraph(f"{section_num}. Cubizon {solution_type}", style="Heading 2")
    

    img_path = SOLUTION_IMAGES.get(solution_type)
    if img_path and os.path.exists(img_path):
        doc.add_paragraph("")
        doc.add_picture(img_path, width=Inches(6.6))
        p = doc.add_paragraph("Reference Image")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.size = Pt(11)
    else:
        p = doc.add_paragraph(f"[Missing reference image: {img_path}]")
        p.runs[0].font.color.rgb = RGBColor(200, 0, 0)


def add_technical_specifications(doc: Document, solution_type: str, section_num: int):
    doc.add_paragraph(f"{section_num}. Technical Specifications", style="Heading 2")

    spec = TECH_SPECS[solution_type]
    all_rows = []

    def add_section(section_name: str, rows: List[tuple]):
        all_rows.append(("__SECTION__", "", section_name))
        for r in rows:
            all_rows.append(r)

    add_section("SIZE", spec["SIZE"])
    add_section("WEIGHT", spec["WEIGHT"])
    add_section("SHIPMENT SHAPE", spec["SHIPMENT SHAPE"])

    tbl = doc.add_table(rows=1 + len(all_rows), cols=3)
    safe_set_table_style(doc, tbl, "Grid Table 4 - Accent 1")

    hdr = tbl.rows[0].cells
    hdr[0].text = "PROPERTY"
    hdr[1].text = "UNITS"
    hdr[2].text = "VALUE"
    for j in range(3):
        set_cell_shading(hdr[j], COLOR_TABLE_HEADER)
        p = hdr[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=11, bold=True, color=COLOR_TABLE_HEADER_TEXT)

    for i, (prop, unit, val) in enumerate(all_rows, start=1):
        c0, c1, c2 = tbl.rows[i].cells

        if prop == "__SECTION__":
            c0.merge(c1).merge(c2)
            c0.text = val
            set_cell_shading(c0, COLOR_LIGHT_BAND)
            p = c0.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=11, bold=True, color=RGBColor(50, 50, 50))
            continue

        c0.text = prop
        c1.text = unit
        c2.text = val

        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in c0.paragraphs[0].runs:
            set_run_font(run, size=11)
        for run in c1.paragraphs[0].runs:
            set_run_font(run, size=11)
        for run in c2.paragraphs[0].runs:
            set_run_font(run, size=11)


def add_commercial_terms(doc: Document, price_rows: List[Dict[str, str]],
                         bullets_text: List[str], payment_lines: List[str], section_num: int,
                         currency_label: str):
    doc.add_paragraph(f"{section_num}. Commercial Terms", style="Heading 2")
    doc.add_paragraph(f"{section_num}.1 Price Sheet", style="Heading 3")

    tbl = doc.add_table(rows=1 + len(price_rows), cols=4)
    safe_set_table_style(doc, tbl, "Grid Table 4 - Accent 1")

    headers = ["Pos.", "Description", "Total Quantity", currency_label]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, COLOR_TABLE_HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=11, bold=True, color=COLOR_TABLE_HEADER_TEXT)

    for i, row in enumerate(price_rows, start=1):
        vals = [
            row.get("Pos.", ""),
            row.get("Description", ""),
            row.get("Total Quantity", ""),
            row.get(headers[3], ""),
        ]
        for j, v in enumerate(vals):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            parts = str(v).split("\n")
            for k, part in enumerate(parts):
                r = p.add_run(part)
                set_run_font(r, size=11)
                if k < len(parts) - 1:
                    p.add_run().add_break()

            set_cell_shading(cell, COLOR_PRICE_BODY)
            if j == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("The Total Price is to be understood.")

    # bullets - safe (no dependence on "List Bullet")
    for b in bullets_text:
        add_bullet_paragraph(doc, b)

    doc.add_paragraph(
        "Our Price does not include taxes, and duties, import duties, charges, fees, taxes etc. charged by "
        "authorities outside India for deliveries and services. Same shall be borne & paid by client itself."
    )
    doc.add_paragraph("*Material Unloading & Laydown area is in customer scope")

    doc.add_paragraph(f"{section_num}.2 Payment terms", style="Heading 3")
    for ln in payment_lines:
        doc.add_paragraph(ln)


def add_warranty_and_exclusions(doc: Document, section_num: int):
    doc.add_paragraph(f"{section_num}. Warranty Period", style="Heading 2")

    doc.add_paragraph(
        "Falcons offered System comes with a standard warranty of 1 year and post completion of "
        "warranty, customer can opt for Extended warranty.  Warranty covers the following support:"
    )
    for b in [
        "Telephonic, Email and Remote Service Support when required.",
        "Regular Software updates and Bug Fixes.",
        "Supply of Mechanical and Electrical components in case of failure (excluding damages as mentioned in Exclusion Clause)",
    ]:
        add_bullet_paragraph(doc, b)

    doc.add_paragraph("The warranty does not apply to replacement or repair of:")
    add_bullet_paragraph(doc, "Faulty articles:")

    for s in [
        "Failure to comply with the manufacturer's recommendations (logistics documentation, Technical Information Note, retrofit document) and the rules of the trade.",
        "Negligence or abnormal use of equipment.",
        "Anomalies produced by an environment of use, storage or transport that does not comply with the specifications or recommendations of Falcon: packaging, temperature, hygrometry, sector, insulation, etc.",
        "A defect due to a cause external to the supplies and services of Falcon.",
    ]:
        p = doc.add_paragraph(s)
        p.paragraph_format.left_indent = Inches(0.5)

    for b in [
        "Equipment other than that supplied by Falcon.",
        "Items that can be repaired exclusively by Falcon that have been repaired or attempted repairs other than those carried out by Falcon.",
        "Items that fail due to normal wear and tear of one or more of its components or whose tamper evident seals (varnish, strip, etc.) have been broken or whose serial numbers have been removed or modified.",
        "Items damaged during transport to Falcon due to the use of unsuitable packaging.",
    ]:
        add_bullet_paragraph(doc, b)

    doc.add_paragraph(f"{section_num + 1}. Exclusions", style="Heading 2")
    doc.add_paragraph("The scope of supply includes all parts which are defined in the Supplier’s quotation.")
    doc.add_paragraph(
        "All other parts which are not defined in the Supplier’s quotation do not belong to the Supplier’s scope "
        "of supply and are excluded. The following parts are also excluded:"
    )

    exclusions = [
        "Unloading of the system",
        "PC-Desktop",
        "Building infrastructure; building structure, doors, fire exits, levelling devices, building extinguisher and fire alarm system, building heating and lighting system.",
        "Electrical power supply and wiring to the main control cabinets.",
        "UPS for Controls and Drives",
        "Electrical Power for Installation",
        "Emergency/Uninterruptable power supply",
        "Laydown Area",
        "All kind of network incl. Local Area Network (LAN/WLAN), exceeding the scope described in Scope of Supply",
        "Any Kind of Civil work",
        "Any item specifically indicated as not forming part of the subject matter of the Seller's supply in the offer documentation.",
    ]
    for idx, ex in enumerate(exclusions, start=1):
        doc.add_paragraph(f"{chr(64+idx)}. {ex}")


# =========================================================
# Build DOCX
# =========================================================
def build_offer_docx(
    client_logo: Optional[bytes],
    client_name: str,
    solution_type: str,
    offer_ref: str,
    exec_names: List[str],
    contact_name: str,
    contact_title: str,
    contact_email: str,
    contact_mobile: str,
    price_rows: List[Dict[str, str]],
    currency_label: str,
) -> bytes:
    doc = Document(COVER_TEMPLATE_PATH)
    set_doc_defaults(doc)

    # Page 1
    add_cover_page_content(doc, client_logo, client_name, solution_type)

    # New section after cover
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    header_title = f"FALCON’s Proposal to {client_name} for Supply of Cubizon {solution_type}"
    add_header_footer_to_section(doc, section2, header_title, client_logo)

    # Page 2
    add_cover_letter(doc, solution_type, client_name, exec_names, offer_ref,
                     contact_name, contact_title, contact_email, contact_mobile)
    doc.add_page_break()

    # Page 3 - TOC heading as normal paragraph (not Heading style) so it won't appear in TOC
    toc_title = doc.add_paragraph()
    toc_run = toc_title.add_run("Table of Contents")
    set_run_font(toc_run, size=16, bold=True, color=COLOR_HEADING_BLUE)
    insert_toc(doc)
    doc.add_page_break()

    # Page 4
    add_reference_image(doc, solution_type, section_num=1)
    doc.add_page_break()

    # Page 5
    add_technical_specifications(doc, solution_type, section_num=2)
    doc.add_page_break()

    # Page 6
    bullets_text = [
        "Ex Works – Greater Noida, India",
        "Freight charges Extra.",
        "Taxes, Mentioned Above.",
        "Duty unpaid",
        "Installation and Commissioning: Inclusive.",
        "Price is valid for 60 days from the date of proposal.",
    ]
    payment_lines = [
        "Stage 1 – 30% advance along with PO",
        "Stage 2 - 70% + taxes before dispatch",
    ]
    add_commercial_terms(doc, price_rows, bullets_text, payment_lines, section_num=3, currency_label=currency_label)
    doc.add_page_break()

    # Last page
    add_warranty_and_exclusions(doc, section_num=4)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# =========================================================
# Streamlit UI
# =========================================================
DEFAULT_PRICE_ROWS = {
    "Cubizone-R": [
        {"Pos.": "FM1", "Description": "Cubizon-R\nAPI Integration\nWired 1D Barcode Scanner", "Total Quantity": "1", "Amount (INR)": "5,20,000/-"},
        {"Pos.": "FM2", "Description": "Non-Powered Infeed and outfeed Roller Conveyor\n1.5 Mtr Each", "Total Quantity": "1", "Amount (INR)": "60,000/-"},
        {"Pos.": "", "Description": "P&F", "Total Quantity": "1", "Amount (INR)": "15,300/-"},
        {"Pos.": "", "Description": "Total Price", "Total Quantity": "", "Amount (INR)": "5,95,300/-"},
    ],
    "Cubizone-R Eco": [
        {"Pos.": "FM1", "Description": "Cubizon-R Eco\nAPI Integration\nWired 1D Barcode Scanner", "Total Quantity": "1", "Amount (INR)": "4,80,000/-"},
        {"Pos.": "FM2", "Description": "P&F", "Total Quantity": "1", "Amount (INR)": "15,300/-"},
        {"Pos.": "", "Description": "Total Price", "Total Quantity": "", "Amount (INR)": "4,95,300/-"},
    ],
    "Cubizone-R Thru": [
        {"Pos.": "FM1", "Description": "Cubizon-R Thru\nAPI Integration\nWired 1D Barcode Scanner", "Total Quantity": "1", "Amount (INR)": "6,60,000/-"},
        {"Pos.": "FM2", "Description": "Non-Powered Infeed and outfeed Roller Conveyor\n1.5 Mtr Each", "Total Quantity": "1", "Amount (INR)": "60,000/-"},
        {"Pos.": "", "Description": "P&F", "Total Quantity": "1", "Amount (INR)": "15,300/-"},
        {"Pos.": "", "Description": "Total Price", "Total Quantity": "", "Amount (INR)": "7,35,300/-"},
    ],
    "Cubizone-R Cross": [
        {"Pos.": "FM1", "Description": "Cubizon-R Cross\nAPI Integration\nWired 1D Barcode Scanner", "Total Quantity": "1", "Amount (INR)": "5,40,000/-"},
        {"Pos.": "FM2", "Description": "Non-Powered Infeed and outfeed Roller Conveyor\n1.5 Mtr Each", "Total Quantity": "1", "Amount (INR)": "60,000/-"},
        {"Pos.": "", "Description": "P&F", "Total Quantity": "1", "Amount (INR)": "15,300/-"},
        {"Pos.": "", "Description": "Total Price", "Total Quantity": "", "Amount (INR)": "6,15,300/-"},
    ],
    "Cubizone-V Eco": [
        {"Pos.": "FM1", "Description": "Cubizon-V Eco\nAPI Integration\nWireless 1D/2D Scanner\nIncluded PC", "Total Quantity": "1", "Amount (INR)": "5,50,000/-"},
        {"Pos.": "", "Description": "P&F", "Total Quantity": "1", "Amount (INR)": "18,500/-"},
        {"Pos.": "", "Description": "Total Price", "Total Quantity": "", "Amount (INR)": "5,68,500/-"},
    ],
    "Cubizone-V Jumbo": [
        {"Pos.": "FM1", "Description": "Cubizon-V Jumbo\nAPI Integration\nWireless 1D/2D Scanner\nIncluded PC", "Total Quantity": "1", "Amount (INR)": "7,20,000/-"},
        {"Pos.": "", "Description": "P&F", "Total Quantity": "1", "Amount (INR)": "20,000/-"},
        {"Pos.": "", "Description": "Total Price", "Total Quantity": "", "Amount (INR)": "7,40,000/-"},
    ],
}


def render_cubizone_builder(embed_mode: bool = False):
    """Render the Cubizon proposal builder. Set embed_mode=True when called inside another app."""
    if not embed_mode:
        st.set_page_config(page_title="Cubizon Offer Builder", layout="wide")

    # Custom CSS for professional styling
    st.markdown(
        """
        <style>
            .main-header {
                font-size: 2rem;
                font-weight: 600;
                color: #1F4E79;
                margin-bottom: 0.5rem;
            }
            .sub-header {
                font-size: 1rem;
                color: #666;
                margin-bottom: 2rem;
            }
            .stTabs [data-baseweb=\"tab-list\"] {
                gap: 2rem;
            }
            .stTabs [data-baseweb=\"tab\"] {
                padding: 1rem 2rem;
                font-weight: 500;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    
    # Initialize session state for data persistence
    if "solution_type" not in st.session_state:
        st.session_state["solution_type"] = "Cubizone-R Thru"
    if "cubizone_solution_type" not in st.session_state:
        st.session_state["cubizone_solution_type"] = st.session_state["solution_type"]
    if "client_name" not in st.session_state:
        st.session_state["client_name"] = ""
    if "offer_ref" not in st.session_state:
        st.session_state["offer_ref"] = ""
    if "exec_raw" not in st.session_state:
        st.session_state["exec_raw"] = ""
    if "client_logo_file" not in st.session_state:
        st.session_state["client_logo_file"] = None
    if "contact_name" not in st.session_state:
        st.session_state["contact_name"] = "Shantanu Rawat"
    if "contact_title" not in st.session_state:
        st.session_state["contact_title"] = "Sr. Engineer"
    if "contact_email" not in st.session_state:
        st.session_state["contact_email"] = "Shantanu.rawat@falconautoonline.com"
    if "contact_mobile" not in st.session_state:
        st.session_state["contact_mobile"] = "+91 9415992218"
    if "currency" not in st.session_state:
        st.session_state["currency"] = "INR"
    if "discount_percentage" not in st.session_state:
        st.session_state["discount_percentage"] = 0.0

    # Create tabs
    tab1, tab2 = st.tabs(["Basic Information", "Price Sheet & Generation"])

    # Tab 1: Basic Information
    with tab1:
        st.markdown("### Client & Project Details")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### Project Information")
            solution_types = ["Cubizone-R", "Cubizone-R Eco", "Cubizone-R Thru", "Cubizone-R Cross", "Cubizone-V Eco", "Cubizone-V Jumbo"]
            
            # Find the index of the current solution type
            current_index = 0
            if st.session_state.get("solution_type") in solution_types:
                current_index = solution_types.index(st.session_state["solution_type"])
            
            # Use a unique key that won't conflict
            selected_solution = st.selectbox(
                "Solution Type",
                solution_types,
                index=current_index,
                key="cubizone_solution_selectbox",
                help="Select the Cubizon solution type for this proposal",
            )
            
            # Update session state immediately
            st.session_state["solution_type"] = selected_solution
            st.session_state["cubizone_solution_type"] = selected_solution
            
            client_name = st.text_input(
                "Client Name",
                value=st.session_state["client_name"],
                placeholder="Enter client company name",
                help="Official name of the client organization"
            )
            st.session_state["client_name"] = client_name
            
            offer_ref = st.text_input(
                "Offer Reference Number",
                value=st.session_state["offer_ref"],
                placeholder="e.g., FLN-2026-001",
                help="Unique reference number for this offer"
            )
            st.session_state["offer_ref"] = offer_ref
            
            exec_raw = st.text_area(
                "Client Executive Names",
                value=st.session_state["exec_raw"],
                height=120,
                placeholder="Enter names separated by comma or new line\nExample:\nJohn Doe\nJane Smith",
                help="Names of client executives to address in the cover letter"
            )
            st.session_state["exec_raw"] = exec_raw
            
            client_logo_file = st.file_uploader(
                "Client Logo",
                type=["png", "jpg", "jpeg"],
                help="Upload client logo for document header (PNG, JPG, or JPEG format)"
            )
            if client_logo_file:
                st.session_state["client_logo_file"] = client_logo_file
        
        with col2:
            st.markdown("#### Contact Person Details")
            st.info("These details will appear in the cover letter contact section")
            
            contact_name = st.text_input(
                "Name",
                value=st.session_state["contact_name"],
                placeholder="Full name of contact person"
            )
            st.session_state["contact_name"] = contact_name
            
            contact_title = st.text_input(
                "Title",
                value=st.session_state["contact_title"],
                placeholder="Job title or designation"
            )
            st.session_state["contact_title"] = contact_title
            
            contact_email = st.text_input(
                "Email",
                value=st.session_state["contact_email"],
                placeholder="contact@falconautotech.com"
            )
            st.session_state["contact_email"] = contact_email
            
            contact_mobile = st.text_input(
                "Mobile",
                value=st.session_state["contact_mobile"],
                placeholder="+91 XXXXXXXXXX"
            )
            st.session_state["contact_mobile"] = contact_mobile

    # Tab 2: Price Sheet & Generation
    with tab2:
        # Header with currency toggle on top right
        col_header1, col_header2 = st.columns([3, 1])
        with col_header1:
            st.markdown("### Price Sheet Configuration")
        with col_header2:
            # Currency toggle button (True = USD, False = INR)
            is_usd = st.toggle(
                "USD",
                value=(st.session_state["currency"] == "USD"),
                help="Toggle between INR and USD (1 INR = 0.011 USD)"
            )
            selected_currency = "USD" if is_usd else "INR"
            
            # Handle currency change
            if selected_currency != st.session_state["currency"]:
                # Convert existing price data
                if "price_df" in st.session_state and st.session_state["price_df"] is not None:
                    df = st.session_state["price_df"].copy()
                    amount_col = "Amount (INR)" if "Amount (INR)" in df.columns else "Amount (USD)"
                    
                    if amount_col in df.columns:
                        # Convert each amount
                        df[amount_col] = df[amount_col].apply(
                            lambda x: convert_price_string(str(x), st.session_state["currency"], selected_currency)
                        )
                        
                        # Rename column
                        new_col_name = f"Amount ({selected_currency})"
                        df = df.rename(columns={amount_col: new_col_name})
                        st.session_state["price_df"] = df
                
                st.session_state["currency"] = selected_currency
                st.rerun()
        
        # Update price dataframe when solution type changes
        if "price_df" not in st.session_state or st.session_state["price_df"] is None or st.session_state.get("price_df_type") != st.session_state["solution_type"]:
            base_df = pd.DataFrame(DEFAULT_PRICE_ROWS[st.session_state["solution_type"]])
            
            # Convert to selected currency if needed
            if st.session_state["currency"] == "USD":
                base_df["Amount (INR)"] = base_df["Amount (INR)"].apply(
                    lambda x: convert_price_string(str(x), "INR", "USD")
                )
                base_df = base_df.rename(columns={"Amount (INR)": "Amount (USD)"})
            
            st.session_state["price_df"] = base_df
            st.session_state["price_df_type"] = st.session_state["solution_type"]
        
        st.info(f"Editing price sheet for {st.session_state['solution_type']}. You can add, edit, or remove rows as needed.")
        
        # Determine current amount column name
        amount_col = f"Amount ({st.session_state['currency']})"
        
        edited_df = st.data_editor(
            st.session_state["price_df"],
            num_rows="dynamic",
            use_container_width=True,
            key=f"price_editor_{st.session_state['solution_type']}_{st.session_state['currency']}",
            column_config={
                "Pos.": st.column_config.TextColumn("Pos.", width="small"),
                "Description": st.column_config.TextColumn("Description", width="large"),
                "Total Quantity": st.column_config.TextColumn("Total Quantity", width="medium"),
                amount_col: st.column_config.TextColumn(amount_col, width="medium"),
            }
        )

        # Persist edited values so reruns keep the latest data
        st.session_state["price_df"] = edited_df.copy()
        
        # Display computed total
        amounts = []
        for _, r in edited_df.iterrows():
            desc = str(r.get("Description", "")).strip().lower()
            if desc == "total price":
                continue
            val = money_to_int(r.get(amount_col, ""))
            if val is not None:
                amounts.append(val)
        computed_total = sum(amounts) if amounts else 0
        
        # Discount slider
        st.markdown("#### Discount")
        discount_percentage = st.slider(
            "Discount Percentage",
            min_value=0.5,
            max_value=60.0,
            value=st.session_state["discount_percentage"],
            step=0.5,
            format="%.1f%%",
            help="Apply discount to the total price"
        )
        st.session_state["discount_percentage"] = discount_percentage
        
        # Calculate discounted total
        discount_amount = int(computed_total * (discount_percentage / 100))
        final_total = computed_total - discount_amount
        
        # Display totals
        col1, col2, col3 = st.columns(3)
        with col1:
            if computed_total:
                st.metric("Subtotal", int_to_money(computed_total, st.session_state["currency"]))
        with col2:
            if discount_amount:
                st.metric(f"Discount ({discount_percentage}%)", f"-{int_to_money(discount_amount, st.session_state['currency'])}")
        with col3:
            if final_total:
                st.metric("Final Total", int_to_money(final_total, st.session_state["currency"]))
        
        st.markdown("---")
        st.markdown("### Generate Proposal Document")
        
        # Generate button
        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            if st.button("Generate Document", type="primary", use_container_width=True):
                # Validation
                missing = []
                if not os.path.exists(COVER_TEMPLATE_PATH):
                    missing.append(f"Missing cover template: {COVER_TEMPLATE_PATH}")
                img_path = SOLUTION_IMAGES.get(st.session_state["solution_type"])
                if not img_path or not os.path.exists(img_path):
                    missing.append(f"Missing reference image for {st.session_state['solution_type']}: {img_path}")
                
                if missing:
                    st.error("Backend files missing:\n\n" + "\n\n".join(missing))
                    st.stop()
                
                if not st.session_state["client_name"].strip():
                    st.error("Client Name is required. Please fill it in the 'Basic Information' tab.")
                    st.stop()
                if not st.session_state["offer_ref"].strip():
                    st.error("Offer Reference Number is required. Please fill it in the 'Basic Information' tab.")
                    st.stop()
                
                # Process inputs
                exec_names = normalize_exec_names(st.session_state["exec_raw"])
                client_logo_bytes = st.session_state["client_logo_file"].read() if st.session_state.get("client_logo_file") else None
                
                # Convert price rows - always use INR column name for document generation
                price_rows_df = edited_df.copy()
                current_amount_col = f"Amount ({st.session_state['currency']})"

                # Compute discount and final total (recompute to avoid staleness)
                amounts = []
                for _, r in price_rows_df.iterrows():
                    desc = str(r.get("Description", "")).strip().lower()
                    if desc == "total price":
                        continue
                    val = money_to_int(r.get(current_amount_col, ""))
                    if val is not None:
                        amounts.append(val)
                computed_total = sum(amounts) if amounts else 0
                discount_pct = st.session_state.get("discount_percentage", 0.0)
                discount_amount = int(computed_total * (discount_pct / 100))
                final_total = computed_total - discount_amount

                # Add Final Price row
                price_rows_df = price_rows_df.rename(columns={current_amount_col: current_amount_col})
                price_rows_df = price_rows_df.fillna("")
                price_rows_df.loc[len(price_rows_df)] = {
                    "Pos.": "",
                    "Description": "Final Price",
                    "Total Quantity": "",
                    current_amount_col: int_to_money(final_total, st.session_state["currency"]),
                }

                price_rows = price_rows_df.astype(str).to_dict(orient="records")

                currency_label = current_amount_col
                
                # Generate document
                with st.spinner("Generating proposal document..."):
                    try:
                        docx_bytes = build_offer_docx(
                            client_logo=client_logo_bytes,
                            client_name=st.session_state["client_name"].strip(),
                            solution_type=st.session_state["solution_type"],
                            offer_ref=st.session_state["offer_ref"].strip(),
                            exec_names=exec_names,
                            contact_name=st.session_state["contact_name"].strip(),
                            contact_title=st.session_state["contact_title"].strip(),
                            contact_email=st.session_state["contact_email"].strip(),
                            contact_mobile=st.session_state["contact_mobile"].strip(),
                            price_rows=price_rows,
                            currency_label=currency_label,
                        )
                        
                        filename = f"Techno-Commercial Offer - Cubizon {st.session_state['solution_type']} - {st.session_state['client_name']}.docx".replace("/", "-")
                        
                        st.success("Document generated successfully!")
                        st.download_button(
                            "Download Proposal",
                            data=docx_bytes,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Error generating document: {str(e)}")


if __name__ == "__main__":
    render_cubizone_builder()

