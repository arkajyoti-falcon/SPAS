"""
PRODUCTION-READY AGENTIC SYSTEM: Process Flow Generation
==========================================================

FIXES APPLIED:
1. Enhanced DXF categorization (fewer UNCATEGORIZED)
2. Proper section numbering detection from references
3. Clean refinement without duplication
4. Strict component verification
5. No debugging output leakage
"""

import os
import re
import tempfile
import json
import time
import logging
from pathlib import Path
from collections import Counter, defaultdict
from typing import Any, List, Dict

import streamlit as st
from dotenv import load_dotenv
import ezdxf
from groq import Groq
from pinecone import Pinecone
from docx import Document
from dxf_extractor import create_dxf_summary_for_embedding

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# CBS Domain Knowledge Document Path
# Try multiple locations for the CBS knowledge document
CBS_KNOWLEDGE_DOC_PATHS = [
    Path(__file__).parent / "Generalized Falcon CBS Ecosystem.docx",  # Same folder as script
    Path(__file__).parent.parent / "Generalized Falcon CBS Ecosystem.docx",  # Parent folder
    Path("D:/Projects/1. Propsal_Automation/Generalized Falcon CBS Ecosystem.docx"),  # Absolute path
]

def _find_cbs_knowledge_doc() -> Path:
    """Find the CBS knowledge document from multiple possible locations."""
    for path in CBS_KNOWLEDGE_DOC_PATHS:
        if path.exists():
            return path
    return CBS_KNOWLEDGE_DOC_PATHS[0]  # Return first path for error message

CBS_KNOWLEDGE_DOC = _find_cbs_knowledge_doc()


def load_cbs_domain_knowledge() -> str:
    """
    Load the generalized CBS ecosystem knowledge from docx file.
    This provides domain expertise for better process flow generation.
    """
    if not CBS_KNOWLEDGE_DOC.exists():
        logger.warning(f"CBS knowledge document not found: {CBS_KNOWLEDGE_DOC}")
        return ""
    
    try:
        doc = Document(str(CBS_KNOWLEDGE_DOC))
        content = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content.append(text)
        
        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    content.append(row_text)
        
        knowledge = "\n".join(content)
        logger.info(f"Loaded CBS domain knowledge: {len(knowledge)} characters")
        return knowledge
    except Exception as e:
        logger.error(f"Error loading CBS knowledge document: {e}")
        return ""


# Load CBS domain knowledge at startup (cached)
@st.cache_data
def get_cbs_knowledge() -> str:
    """Cached loader for CBS domain knowledge."""
    return load_cbs_domain_knowledge()

# CONFIG
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY","pcsk_8akoe_FxzXaW2zvAsEd1uiHqxiMrosvumSujgFyrWAB9vyqG87DGWpnDc6rSxaDYrkP3v")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME", "spas-dxf-samples")
EMBED_MODEL = "llama-text-embed-v2"
EMBED_NAMESPACE = "v1-dxf"

st.set_page_config(page_title="Process Flow Agent", layout="wide")

if not GROQ_API_KEY:
    st.error("❌ GROQ_API_KEY not found")
    st.stop()

groq_client = Groq(api_key=GROQ_API_KEY)

# ENHANCED COMPONENT PATTERNS
COMPONENT_PATTERNS = {
    "AUTO_INDUCT": [r"fal.*fs\d+", r"fal.*feed", r"feedline", r"transfer.*plate"],
    "CONVEYOR_INFEED": [r"telescopic", r"infeed.*conv", r"inclined"],
    "VDS_BUFFER": [r"vds", r"distribution", r"buffer", r"arm.*vds", r"fal.*s013"],
    "OPERATOR_STATION": [r"operator(?!.*safety)", r"manual.*station"],
    "CHUTE": [r"chute", r"slide", r"sliding", r"irregular", r"parcel.*chute", 
              r"non.*sort", r"rejection", r"sortfail"],  # Enhanced!
    "PTL": [r"ptl", r"put.*to.*light", r"light.*rack", r"pallet.*setup"],
    "BAG_SYSTEM": [r"bag", r"bagging", r"takeaway", r"trolley"],  # Added trolley
    "RECIRCULATION": [r"recirculation", r"refeed"],
    "CBS_SORTER": [r"cbs", r"cross.*belt", r"sorter"],
    "COLLECTION": [r"collection.*bin", r"pallet(?!.*setup)"],  # New category
}

STRUCTURAL_PATTERNS = [r"leg.*guard", r"guard(?!.*operator)", r"fenc", 
                       r"safety.*(?!operator)", r"crossover.*main", r"end.*joint"]

def _is_noise_block(name: str) -> bool:
    n = name.strip()
    return (re.match(r"^\*[UDXATE]\d+$", n, re.IGNORECASE) or 
            n.startswith("*") or n.startswith("~") or n.startswith("A$C"))

def _is_structural(name: str) -> bool:
    return any(re.search(p, name.lower()) for p in STRUCTURAL_PATTERNS)

def _categorize_component(name: str) -> str:
    n_lower = name.lower()
    for category, patterns in COMPONENT_PATTERNS.items():
        if any(re.search(p, n_lower) for p in patterns):
            return category
    return "UNCATEGORIZED"

def _normalize_group_name(name: str) -> str:
    n = name.strip()
    if "|" in n:
        n = n.split("|")[-1]
    if not n.startswith("FAL"):
        n = re.sub(r"[_\-]+", " ", n)
    n = re.sub(r"\s+", " ", n).strip()
    if not re.search(r"V\d+$", n, re.IGNORECASE):
        n = re.sub(r"\s*\(?\d+\)?$", "", n).strip()
    return n.lower()

def _detect_cbs_type(project_name: str) -> str:
    return "Linear CBS" if "linear" in project_name.lower() else "Loop CBS"

def _analyze_chute_types(components: dict) -> dict:
    """Enhanced chute analysis"""
    chute_analysis = {"total": 0, "by_type": defaultdict(int), "has_type_info": False}
    
    for comp_name, count in components.items():
        n_lower = comp_name.lower()
        if "chute" in n_lower:
            chute_analysis["total"] += count
            
            # Detect types
            if "sliding" in n_lower or "slide" in n_lower:
                chute_analysis["by_type"]["sliding"] += count
                chute_analysis["has_type_info"] = True
            elif "irregular" in n_lower:
                chute_analysis["by_type"]["irregular"] += count
                chute_analysis["has_type_info"] = True
            elif "non-sort" in n_lower or "nonsort" in n_lower:
                chute_analysis["by_type"]["non_sort"] += count
                chute_analysis["has_type_info"] = True
            elif "reject" in n_lower or "sortfail" in n_lower:
                chute_analysis["by_type"]["rejection"] += count
                chute_analysis["has_type_info"] = True
            elif "big parcel" in n_lower or "parcel" in n_lower:
                chute_analysis["by_type"]["big_parcel"] += count
                chute_analysis["has_type_info"] = True
            elif "gravity" in n_lower or "collection" in n_lower:
                chute_analysis["by_type"]["gravity"] += count
                chute_analysis["has_type_info"] = True
            elif "live" in n_lower:
                chute_analysis["by_type"]["live"] += count
                chute_analysis["has_type_info"] = True
    
    return chute_analysis

def extract_dxf_components(dxf_path: Path, project_name: str = "") -> dict:
    """Extract with enhanced categorization"""
    doc = ezdxf.readfile(str(dxf_path))
    msp = doc.modelspace()
    hdr = doc.header

    units_code = hdr.get("$INSUNITS", None)
    try:
        units_code = int(units_code) if units_code is not None else None
    except:
        units_code = None

    raw_counts: Counter[str] = Counter()
    for e in msp:
        try:
            if e.dxftype() == "INSERT":
                bname = e.dxf.name
                if not _is_noise_block(bname) and not _is_structural(bname):
                    raw_counts[bname] += 1
        except:
            continue

    categorized = defaultdict(lambda: defaultdict(lambda: {"count": 0, "examples": []}))
    for raw_name, cnt in raw_counts.items():
        category = _categorize_component(raw_name)
        gname = _normalize_group_name(raw_name)
        categorized[category][gname]["count"] += cnt
        categorized[category][gname]["examples"].append(raw_name)

    cbs_type = _detect_cbs_type(project_name or dxf_path.name)
    chute_analysis = _analyze_chute_types(raw_counts)

    has_auto = len(categorized.get("AUTO_INDUCT", {})) > 0
    has_ops = len(categorized.get("OPERATOR_STATION", {})) > 0
    has_vds = len(categorized.get("VDS_BUFFER", {})) > 0

    induction_type = ("MIXED (Auto + Manual)" if has_auto and has_ops else
                     "AUTO" if has_auto else "MANUAL" if has_ops else "UNKNOWN")

    category_summary = {cat: sum(item["count"] for item in items.values())
                       for cat, items in categorized.items()}
    
    total_components = sum(category_summary.values())

    return {
        "file": dxf_path.name,
        "cbs_type": cbs_type,
        "induction_type": induction_type,
        "has_vds": has_vds,
        "total_components": total_components,
        "category_summary": dict(category_summary),
        "categorized_components": {cat: {name: data["count"] 
                                         for name, data in items.items()}
                                  for cat, items in categorized.items()},
        "chute_analysis": chute_analysis,
        "raw_block_counts": {k: int(v) for k, v in raw_counts.items()},
    }

def create_dxf_summary(dxf_json: dict) -> str:
    lines = ["=" * 70,
             "DXF COMPONENT ANALYSIS",
             "=" * 70, "",
             f"FILE: {dxf_json['file']}",
             f"CBS TYPE: {dxf_json['cbs_type']}",
             f"INDUCTION: {dxf_json['induction_type']}",
             f"VDS/BUFFER: {'YES' if dxf_json['has_vds'] else 'NO'}",
             ""]
    
    cats = dxf_json.get("category_summary", {})
    if cats:
        lines.append("COMPONENTS:")
        for cat in ["AUTO_INDUCT", "OPERATOR_STATION", "VDS_BUFFER", "CHUTE",
                    "PTL", "BAG_SYSTEM", "COLLECTION", "CBS_SORTER", "UNCATEGORIZED"]:
            if cat in cats and cats[cat] > 0:
                lines.append(f"  • {cat}: {cats[cat]} units")
        lines.append("")
    
    chute = dxf_json.get("chute_analysis", {})
    if chute.get("total", 0) > 0:
        lines.append(f"CHUTES: {chute['total']} total")
        if chute.get("by_type"):
            for ct, cnt in sorted(chute["by_type"].items(), key=lambda x: -x[1]):
                lines.append(f"  • {ct.replace('_', ' ').title()}: {cnt}")
        lines.append("")
    
    lines.extend(["=" * 70])
    return "\n".join(lines)

# Pinecone functions (same as before, keeping them brief)
def get_pinecone_index():
    pc = Pinecone(api_key=PINECONE_API_KEY)
    return pc, pc.Index(PINECONE_INDEX_NAME)


def load_actual_references(base_dir: str = None) -> List[Dict]:
    """Load actual process-flow docx files from the V6_Actual folder as high-priority references.

    Looks for any .docx files under `V6_Actual_Vs_Generated_Processflow_15.12` and returns
    a list of dicts with `client` and `process_flow` keys.
    """
    refs: List[Dict] = []
    try:
        base = Path(base_dir) if base_dir else Path(__file__).parent / ".." / "V6_Actual_Vs_Generated_Processflow_15.12"
        base = base.resolve()
        if not base.exists():
            return refs

        for p in base.rglob("*.docx"):
            # prefer files named Actual/actual/Actual.docx, but accept any .docx
            try:
                client_name = p.parent.name
                doc = Document(str(p))
                paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                content = "\n".join(paragraphs)
                if content:
                    refs.append({"client": client_name, "process_flow": content})
            except Exception:
                logger.debug(f"Failed to read docx {p}", exc_info=True)
                continue
    except Exception:
        logger.debug("Error while scanning for actual reference docx files", exc_info=True)
    return refs

def embed_text(pc, text: str) -> List[float]:
    resp = pc.inference.embed(model=EMBED_MODEL, inputs=[{"text": text}],
                             parameters={"input_type": "passage", "truncate": "END"})
    return list(resp.data[0].values if hasattr(resp.data[0], "values") 
               else resp.data[0]["values"])

def query_similar_flows(pc, index, dxf_summary: str, dxf_json: dict,
                       top_k: int = 2, threshold: float = 0.80) -> List[Dict]:
    """
    HYBRID RETRIEVAL: Embedding similarity + Component matching
    
    Stage 1: Get candidates by embedding (top_k * 5)
    Stage 2: Re-rank by component similarity
    Stage 3: Return only matches above threshold
    
    Args:
        pc: Pinecone client
        index: Pinecone index
        dxf_summary: Component-focused embedding text
        dxf_json: Extracted DXF metadata for component matching
        top_k: Number of results to return
        threshold: Combined score threshold (0-1)
    """
    from dxf_extractor import calculate_component_similarity
    
    # Stage 1: Embed and get candidates
    vec = embed_text(pc, dxf_summary)
    candidates_k = min(top_k * 5, 20)
    
    resp = index.query(
        vector=vec,
        top_k=candidates_k,
        include_metadata=True,
        namespace=EMBED_NAMESPACE,
        include_values=False
    )
    
    # Stage 2: Re-rank by component similarity
    query_cats = dxf_json.get("category_summary", {})
    query_chute = dxf_json.get("chute_analysis", {})
    
    results = []
    for m in (resp.matches if hasattr(resp, "matches") else resp.get("matches", [])):
        embedding_score = getattr(m, "score", None) or m.get("score", 0)
        meta = getattr(m, "metadata", {}) or m.get("metadata", {})
        pf = meta.get("process_flow") or ""
        
        if not pf:
            continue
        
        # Rebuild stored component summary from metadata (handles flattened + json)
        stored_cats = meta.get("category_summary", {}) or {}

        if not stored_cats and meta.get("category_summary_json"):
            try:
                stored_cats = json.loads(meta.get("category_summary_json")) or {}
            except Exception:
                stored_cats = {}

        # Populate from flattened fields if still empty
        if not stored_cats:
            flat_keys = [
                "AUTO_INDUCT", "OPERATOR_STATION", "CONVEYOR_INFEED",
                "VDS_BUFFER", "CHUTE", "RECIRCULATION", "PTL",
                "BAG_SYSTEM", "SCANNER", "CBS_SORTER"
            ]
            rebuilt = {}
            for k in flat_keys:
                val = meta.get(f"cat_{k}")
                if isinstance(val, (int, float)):
                    rebuilt[k] = int(val)
            stored_cats = rebuilt
        
        if not stored_cats:
            # If we cannot reconstruct categories, skip
            continue

        # Rebuild stored chute info for chute-type similarity
        stored_chute = meta.get("chute_analysis", {}) or {}
        if not stored_chute and meta.get("chute_analysis_json"):
            try:
                stored_chute = json.loads(meta.get("chute_analysis_json")) or {}
            except Exception:
                stored_chute = {}
        if not stored_chute:
            stored_chute = {
                "total": meta.get("chute_total", 0) or 0,
                "by_type": {
                    "live": meta.get("chute_live", 0) or 0,
                    "collection": meta.get("chute_collection", 0) or 0,
                    "rejection": meta.get("chute_rejection", 0) or 0,
                    "sliding": meta.get("chute_sliding", 0) or 0,
                    "mini_gravity": meta.get("chute_mini_gravity", 0) or 0,
                    "bulk": meta.get("chute_bulk", 0) or 0,
                    "big_parcel": meta.get("chute_big_parcel", 0) or 0,
                    "gravity": meta.get("chute_gravity", 0) or 0,
                }
            }
        
        # Calculate component similarity
        component_sim = calculate_component_similarity(
            query_cats, stored_cats, query_chute, {"chute_analysis": stored_chute}
        )
        
        # Combined score (70% component match, 30% embedding)
        combined_score = (component_sim * 0.7) + (embedding_score * 0.3)
        
        # Only include if above threshold
        if combined_score >= threshold:
            results.append({
                "id": getattr(m, "id", None) or m.get("id"),
                "embedding_score": embedding_score,
                "component_similarity": component_sim,
                "combined_score": combined_score,
                "process_flow": pf,
                "client": meta.get("client", "Unknown"),
                "category_summary": stored_cats,
                "cbs_type": meta.get("cbs_type"),
                "induction_type": meta.get("induction_type"),
            })
    
    # Sort by combined score
    results.sort(key=lambda x: x["combined_score"], reverse=True)
    
    return results[:top_k]

def call_groq(messages: List[Dict], temp: float = 0.2, max_tok: int = 2000) -> str:
    delay = 2
    for attempt in range(5):
        try:
            resp = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=messages,
                temperature=temp,
                max_tokens=max_tok,
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            if "rate_limit" in str(e).lower() or "429" in str(e):
                if attempt < 4:
                    time.sleep(delay)
                    delay *= 2
                else:
                    raise
            else:
                raise

def detect_numbering_style(reference_flows: List[Dict]) -> bool:
    """Detect if references use section numbering"""
    for ref in reference_flows:
        flow = ref["process_flow"]
        # Check for patterns like "1. Infeed", "2. Inducts"
        if re.search(r'^\d+\.\s+\w', flow, re.MULTILINE):
            return True
    return False

def generate_initial_flow(client_name: str, dxf_json: dict, 
                          reference_flows: List[Dict]) -> str:
    """Generate with comprehensive prompt and reference style matching"""
    
    dxf_summary = create_dxf_summary(dxf_json)
    use_numbering = detect_numbering_style(reference_flows)
    
    # Load CBS domain knowledge for better understanding
    cbs_knowledge = get_cbs_knowledge()
    
    # Build reference context
    ref_context = ""
    if reference_flows:
        ref_context = "\n=== REFERENCE EXAMPLES (for style) ===\n"
        for i, ref in enumerate(reference_flows[:2], 1):
            ref_context += f"\nExample {i} ({ref['client']}):\n"
            ref_context += ref["process_flow"][:5000] + "...\n"
    
    # Build domain knowledge context
    domain_context = ""
    if cbs_knowledge:
        domain_context = f"""
=== CBS DOMAIN KNOWLEDGE (Falcon CBS Ecosystem) ===
Use this knowledge to understand how CBS systems work and write accurate process flows:

{cbs_knowledge}

=== END DOMAIN KNOWLEDGE ==="""
    
    system_prompt = """## ROLE
You are a senior SALES engineer presenting the "Process Flow of the System" to a potential client. You're not just describing—you're SELLING how this solution transforms their operations.

## SALES-FIRST MINDSET (CRITICAL)
- This is a SALES document, not a technical manual
- Every step should answer: "Why does this matter to the client?"
- Highlight BENEFITS: speed, accuracy, efficiency, reduced errors, labor savings
- Make the client visualize parcels flowing SMOOTHLY through their new system

## WHY + WHAT (Always explain WHY, not just WHAT)
- DON'T: "Parcels are inducted onto the sorter"
- DO: "Parcels are smoothly inducted, ensuring zero jams and maximum throughput"

---

## � STORYTELLING NARRATIVE REQUIREMENT (MOST IMPORTANT)

**This is NOT a numbered list. Write it like a human expert describing a journey.**

### ✅ HOW TO WRITE LIKE A HUMAN:
1. **CONNECT Each Step to the Next**: Each section should naturally flow into the next, showing how material moves through the system
2. **Use Transitional Language**: "From there...", "Subsequently...", "Once the parcels reach...", "They are then directed to...", "Finally..."
3. **Tell a STORY**: Imagine explaining the system to a client. How would a human do it? With narrative flow, not bullet points.
4. **Meaningful Complete Sentences**: Each section should be a complete thought that builds upon the previous one
5. **Show the JOURNEY**: Parcels/shipments arrive → move through system → get sorted → leave. Make the reader follow this journey.

### ✅ GOOD EXAMPLE (Storytelling):
```
Process Flow

Incoming shipments are dumped in bulk onto the infeed conveyor, where they begin their journey through the system. From there, the shipments ascend to a higher level via the inclined conveyor, arriving at the induction zone. Here, they are positioned by the operator onto the feedlines. Once on the feedlines, parcels are automatically inducted into the Loop CBS, where the sorting operation begins. The system efficiently sorts the shipments into their respective output chutes based on destination data provided by [Client]. Finally, the sorted packages are discharged into the collection chutes for further processing.
```

### ❌ BAD EXAMPLE (Disconnected):
```
Process Flow
Infeed System: Incoming shipments
Auto Induct Line: Parcels are inducted
Loop CBS: Sorting happens
Output Chutes: 51 chutes
```

---

## �🚨 CRITICAL FORMATTING RULES (NON-NEGOTIABLE)

### Rule 1: NO SECTION NUMBERING
- ❌ WRONG: `1. Infeed System:` `2. Auto Induct Line:`
- ✅ CORRECT: `Infeed System:` `Auto Induct Line:`
- **This is the #1 cause of low structural_coherence scores**

### Rule 2: NO SUB-POINT PARENT NUMBERING  
- ❌ WRONG: `5. a. Live Chutes` `6. b. Collection Chutes`
- ✅ CORRECT: `a. Live Chutes` `b. Collection Chutes`

### Rule 3: Start with "Process Flow" Header Only
- ✅ CORRECT: `Process Flow` (standalone line)
- ❌ WRONG: `Process flow of the Loop CBS System-`
- ❌ WRONG: `1. Process Flow`

### Output Structure Template:
```
Process Flow
Infeed System: - <Description>

Auto Induct Line: <Description>

Loop CBS: - <Description>

Output Chutes: - <Description>
a. <Type> - <Description>
b. <Type> - <Description>

<Conditional Section>: <Description>
```

---

## 📊 DATA PRIORITY HIERARCHY

### Priority 1: METADATA (Highest)
If metadata explicitly states something, use that **exact wording**:
- "existing conveyor" → use "existing conveyor"
- "lengthwise orientation" → use "lengthwise orientation"
- "based on dimensions and weight" → use "based on their dimensions and weight"
- Client name "Amazon" → use "using data provided by Amazon"
- "Falcon's fully automatic induct line" → use this exact phrase

### Priority 2: DXF GUIDANCE
Follow the DXF analysis guidance section provided.

### Priority 3: Generic Functional Descriptions
When data is sparse, use simple, generic descriptions based on system type.

### Priority 4: NEVER INVENT
DO NOT add details not supported by metadata or DXF guidance.

---

## 🚫 CRITICAL "DO NOT" RULES

### DO NOT Invent Scanner Details
- ❌ "Barcode scanners are positioned along the infeed line"
- ❌ "Top-side barcode scanner reads each parcel's barcode"
- ❌ "Side barcode scanners are mounted on the infeed conveyors"
- ✅ ONLY mention scanners if metadata explicitly describes them

### DO NOT Expose CAD Codes
- ❌ "Falcon FS002V02 auto-induct units"
- ❌ "fal_fs002v02 and feedline transfer plates"
- ❌ "supported by transfer plates"
- ✅ Use generic: "auto induct lines", "automatically inducted"

### DO NOT Add Unnecessary Technical Details
- ❌ "ensuring accurate identification before induction"
- ❌ "without operator intervention"
- ❌ "rapidly routes each shipment"
- ✅ Keep it simple and direct like actual examples

### DO NOT Invent Metadata Terms
- ❌ "existing conveyor" (unless metadata says this)
- ❌ "highway line" (unless metadata says this)
- ❌ "lengthwise orientation" (unless metadata says this)

### DO NOT Describe What Happens Inside Equipment
- ❌ "The parcels are lifted via an inclined conveyor"
- ✅ "The shipments ascend to a higher level via an inclined conveyor"

---

## 🎯 LANGUAGE & STYLE MATCHING

### Use These Sentence Patterns:

**Arrivals:**
- "Boxes and totes are loaded onto..."
- "Shipments are placed on..."
- "Bags containing shipments are unsealed and dumped..."
- "Shipments from [source] are dumped in bulk onto..."

**Movement:**
- "The shipments ascend to a higher level..."
- "travel from lower level to Mezzanine level"
- "From there, they are directed to..."

**Transitions:**
- "Once the [items] are [state], they [action]..."
- "Upon arrival at the induct zone..."
- "After the shipments are collected..."

**Operations:**
- "The operator picks and positions each shipment..."
- "Feedlines automatically induct the parcels..."
- "efficiently sorts the shipments into their respective output chutes"
- "by utilizing the data provided by [Client]'s sorting logic"

**Counts:**
- "there are [X] chutes present in the system"
- "A total of [X] chutes are designed to..."
- "Within the system, there are a total of [X]..."

### Tone Characteristics:
- **Direct and factual**, not flowery
- **Active voice** preferred
- **Present tense**
- **Specific over generic** when data available
- **Natural flow** with transitions

---

## ✅ QUALITY ASSURANCE CHECKLIST

### Before Generating Output, Verify:

**Structure (structural_coherence):**
- [ ] NO section numbering (`1.`, `2.`, `3.`)
- [ ] NO sub-point parent numbers (`5. a.`, `6. b.`)
- [ ] Starts with "Process Flow" header only
- [ ] Proper blank lines between sections
- [ ] Lettered sub-points only in Output Chutes section

**Content Accuracy (content_coverage & numeric_accuracy):**
- [ ] Used exact client name from metadata
- [ ] Included all sections specified in DXF guidance
- [ ] Used exact chute counts from DXF/metadata
- [ ] NO invented scanner details at infeed
- [ ] NO exposed CAD codes (FAL_FS002V02, etc.)
- [ ] Included VDS if guidance specified
- [ ] Used metadata terminology exactly where provided

**Language Quality (semantic_similarity & domain_tone):**
- [ ] Matches sentence patterns from actual examples
- [ ] Natural transitions between sections
- [ ] Professional but accessible tone
- [ ] No unnecessary technical elaboration
- [ ] Active voice predominant

---

## 🎯 FINAL REMINDER

**Your output will be evaluated on:**
- **semantic_similarity**: Match actual language patterns and phrasing
- **content_coverage**: Include all relevant details from metadata/DXF, no invented content
- **structural_coherence**: Perfect formatting (no section numbers, proper sub-points)
- **numeric_accuracy**: Exact counts from DXF/metadata
- **domain_tone**: Professional, direct style matching actual examples

**Keys to Success:**
1. **NO section numbering** - this alone will boost structural_coherence by 20+ points
2. **NO invented scanners** - boosts content_coverage and numeric_accuracy
3. **Use metadata verbatim** - boosts semantic_similarity
4. **Match sentence patterns** - boosts semantic_similarity and domain_tone
5. **Follow DXF guidance exactly** - boosts content_coverage

**Example Process flow"

**Generate output that a human expert would write, not a template filler.**"""
    
    user_prompt = f"""
=== DXF COMPONENTS ===
{dxf_summary}
{domain_context}
{ref_context}

=== GENERATION INSTRUCTIONS ===

1. STRUCTURE:
   {"- Use section numbering (1., 2., 3.)" if use_numbering else "- NO section numbering"}
   - Section titles: "<Title>:" or "<Title>: -"
   - Sub-points (chutes only): "a.", "b.", "c."
   - Start with "Process Flow" header

2. SECTIONS TO INCLUDE (based on DXF):
   {_build_section_plan(dxf_json)}

3. CRITICAL RULES:
   - Use ONLY components from DXF above
   - Use exact counts from DXF
   - Apply CBS DOMAIN KNOWLEDGE for accurate technical descriptions
   - Match reference style and language
   - Client: {client_name}
   - CBS Type: {dxf_json['cbs_type']}

4. DO NOT:
   - Invent components not in DXF
   - Add scanner details unless in DXF
   - Expose CAD codes
   - Add UNCATEGORIZED sections

Generate clean process flow now following all rules from the system prompt."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    return call_groq(messages, temp=0.2, max_tok=5000)

def _build_section_plan(dxf_json: dict) -> str:
    """Build section guidance"""
    cats = dxf_json.get("category_summary", {})
    chute = dxf_json.get("chute_analysis", {})
    
    plan = []
    
    # Infeed
    if cats.get("CONVEYOR_INFEED", 0) > 0 or cats.get("AUTO_INDUCT", 0) > 0:
        plan.append("- Infeed System (how shipments arrive, ascend)")
    
    # VDS
    if dxf_json.get("has_vds"):
        plan.append("- VDS/Buffer in Infeed (distribution before induction)")
    
    # Induction
    if dxf_json.get("induction_type") == "MIXED (Auto + Manual)":
        plan.append("- Inducts section (describe both auto and manual)")
    elif "AUTO" in dxf_json.get("induction_type", ""):
        plan.append("- Auto Induct Line")
    elif "MANUAL" in dxf_json.get("induction_type", ""):
        plan.append("- Manual Induct Station")
    
    # CBS
    plan.append(f"- {dxf_json['cbs_type']} (sorting operation)")
    
    # Output Chutes
    if chute.get("total", 0) > 0:
        plan.append(f"- Output Chutes ({chute['total']} total)")
        if chute.get("by_type"):
            for ct, cnt in chute["by_type"].items():
                plan.append(f"    a. {ct.title()} - {cnt} chutes")
    
    # PTL
    if cats.get("PTL", 0) > 0:
        plan.append(f"- Put To Light System ({cats['PTL']} locations)")
    
    # Bag System
    if cats.get("BAG_SYSTEM", 0) > 0:
        plan.append("- Bag Takeaway Conveyor")
    
    return "\n   ".join(plan)


def aggressive_language_cleanup(flow: str) -> str:
    """Ultra-aggressive cleanup of technical terms to ensure natural language."""
    
    # Direct category name replacements
    replacements = [
        # Remove entire problematic phrases first
        (r',?\s*utilizing \d+\s+(VDS_BUFFER|AUTO_INDUCT|OPERATOR_STATION|BAG_SYSTEM|COLLECTION)\s+units?,?', ''),
        (r',?\s*consisting of \d+\s+(VDS_BUFFER|AUTO_INDUCT|OPERATOR_STATION|BAG_SYSTEM|COLLECTION)\s+units?,?', ''),
        (r',?\s*which consists? of \d+\s+(VDS_BUFFER|AUTO_INDUCT|OPERATOR_STATION|BAG_SYSTEM|COLLECTION)\s+units?,?', ''),
        (r',?\s*comprises? \d+\s+(VDS_BUFFER|AUTO_INDUCT|OPERATOR_STATION|BAG_SYSTEM|COLLECTION)\s+units?,?', ''),
        
        # Fix "mixed induction" phrases
        (r'utilizes a mixed induction method[^.]*', 'uses both manual positioning by operators and automatic feedline induction'),
        (r'combining both auto and manual induction[^.]*', ''),
        (r'allowing for both automated and manual sorting[^.]*', ''),
        
        # Standalone category names with units
        (r'\bVDS_BUFFER\s+units?\b', 'VDS loop system'),
        (r'\bAUTO_INDUCT\s+units?\b', 'feedlines'),
        (r'\bOPERATOR_STATION\s+units?\b', 'manual induction stations'),
        (r'\bBAG_SYSTEM\s+units?\b', 'bagging system'),
        (r'\bCOLLECTION\s+units?\b', 'collection points'),
        
        # Bare category names
        (r'\bVDS_BUFFER\b', 'VDS loop system'),
        (r'\bAUTO_INDUCT\b', 'feedlines'),
        (r'\bOPERATOR_STATION\b', 'manual induction stations'),
        (r'\bBAG_SYSTEM\b', 'bagging system'),
        (r'\bCOLLECTION\b', 'collection points'),
        (r'\bCONVEYOR_INFEED\b', 'infeed conveyors'),
        
        # "X CATEGORY units" patterns
        (r'(\d+)\s+VDS_BUFFER\s+units?', r''),
        (r'(\d+)\s+AUTO_INDUCT\s+units?', r''),
        (r'(\d+)\s+OPERATOR_STATION\s+units?', r''),
        (r'(\d+)\s+BAG_SYSTEM\s+units?', r''),
        (r'(\d+)\s+COLLECTION\s+units?', r''),
        
        # Clean up extra spaces
        (r'\s+', ' '),
        (r'\n{3,}', '\n\n'),
        (r',\s*,', ','),  # Remove double commas
        (r'\.\s*\.', '.'),  # Remove double periods
    ]
    
    for pattern, replacement in replacements:
        flow = re.sub(pattern, replacement, flow, flags=re.IGNORECASE)
    
    return flow.strip()


def enforce_induction_paragraph(flow: str, dxf_json: dict) -> str:
    """Enforce canonical induction wording.

    Replaces any 'Inducts' / 'Induction' section that contains forbidden phrasing
    (e.g., 'auto induct', 'AUTO_INDUCT', 'mixed induction') with a canonical
    operator + feedlines paragraph. Keeps counts if available, expressed naturally.
    """
    try:
        # Canonical paragraph (always mention barcode orientation)
        canonical = (
            "Inducts: - The operator picks and positions each shipment on the induct line, "
            "ensuring that the shipment is properly aligned and that its barcode is facing upwards. "
            "The feedlines then automatically induct the shipments onto the Cross-Belt Sorter Loop."
        )

        # Replace any existing Induct(s)/Induction section
        pattern = r'(?mi)^(Inducts?|Induction)\s*:\s*-.*?(?=\n^[A-Z][A-Za-z0-9 _\-]{0,80}:\s*-|\Z)'
        if re.search(pattern, flow):
            flow = re.sub(pattern, canonical, flow, flags=re.MULTILINE | re.DOTALL)
            return flow

        # If there is mention of AUTO_INDUCT or 'auto induct' inline, try to insert section after Infeed/VDS
        if re.search(r'auto\s*induct|AUTO_INDUCT|mixed\s+induction', flow, flags=re.IGNORECASE):
            # try to insert after 'Infeed System' or 'VDS/Buffer' sections
            insert_after = None
            m = re.search(r'(?mi)^(VDS/Buffer|VDS BUFFER|VDS loop system|Infeed System)\s*:\s*-.*?(?=\n^[A-Z][A-Za-z0-9 _\-]{0,80}:\s*-|\Z)', flow, flags=re.MULTILINE | re.DOTALL)
            if m:
                insert_after = m.end()
            if insert_after:
                flow = flow[:insert_after] + "\n\n" + canonical + flow[insert_after:]
                return flow

    except Exception:
        logger.debug("enforce_induction_paragraph failed", exc_info=True)

    return flow


def final_forbidden_pass(flow: str, dxf_json: dict) -> str:
    """Final pass to replace any remaining raw category tokens with natural phrases.

    Also writes a violations file if any forbidden tokens remain after replacement.
    """
    mapping = {
        r'\bVDS_BUFFER\b': 'VDS loop system',
        r'\bAUTO_INDUCT\b': 'feedlines',
        r'\bOPERATOR_STATION\b': 'manual induction stations',
        r'\bBAG_SYSTEM\b': 'bagging system',
        r'\bCOLLECTION\b': 'collection points',
        r'\bCONVEYOR_INFEED\b': 'infeed conveyors',
    }

    for pat, repl in mapping.items():
        flow = re.sub(pat, repl, flow, flags=re.IGNORECASE)

    # Replace patterns like '24 AUTO_INDUCT units' -> '24 feedlines'
    flow = re.sub(r"(\d+)\s+AUTO_INDUCT\s+units?", r"\1 feedlines", flow, flags=re.IGNORECASE)
    flow = re.sub(r"(\d+)\s+VDS_BUFFER\s+units?", r"\1 VDS loop system", flow, flags=re.IGNORECASE)
    flow = re.sub(r"(\d+)\s+BAG_SYSTEM\s+units?", r"\1 bagging stations", flow, flags=re.IGNORECASE)

    # Detect any remaining uppercase category tokens
    violations = re.findall(r'\b(VDS_BUFFER|AUTO_INDUCT|OPERATOR_STATION|BAG_SYSTEM|COLLECTION|CONVEYOR_INFEED)\b', flow)
    try:
        dump_dir = Path(__file__).parent / "tmp_stage_outputs"
        dump_dir.mkdir(parents=True, exist_ok=True)
        (dump_dir / "stage3_final_pass.txt").write_text(flow or "", encoding="utf-8")
        if violations:
            (dump_dir / "stage3_violations.txt").write_text("\n".join(violations), encoding="utf-8")
    except Exception:
        logger.debug("Failed to write final pass dumps", exc_info=True)

    return flow
        


def refine_with_references(initial_flow: str, reference_flows: List[Dict],
                          dxf_json: dict, client_name: str) -> str:
    """
    ALWAYS refine and transform to natural language.
    - Stage 1: Language transformation (remove technical terms, use natural language)
    - Stage 2: Add missing sections from DXF
    - Stage 3: Final polish and narrative flow
    
    Works WITH or WITHOUT references.
    """
    
    dxf_summary = create_dxf_summary(dxf_json)
    dxf_cats = dxf_json.get("category_summary", {})
    
    # Build reference context - prefer actual 'Actual.docx' references from the V6_Actual folder,
    # then include any provided reference_flows. If none exist, fall back to a short default.

    
    # Build detailed reference context
    ref_context = ""
    for i, ref in enumerate(reference_flows[:2], 1):
        display_score = ref.get("combined_score") or ref.get("score") or 0.0
        ref_context += f"\n=== REFERENCE {i}: {ref['client']} (Score: {display_score:.3f}) ===\n"
        ref_context += "Components in this reference:\n"
        
        ref_cats = ref.get("category_summary", {})
        for cat in ["AUTO_INDUCT", "OPERATOR_STATION", "VDS_BUFFER", "CHUTE", 
                    "PTL", "BAG_SYSTEM", "RECIRCULATION", "CBS_SORTER"]:
            if cat in ref_cats and ref_cats[cat] > 0:
                ref_context += f"  • {cat}: {ref_cats[cat]} units\n"
        
        ref_context += "\nProcess Flow Excerpt:\n"
        ref_context += ref["process_flow"][:1500] + "...\n"
    
    system_prompt = """You are a senior proposal engineer refining a CBS Process Flow section to match the quality and style of actual winning proposals.

# CRITICAL SUCCESS FACTORS

## 1. 🎯 LANGUAGE TRANSFORMATION (Highest Priority)

Your PRIMARY goal is to transform technical/robotic language into natural, professional proposal language.

### âŒ WRONG (Technical/Robotic):
- "The Loop CBS system combines auto and manual induction with 24 AUTO_INDUCT units"
- "VDS/Buffer area consists of 2 VDS_BUFFER units"
- "conveyor consists of 104 BAG_SYSTEM units"
- "Shipments are loaded onto the infeed conveyor"

### âœ… CORRECT (Natural/Professional):
- "Operators manually position shipments on the induct line"
- "Once within the loop, shipments are manually picked and fed into the inducts"
- "After secondary sorting and bagging, shipments are placed in bags"
- "Shipments from FC and marketplace are dumped in bulk onto the infeed lines"

### Language Transformation Rules:
1. **NEVER expose category names** (VDS_BUFFER, AUTO_INDUCT, BAG_SYSTEM, etc.)
2. **Use descriptive verbs**: "dumped in bulk", "manually picked", "ascend to", "discharged into"
3. **Focus on the shipment journey**, not equipment specifications
4. **Use natural transitions**: "Once within...", "After...", "From there..."
5. **Prefer active descriptions**: "Operators position shipments" vs "Shipments are positioned"

---

## 2. 📖 NARRATIVE FLOW (Tell the Story)

Each section should connect to create a shipment's journey through the system.

### Story Arc Template:
```
Arrival → Distribution → Preparation → Induction → Sorting → Collection → Dispatch
```

### Connection Phrases to Use:
- "Once the shipments..." / "Once within..."
- "After [action], they..."
- "From there, they..."
- "The shipments then..."
- "Following [process]..."

### Example of Good Flow:
"Shipments from FC and marketplace are dumped in bulk onto the infeed lines, where they **ascend to a higher level** and **enter the VDS loop system**. **Once within the loop**, shipments are manually picked and fed into the inducts."

---

## 3. 🔢 COMPONENT INTEGRATION RULES

### Rule A: Add Missing Sections
IF component exists in BOTH:
- DXF data (count > 0)
- Reference flows (described/mentioned)

AND component is missing from initial flow
→ ADD it using reference language

### Rule B: Component Count Usage
- **DXF counts are sacred** - use them exactly as provided
- Reference counts are for language/style ONLY
- When adding sections, extract the count from DXF, not references
- 

### Rule C: Translation of Technical Categories
When you see these in DXF, translate naturally:

| DXF Category | Natural Language Options |
|--------------|-------------------------|
| VDS_BUFFER | "VDS loop system", "distribution loop", "buffer system" |
| AUTO_INDUCT | "automatically induct", "feedlines" |
| OPERATOR_STATION | "manually position", "operators pick and place", "manual induct" |
| CHUTE (by type) | "sliding chutes", "collection chutes", "rejection chutes" |
| PTL | "PTL racks", "PTL setup", "Put-to-Light system" |
| BAG_SYSTEM | "bags", "bagging process", "roller cage trolleys" |
| COLLECTION | "trolleys", "pallets", "collection points" |

---

## 4. 🎨 STYLE MATCHING FROM REFERENCES

Extract and apply these elements from reference flows:

### A. Sentence Structures
Study how references construct sentences:
- "Shipments from [source] are dumped in bulk onto..."
- "The operator picks and positions each shipment on..."
- "Once the shipments have entered..., [equipment] efficiently sorts..."
- "Within the system, there are [X] chutes..."

### B. Technical Detail Level
- References balance detail with readability
- Count mentions: "50 chutes per zone", "13 chutes per zone"
- BUT avoid: "consisting of X units of Y type"

### C. Section Structure
From references, maintain:
- Section numbering IF references use it
- Colon after section title
- Sub-points (a., b., c.) for chute types
- Descriptive details after the count

---

## 5. 🚫 CRITICAL "DO NOT" RULES

### NEVER Do These:
1. âŒ Expose category names: "VDS_BUFFER units", "AUTO_INDUCT units"
2. âŒ Use technical equipment specs: "consists of X BAG_SYSTEM units"
3. âŒ Write disconnected sections without transitions
4. âŒ Add components that have 0 count in DXF
5. âŒ Use reference counts (only DXF counts)
6. âŒ Create duplicate sections
7. âŒ Remove sections from initial flow
8. âŒ Add debug output or category labels
9. âŒ Expose CAD codes (FS002V02, etc.)
10. âŒ Invent details not in DXF or references

### ALWAYS Do These:
1. âœ… Use natural, professional language from references
2. âœ… Create narrative flow with transitions
3. âœ… Use exact DXF counts
4. âœ… Translate category names naturally
5. âœ… Match reference tone and style
6. âœ… Add missing sections if in DXF + references
7. âœ… Keep client-specific terminology (e.g., "Noon's sorting logic")

---

## 6. 📋 SECTION ENHANCEMENT GUIDE

For each section type, here's how to refine:

### Infeed System:
- Start with shipment arrival: "dumped in bulk", "loaded onto"
- Describe movement: "ascend to", "enter the", "arrive at"
- Include distribution if VDS present: "Once within the loop"

### Inducts/Induction:
- Focus on operator action: "Operators manually position"
- Mention barcode orientation: "ensuring proper alignment and barcode visibility"
- Describe automation: "feedlines automatically induct"

### CBS (Loop or Linear):
- Keep it concise
- Emphasize sorting logic: "using [Client]'s sorting logic"
- Describe efficiency: "efficiently sorts shipments into"

### Output Chutes:
- Lead with discharge action: "Shipments are discharged into"
- Use sub-points (a., b., c.) for types
- Include count AND purpose for each type
- Natural descriptions: "50 chutes per zone collect shipments in roller cage trolleys"

### PTL System:
- Focus on function, not just count
- "PTL racks for consolidation", "further sorted via PTL setup"

### Bag Takeaway:
- Describe the post-sorting journey
- "After secondary sorting and bagging"
- Include destination: "transports them to the outbound docks"

---

## 7. ✅ QUALITY CHECKLIST

Before finalizing, verify:

**Language Quality:**
- [ ] NO category names exposed (VDS_BUFFER, AUTO_INDUCT, etc.)
- [ ] Natural verbs and descriptions used throughout
- [ ] Reads like it was written by a human expert
- [ ] Matches reference tone and phrasing

**Narrative Flow:**
- [ ] Each section connects to the next
- [ ] Tells the shipment journey from arrival to dispatch
- [ ] Transition phrases used between sections
- [ ] Logical progression maintained

**Content Accuracy:**
- [ ] All DXF counts used exactly as provided
- [ ] No reference counts used
- [ ] Missing sections added if in DXF + references
- [ ] No invented components
- [ ] Client name used correctly

**Structure:**
- [ ] Section numbering matches references (if applicable)
- [ ] Sub-points only for chute types
- [ ] Proper spacing between sections
- [ ] No duplicate sections

---

## 8. 🎯 FINAL REMINDER

Your output should read like this actual example:

**GOOD:**
"Shipments from FC and marketplace are dumped in bulk onto the infeed lines, where they ascend to a higher level and enter the VDS loop system. Once within the loop, shipments are manually picked and fed into the inducts."

**NOT like this:**
"The infeed system consists of 5 CONVEYOR_INFEED units. Shipments are loaded onto the conveyor and directed to the VDS/Buffer area, which consists of 2 VDS_BUFFER units."

Transform technical data into professional narrative.
Use DXF for accuracy, references for style.
Tell the shipment's story.
**Sample Examples (Only to understand style, do NOT copy content):** 
Process Flow

Infeed System: - Shipments from FC and Market place are dumped in bulk onto the infeed
lines . The shipments ascend to a higher level and arrive at the VDS loop system. Once the
shipments are within the VDS loop, they are picked manually and are fed among all inducts.
Inducts : - The operator picks and positions each shipment on the induct line, ensuring that
the shipment is properly aligned and that its barcode is facing upwards. The feedlines then
automatically induct the shipments onto the Cross -Belt Sorter Loop.
Loop CBS: - Once the shipments have entered the main loop, the Cross -Belt Sorter (CBS)
efficiently sorts the shipments into their respective output chutes by utilizing the data
provided by Noon 's sorting logic.
Output Chutes : - The shipments are discharged into two types of chutes.
a. Sliding Chutes - Within the loop CBS system, there are a total of 50 Sliding chutes for
each zone . The Shipments collected in Roller Cage trolleys, then they are
consolidated into bags using bagging type PTL racks .
b. Non -Sort Chutes - Within the loop CBS system, there are a total of 13 Non-Sort
Chutes per zone . Shipments collected within these chutes further undergo sortation
via PTL setup into Pallets.
c. Rejection Chute s- Two Rejection Chutes per zone are present to handle rejected
Shipments .
Bag Takeaway Conveyor: - Following the Secondary Sorting process of bagging type PTL , the
shipments are placed into bags and then manually loaded onto a bag takeaway conveyor
located beneath the CBS loop. This conveyor transports the bags out of shipment sorter area
to outbound docks.


Process Flow
Infeed System:
Bags containing shipments are unsealed and dumped in bulk onto the infeed lines of the Cross Belt Sorter equipped with Telescopic Belt Conveyor. The shipments ascend to a higher level and arrive at the VDS system. Once the shipments are within the VDS loop, they are evenly distributed among all inducts using Arm VDS technology.
Inducts:
After the shipments are collected in the VDS chute, an operator picks and positions each shipment on the induct line, ensuring that the shipment is properly aligned and that its barcode is facing upwards. The feed lines then automatically induct the shipments onto the Cross-Belt Sorter Loop.
Loop CBS:
Once the shipments have entered the main loop, the Cross-Belt Sorter (CBS) efficiently sorts the shipments into their respective output chutes by utilizing the data provided by Shadowfax's sorting logic.
Output Chutes:
The shipments are discharged into two types of chutes.
•	Direct Bagging Chutes (L-type):
Within a double decker loop CBS system, there are a total of 104 L-Type direct bagging chutes. Shipments collected within these direct bagging chutes are bagged and will be treated as high volume chutes.
•	Secondary Chutes (L-type):
Within a double decker loop CBS system, there are a total of 100 L-Type Secondary chutes. Shipments collected within these secondary chutes further undergo sortation via PTL setup placed at two levels.
•	Rejection Chute:
Four rejection chutes are present to handle rejected shipments.
Put To Light System:
In the system there are 3000 PTL locations. Each secondary chute is linked to 30 PTL locations. The PTL racks are placed in L-Shape double decker arrangement.
Bag Takeaway Conveyor:
Following the direct bagging process and secondary sorting process, the shipments are placed into bags and then manually loaded onto a bag takeaway conveyor located beneath the CBS loop. This conveyor transports the bags out of shipment sorter to outbound sorter located beneath base mezzanine in the approximate centre of the Loop CBS.


Process Flow
Infeed System: Boxes and totes are loaded onto the existing conveyor in a lengthwise orientation. From there, they are directed to their assigned highway line, which transports them to the CBS induct zone in a singulated manner.
Inducts: Upon arrival at the induct zone, Falcon's fully automatic induct line accurately and smoothly inducts the parcels onto the Linear CBS, based on their dimensions and weight.
Linear CBS: Once the parcels enter the main Linear CBS, the Cross-Belt Sorter (CBS) capture the barcode details & volume data after which it efficiently sorts the boxes and totes into their designated output chutes using data provided by Amazon.
Output Chutes: The Totes/Boxes are discharged into below output chutes.
a. Live Chutes - There are 9 sliding-type live chutes within the Linear CBS system, integrated with PVC belt conveyors and TBCs for live loading.
b. Collection chute – A total of 20 friction roller-based chutes are designed to collect and gradually accumulate the parcels.
c. Rejection Chute- One friction roller-based chute handles rejected shipments.
Recirculation Line: A recirculation line is available to automatically feed sortfail parcels back into the Linear CBS. It is also integrated with a manual loading point for reprocessed boxes and totes collected from the rejection chute.

"""

    user_prompt = f"""
# REFINEMENT TASK

=== INITIAL PROCESS FLOW (Base to Refine) ===
{initial_flow}

=== DXF COMPONENTS (Source of Truth for Counts) ===
{dxf_summary}

**DXF Categories Present:**
{json.dumps(dxf_cats, indent=2)}

=== REFERENCE FLOWS (Source of Language & Style) ===
{ref_context}

---

# YOUR TASK: Refine the Initial Flow

## Step 1: Language Transformation
Go through each section of the initial flow:
- Remove any exposed category names (VDS_BUFFER, AUTO_INDUCT, etc.)
- Replace robotic language with natural descriptions from references
- Add transition phrases to connect sections

## Step 2: Identify Missing Components
Check for components that are:
- Present in DXF (count > 0)
- Described in reference flows
- Missing from initial flow

Examples to check:
- VDS/Buffer system ({dxf_cats.get('VDS_BUFFER', 0)} in DXF)
- PTL System ({dxf_cats.get('PTL', 0)} in DXF)
- Bag Takeaway ({dxf_cats.get('BAG_SYSTEM', 0)} in DXF)
- Recirculation ({dxf_cats.get('RECIRCULATION', 0)} in DXF)

## Step 3: Add Missing Sections
For each missing component:
- Extract language pattern from references
- Use DXF count exactly
- Insert in logical position (follow shipment journey order)
- Integrate with transition phrases

## Step 4: Create Narrative Flow
Ensure the entire flow tells a coherent story:
- Arrival → Distribution → Induction → Sorting → Collection → Dispatch
- Each section flows naturally to the next
- Use transition phrases between sections

## Step 5: Final Polish
- Remove any technical jargon
- Ensure client name is used: "{client_name}"
- Verify CBS type is correct: "{dxf_json['cbs_type']}"
- Check all counts match DXF exactly
- Confirm no category names are exposed

---

# CONSTRAINTS

**MUST USE from DXF:**
- All component counts
- CBS type ({dxf_json['cbs_type']})
- Induction type ({dxf_json.get('induction_type', 'Unknown')})

**MUST ADAPT from References:**
- Language patterns and phrasing
- Sentence structures
- Transition phrases
- Professional tone

**MUST NOT:**
- Add components with 0 count in DXF
- Use reference counts instead of DXF counts
- Expose category names (VDS_BUFFER, etc.)
- Create duplicate sections
- Add debug output

---

Generate the refined process flow now. Output ONLY the final flow text (plain text, ready for proposal document)."""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    refined = call_groq(messages, temp=0.15, max_tok=2500)
    
    # Clean any leaked technical terms or debug output
    refined = re.sub(r'\n+UNCATEGORIZED:.*$', '', refined, flags=re.DOTALL)
    refined = re.sub(r'\n+- [A-Z_]+:.*$', '', refined, flags=re.DOTALL | re.MULTILINE)
    refined = re.sub(r'\n+DXF.*:.*$', '', refined, flags=re.DOTALL | re.MULTILINE)
    
    # Additional cleaning for exposed category names
    category_fixes = {
        r'VDS_BUFFER units?': 'distribution loop',
        r'AUTO_INDUCT units?': 'feedlines',
        r'OPERATOR_STATION units?': 'manual induct stations',
        r'BAG_SYSTEM units?': 'bagging system',
        r'COLLECTION units?': 'collection points',
        r'\b\d+\s+(VDS_BUFFER|AUTO_INDUCT|OPERATOR_STATION|BAG_SYSTEM|COLLECTION)\b': lambda m: m.group(0).split()[0],
    }
    
    for pattern, replacement in category_fixes.items():
        refined = re.sub(pattern, replacement, refined, flags=re.IGNORECASE)
    
    return refined.strip()

def verify_flow(flow: str, dxf_json: dict) -> Dict:
    """Verify generated flow"""
    cats = dxf_json.get("category_summary", {})
    flow_lower = flow.lower()
    
    issues = []
    
    # Check for debug leakage
    if "UNCATEGORIZED:" in flow or re.search(r'- [A-Z_]+:', flow):
        issues.append("Debug output leaked into flow")
    
    # Check for duplicates
    lines = flow.split('\n')
    titles = [l.strip() for l in lines if re.match(r'^[\d\.]*\s*[A-Z].*:', l)]
    if len(titles) != len(set(titles)):
        issues.append(f"Duplicate sections: {[t for t in titles if titles.count(t) > 1]}")
    
    # Check component coverage
    coverage = {}
    keywords = {
        "AUTO_INDUCT": ["feedline", "feedlines", "automatic"],
        "OPERATOR_STATION": ["operator", "manual station", "manual induct"],
        "VDS_BUFFER": ["vds", "buffer"],
        "CHUTE": ["chute", "output"],
        "PTL": ["ptl", "put to light"],
        "BAG_SYSTEM": ["bag", "takeaway"],
    }
    
    for cat, count in cats.items():
        if count > 0 and cat in keywords:
            kws = keywords[cat]
            found = any(kw in flow_lower for kw in kws)
            coverage[cat] = found
            if not found:
                issues.append(f"Missing: {cat} ({count} units)")
    
    return {
        "is_valid": len(issues) == 0,
        "issues": issues,
        "coverage": coverage
    }

def run_agent(dxf_path: Path, client_name: str, project_name: str) -> Dict:
    results = {"status": "processing"}
    
    try:
        # Extract
        with st.status("📊 Extracting DXF...") as status:
            dxf_json = extract_dxf_components(dxf_path, project_name)
            results["dxf"] = dxf_json
            st.write(f"✅ {dxf_json['total_components']} components")
            status.update(label="✅ DXF extracted", state="complete")
        
        # Query
        with st.status("🔍 Finding references...") as status:
            pc, index = get_pinecone_index()
            # CRITICAL: Use SAME embedding format as push.py for consistency!
            dxf_summary = create_dxf_summary_for_embedding(dxf_json)
            refs = query_similar_flows(pc, index, dxf_summary, dxf_json, top_k=2, threshold=0.80)
            results["references"] = refs
            st.write(f"✅ Found {len(refs)} matches")
            if refs:
                for i, ref in enumerate(refs, 1):
                    st.write(f"  {i}. {ref['client']} - Score: {ref['combined_score']:.3f} (Embedding: {ref['embedding_score']:.3f}, Components: {ref['component_similarity']:.3f})")
            status.update(label="✅ References found", state="complete")
        
        # Generate
        with st.status("✍️  Generating flow...") as status:
            initial = generate_initial_flow(client_name, dxf_json, refs)
            results["initial"] = initial
            st.write("✅ Initial flow generated")
            status.update(label="✅ Generated", state="complete")
        
        # Refine
        with st.status("🔧 Refining...") as status:
            refined = refine_with_references(initial, refs, dxf_json, client_name)
            results["refined"] = refined
            st.write("✅ Refined")
            status.update(label="✅ Refined", state="complete")
        
        # Verify
        with st.status("🔍 Verifying...") as status:
            verification = verify_flow(refined, dxf_json)
            results["verification"] = verification
            
            if verification["is_valid"]:
                st.write("✅ Valid")
                results["final"] = refined
            else:
                st.write(f"⚠️  Issues: {', '.join(verification['issues'])}")
                results["final"] = initial  # Fall back
            
            status.update(label="✅ Verified", state="complete")
        
        results["status"] = "success"
        
    except Exception as e:
        logger.error(f"Failed: {e}", exc_info=True)
        results["status"] = "error"
        results["error"] = str(e)
    
    return results

def main():
    st.title("🤖 Process Flow Generator")
    
    uploaded = st.file_uploader("Upload DXF", type=["dxf"])
    col1, col2 = st.columns(2)
    with col1:
        client = st.text_input("Client Name", "Noon")
    with col2:
        project = st.text_input("Project Name", "")
    
    if st.button("Generate", type="primary") and uploaded:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".dxf") as tmp:
            tmp.write(uploaded.read())
            tmp_path = Path(tmp.name)
        
        results = run_agent(tmp_path, client, project or uploaded.name)
        tmp_path.unlink()
        
        if results["status"] == "success":
            st.success("✅ Generated!")
            
            tab1, tab2, tab3 = st.tabs(["Final Flow", "DXF Analysis", "References"])
            
            with tab1:
                final = results.get("final", "")
                st.text_area("Process Flow", final, height=500)
                st.download_button("Download", final, 
                                  f"{Path(uploaded.name).stem}_flow.txt")
                
                # Verification
                verif = results.get("verification", {})
                if not verif.get("is_valid"):
                    st.warning(f"Issues: {', '.join(verif.get('issues', []))}")
            
            with tab2:
                dxf = results["dxf"]
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("CBS", dxf["cbs_type"])
                with col2:
                    st.metric("Induction", dxf["induction_type"])
                with col3:
                    st.metric("Components", dxf["total_components"])
                
                st.write("**Categories:**")
                for cat, cnt in sorted(dxf["category_summary"].items(), key=lambda x: -x[1]):
                    st.write(f"• {cat}: {cnt}")
                
                chute = dxf.get("chute_analysis", {})
                if chute.get("total"):
                    st.write(f"\n**Chutes:** {chute['total']} total")
                    for ct, cnt in chute.get("by_type", {}).items():
                        st.write(f"  • {ct.title()}: {cnt}")
            
            with tab3:
                refs = results.get("references", [])
                if refs:
                    for i, ref in enumerate(refs, 1):
                        display_score = ref.get("combined_score") or ref.get("score") or 0.0
                        with st.expander(f"{ref['client']} ({display_score:.3f})"):
                            st.code(ref["process_flow"][:1000], language="text")
                else:
                    st.info("No references found")
        else:
            st.error(f"Error: {results.get('error')}")

if __name__ == "__main__":
    main()

