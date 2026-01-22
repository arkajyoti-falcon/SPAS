# sys_desc_latest_streamlit.py
# Streamlit app: DXF + (optional) Costing Excel -> System Description (DOCX)
# Uses GROQ (openai/gpt-oss-120b) for detection + writing + judge pass.
# Inserts ONLY Conveyor BOQ table at [[CONVEYOR_BOQ_TABLE]] marker.
#
# CHANGES implemented (latest):
# 1) Induct subcomponent order (ONLY if found in DXF):
#    - Orientation / Loading Conveyor
#    - Buffer Conveyors
#    - Weighing Conveyor
#    - Intelligent Merge Conveyor
#
# 2) Infeed Conveyors can also have subcomponents (ONLY if found in DXF):
#    - Straight and Inclined conveyor
#    - Plastic Modular conveyor
#    - Curve conveyor
#    - Buffer conveyor  (can be in infeed OR induct; we classify to avoid duplicates)
#    - Alligning conveyor
#    - Belt merge
#
# 3) Conveyor BOQ: extract ONLY columns (in order):
#    S No. | Name | EL_1 | EL_2 | Length (m) | width (mm) | Set
#
# 4) DOCX: fixed images replace placeholders for induct subcomponents:
#    FIXED_IMAGE/weigh_conv.PNG
#    FIXED_IMAGE/buffer_conv.PNG
#    FIXED_IMAGE/oriant_conv.PNG
#    FIXED_IMAGE/merge_conv.PNG

import os
import re
import json
import tempfile
from pathlib import Path
from collections import Counter
from typing import Dict, Any, List, Tuple, Optional

import requests
import ezdxf
import streamlit as st
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from dotenv import load_dotenv

# New dynamic system description pipeline
from Supportive_Functions.costing_sheet_mapper import load_component_sheets
from Supportive_Functions.dynamic_system_description import (
    DXFComponent,
    generate_dynamic_system_description,
    load_catalog,
    normalize_name,
)

load_dotenv()
st.cache_data.clear()
# -----------------------------
# CONFIG
# -----------------------------
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_BASE_URL = os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1/chat/completions")
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")

TEMPLATE_PATH = os.getenv("CBS_TEMPLATE_PATH", "CBS_SYSTEM_DESC.txt")

UNITS = {
    0: "Unitless", 1: "inches", 2: "feet", 3: "miles",
    4: "millimeters", 5: "centimeters", 6: "meters", 7: "kilometers"
}

SCRIPT_DIR = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()
FIXED_IMAGE_DIR = SCRIPT_DIR.parent / "assests" / "Images"

# Fixed images for system description components
FIXED_IMAGE_MAP = {
    # Induct subcomponents
    "weighing conveyor": FIXED_IMAGE_DIR / "weigh_conv.PNG",
    "buffer conveyors": FIXED_IMAGE_DIR / "buffer_conv.PNG",
    "buffer conveyor": FIXED_IMAGE_DIR / "buffer_conv.PNG",
    "orientation / loading conveyor": FIXED_IMAGE_DIR / "oriant_conv.PNG",
    "orientation conveyor": FIXED_IMAGE_DIR / "oriant_conv.PNG",
    "orientation loading conveyor": FIXED_IMAGE_DIR / "oriant_conv.PNG",
    "intelligent merge conveyor": FIXED_IMAGE_DIR / "merge_conv.PNG",
    "merge conveyor": FIXED_IMAGE_DIR / "merge_conv.PNG",
    
    # Main Loop / CBS
    "main loop": FIXED_IMAGE_DIR / "CROS_BELT_SORTER.PNG",
    "main linear cbs": FIXED_IMAGE_DIR / "CROS_BELT_SORTER.PNG",
    "cross belt sorter": FIXED_IMAGE_DIR / "CROS_BELT_SORTER.PNG",
    "cbs": FIXED_IMAGE_DIR / "CROS_BELT_SORTER.PNG",
    
    # Carrier/Drive system
    "carrier wheel drive": FIXED_IMAGE_DIR / "CAREER_WHEEL.PNG",
    "carrier": FIXED_IMAGE_DIR / "CBS_CAREER.PNG",
    "friction wheel drive": FIXED_IMAGE_DIR / "FRINCTION_WHEEL_DRIVE.PNG",
    "linear motor drive": FIXED_IMAGE_DIR / "LINEAR_MOTOR_DRIVE.PNG",
    "servo roller": FIXED_IMAGE_DIR / "SERVO_ROLLER.PNG",
    
    # Infeed/Conveyors
    "infeed conveyor": FIXED_IMAGE_DIR / "buffer_conv.PNG",
    "infeed conveyors": FIXED_IMAGE_DIR / "buffer_conv.PNG",
    "straight and inclined conveyor": FIXED_IMAGE_DIR / "buffer_conv.PNG",
    "curve conveyor": FIXED_IMAGE_DIR / "buffer_conv.PNG",
}

CONVEYOR_BOQ_REQUIRED_COLS = ["S No.", "Name", "EL_1", "EL_2", "Length (m)", "width (mm)", "Set"]

# -----------------------------
# NOISE FILTERS (*U### / U### etc)
# -----------------------------
def is_noise_block(name: str) -> bool:
    if not name:
        return True
    n = name.strip()
    if re.match(r"^\*?U\d+$", n, flags=re.IGNORECASE):
        return True
    if re.match(r"^\*[UDXATE]\d+$", n, flags=re.IGNORECASE):
        return True
    if n.startswith("*") or n.startswith("~") or n.startswith("A$C"):
        return True
    return False


def safe_int(x) -> Optional[int]:
    try:
        return int(x)
    except Exception:
        return None


def load_template_text() -> str:
    if os.path.exists(TEMPLATE_PATH):
        with open(TEMPLATE_PATH, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    return ""

# -----------------------------
# DXF FULL EXTRACTION (helpers)
# -----------------------------
def _clean_text_snip(s: str, max_len: int = 80) -> str:
    s = re.sub(r"\s+", " ", (s or "").strip())
    return s[:max_len] if s else ""


def _get_doc_extents(doc) -> Optional[Dict[str, List[float]]]:
    try:
        extmin = doc.header.get("$EXTMIN")
        extmax = doc.header.get("$EXTMAX")
        if extmin is not None and extmax is not None:
            mn = extmin.value
            mx = extmax.value
            if mn and mx:
                return {
                    "min": [float(mn.x), float(mn.y), float(getattr(mn, "z", 0.0))],
                    "max": [float(mx.x), float(mx.y), float(getattr(mx, "z", 0.0))]
                }
    except Exception:
        pass
    return None


def _collect_space_signals(space, top_inserts: Counter, layers: Counter, text_snips: Counter, entity_types: Counter):
    for e in space:
        try:
            et = e.dxftype()
            entity_types[et] += 1
        except Exception:
            pass

        try:
            layer = getattr(e.dxf, "layer", None)
            if layer:
                layers[str(layer)] += 1
        except Exception:
            pass

        try:
            if e.dxftype() == "INSERT":
                nm = getattr(e.dxf, "name", "") or ""
                if nm and not is_noise_block(nm):
                    top_inserts[nm] += 1
        except Exception:
            pass

        try:
            if e.dxftype() == "TEXT":
                t = getattr(e.dxf, "text", "") or ""
                t = _clean_text_snip(t)
                if t:
                    text_snips[t] += 1
            elif e.dxftype() == "MTEXT":
                t = getattr(e, "text", "") or ""
                t = _clean_text_snip(t)
                if t:
                    text_snips[t] += 1
        except Exception:
            pass


def _expand_nested_inserts(doc, top_inserts: Counter) -> Counter:
    nested = Counter()
    visited = set()

    def walk_block(block_name: str, mult: int, depth: int = 0):
        if not block_name or is_noise_block(block_name):
            return
        if depth > 10:
            return
        key = (block_name, depth)
        if key in visited:
            return
        visited.add(key)

        try:
            blk = doc.blocks.get(block_name)
        except Exception:
            return
        if blk is None:
            return

        for e in blk:
            try:
                if e.dxftype() == "INSERT":
                    nm = getattr(e.dxf, "name", "") or ""
                    if nm and not is_noise_block(nm):
                        nested[nm] += int(mult)
                        walk_block(nm, mult, depth + 1)
            except Exception:
                continue

    for nm, cnt in top_inserts.items():
        walk_block(nm, int(cnt), 0)

    return nested


# -----------------------------
# DXF FULL EXTRACTION (kept name/signature)
# -----------------------------
def _collect_from_space(space, top_inserts: Counter, layers: Counter, text_snips: Counter, entity_types: Counter):
    # Previously placeholder. Kept same signature; now delegates to signal collector.
    _collect_space_signals(space, top_inserts, layers, text_snips, entity_types)


def extract_dxf_full_json(dxf_path: Path) -> Dict[str, Any]:
    try:
        doc = ezdxf.readfile(str(dxf_path))
    except Exception:
        doc, _ = ezdxf.recover.readfile(str(dxf_path))

    units_code = getattr(doc, "units", 0) or 0
    units_name = UNITS.get(int(units_code), "unknown")

    top_inserts = Counter()
    layers = Counter()
    text_snips = Counter()
    entity_types = Counter()

    try:
        _collect_from_space(doc.modelspace(), top_inserts, layers, text_snips, entity_types)
    except Exception:
        pass

    try:
        for layout in doc.layouts:
            try:
                if layout.name.lower() != "model":
                    _collect_from_space(layout, top_inserts, layers, text_snips, entity_types)
            except Exception:
                continue
    except Exception:
        pass

    nested_inserts = _expand_nested_inserts(doc, top_inserts)

    block_defs = []
    try:
        for blk in doc.blocks:
            bn = getattr(blk, "name", "") or ""
            if bn and not is_noise_block(bn):
                block_defs.append(bn)
    except Exception:
        pass

    extents = _get_doc_extents(doc) or {"min": [], "max": []}

    return {
        "file": str(dxf_path),
        "units_code": int(units_code),
        "units_name": units_name,
        "extents": extents,
        "top_inserts": dict(top_inserts),
        "nested_inserts": dict(nested_inserts),
        "layers": dict(layers),
        "text_snips": dict(text_snips),
        "block_definitions": block_defs,
        "entity_types": dict(entity_types),
    }


def _resolve_stage_from_catalog(name: str, catalog: Dict[str, Any]) -> Optional[str]:
    """Return stage for a component name using catalog exact/alias lookup."""
    norm_name = normalize_name(name)
    if norm_name in catalog:
        return catalog[norm_name].get("stage")

    for base_name, data in catalog.items():
        aliases = [normalize_name(a) for a in data.get("aliases", [])]
        if norm_name in aliases:
            return data.get("stage")
    return None


def build_dxf_components(full_dxf: Dict[str, Any], catalog: Dict[str, Any]) -> List[DXFComponent]:
    """Build DXFComponent list from DXF extraction using catalog for stage mapping."""
    components: List[DXFComponent] = []
    seen: set = set()

    def _maybe_add(raw_name: str):
        if not raw_name:
            return
        if is_noise_block(raw_name):
            return
        norm = normalize_name(raw_name)
        if not norm or norm in seen:
            return
        stage = _resolve_stage_from_catalog(norm, catalog)
        components.append(DXFComponent(name=raw_name, stage=stage, raw=raw_name))
        seen.add(norm)

    for name in (full_dxf.get("nested_inserts") or {}).keys():
        _maybe_add(name)
    for name in (full_dxf.get("top_inserts") or {}).keys():
        _maybe_add(name)
    for name in (full_dxf.get("block_definitions") or []):
        _maybe_add(name)
    for name in (full_dxf.get("layers") or {}).keys():
        _maybe_add(name)
    for name in (full_dxf.get("text_snips") or {}).keys():
        _maybe_add(name)

    return components


def _search_counts_multi(full: Dict[str, Any], patterns: List[str]) -> int:
    rx = [re.compile(p, flags=re.IGNORECASE) for p in patterns]

    nested = full.get("nested_inserts", {}) or {}
    top = full.get("top_inserts", {}) or {}
    layers = full.get("layers", {}) or {}
    texts = full.get("text_snips", {}) or {}
    block_defs = full.get("block_definitions", []) or []

    total = 0

    for name, cnt in nested.items():
        try:
            if any(r.search(str(name)) for r in rx):
                total += int(cnt)
        except Exception:
            continue

    for name, cnt in top.items():
        try:
            if any(r.search(str(name)) for r in rx):
                total += int(cnt)
        except Exception:
            continue

    if total == 0:
        for lname in layers.keys():
            if any(r.search(str(lname)) for r in rx):
                return 1
        for t in texts.keys():
            if any(r.search(str(t)) for r in rx):
                return 1
        for bn in block_defs:
            if any(r.search(str(bn)) for r in rx):
                return 1

    return total


def extract_angle_degrees(d: Dict[str, int]) -> List[int]:
    degrees = set()
    for name in (d or {}).keys():
        s = str(name).lower()

        for m in re.findall(r"(\d{1,3})\s*¬∞", s):
            degrees.add(int(m))

        for m in re.findall(r"(\d{1,3})\s*deg", s):
            degrees.add(int(m))
        for m in re.findall(r"_(\d{1,3})_deg", s):
            degrees.add(int(m))

        for m in re.findall(r"(\d{1,3})\s*deg\s*turn", s):
            degrees.add(int(m))

    return sorted(degrees)


def compute_dxf_metrics(full: Dict[str, Any]) -> Dict[str, Any]:
    units_name = full.get("units_name") or "unknown"
    fname = (full.get("file") or "").lower()

    cbs_type = "Linear CBS" if "linear" in fname else "Loop CBS"

    feedline_count = _search_counts_multi(full, [
        r"feedline", r"feed\s*line", r"fal.*feed", r"\bfs\d{3,4}\b", r"\bfs0\d+\b", r"induct"
    ])

    fs002_without_weighing = _search_counts_multi(full, [
        r"fal[_\-\s]*fs002v02", r"\bfs002v02\b", r"without\s*weigh"
    ])

    infeed_telescopic = _search_counts_multi(full, [r"telescopic", r"telescopico", r"\btbc\b"])
    infeed_generic = _search_counts_multi(full, [
        r"infeed", r"ingresso", r"receiving", r"highway", r"singulat", r"pvc\s*belt"
    ])

    degrees = extract_angle_degrees(full.get("nested_inserts", {}) or {})

    # VDS Loop detection - STRICT: only if explicit VDS pattern is found
    # Removed patterns like "fal[_\-\s]*f001", "vipacsystem", "return\s*line" which cause false positives
    vds_loop = _search_counts_multi(full, [
        r"\bvds\b(?!.*chute)",  # VDS but not VDS chute
        r"vds\s*loop",
        r"vds\s*conveyor",
        r"distribution\s*loop",
        r"\bvds[_\-\s]*return\b",
    ])
    has_vds_loop = vds_loop > 0

    recirc_count = _search_counts_multi(full, [
        r"recirculation", r"recirculate", r"refeed", r"return.*conv", r"loop.*back", r"re[-\s]*cir"
    ])
    has_recirculation = recirc_count > 0

    rejection_chute_count = _search_counts_multi(full, [r"reject", r"rejection", r"sortfail", r"exception"])
    dispersion_chute_count = _search_counts_multi(full, [r"disperson", r"dispersion"])
    collection_chute_count = _search_counts_multi(full, [r"collection", r"friction", r"accumulation"])
    gravity_chute_count = _search_counts_multi(full, [r"\bgravity\b"])
    mini_gravity_count = _search_counts_multi(full, [r"mini.*gravity", r"chutes\$0\$mini\s*gravity"])
    direct_bagging_chute_count = _search_counts_multi(full, [r"direct.*bagging", r"bagging.*chute", r"bag.*chute"])
    sliding_chute_count = _search_counts_multi(full, [r"sliding", r"slide.*chute"])
    secondary_chute_count = _search_counts_multi(full, [r"secondary", r"l[-\s]*type"])
    bulk_chute_count = _search_counts_multi(full, [r"\bbulk\b"])
    # Generic chute count (for fallback when no specific type detected)
    generic_chute_count = _search_counts_multi(full, [r"\bchute\b"])

    scan_count = _search_counts_multi(full, [
        r"scanner", r"\bscan\b", r"barcode", r"\bdws\b", r"dimension", r"reader"
    ])
    weighing_signal = _search_counts_multi(full, [r"weigh", r"weight", r"scale"])
    has_scanner = (scan_count > 0) or (weighing_signal > 0)

    manual_station_count = _search_counts_multi(full, [
        r"manual", r"operator", r"workstation", r"induct.*station"
    ])
    has_manual = manual_station_count > 0

    # Signals used for induct / infeed classification
    loading_sig = _search_counts_multi(full, [r"\bloading\b", r"\bspacing\b", r"\bspacer\b", r"gap\s*optimizer"])
    buffer_sig = _search_counts_multi(full, [r"\bbuffer\b", r"accumulation"])
    merge_sig = _search_counts_multi(full, [r"intelligent\s*merge", r"angle\s*merge", r"\bmerge\b", r"fal[_\-\s]*f012", r"\bf012\b"])
    orientation_sig = _search_counts_multi(full, [r"\borientation\b", r"\borient\b", r"orientat", r"oriant"])
    belt_merge_sig = _search_counts_multi(full, [r"belt\s*merge"])

    return {
        "UNITS": units_name,
        "TYPE OF CBS": cbs_type,

        "FEEDLINE COUNT": feedline_count,
        "FS002 WITHOUT WEIGHING COUNT": fs002_without_weighing,

        "TEL. BELT CONVEYOR COUNT": infeed_telescopic,
        "INFEED CONVEYOR COUNT": infeed_generic,

        "VDS LOOP COUNT": vds_loop,
        "HAS_VDS_LOOP": has_vds_loop,

        "HAS_REIRCULATION": has_recirculation,
        "RECIRCULATION COUNT": recirc_count,

        "HAS_MANUAL_INDUCT": has_manual,
        "MANUAL STATION COUNT": manual_station_count,

        "HAS_SCANNER": has_scanner,
        "SCANNER SIGNAL COUNT": scan_count,
        "WEIGHING SIGNAL COUNT": weighing_signal,

        "COUNT OF REJECTION CHUTE": rejection_chute_count,
        "COUNT OF DISPERSION CHUTE": dispersion_chute_count,
        "COUNT OF COLLECTION CHUTE": collection_chute_count,

        "GRAVITY CHUTE COUNT": gravity_chute_count,
        "MINI GRAVITY CHUTE COUNT": mini_gravity_count,
        "DIRECT BAGGING CHUTE COUNT": direct_bagging_chute_count,
        "SLIDING CHUTE COUNT": sliding_chute_count,
        "SECONDARY CHUTE COUNT": secondary_chute_count,
        "BULK CHUTE COUNT": bulk_chute_count,
        "GENERIC CHUTE COUNT": generic_chute_count,

        "DEGREE OF ANGLE MERGE": ", ".join(str(x) for x in degrees) if degrees else "",

        # raw signals (for debugging / optional use)
        "SIG_LOADING": loading_sig > 0,
        "SIG_BUFFER": buffer_sig > 0,
        "SIG_ORIENTATION": orientation_sig > 0,
        "SIG_MERGE": (merge_sig > 0) or bool(degrees),
        "SIG_BELT_MERGE": belt_merge_sig > 0,
    }


def build_detected_json(metrics: Dict[str, Any]) -> Dict[str, Any]:
    """
    Deterministic fallback detected JSON (present-only).
    NOTE: We will further enforce induct/infeed subcomponents strictly from DXF later.
    """
    det: Dict[str, Any] = {}

    # Infeed System (base)
    infeed: Dict[str, Any] = {}
    if metrics.get("TEL. BELT CONVEYOR COUNT", 0) > 0:
        infeed["Telescopic Belt Conveyor"] = int(metrics["TEL. BELT CONVEYOR COUNT"])
    if metrics.get("INFEED CONVEYOR COUNT", 0) > 0:
        infeed["Infeed Conveyors"] = int(metrics["INFEED CONVEYOR COUNT"])
    if metrics.get("HAS_VDS_LOOP"):
        infeed["VDS Loop Conveyor"] = max(1, int(metrics.get("VDS LOOP COUNT", 1)))
    if infeed:
        det["Infeed System"] = infeed

    # Inducts
    if metrics.get("FEEDLINE COUNT", 0) > 0:
        feedlines: Dict[str, Any] = {"Feedline Count": int(metrics["FEEDLINE COUNT"]), "Subcomponents": {}}
        det["Parcel Inducts / Induction to Sorter"] = {"Feedlines": feedlines}

    if metrics.get("HAS_MANUAL_INDUCT"):
        det.setdefault("Parcel Inducts / Induction to Sorter", {})
        det["Parcel Inducts / Induction to Sorter"]["Manual Induct Stations"] = {
            "Manual Induct Station Count": int(metrics.get("MANUAL STATION COUNT", 1))
        }

    # CBS
    det["CBS"] = {"Type": metrics.get("TYPE OF CBS", "")}

    # Barcode scanning
    if metrics.get("HAS_SCANNER"):
        det["Barcode Scanning System"] = {"Present": True}

    # Output Chutes (counts)
    out: Dict[str, Any] = {}
    if metrics.get("MINI GRAVITY CHUTE COUNT", 0) > 0:
        out["Mini-Gravity Chutes"] = int(metrics["MINI GRAVITY CHUTE COUNT"])
    if metrics.get("GRAVITY CHUTE COUNT", 0) > 0:
        out["Gravity Chutes"] = int(metrics["GRAVITY CHUTE COUNT"])
    if metrics.get("COUNT OF COLLECTION CHUTE", 0) > 0:
        out["Collection Chutes"] = int(metrics["COUNT OF COLLECTION CHUTE"])
    if metrics.get("COUNT OF DISPERSION CHUTE", 0) > 0:
        out["Dispersion Chutes"] = int(metrics["COUNT OF DISPERSION CHUTE"])
    if metrics.get("COUNT OF REJECTION CHUTE", 0) > 0:
        out["Rejection Chutes"] = int(metrics["COUNT OF REJECTION CHUTE"])
    if metrics.get("DIRECT BAGGING CHUTE COUNT", 0) > 0:
        out["Direct Bagging Chutes"] = int(metrics["DIRECT BAGGING CHUTE COUNT"])
    if metrics.get("SLIDING CHUTE COUNT", 0) > 0:
        out["Sliding Chutes"] = int(metrics["SLIDING CHUTE COUNT"])
    if metrics.get("SECONDARY CHUTE COUNT", 0) > 0:
        out["Secondary Chutes"] = int(metrics["SECONDARY CHUTE COUNT"])
    if metrics.get("BULK CHUTE COUNT", 0) > 0:
        out["Bulk Chutes"] = int(metrics["BULK CHUTE COUNT"])
    # Fallback: if no specific chutes found but generic chute exists, add as "Output Chutes"
    if not out and metrics.get("GENERIC CHUTE COUNT", 0) > 0:
        out["Chutes"] = int(metrics["GENERIC CHUTE COUNT"])
    if out:
        det["Output Chutes"] = out

    # Recirculation (kept as-is)
    if metrics.get("HAS_REIRCULATION"):
        det["Recirculation"] = {"Recirculation Count": int(metrics.get("RECIRCULATION COUNT", 1))}

    return det

# -----------------------------
# COSTING EXCEL EXTRACTION
# -----------------------------
def _normalize(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).lower()


def read_block_table(ws, header_row: int, ncols: int = None, max_rows: int = 200) -> List[List[str]]:
    table: List[List[str]] = []
    hdr: List[str] = []

    if ncols is None:
        ncols = 1
        for c in range(1, 50):
            v = ws.cell(header_row, c).value
            if v is not None and str(v).strip() != "":
                ncols = c
            elif ncols > 1 and v is None:
                break

    for c in range(1, ncols + 1):
        v = ws.cell(header_row, c).value
        hdr.append("" if v is None else str(v).strip())
    table.append(hdr)

    for r in range(header_row + 1, header_row + 1 + max_rows):
        v0 = ws.cell(r, 1).value
        if v0 is None or str(v0).strip() == "":
            break
        row: List[str] = []
        for c in range(1, ncols + 1):
            v = ws.cell(r, c).value
            row.append("" if v is None else str(v).strip())
        table.append(row)

    return table


def extract_costing_tables(xlsx_path: Path) -> Dict[str, List[List[str]]]:
    wb = load_workbook(filename=str(xlsx_path), data_only=True)
    tables: Dict[str, List[List[str]]] = {}

    if "Conveyors" in wb.sheetnames:
        ws = wb["Conveyors"]
        tables["Conveyor BOQ"] = read_block_table(ws, header_row=2, ncols=None)

        full = tables["Conveyor BOQ"]
        hdr = full[0] if full else []
        name_idx = hdr.index("Name") if "Name" in hdr else 1
        bag_rows = [hdr] if hdr else []
        sno = 1
        for row in full[1:]:
            nm = (row[name_idx] or "").lower() if name_idx < len(row) else ""
            if "bagging" in nm:
                row2 = row[:]
                if row2:
                    row2[0] = str(sno)
                sno += 1
                bag_rows.append(row2)
        if len(bag_rows) > 1:
            tables["Bagging Conveyor BOQ"] = bag_rows

    if "Conveyor BOQ" not in tables:
        for sname in wb.sheetnames:
            if "convey" in sname.lower():
                ws = wb[sname]
                for hr in range(1, 11):
                    row_vals = [ws.cell(hr, c).value for c in range(1, 10)]
                    row_txt = " | ".join([str(v).strip().lower() for v in row_vals if v is not None])
                    if "name" in row_txt and ("qty" in row_txt or "quantity" in row_txt):
                        tables["Conveyor BOQ"] = read_block_table(ws, header_row=hr, ncols=7)
                        break
            if "Conveyor BOQ" in tables:
                break

    if "Inducts" in wb.sheetnames:
        tables["Feedlines BOQ"] = read_block_table(wb["Inducts"], header_row=2, ncols=7)

    if "Bag Sorter Induct" in wb.sheetnames:
        tables["Bag Induct BOQ"] = read_block_table(wb["Bag Sorter Induct"], header_row=2, ncols=7)

    if "Destinations" in wb.sheetnames:
        tables["Output Destinations"] = read_block_table(wb["Destinations"], header_row=2, ncols=6)

    if "Weighing Conveyors" in wb.sheetnames:
        tables["Weighing Conveyor BOQ"] = read_block_table(wb["Weighing Conveyors"], header_row=2, ncols=7)

    return tables


def extract_costing_values(xlsx_path: Path) -> Dict[str, str]:
    wb = load_workbook(filename=str(xlsx_path), data_only=True)

    ws = None
    if "Loop CBS" in wb.sheetnames:
        ws = wb["Loop CBS"]
    else:
        for s in wb.sheetnames:
            if "loop cbs" in s.lower():
                ws = wb[s]
                break
    if ws is None:
        ws = wb[wb.sheetnames[0]]

    carrier_pitch = ""
    sorter_height = ""
    sorter_speed = ""
    belt_speed = ""
    throughput = ""
    total_chutes = ""
    feedline_count = ""

    for r in range(1, ws.max_row + 1):
        b = ws.cell(r, 2).value
        b_str = str(b).strip().lower() if b else ""
        
        if b_str == "carrier pitch":
            v = ws.cell(r, 3).value
            if v is not None:
                carrier_pitch = str(v).strip()
                
        elif "select the sorter height" in b_str:
            v = ws.cell(r, 4).value or ws.cell(r, 3).value
            if v is not None:
                sorter_height = str(v).strip()
                
        elif "sorter speed" in b_str or "carrier speed" in b_str:
            v = ws.cell(r, 3).value or ws.cell(r, 4).value
            if v is not None:
                sorter_speed = str(v).strip()
                
        elif "belt speed" in b_str:
            v = ws.cell(r, 3).value or ws.cell(r, 4).value
            if v is not None:
                belt_speed = str(v).strip()
                
        elif "throughput" in b_str and "per hour" in b_str:
            v = ws.cell(r, 3).value or ws.cell(r, 4).value
            if v is not None:
                throughput = str(v).strip()
                
        elif "total chutes" in b_str or "number of chutes" in b_str:
            v = ws.cell(r, 3).value or ws.cell(r, 4).value
            if v is not None:
                total_chutes = str(v).strip()
                
        elif "feedline" in b_str and "count" in b_str:
            v = ws.cell(r, 3).value or ws.cell(r, 4).value
            if v is not None:
                feedline_count = str(v).strip()

    return {
        "COSTING_SHEET_NAME": ws.title,
        "CBS HEIGHT FROM GROUND": sorter_height,
        "PITCH LENGTH": carrier_pitch,
        "SORTER SPEED": sorter_speed,
        "BELT SPEED": belt_speed,
        "THROUGHPUT": throughput,
        "TOTAL CHUTES": total_chutes,
        "FEEDLINE COUNT": feedline_count,
    }

# -----------------------------
# Conveyor BOQ filter (ONLY requested columns)
# -----------------------------
def filter_conveyor_boq_columns(table: List[List[str]]) -> List[List[str]]:
    if not table or len(table) < 1:
        return [CONVEYOR_BOQ_REQUIRED_COLS]

    hdr = table[0]
    norm_to_idx = {}
    for i, h in enumerate(hdr):
        nh = _normalize(h)
        if nh:
            norm_to_idx[nh] = i

    def idx_for(col: str) -> Optional[int]:
        c = _normalize(col)

        aliases = [c]
        if c in ("s no.", "s no", "sno", "sr no", "sr. no", "s.no", "s.no."):
            aliases += ["s no", "s no.", "sno", "sr no", "sr. no", "s.no", "s.no.", "s#"]
        if c == "length (m)":
            aliases += ["length", "length(m)", "length m", "len (m)", "len(m)", "l (m)"]
        if c == "width (mm)":
            aliases += ["width", "width(mm)", "width mm", "w (mm)", "w(mm)"]
        if c == "el_1":
            aliases += ["el1", "el 1", "elev 1", "elevation 1"]
        if c == "el_2":
            aliases += ["el2", "el 2", "elev 2", "elevation 2"]
        if c == "set":
            aliases += ["set", "sets"]
        if c == "name":
            aliases += ["name", "item", "description"]

        for a in aliases:
            if a in norm_to_idx:
                return norm_to_idx[a]

        for nh, ii in norm_to_idx.items():
            if c in nh:
                return ii
        return None

    desired_idxs = [idx_for(c) for c in CONVEYOR_BOQ_REQUIRED_COLS]

    out = [CONVEYOR_BOQ_REQUIRED_COLS]
    for row in table[1:]:
        new_row = []
        for ii in desired_idxs:
            if ii is None or ii >= len(row):
                new_row.append("")
            else:
                new_row.append(row[ii])
        out.append(new_row)
    return out

# -----------------------------
# GROQ CALL
# -----------------------------
def groq_chat(messages: List[Dict[str, str]], temperature: float = 0.1, max_tokens: int = 6000) -> str:
    if not GROQ_API_KEY:
        raise RuntimeError("GROQ_API_KEY is not set.")

    headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
    body = {
        "model": GROQ_MODEL,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }

    r = requests.post(GROQ_BASE_URL, headers=headers, data=json.dumps(body), timeout=180)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]


def extract_json(txt: str) -> Dict[str, Any]:
    if not txt:
        return {}
    txt = txt.strip()
    try:
        return json.loads(txt)
    except Exception:
        pass
    a = txt.find("{")
    b = txt.rfind("}")
    if a != -1 and b != -1 and b > a:
        try:
            return json.loads(txt[a:b + 1])
        except Exception:
            return {}
    return {}


def normalize_detected(d: Dict[str, Any], fallback: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(d, dict) or not d:
        return fallback

    if "detected" in d and isinstance(d["detected"], dict):
        d = d["detected"]

    has_any = any(k in d for k in ["Infeed System", "CBS", "Parcel Inducts / Induction to Sorter", "Output Chutes"])
    if not has_any:
        return fallback

    return d

# -----------------------------
# PROMPTS
# -----------------------------
def prompt_detect_components(full_dxf: Dict[str, Any], metrics: Dict[str, Any]) -> List[Dict[str, str]]:
    system = r"""
ROLE
You are a DXF component detector for Cross-Belt Sorter (CBS) systems.

INPUTS YOU WILL RECEIVE
- DXF FULL JSON: insert names + nested inserts + layers + text snippets + block definitions.
- COMPUTED METRICS: precomputed counts/presence hints.

YOUR TASK
Return a single JSON object listing ONLY the components/subcomponents that are PRESENT in the system, with best-effort counts.

STRICT OUTPUT RULES
- Output MUST be a single VALID JSON object (no markdown, no commentary).
- Include ONLY PRESENT items. Do NOT output "Not present".
- Use the CANONICAL OUTPUT SCHEMA exactly as below.
- For feedline subcomponents, set count = 1 (module presence) unless DXF clearly indicates >1.
- Enforce mandatory mappings:
  1) If any block/text/layer contains "FAL_FS002V02" AND "(Without weighing)" OR "without weigh",
     then detect Feedlines -> Buffer Conveyor as present (count=1).
  2) If any block/text/layer indicates "VDS" or "distribution loop" OR matches common VDS patterns,
     then detect Infeed System -> VDS Loop Conveyor as present (count>=1).
  3) Detect CBS type:
     - If DXF indicates "Linear" -> Type="Linear CBS" else if "Loop" -> Type="Loop CBS".

CANONICAL OUTPUT SCHEMA (ONLY THESE KEYS)
{
  "Infeed System": {
    "<subcomponent name>": <count>
  },
  "Parcel Inducts / Induction to Sorter": {
    "Feedlines": {
      "Feedline Count": <int>,
      "Subcomponents": {
        "<subcomponent name>": <int>
      },
      "Angle Merge Degrees": "<string optional>"
    },
    "Manual Induct Stations": {
      "Manual Induct Station Count": <int>
    }
  },
  "CBS": { "Type": "Loop CBS or Linear CBS" },
  "Barcode Scanning System": { "Present": true },
  "Output Chutes": {
    "<chute type>": <int>
  },
  "Recirculation": { "Recirculation Count": <int> }
}

NOTES
- You may omit any top-level key if nothing is present for it.
- Prefer counts from nested inserts when available.
- If you only have presence evidence from layers/text/definitions, use count=1.

NOW RETURN THE JSON ONLY.
"""
    user = f"""DXF FULL JSON:
{json.dumps(full_dxf, ensure_ascii=False)}

COMPUTED METRICS:
{json.dumps(metrics, ensure_ascii=False)}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


# Kept as-is (older fixed-flow prompt). Not used in new pipeline.
def prompt_generate_system_description(template_text: str,
                                       detected: Dict[str, Any],
                                       variables: Dict[str, str]) -> List[Dict[str, str]]:
    system = """
ROLE
You are a senior solution engineer writing the ‚ÄúSystem Description‚Äù section for a Cross-Belt Sorter (CBS) proposal.

INPUTS YOU WILL RECEIVE
1) TEMPLATE TEXT (canonical wording library; may contain [VAR] placeholders)
2) DETECTED JSON (the single source of truth for which components/subcomponents are PRESENT and their counts/types)
3) VARIABLES MAP (values for placeholders; some may be missing)

PRIMARY OBJECTIVE
Generate the final ‚ÄúSystem Description‚Äù using the TEMPLATE TEXT as the base wording, but include ONLY those sections and subcomponents that are PRESENT in DETECTED JSON.

NON-NEGOTIABLE HARD RULES
1) TEMPLATE ANCHORING
- Use TEMPLATE TEXT as the canonical description wording.
- Keep descriptions the SAME in meaning and style as the template.
- You may reorder sentences slightly ONLY to fit the fixed flow or to insert detected counts.
- Do NOT introduce new claims/specifications not supported by template or detected JSON.

2) PRESENCE FILTER
- Include ONLY components/subcomponents that exist in DETECTED JSON.
- Never mention missing components.
- Never write ‚ÄúNot present / Not available / not detected‚Äù.

3) PLACEHOLDERS
- Replace placeholders like [VAR] using VARIABLES MAP.
- If any placeholder is missing, keep it exactly as [VAR] (do not guess).
- Do not invent numbers or values.

4) IMAGE PLACEHOLDERS
- After each component/subcomponent description block, add on a new line:
  [IMAGE PLACEHOLDER: <exact component/subcomponent name>]

5) OUTPUT FORMAT
- Output plain text only. No JSON. No Markdown.

FIXED FLOW (MUST FOLLOW EXACT ORDER)
1. Infeed System
2. Induction to Sorter
   2.1 Feedlines
   2.2 Manual Induct Stations (if present)
3. Loop CBS / Linear CBS
4. Barcode Scanning System (if present)
5. Output Chutes
6. Exception Handling Area (if present)
7. Recirculation & Manual feedline (if present)
"""
    user = f"""TEMPLATE TEXT:
{template_text}

DETECTED JSON:
{json.dumps(detected, ensure_ascii=False)}

VARIABLES (already resolved):
{json.dumps(variables, ensure_ascii=False)}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def prompt_generate_system_description_dynamic(template_text: str,
                                               detected: Dict[str, Any],
                                               variables: Dict[str, str]) -> List[Dict[str, str]]:
    system = r"""
ROLE
You are a senior solution engineer writing the ‚ÄúSystem Description‚Äù section for a Cross-Belt Sorter (CBS) proposal.

INPUTS
1) TEMPLATE TEXT - Contains EXACT descriptions to use for each component. Look for lines starting with "# Component Name -"
2) DETECTED JSON (ONLY source of truth for which components are present)
3) VARIABLES MAP

üö®üö®üö® CRITICAL INSTRUCTION - NO PARAPHRASING ALLOWED üö®üö®üö®
YOU ARE A TEXT COPIER, NOT A TEXT WRITER
- Your ONLY job is to COPY text from TEMPLATE TEXT exactly as written
- Find the component section in TEMPLATE TEXT (marked with "# Component Name -")
- Copy EVERY WORD, EVERY SENTENCE, EVERY PARAGRAPH exactly as it appears
- Replace ONLY [PLACEHOLDER] values with actual numbers from VARIABLES MAP
- FORBIDDEN EXAMPLES:
  ‚ùå Changing "Each Feedline has the following Conveyor Modules:" to "The feedlines consist of..."
  ‚ùå Changing "‚Ä¢ Loading Conveyor- 7 No." to "seven loading conveyors"
  ‚úÖ Copy EXACTLY: "‚Ä¢ Loading Conveyor- 7 No."
- REQUIRED: Copy the EXACT formatting including bullets (‚Ä¢), line breaks, and numbering

ABSOLUTE RULE - COPY DESCRIPTIONS VERBATIM (THIS IS THE MOST IMPORTANT RULE)
- For EVERY component, you MUST find its description in TEMPLATE TEXT and COPY IT WORD-FOR-WORD
- The template has multi-line, multi-paragraph descriptions - USE THE ENTIRE TEXT
- DO NOT SUMMARIZE. DO NOT PARAPHRASE. DO NOT SHORTEN. DO NOT REWRITE IN YOUR OWN WORDS.
- ONLY replace placeholders like [COUNT], [CLIENT'S NAME], [PITCH LENGTH], etc. with values from VARIABLES MAP
- If a template description has 5 lines, your output must have those same 5 lines

EXAMPLE OF WHAT YOU MUST DO:
Template has for "# Buffer Conveyor-":
"A buffer conveyor, also known as a buffering conveyor or accumulation conveyor, is a type of conveyor system used to temporarily store or hold items in a controlled manner. Its primary purpose is to manage the flow of items between different stages of a production or handling process when there is a mismatch in the speeds or capacities of the upstream and downstream equipment. This Conveyors required to maintain the Throughput of Line."

Your output MUST be:
**2.2 Buffer Conveyor**
A buffer conveyor, also known as a buffering conveyor or accumulation conveyor, is a type of conveyor system used to temporarily store or hold items in a controlled manner. Its primary purpose is to manage the flow of items between different stages of a production or handling process when there is a mismatch in the speeds or capacities of the upstream and downstream equipment. This Conveyors required to maintain the Throughput of Line.
[IMAGE PLACEHOLDER: Buffer Conveyor]

WRONG (DO NOT DO THIS):
**2.2 Buffer Conveyor**
Temporarily stores parcels to balance flow between upstream and downstream equipment.
[IMAGE PLACEHOLDER: Buffer Conveyor]

The WRONG example above summarizes the description - THIS IS FORBIDDEN.

CRITICAL FORMATTING RULES

1) HEADING FORMAT:
   - Main section headings: Use format "<N>. <Title>" where N is auto-incremented (1, 2, 3...)
   - Subsection headings: Use format "<N.M> <Title>" where M is auto-incremented within the section
   - Sub-subsection headings: Use format "<N.M.P> <Title>" where P is auto-incremented
   - IMPORTANT: Number sections sequentially based on what is ACTUALLY PRESENT, not hardcoded numbers
   - Example: If only "Infeed Conveyors" and "VDS Loop Conveyor" exist under Infeed System,
     they become 1.1 and 1.2 (NOT 1.1 and 1.3)

2) HEADING STYLE:
   - ALL headings must be prefixed with **bold markers**: **<N>. <Title>** or **<N.M> <Title>**
   - This applies to ALL section, subsection, and sub-subsection headings

HARD OUTPUT STRUCTURE (USE ONLY PRESENT ITEMS, AUTO-NUMBER SEQUENTIALLY)

**1. Infeed System**
  **1.1 Infeed Conveyors** (if DETECTED JSON -> Infeed System contains any infeed conveyors)
      - If the following infeed subcomponents exist in DETECTED JSON -> Infeed System, add them
        in this exact order (ONLY if present), numbered sequentially:
        - Straight and Inclined conveyor -> use "# Straight and Inclined Powered Belt Conveyor" from TEMPLATE TEXT
        - Plastic Modular conveyor -> COPY VERBATIM from "# Plastic Modular Conveyor" in TEMPLATE TEXT
        - Curve conveyor -> COPY VERBATIM from "# Curve Conveyor" in TEMPLATE TEXT
        - Buffer conveyor -> COPY VERBATIM from "# Buffer Conveyor" in TEMPLATE TEXT
        - Alligning conveyor -> COPY VERBATIM from "# Aligning Conveyor" in TEMPLATE TEXT
        - Belt merge -> COPY VERBATIM from "# Belt Merge" in TEMPLATE TEXT
  **1.x VDS Loop Conveyor** (if present) -> COPY VERBATIM from "# VDS Loop" in TEMPLATE TEXT
  **1.x Conveyor BOQ** (MUST be the LAST subsection under Infeed System)
      [[CONVEYOR_BOQ_TABLE]]

**2. Induct**
  - This section is derived from:
    DETECTED JSON -> "Parcel Inducts / Induction to Sorter"
  - CRITICAL - USE EXACT TEXT FROM TEMPLATE (NO PARAPHRASING):
    Step 1: Copy the introductory paragraph from "# Induction to Sorter: Feedlines (variable)" section
      - If VDS exists in DETECTED JSON -> Infeed System, use the paragraph after "! If VDS is present -"
      - If VDS does NOT exist, use the paragraph after "! If VDS is not present -"
      - COPY EVERY WORD EXACTLY - do not rewrite or paraphrase
    
    Step 2: Copy the "# Feedlines -" section EXACTLY as written in template:
      - Line 1 MUST start with: "[Type of CBS] CBS has a total [Feedlines Count] Feedlines with [Conveyor Module Count] conveyor modules."
      - Line 2 MUST be: "Each Feedline has the following Conveyor Modules:"
      - Then list the conveyor modules with placeholders: [Loading Conveyor Count], [Buffer Conveyor Count], [Intelligent Merge Count]
      - Lines 2 and 3 about sensors and transfer MUST be copied word-for-word
      - Replace ONLY the [PLACEHOLDER] values with actual numbers from VARIABLES MAP
      - DO NOT change "Each Feedline has the following Conveyor Modules:" to "The feedlines consist of..."
      - DO NOT paraphrase any part of this section
  
  - After the feedlines paragraph, create numbered sub-sections for components:
    - Orientation / Loading Conveyor -> COPY VERBATIM from "# Orientation Conveyor" in TEMPLATE TEXT
    - Buffer Conveyors -> COPY VERBATIM from "# Buffer Conveyor-" in TEMPLATE TEXT
    - Weighing Conveyor -> COPY VERBATIM from "# Weighing Conveyor-" in TEMPLATE TEXT
    - Intelligent Merge Conveyor -> COPY VERBATIM from "# Intelligent merge conveyor" in TEMPLATE TEXT
  - IMPORTANT: COPY THE ENTIRE MULTI-LINE DESCRIPTION from TEMPLATE TEXT - DO NOT SUMMARIZE INTO ONE LINE

**3. Main Loop** OR **3. Main Linear CBS**
  - If CBS Type is "Loop CBS" => heading must be exactly: **3. Main Loop**
  - If CBS Type is "Linear CBS" => heading must be exactly: **3. Main Linear CBS**
  - CRITICAL: COPY THE EXACT BULLET POINTS from "# Main Loop -" section in TEMPLATE TEXT
  - Replace [CBS HEIGHT FROM GROUND], [PITCH LENGTH] etc. with values from VARIABLES MAP
  - Use - prefix for bullet lines

**4. Output Chutes** (MANDATORY if any chutes exist in DETECTED JSON -> "Output Chutes")
  - This section MUST appear after Main Loop if ANY chute types exist
  - Create sub-sections for each chute type present in DETECTED JSON -> "Output Chutes"
  - Use this preferred order if present (auto-number sequentially):
    - Direct Bagging Chutes -> COPY VERBATIM from "# Direct Bagging Chute -" in TEMPLATE TEXT
    - Gravity Chutes -> COPY VERBATIM from "# Gravity Chute -" in TEMPLATE TEXT
    - Mini-Gravity Chutes -> COPY VERBATIM from "# Mini-Gravity Chute -" in TEMPLATE TEXT
    - Collection Chutes -> COPY VERBATIM from "# Collection Chute -" in TEMPLATE TEXT
    - Dispersion Chutes -> COPY VERBATIM from "# Dispersion Chute -" in TEMPLATE TEXT
    - Rejection Chutes -> COPY VERBATIM from "# Rejection Chute -" in TEMPLATE TEXT
    - Bulk Chutes -> use "# Bulk Chute -" description from TEMPLATE TEXT
    - Sliding Chutes -> COPY VERBATIM from "# Sliding Chute -" in TEMPLATE TEXT
    - Secondary Chutes -> COPY VERBATIM from "#Secondary Chute (L-Type)" in TEMPLATE TEXT
  - IMPORTANT: COPY THE ENTIRE MULTI-LINE DESCRIPTION from TEMPLATE TEXT - DO NOT SUMMARIZE INTO ONE LINE

DESCRIPTION RULE (ABSOLUTELY CRITICAL - READ THIS CAREFULLY)
- For EVERY component and subcomponent, find its description in TEMPLATE TEXT (marked with # Component Name -)
- COPY THE FULL MULTI-LINE DESCRIPTION from TEMPLATE TEXT - NOT A ONE-LINE SUMMARY
- Replace any [PLACEHOLDER] with values from VARIABLES MAP
- DO NOT paraphrase, summarize, or shorten the descriptions
- The description MUST appear in the output text - it should come RIGHT AFTER the heading line and BEFORE the image placeholder
- Structure for each component:
  1. **Heading line** (bold with section number)
  2. **Full description paragraph(s) from template** (THIS IS MANDATORY - COPY ALL LINES, NOT JUST ONE)
  3. **[IMAGE PLACEHOLDER: ...]** at the end

PRESENCE RULE
- Include ONLY items present in DETECTED JSON. Never mention missing items.
- Number sections/subsections SEQUENTIALLY based on what is actually present.

PLACEHOLDERS
- Replace [VAR] using VARIABLES MAP when available; otherwise keep [VAR] unchanged.

IMAGE PLACEHOLDERS (MANDATORY)
- For EACH component/subcomponent, the structure MUST be:
  1. **Heading** (e.g., **2.2 Buffer Conveyors**)
  2. **Full description text from TEMPLATE** (the complete multi-line paragraph - DO NOT SUMMARIZE INTO ONE LINE)
  3. **Image placeholder** at the END: [IMAGE PLACEHOLDER: <heading text without numbering>]

- CORRECT Example for "Buffer Conveyors":
  **2.2 Buffer Conveyors**
  A buffer conveyor, also known as a buffering conveyor or accumulation conveyor, is a type of conveyor system used to temporarily store or hold items in a controlled manner. Its primary purpose is to manage the flow of items between different stages of a production or handling process when there is a mismatch in the speeds or capacities of the upstream and downstream equipment. This Conveyors required to maintain the Throughput of Line.
  [IMAGE PLACEHOLDER: Buffer Conveyors]

- WRONG Example (DO NOT DO THIS):
  **2.2 Buffer Conveyors**
  Temporarily stores parcels to balance flow.
  [IMAGE PLACEHOLDER: Buffer Conveyors]

- CORRECT Example for "Weighing Conveyor":
  **2.3 Weighing Conveyor**
  A weighing conveyor, also known as a weigh belt conveyor, is a type of conveyor system specifically designed to measure the weight of materials as they move along the conveyor belt. It combines the functions of conveying and weighing into a single integrated process. Weighing Conveyors equipped with high precision Load Cells to capture the weight of Parcels.
  [IMAGE PLACEHOLDER: Weighing Conveyor]

- WRONG Example (DO NOT DO THIS):
  **2.3 Weighing Conveyor**
  Measures parcel weight while conveying them to the sorter.
  [IMAGE PLACEHOLDER: Weighing Conveyor]

- CRITICAL: The description text MUST appear BETWEEN the heading and the image placeholder. Never put an image placeholder immediately after a heading without the full multi-line description from the template.

TABLE RULE
- The ONLY table marker allowed is [[CONVEYOR_BOQ_TABLE]] exactly once.

OUTPUT
- Plain text only. No JSON.
- Use **bold** markers for headings.
- Use - bullet prefix for Main Loop content lines.
- REMEMBER: Every component description must be the FULL TEXT from the template, NOT a one-line summary.
"""
    user = f"""TEMPLATE TEXT:
{template_text}

DETECTED JSON:
{json.dumps(detected, ensure_ascii=False)}

VARIABLES:
{json.dumps(variables, ensure_ascii=False)}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def prompt_judge_fix_system_description(detected: Dict[str, Any], draft_text: str) -> List[Dict[str, str]]:
    system = r"""
ROLE
You are a strict QA judge for CBS System Description text.

INPUTS
- DETECTED JSON (truth)
- DRAFT TEXT

YOUR JOB
Rewrite the draft to comply with ALL rules below. If already compliant, return unchanged.

RULES (STRICT)

A) HEADING FORMAT & STYLE
- ALL headings MUST be wrapped in **bold** markers: **1. Title**, **1.1 Subtitle**, **1.1.1 Sub-subtitle**
- Numbers MUST be sequential based on what is ACTUALLY PRESENT (not hardcoded)
- If only 2 items exist under a section, number them 1.1 and 1.2 (NOT 1.1 and 1.3)

B) STRUCTURE (use only present items, auto-number sequentially)

**1. Infeed System**
  **1.1 Infeed Conveyors** (if present)
    Sub-subsections numbered sequentially (1.1.1, 1.1.2, etc.) for:
    - Straight and Inclined conveyor (if present)
    - Plastic Modular conveyor (if present)
    - Curve conveyor (if present)
    - Buffer conveyor (if present)
    - Alligning conveyor (if present)
    - Belt merge (if present)
  **1.x VDS Loop Conveyor** (if present, x = next available number)
  **1.x Conveyor BOQ** (MUST be LAST under Infeed System)
      [[CONVEYOR_BOQ_TABLE]]

**2. Induct**
  CRITICAL FORMAT REQUIREMENTS:
  - MUST start with introductory paragraph about feedline purpose (conditional on VDS)
  - Then MUST include the exact numbered structure from template "# Feedlines -":
    1. MUST start with: "[Type of CBS] CBS has a total [X] Feedlines with [Y] conveyor modules."
       Followed by: "Each Feedline has the following Conveyor Modules:"
       Then bullet list with actual counts: "‚Ä¢ Loading Conveyor- [X] No."
    2. Sensors paragraph (exact text from template about position and dimensions)
    3. Transfer paragraph (exact text from template about spacing and smooth transfer)
  - After these paragraphs, add subsections numbered sequentially for present components:
    - Orientation / Loading Conveyor (if present)
    - Buffer Conveyors (if present)
    - Weighing Conveyor (if present)
    - Intelligent Merge Conveyor (if present)
  
  FORBIDDEN PARAPHRASING:
  - DO NOT change "Each Feedline has the following Conveyor Modules:" to "The feedlines consist of..."
  - DO NOT change "‚Ä¢ Loading Conveyor- [X] No." to generic descriptions
  - MUST preserve the exact bullet format with counts

**3. Main Loop** OR **3. Main Linear CBS** (based on DETECTED JSON -> CBS -> Type)
  CRITICAL: Content MUST be exactly 3 bullet points using "-" prefix:
  - First bullet: Installation level and sorter height
  - Second bullet: Belt carrier pitch and belt size
  - Third bullet: Barcode scanning and chute assignment

**4. Output Chutes** (MANDATORY if any chutes exist in DETECTED JSON)
  Subsections numbered sequentially for present chute types:
    - Direct Bagging Chutes (if present)
    - Gravity Chutes (if present)
    - Mini-Gravity Chutes (if present)
    - Collection Chutes (if present)
    - Dispersion Chutes (if present)
    - Rejection Chutes (if present)
    - Bulk Chutes (if present)
    - Sliding Chutes (if present)
    - Secondary Chutes (if present)

C) PRESENCE FILTER
- Include ONLY items present in DETECTED JSON.

D) NO DUPLICATES
- No subcomponent heading should appear twice.

E) TABLE MARKER
- Only one marker allowed: [[CONVEYOR_BOQ_TABLE]] under Conveyor BOQ only.

F) IMAGE PLACEHOLDERS
- After every component heading AND its full description, add exactly one line:
  [IMAGE PLACEHOLDER: <heading text without numbering>]
- IMPORTANT: There must be a multi-line description paragraph BETWEEN the heading and the image placeholder

G) MAIN LOOP/LINEAR CBS BULLET FORMAT
- The Main Loop or Main Linear CBS section content MUST use exactly 3 bullet points
- Each bullet starts with "- " on its own line
- Do NOT use paragraph format for this section

H) DESCRIPTION COMPLETENESS (CRITICAL)
- Each component MUST have its full description from the template
- Descriptions should be 2-4 sentences minimum, NOT one-line summaries
- If a description looks like a single short sentence summary, it is WRONG
- The description must come AFTER the heading and BEFORE the [IMAGE PLACEHOLDER]

OUTPUT
- Plain text with **bold** markers for headings.
- Use - bullets for Main Loop content.
- No JSON. No explanations.
"""
    user = f"""DETECTED JSON:
{json.dumps(detected, ensure_ascii=False)}

DRAFT TEXT:
{draft_text}
"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]

# -----------------------------
# Enforce Induct/Infeed subcomponents ONLY if DXF evidence exists,
# and classify "Buffer" to avoid wrong duplication.
# -----------------------------
def _classify_buffer_location(metrics: Dict[str, Any], full_dxf: Dict[str, Any]) -> str:
    """
    Returns: 'induct' | 'infeed' | 'none'
    Heuristic:
      - If FS002(without weighing) present => induct buffer
      - Else if explicit infeed-context buffer signals exist => infeed buffer
      - Else if feedlines exist and generic 'buffer' exists => induct buffer
      - Else if generic 'buffer' exists and no feedlines => infeed buffer
      - Else none
    """
    fs002 = int(metrics.get("FS002 WITHOUT WEIGHING COUNT", 0) or 0)
    feedline_cnt = int(metrics.get("FEEDLINE COUNT", 0) or 0)

    has_generic_buffer = _search_counts_multi(full_dxf, [r"\bbuffer\b", r"accumulation"]) > 0
    has_infeed_context_buffer = _search_counts_multi(full_dxf, [
        r"infeed.*buffer", r"buffer.*infeed",
        r"receiving.*buffer", r"buffer.*receiving",
        r"highway.*buffer", r"buffer.*highway",
        r"telescopic.*buffer", r"buffer.*telescopic"
    ]) > 0

    if fs002 > 0:
        return "induct"

    if has_infeed_context_buffer:
        return "infeed"

    if has_generic_buffer and feedline_cnt > 0:
        return "induct"

    if has_generic_buffer and feedline_cnt == 0:
        return "infeed"

    return "none"


def enforce_induct_subcomponents_from_dxf(detected: Dict[str, Any],
                                         metrics: Dict[str, Any],
                                         full_dxf: Dict[str, Any]) -> Dict[str, Any]:
    """
    Induct subcomponents are placed ONLY under:
    detected["Parcel Inducts / Induction to Sorter"]["Feedlines"]["Subcomponents"]

    Required output names (for LLM structure):
      - Orientation / Loading Conveyor
      - Buffer Conveyors
      - Weighing Conveyor
      - Intelligent Merge Conveyor
    """
    d = detected if isinstance(detected, dict) else {}

    pi = d.get("Parcel Inducts / Induction to Sorter")
    if not isinstance(pi, dict):
        pi = {}
        d["Parcel Inducts / Induction to Sorter"] = pi

    feedlines = pi.get("Feedlines")
    if not isinstance(feedlines, dict):
        feedlines = {}
        pi["Feedlines"] = feedlines

    subs = feedlines.get("Subcomponents")
    if not isinstance(subs, dict):
        subs = {}
        feedlines["Subcomponents"] = subs

    # Strict evidence
    has_orientation = _search_counts_multi(full_dxf, [r"\borientation\b", r"\borient\b", r"orientat", r"oriant"]) > 0
    has_loading = _search_counts_multi(full_dxf, [r"\bloading\b", r"\bspacing\b", r"\bspacer\b", r"gap\s*optimizer"]) > 0
    has_orient_or_loading = has_orientation or has_loading

    buffer_loc = _classify_buffer_location(metrics, full_dxf)
    has_buffer_induct = (buffer_loc == "induct")

    has_weighing = _search_counts_multi(full_dxf, [r"\bweigh\b", r"weight", r"scale"]) > 0

    has_intelligent_merge = (
        _search_counts_multi(full_dxf, [
            r"intelligent\s*merge", r"angle\s*merge", r"fal[_\-\s]*f012", r"\bf012\b"
        ]) > 0
    ) or bool(metrics.get("DEGREE OF ANGLE MERGE"))

    # Apply: add only if present, remove if not
    if has_orient_or_loading:
        subs["Orientation / Loading Conveyor"] = 1
    else:
        subs.pop("Orientation / Loading Conveyor", None)

    if has_buffer_induct:
        subs["Buffer Conveyors"] = 1
    else:
        subs.pop("Buffer Conveyors", None)

    if has_weighing:
        subs["Weighing Conveyor"] = 1
    else:
        subs.pop("Weighing Conveyor", None)

    if has_intelligent_merge:
        subs["Intelligent Merge Conveyor"] = 1
        subs.pop("Merge Conveyor", None)
        subs.pop("Intelligent / Angle Merge Conveyor", None)
    else:
        subs.pop("Intelligent Merge Conveyor", None)
        subs.pop("Merge Conveyor", None)
        subs.pop("Intelligent / Angle Merge Conveyor", None)

    # Ensure these names do NOT appear under Infeed System
    infeed = d.get("Infeed System")
    if isinstance(infeed, dict):
        for k in ["Orientation / Loading Conveyor", "Buffer Conveyors", "Weighing Conveyor", "Intelligent Merge Conveyor"]:
            infeed.pop(k, None)

    return d


def enforce_infeed_subcomponents_from_dxf(detected: Dict[str, Any],
                                         metrics: Dict[str, Any],
                                         full_dxf: Dict[str, Any]) -> Dict[str, Any]:
    """
    Adds infeed conveyor subcomponents as keys under detected["Infeed System"] ONLY if DXF evidence exists.

    Required names (for LLM structure):
      - Straight and Inclined conveyor
      - Plastic Modular conveyor
      - Curve conveyor
      - Buffer conveyor
      - Alligning conveyor
      - Belt merge

    Buffer conveyor is classified to avoid putting it into infeed if it belongs to induct.
    """
    d = detected if isinstance(detected, dict) else {}

    infeed = d.get("Infeed System")
    if not isinstance(infeed, dict):
        infeed = {}
        d["Infeed System"] = infeed

    # Base infeed conveyors presence (keep existing if any)
    if "Infeed Conveyors" not in infeed:
        # If any infeed-ish signals exist, mark present
        if int(metrics.get("TEL. BELT CONVEYOR COUNT", 0) or 0) > 0 or int(metrics.get("INFEED CONVEYOR COUNT", 0) or 0) > 0:
            infeed["Infeed Conveyors"] = max(1, int(metrics.get("INFEED CONVEYOR COUNT", 1) or 1))

    # Subcomponents
    has_straight_inclined = _search_counts_multi(full_dxf, [
        r"straight.*conveyor", r"\bstraight\b.*\bconv\b",
        r"\binclined\b", r"\bincline\b", r"\bincl\b", r"\buphill\b", r"\bdownhill\b", r"\bslope\b"
    ]) > 0

    has_plastic_modular = _search_counts_multi(full_dxf, [
        r"plastic\s*modular", r"\bmodular\b.*\bplastic\b", r"\bplastic\b.*\bmodular\b", r"\bmodular\b.*\bconv\b"
    ]) > 0

    has_curve = _search_counts_multi(full_dxf, [
        r"\bcurve\b", r"\bcurved\b", r"\bbend\b", r"\bradius\b", r"curve.*conveyor", r"curved.*conveyor"
    ]) > 0

    buffer_loc = _classify_buffer_location(metrics, full_dxf)
    has_buffer_infeed = (buffer_loc == "infeed")

    has_aligning = _search_counts_multi(full_dxf, [
        r"aligning\s*conveyor", r"align.*conveyor", r"\ballign\s*conv\b", r"alligning"
    ]) > 0

    has_belt_merge = _search_counts_multi(full_dxf, [
        r"belt\s*merge", r"merge\s*belt"
    ]) > 0

    if has_straight_inclined:
        infeed["Straight and Inclined conveyor"] = 1
    else:
        infeed.pop("Straight and Inclined conveyor", None)

    if has_plastic_modular:
        infeed["Plastic Modular conveyor"] = 1
    else:
        infeed.pop("Plastic Modular conveyor", None)

    if has_curve:
        infeed["Curve conveyor"] = 1
    else:
        infeed.pop("Curve conveyor", None)

    if has_buffer_infeed:
        infeed["Buffer conveyor"] = 1
    else:
        infeed.pop("Buffer conveyor", None)

    if has_aligning:
        infeed["Alligning conveyor"] = 1
    else:
        infeed.pop("Alligning conveyor", None)

    if has_belt_merge:
        infeed["Belt merge"] = 1
    else:
        infeed.pop("Belt merge", None)

    return d

# -----------------------------
# VARIABLES MAP
# -----------------------------
def build_variables_map(metrics: Dict[str, Any], costing_values: Dict[str, str]) -> Dict[str, str]:
    vars_map: Dict[str, str] = {}
    
    # Copy all costing values
    for k, v in (costing_values or {}).items():
        if v is None or str(v).strip() == "":
            continue
        vars_map[k] = str(v)

    # Copy relevant metrics
    metric_keys = [
        "UNITS", "TYPE OF CBS", "PITCH LENGTH", "CBS HEIGHT FROM GROUND", 
        "DEGREE OF ANGLE MERGE", "FEEDLINE COUNT", "SORTER SPEED", "BELT SPEED",
        "THROUGHPUT", "TOTAL CHUTES"
    ]
    for k in metric_keys:
        if k in metrics and metrics[k] is not None and str(metrics[k]).strip() != "":
            vars_map[k] = str(metrics[k])
    
    # ==================== ENHANCED INDUCT SUBCOMPONENT ANALYSIS ====================
    # Analyze actual DXF components to determine induct structure
    feedline_count = int(metrics.get("FEEDLINE COUNT", 0) or 0)
    
    if feedline_count > 0:
        # Get actual component counts from metrics
        loading_count = 0
        buffer_count = 0
        merge_count = 0
        weighing_count = 0
        spacing_count = 0
        positioning_count = 0
        
        # Extract from metrics (these should be populated from DXF analysis)
        for key, value in metrics.items():
            key_lower = key.lower()
            
            # Loading/Receiving conveyors
            if any(term in key_lower for term in ["receiving", "loading", "infeed conveyor"]):
                loading_count += int(value or 0)
            
            # Buffer conveyors
            elif "buffer" in key_lower and "conveyor" in key_lower:
                buffer_count += int(value or 0)
            
            # Intelligent merge / Angle merge
            elif any(term in key_lower for term in ["intelligent merge", "angle merge", "merge"]):
                if "conveyor" in key_lower or "merge" in key_lower:
                    # Skip if value is a comma-separated string (e.g. "30, 60" for degrees)
                    if isinstance(value, str) and "," in str(value):
                        continue
                    try:
                        merge_count += int(value or 0)
                    except (ValueError, TypeError):
                        pass  # Skip non-numeric values
            
            # Weighing conveyors
            elif "weighing" in key_lower and "conveyor" in key_lower:
                weighing_count += int(value or 0)
            
            # Spacing/Positioning conveyors
            elif "spacing" in key_lower or "positioning" in key_lower:
                if "conveyor" in key_lower:
                    spacing_count += int(value or 0)
                    positioning_count += int(value or 0)
        
        # Calculate total conveyor modules (all subcomponents)
        total_modules = loading_count + buffer_count + merge_count + weighing_count + spacing_count
        
        # If no specific counts found, use default of 3 modules per feedline
        if total_modules == 0:
            # Standard configuration: 3 basic modules per feedline
            total_modules = feedline_count * 3
            loading_count = feedline_count  # 1 per feedline
            buffer_count = feedline_count   # 1 per feedline
            merge_count = feedline_count    # 1 per feedline
        
        # Build the induct module structure text
        vars_map["Conveyor Module Count"] = str(total_modules)
        vars_map["Feedlines Count"] = str(feedline_count)
        vars_map["Loading Conveyor Count"] = str(loading_count)
        vars_map["Buffer Conveyor Count"] = str(buffer_count)
        vars_map["Intelligent Merge Count"] = str(merge_count)
        
        if weighing_count > 0:
            vars_map["Weighing Conveyor Count"] = str(weighing_count)
        if spacing_count > 0:
            vars_map["Spacing Conveyor Count"] = str(spacing_count)
        if positioning_count > 0:
            vars_map["Positioning Conveyor Count"] = str(positioning_count)

    return vars_map

# -----------------------------
# DOCX WRITER
# -----------------------------
def add_docx_table(doc: Document, table_data: List[List[str]]):
    if not table_data or len(table_data) < 2:
        return
    rows = len(table_data)
    cols = len(table_data[0])
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    for r in range(rows):
        for c in range(cols):
            t.cell(r, c).text = table_data[r][c]
    doc.add_paragraph("")


def _try_add_fixed_image(doc: Document, placeholder_heading: str) -> bool:
    h = (placeholder_heading or "").strip().lower()

    # direct match
    if h in FIXED_IMAGE_MAP:
        img = FIXED_IMAGE_MAP[h]
        if img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
        return False

    # tolerant matching for induct subcomponents
    if "weigh" in h:
        img = FIXED_IMAGE_MAP.get("weighing conveyor")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    if "buffer" in h:
        img = FIXED_IMAGE_MAP.get("buffer conveyors") or FIXED_IMAGE_MAP.get("buffer conveyor")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    if "orientation" in h or "loading" in h or "oriant" in h:
        img = FIXED_IMAGE_MAP.get("orientation / loading conveyor") or FIXED_IMAGE_MAP.get("orientation conveyor")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    if "merge" in h:
        img = FIXED_IMAGE_MAP.get("intelligent merge conveyor") or FIXED_IMAGE_MAP.get("merge conveyor")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    
    # Main Loop / CBS matching
    if "main loop" in h or "main linear" in h or "cross belt" in h:
        img = FIXED_IMAGE_MAP.get("main loop")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    
    # Carrier/Drive matching
    if "carrier" in h:
        img = FIXED_IMAGE_MAP.get("carrier")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    if "friction" in h and "drive" in h:
        img = FIXED_IMAGE_MAP.get("friction wheel drive")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    if "linear motor" in h:
        img = FIXED_IMAGE_MAP.get("linear motor drive")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    if "servo" in h:
        img = FIXED_IMAGE_MAP.get("servo roller")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    
    # Infeed conveyor matching
    if "infeed" in h or "straight" in h or "inclined" in h:
        img = FIXED_IMAGE_MAP.get("infeed conveyor")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True
    if "curve" in h:
        img = FIXED_IMAGE_MAP.get("curve conveyor")
        if img and img.exists():
            doc.add_picture(str(img), width=Inches(5.8))
            doc.add_paragraph("")
            return True

    return False


def build_docx(system_description_text: str,
               out_path: Path,
               title: str = "System Description",
               tables: Optional[Dict[str, List[List[str]]]] = None,
               detected: Optional[Dict[str, Any]] = None) -> None:

    tables = tables or {}
    detected = detected or {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("")

    conveyor_boq = tables.get("Conveyor BOQ", [])

    ph_rx = re.compile(r"^\s*\[IMAGE PLACEHOLDER:\s*(.+?)\s*\]\s*$", flags=re.IGNORECASE)
    # Regex to detect **bold** markers in text
    bold_rx = re.compile(r"\*\*(.+?)\*\*")
    # Regex to detect heading lines (numbered sections like "**1. Title**" or "**1.1 Title**")
    heading_rx = re.compile(r"^\s*\*\*\s*(\d+(?:\.\d+)*\.?)\s*(.+?)\s*\*\*\s*$")
    # Regex for bullet points (-, *, ‚Ä¢)
    bullet_rx = re.compile(r"^\s*[-*‚Ä¢]\s+(.+)$")

    for line in system_description_text.splitlines():
        stripped = line.strip()
        
        # Handle table marker
        if stripped == "[[CONVEYOR_BOQ_TABLE]]":
            add_docx_table(doc, conveyor_boq)
            continue

        # Handle image placeholders
        m = ph_rx.match(stripped)
        if m:
            heading = m.group(1).strip()
            # Replace placeholder with fixed images for induct subcomponents
            if _try_add_fixed_image(doc, heading):
                continue
            # else skip placeholder line (don't show in final doc)
            continue

        # Handle empty lines
        if not stripped:
            doc.add_paragraph("")
            continue

        # Handle heading lines with **bold** markers
        heading_match = heading_rx.match(stripped)
        if heading_match:
            number = heading_match.group(1).strip()
            heading_text = heading_match.group(2).strip()
            p = doc.add_paragraph()
            run = p.add_run(f"{number} {heading_text}")
            run.bold = True
            run.font.size = Pt(12)
            continue

        # Handle bullet points
        bullet_match = bullet_rx.match(stripped)
        if bullet_match:
            bullet_text = bullet_match.group(1).strip()
            p = doc.add_paragraph(style='List Bullet')
            # Check for any **bold** text within the bullet
            if "**" in bullet_text:
                parts = bold_rx.split(bullet_text)
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = p.add_run(part)
                    if i % 2 == 1:  # Bold parts are at odd indices after split
                        run.bold = True
            else:
                p.add_run(bullet_text)
            continue

        # Handle regular text with possible **bold** markers
        if "**" in stripped:
            p = doc.add_paragraph()
            parts = bold_rx.split(stripped)
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                if i % 2 == 1:  # Bold parts are at odd indices after split
                    run.bold = True
        else:
            # Regular paragraph
            doc.add_paragraph(line)

    doc.save(str(out_path))

# -----------------------------
# Generation pipeline
# -----------------------------
def generate_system_description(
    dxf_path: Path,
    xlsx_path: Optional[Path] = None,
    template_override_text: Optional[str] = None,
    title: str = "System Description",
    temperature_detect: float = 0.1,
    temperature_write: float = 0.1,
) -> Tuple[str, Dict[str, Any], Dict[str, Any], Dict[str, List[List[str]]]]:
    # New dynamic pipeline: no hallucinated components
    catalog = load_catalog()

    full_dxf = extract_dxf_full_json(dxf_path)
    metrics = compute_dxf_metrics(full_dxf)

    # Build DXF components list (raw names only from DXF)
    dxf_components = build_dxf_components(full_dxf, catalog)

    # Load Excel registry (normalized sheet names ‚Üí table/key_values)
    excel_registry: Dict[str, Dict[str, Any]] = {}
    if xlsx_path is not None and xlsx_path.exists():
        excel_registry = load_component_sheets(str(xlsx_path))

    # Generate dynamic system description
    use_polish = temperature_write > 0.5
    final_txt, diagnostics = generate_dynamic_system_description(
        dxf_components=dxf_components,
        excel_registry=excel_registry,
        catalog=catalog,
        use_lm_polish=use_polish,
        groq_api_key=GROQ_API_KEY if use_polish else None,
    )

    # Build tables map for matched components (only one per component)
    tables: Dict[str, List[List[str]]] = {}
    for reg_name, reg_entry in excel_registry.items():
        table = reg_entry.get("table")
        if table is not None:
            tables[reg_entry.get("sheet_name", reg_name)] = table

    # Mandatory diagnostics (rendered, missing mappings)
    diagnostics.setdefault("dxf_components_count", len(dxf_components))
    diagnostics.setdefault("rendered_components_count", diagnostics.get("total_components", 0))

    return final_txt, diagnostics, metrics, tables
